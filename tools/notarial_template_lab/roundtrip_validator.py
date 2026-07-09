from __future__ import annotations

import sys
from pathlib import Path

from docx import Document

from tools.notarial_template_lab.models import ValidationResult


class RoundtripValidator:
    def validate(self, draft_docx: str | Path) -> ValidationResult:
        path = Path(draft_docx)
        errors: list[str] = []
        opens = False
        contains_markers = False
        fields: list[dict] = []

        try:
            doc = Document(str(path))
            opens = True
            contains_markers = "{{" in "\n".join(iter_all_text(doc))
        except Exception as exc:
            errors.append(f"No fue posible abrir el DOCX experimental: {exc}")
            doc = None

        if opens:
            try:
                detect_marked_template = self._load_marked_template_detector()
                detected = detect_marked_template(path)
                fields = detected.get("fields", []) if isinstance(detected, dict) else []
            except Exception as exc:
                errors.append(f"marked_template_detector fallo: {exc}")

        passed = opens and contains_markers and len(fields) > 0 and not errors
        if opens and not contains_markers:
            errors.append("El DOCX experimental no contiene marcadores {{...}}.")
        if opens and contains_markers and not fields:
            errors.append("marked_template_detector no devolvio campos detectados.")

        return ValidationResult(
            passed=passed,
            output_path=str(path),
            opens=opens,
            contains_markers=contains_markers,
            marked_fields_count=len(fields),
            marked_fields=fields,
            errors=errors,
        )

    def _load_marked_template_detector(self):
        repo_root = Path(__file__).resolve().parents[2]
        backend_path = repo_root / "backend"
        if str(backend_path) not in sys.path:
            sys.path.insert(0, str(backend_path))
        from app.services.minuta.marked_template_detector import detect_marked_template

        return detect_marked_template


def iter_all_text(doc):
    for paragraph in doc.paragraphs:
        yield paragraph.text or ""
    for table in doc.tables:
        yield from iter_table_text(table)
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            yield paragraph.text or ""
        for table in section.header.tables:
            yield from iter_table_text(table)
        for paragraph in section.footer.paragraphs:
            yield paragraph.text or ""
        for table in section.footer.tables:
            yield from iter_table_text(table)


def iter_table_text(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                yield paragraph.text or ""
            for nested in cell.tables:
                yield from iter_table_text(nested)
