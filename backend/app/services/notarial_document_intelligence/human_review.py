from __future__ import annotations

import hashlib
import json
import re
import tempfile
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import RGBColor
from sqlalchemy import func
from sqlalchemy.orm import Session

from app.models.legal_entity import LegalEntity
from app.models.notarial_document_intelligence import (
    NotarialDocument,
    NotarialDocumentBlock,
    NotarialDocumentDecision,
    NotarialDocumentEntity,
    NotarialDocumentFamily,
    NotarialHumanFieldReview,
    NotarialHumanReviewAudit,
    NotarialHumanReviewSession,
    NotarialTemplateLibraryItem,
    NotarialTemplateVersion,
)
from app.services.minuta.inverse_conversion_catalog.models import FieldDefinition


VALID_ACTIONS = {"accept", "reject", "correct", "propose_new"}
VALID_SCOPES = {"single", "all"}
VALID_FIXED_LABELS = {"fixed", "variable", "optional", "unknown"}
RED = RGBColor(0xFF, 0x00, 0x00)


@dataclass(frozen=True)
class GeneratedTemplateDocx:
    path: Path
    filename: str
    content_type: str = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


class NotarialHumanReviewService:
    def __init__(self, db: Session, *, notary_id: int, actor_user_id: int | None = None) -> None:
        self.db = db
        self.notary_id = notary_id
        self.actor_user_id = actor_user_id

    def create_or_get_session(self, decision_id: int, *, idempotency_key: str | None = None) -> NotarialHumanReviewSession:
        decision = self._decision(decision_id)
        existing = (
            self.db.query(NotarialHumanReviewSession)
            .filter(
                NotarialHumanReviewSession.decision_id == decision.id,
                NotarialHumanReviewSession.notary_id == self.notary_id,
                NotarialHumanReviewSession.status.in_(["open", "approved"]),
            )
            .order_by(NotarialHumanReviewSession.version.desc())
            .first()
        )
        if existing is not None:
            self._audit("review_session", existing.id, "create_idempotent", idempotency_key, {"decision_id": decision_id})
            self.db.commit()
            return existing

        version = (
            self.db.query(func.coalesce(func.max(NotarialHumanReviewSession.version), 0))
            .filter(NotarialHumanReviewSession.decision_id == decision.id)
            .scalar()
            or 0
        ) + 1
        session = NotarialHumanReviewSession(
            notary_id=self.notary_id,
            decision_id=decision.id,
            run_id=_json_object(decision.metadata_json).get("intelligence_run_id"),
            document_id=decision.document_id,
            parse_run_id=decision.parse_run_id,
            version=version,
            status="open",
            reviewer_user_id=self.actor_user_id,
            locked_by_user_id=self.actor_user_id,
            lock_token=_lock_token(decision.id, version, self.actor_user_id),
            metadata_json=json.dumps({"source": "hybrid_decision", "idempotency_key": idempotency_key}, ensure_ascii=False, sort_keys=True),
        )
        self.db.add(session)
        self.db.flush()
        self._seed_field_reviews(session, decision)
        self._audit("review_session", session.id, "created", idempotency_key, {"decision_id": decision_id, "version": version})
        self.db.commit()
        return session

    def detail(self, session_id: int) -> dict[str, Any]:
        session = self._session(session_id)
        decision = self._decision(session.decision_id)
        fields = self._field_rows(session.id)
        return {
            "session": _session_payload(session),
            "decision": {
                "id": decision.id,
                "status": decision.status,
                "document_id": decision.document_id,
                "parse_run_id": decision.parse_run_id,
                "hybrid": _json_object(decision.hybrid_decision_json),
                "human": _json_object(decision.human_decision_json),
            },
            "fields": [_field_payload(row) for row in fields],
            "visual_review": self._visual_review(session, fields),
        }

    def apply_field_decision(
        self,
        session_id: int,
        field_review_id: int,
        *,
        action: str,
        proposed_value: str | None = None,
        proposed_field_code: str | None = None,
        apply_scope: str = "single",
        fixed_variable_label: str | None = None,
        idempotency_key: str | None = None,
        reason: str | None = None,
    ) -> list[NotarialHumanFieldReview]:
        session = self._session(session_id)
        if session.status != "open":
            raise ValueError("La sesion de revision no esta abierta.")
        if action not in VALID_ACTIONS:
            raise ValueError(f"Accion de revision invalida: {action}.")
        if apply_scope not in VALID_SCOPES:
            raise ValueError(f"Alcance invalido: {apply_scope}.")
        if fixed_variable_label is not None and fixed_variable_label not in VALID_FIXED_LABELS:
            raise ValueError(f"Etiqueta fija/variable invalida: {fixed_variable_label}.")
        if self._audit_seen(idempotency_key):
            return self._field_rows(session.id)

        row = (
            self.db.query(NotarialHumanFieldReview)
            .filter(NotarialHumanFieldReview.id == field_review_id, NotarialHumanFieldReview.session_id == session.id)
            .first()
        )
        if row is None:
            raise ValueError("Campo de revision no encontrado.")
        if action == "propose_new":
            normalized_code = _normalize_code(proposed_field_code)
            proposal = NotarialHumanFieldReview(
                session_id=session.id,
                notary_id=self.notary_id,
                decision_id=session.decision_id,
                document_id=session.document_id,
                parse_run_id=session.parse_run_id,
                block_id=row.block_id,
                entity_id=None,
                occurrence_key="proposal:" + hashlib.sha1(f"{session.id}:{field_review_id}:{normalized_code}:{proposed_value}".encode("utf-8")).hexdigest()[:24],
                field_code=None,
                proposed_field_code=normalized_code,
                original_value=None,
                proposed_value=proposed_value,
                action="propose_new",
                apply_scope="single",
                fixed_variable_label=fixed_variable_label or "unknown",
                is_new_field_proposal=True,
                status="proposed_catalog_review",
                decided_by_user_id=self.actor_user_id,
                decided_at=_now(),
                metadata_json=json.dumps(
                    {
                        "reason": reason,
                        "based_on_field_review_id": field_review_id,
                        "known_field_code": self._known_field_code(normalized_code) if normalized_code else False,
                        "no_catalog_autoapproval": True,
                    },
                    ensure_ascii=False,
                    sort_keys=True,
                ),
            )
            self.db.add(proposal)
            self.db.flush()
            self._sync_decision_human_payload(session)
            self._audit("field_review", proposal.id, "new_field_proposed", idempotency_key, {"based_on_field_review_id": field_review_id})
            self.db.commit()
            return [proposal]
        targets = [row]
        if apply_scope == "all":
            targets = (
                self.db.query(NotarialHumanFieldReview)
                .filter(
                    NotarialHumanFieldReview.session_id == session.id,
                    NotarialHumanFieldReview.field_code == row.field_code,
                    NotarialHumanFieldReview.original_value == row.original_value,
                )
                .all()
            )

        normalized_code = _normalize_code(proposed_field_code or row.proposed_field_code or row.field_code)
        is_known_code = self._known_field_code(normalized_code) if normalized_code else False
        now = _now()
        for target in targets:
            target.action = action
            target.apply_scope = apply_scope
            target.proposed_value = proposed_value if proposed_value is not None else target.proposed_value
            target.proposed_field_code = normalized_code
            target.fixed_variable_label = fixed_variable_label or target.fixed_variable_label
            target.decided_by_user_id = self.actor_user_id
            target.decided_at = now
            if action == "reject":
                target.status = "rejected"
            elif action == "propose_new" or (normalized_code and not is_known_code):
                target.status = "proposed_catalog_review"
                target.is_new_field_proposal = True
            else:
                target.status = "approved"
                self._apply_supervised_label(target)
            target.metadata_json = _merge_json(
                target.metadata_json,
                {"reason": reason, "known_field_code": is_known_code, "no_catalog_autoapproval": target.is_new_field_proposal},
            )
        self._sync_decision_human_payload(session)
        self._audit(
            "field_review",
            field_review_id,
            "field_decision_applied",
            idempotency_key,
            {"action": action, "apply_scope": apply_scope, "targets": [item.id for item in targets]},
        )
        self.db.commit()
        return targets

    def approve_session(
        self,
        session_id: int,
        *,
        template_name: str,
        template_kind: str = "individual",
        idempotency_key: str | None = None,
    ) -> dict[str, Any]:
        session = self._session(session_id)
        if self._audit_seen(idempotency_key):
            item = self._library_item_for_session(session)
            return {"session": _session_payload(session), "library_item": _library_payload(item), "version": _version_payload(self._latest_version(item))}
        if session.status != "open":
            raise ValueError("La sesion de revision no esta abierta.")
        fields = self._field_rows(session.id)
        if any(field.status == "pending" for field in fields):
            raise ValueError("No se puede aprobar una sesion con campos pendientes.")
        session.status = "approved"
        session.approved_at = _now()
        session.reviewer_user_id = self.actor_user_id
        decision = self._decision(session.decision_id)
        decision.status = "human_approved"
        self._sync_decision_human_payload(session)
        item = self._ensure_library_item(session, decision, template_name=template_name, template_kind=template_kind)
        version = self._create_template_version(item, session, decision, status="approved")
        item.latest_version_id = version.id
        item.approved_version_id = version.id
        item.status = "approved"
        version.approved_by_user_id = self.actor_user_id
        version.approved_at = _now()
        self._audit("review_session", session.id, "approved", idempotency_key, {"library_item_id": item.id, "version_id": version.id})
        self.db.commit()
        return {"session": _session_payload(session), "library_item": _library_payload(item), "version": _version_payload(version)}

    def list_library(self, *, act_code: str | None = None, bank_name: str | None = None, project_name: str | None = None, legal_entity_id: int | None = None) -> list[dict[str, Any]]:
        query = self.db.query(NotarialTemplateLibraryItem).filter(NotarialTemplateLibraryItem.notary_id == self.notary_id)
        if act_code:
            query = query.filter(NotarialTemplateLibraryItem.act_code == act_code)
        if bank_name:
            query = query.filter(NotarialTemplateLibraryItem.bank_name == bank_name)
        if project_name:
            query = query.filter(NotarialTemplateLibraryItem.project_name == project_name)
        if legal_entity_id is not None:
            query = query.filter(NotarialTemplateLibraryItem.legal_entity_id == legal_entity_id)
        return [_library_payload(item) for item in query.order_by(NotarialTemplateLibraryItem.updated_at.desc()).all()]

    def link_bank(self, library_item_id: int, legal_entity_id: int | None) -> dict[str, Any]:
        """Enlaza (o desvincula) un modelo de minuta aprendido con un banco del
        registro de entidades. legal_entity_id=None quita el vinculo."""
        item = (
            self.db.query(NotarialTemplateLibraryItem)
            .filter(
                NotarialTemplateLibraryItem.notary_id == self.notary_id,
                NotarialTemplateLibraryItem.id == library_item_id,
            )
            .first()
        )
        if item is None:
            raise ValueError(f"Modelo de minuta {library_item_id} no existe para esta notaria")
        if legal_entity_id is not None:
            entity = self.db.query(LegalEntity).filter(LegalEntity.id == legal_entity_id).first()
            if entity is None:
                raise ValueError(f"Entidad (banco) {legal_entity_id} no existe")
        item.legal_entity_id = legal_entity_id
        self.db.commit()
        self.db.refresh(item)
        return _library_payload(item)

    def rollback_template(self, library_item_id: int, target_version_id: int, *, idempotency_key: str | None = None) -> dict[str, Any]:
        item = self._library_item(library_item_id)
        target = self._version(target_version_id, item.id)
        if self._audit_seen(idempotency_key):
            return {"library_item": _library_payload(item), "version": _version_payload(self._latest_version(item))}
        version = NotarialTemplateVersion(
            library_item_id=item.id,
            notary_id=self.notary_id,
            version_number=self._next_template_version(item.id),
            status="approved",
            source_decision_id=target.source_decision_id,
            review_session_id=target.review_session_id,
            source_document_id=target.source_document_id,
            content_json=target.content_json,
            placeholder_map_json=target.placeholder_map_json,
            provenance_json=_merge_json(target.provenance_json, {"rollback_from_version_id": target.id}),
            storage_path=target.storage_path,
            created_by_user_id=self.actor_user_id,
            approved_by_user_id=self.actor_user_id,
            approved_at=_now(),
            rollback_of_version_id=target.id,
            metadata_json=json.dumps({"rollback": True}, ensure_ascii=False, sort_keys=True),
        )
        self.db.add(version)
        self.db.flush()
        item.latest_version_id = version.id
        item.approved_version_id = version.id
        item.status = "approved"
        self._audit("template_library_item", item.id, "rolled_back", idempotency_key, {"target_version_id": target.id, "new_version_id": version.id})
        self.db.commit()
        return {"library_item": _library_payload(item), "version": _version_payload(version)}

    def generate_template_docx(self, version_id: int) -> GeneratedTemplateDocx:
        version = (
            self.db.query(NotarialTemplateVersion)
            .filter(NotarialTemplateVersion.id == version_id, NotarialTemplateVersion.notary_id == self.notary_id)
            .first()
        )
        if version is None:
            raise ValueError("Version de plantilla no encontrada.")
        content = _json_object(version.content_json)
        paragraphs = content.get("paragraphs") if isinstance(content.get("paragraphs"), list) else []
        placeholders = _json_object(version.placeholder_map_json)
        doc = Document()
        for item in paragraphs:
            text = str(item.get("text") if isinstance(item, dict) else item)
            paragraph = doc.add_paragraph()
            remaining = text
            for marker in sorted(placeholders, key=len, reverse=True):
                if marker not in remaining:
                    continue
                before, remaining = remaining.split(marker, 1)
                if before:
                    paragraph.add_run(before)
                run = paragraph.add_run(marker)
                run.font.color.rgb = RED
            if remaining:
                paragraph.add_run(remaining)
        tmp = tempfile.NamedTemporaryFile(prefix="easypro2-template-", suffix=".docx", delete=False)
        tmp.close()
        path = Path(tmp.name)
        doc.save(str(path))
        return GeneratedTemplateDocx(path=path, filename=f"plantilla-{version.id}.docx")

    def _seed_field_reviews(self, session: NotarialHumanReviewSession, decision: NotarialDocumentDecision) -> None:
        entities = (
            self.db.query(NotarialDocumentEntity)
            .filter(
                NotarialDocumentEntity.notary_id == self.notary_id,
                NotarialDocumentEntity.document_id == decision.document_id,
                NotarialDocumentEntity.parse_run_id == decision.parse_run_id,
            )
            .order_by(NotarialDocumentEntity.id.asc())
            .all()
        )
        for entity in entities:
            self._add_field_review(
                session,
                decision,
                occurrence_key=f"entity:{entity.id}",
                entity_id=entity.id,
                block_id=entity.block_id,
                field_code=entity.canonical_field_code,
                original_value=entity.text,
                confidence=entity.confidence,
            )

        llm = _json_object(decision.llm_decision_json)
        for index, field in enumerate(llm.get("fields") or []):
            if not isinstance(field, dict):
                continue
            evidence_ids = field.get("evidence_block_ids") if isinstance(field.get("evidence_block_ids"), list) else []
            self._add_field_review(
                session,
                decision,
                occurrence_key=f"llm:{index}:{field.get('field_code')}",
                block_id=int(evidence_ids[0]) if evidence_ids else None,
                field_code=field.get("field_code"),
                original_value=field.get("value"),
                confidence=float(field.get("confidence") or 0.0),
                metadata={"source": "llm", "raw": field},
            )

    def _add_field_review(
        self,
        session: NotarialHumanReviewSession,
        decision: NotarialDocumentDecision,
        *,
        occurrence_key: str,
        entity_id: int | None = None,
        block_id: int | None = None,
        field_code: str | None,
        original_value: str | None,
        confidence: float,
        metadata: dict[str, Any] | None = None,
    ) -> None:
        code = _normalize_code(field_code)
        self.db.add(
            NotarialHumanFieldReview(
                session_id=session.id,
                notary_id=self.notary_id,
                decision_id=decision.id,
                document_id=decision.document_id,
                parse_run_id=decision.parse_run_id,
                block_id=block_id,
                entity_id=entity_id,
                occurrence_key=occurrence_key,
                field_code=code,
                proposed_field_code=code,
                original_value=original_value,
                proposed_value=original_value,
                action="pending",
                fixed_variable_label="unknown",
                status="pending",
                metadata_json=json.dumps({"confidence": confidence, **(metadata or {})}, ensure_ascii=False, sort_keys=True),
            )
        )

    def _apply_supervised_label(self, field: NotarialHumanFieldReview) -> None:
        if field.block_id is None:
            return
        block = self.db.get(NotarialDocumentBlock, field.block_id)
        if block is not None and block.notary_id == self.notary_id:
            block.fixed_variable_label = field.fixed_variable_label

    def _sync_decision_human_payload(self, session: NotarialHumanReviewSession) -> None:
        decision = self._decision(session.decision_id)
        fields = self._field_rows(session.id)
        decision.human_decision_json = json.dumps(
            {
                "review_session_id": session.id,
                "review_status": session.status,
                "fields": [_field_payload(field) for field in fields],
                "catalog_autoapproved": False,
                "corpus_autoapproved": False,
            },
            ensure_ascii=False,
            sort_keys=True,
        )

    def _ensure_library_item(self, session: NotarialHumanReviewSession, decision: NotarialDocumentDecision, *, template_name: str, template_kind: str) -> NotarialTemplateLibraryItem:
        document = self.db.get(NotarialDocument, decision.document_id) if decision.document_id is not None else None
        act_code = _first_act(decision)
        raw_key = "|".join(
            [
                str(self.notary_id),
                template_kind,
                act_code or "",
                str(document.family_id if document else ""),
                (document.bank_name if document else "") or "",
                (document.project_name if document else "") or "",
                template_name,
            ]
        )
        key = "tpl_" + hashlib.sha1(raw_key.encode("utf-8")).hexdigest()[:16]
        item = (
            self.db.query(NotarialTemplateLibraryItem)
            .filter(NotarialTemplateLibraryItem.notary_id == self.notary_id, NotarialTemplateLibraryItem.library_key == key)
            .first()
        )
        if item is None:
            item = NotarialTemplateLibraryItem(
                notary_id=self.notary_id,
                library_key=key,
                name=template_name,
                template_kind=template_kind,
                status="draft",
                act_code=act_code,
                document_type=document.document_type if document else None,
                family_id=document.family_id if document else None,
                bank_name=document.bank_name if document else None,
                project_name=document.project_name if document else None,
                source_document_id=document.id if document else None,
                metadata_json=json.dumps({"review_session_id": session.id}, ensure_ascii=False, sort_keys=True),
            )
            self.db.add(item)
            self.db.flush()
        return item

    def _create_template_version(self, item: NotarialTemplateLibraryItem, session: NotarialHumanReviewSession, decision: NotarialDocumentDecision, *, status: str) -> NotarialTemplateVersion:
        blocks = (
            self.db.query(NotarialDocumentBlock)
            .filter(
                NotarialDocumentBlock.notary_id == self.notary_id,
                NotarialDocumentBlock.document_id == decision.document_id,
                NotarialDocumentBlock.parse_run_id == decision.parse_run_id,
            )
            .order_by(NotarialDocumentBlock.block_index.asc())
            .all()
        )
        fields = [field for field in self._field_rows(session.id) if field.status == "approved"]
        paragraphs, placeholders = _template_paragraphs(blocks, fields)
        version = NotarialTemplateVersion(
            library_item_id=item.id,
            notary_id=self.notary_id,
            version_number=self._next_template_version(item.id),
            status=status,
            source_decision_id=decision.id,
            review_session_id=session.id,
            source_document_id=decision.document_id,
            content_json=json.dumps({"paragraphs": paragraphs}, ensure_ascii=False, sort_keys=True),
            placeholder_map_json=json.dumps(placeholders, ensure_ascii=False, sort_keys=True),
            provenance_json=json.dumps(
                {
                    "review_session_id": session.id,
                    "decision_id": decision.id,
                    "document_id": decision.document_id,
                    "parse_run_id": decision.parse_run_id,
                    "field_review_ids": [field.id for field in fields],
                },
                ensure_ascii=False,
                sort_keys=True,
            ),
            created_by_user_id=self.actor_user_id,
            metadata_json=json.dumps({"writer": "safe-red-placeholder-v1"}, ensure_ascii=False, sort_keys=True),
        )
        self.db.add(version)
        self.db.flush()
        return version

    def _known_field_code(self, code: str | None) -> bool:
        if not code:
            return False
        return self.db.query(FieldDefinition.id).filter(func.upper(FieldDefinition.field_code) == code).first() is not None

    def _decision(self, decision_id: int) -> NotarialDocumentDecision:
        decision = self.db.get(NotarialDocumentDecision, decision_id)
        if decision is None or decision.notary_id != self.notary_id:
            raise ValueError("Decision no encontrada.")
        return decision

    def _session(self, session_id: int) -> NotarialHumanReviewSession:
        session = self.db.get(NotarialHumanReviewSession, session_id)
        if session is None or session.notary_id != self.notary_id:
            raise ValueError("Sesion de revision no encontrada.")
        return session

    def _library_item(self, item_id: int) -> NotarialTemplateLibraryItem:
        item = self.db.get(NotarialTemplateLibraryItem, item_id)
        if item is None or item.notary_id != self.notary_id:
            raise ValueError("Plantilla de biblioteca no encontrada.")
        return item

    def _version(self, version_id: int, item_id: int) -> NotarialTemplateVersion:
        version = self.db.get(NotarialTemplateVersion, version_id)
        if version is None or version.notary_id != self.notary_id or version.library_item_id != item_id:
            raise ValueError("Version de plantilla no encontrada.")
        return version

    def _latest_version(self, item: NotarialTemplateLibraryItem) -> NotarialTemplateVersion:
        version_id = item.latest_version_id or item.approved_version_id
        version = self.db.get(NotarialTemplateVersion, version_id) if version_id else None
        if version is None:
            raise ValueError("La plantilla no tiene versiones.")
        return version

    def _library_item_for_session(self, session: NotarialHumanReviewSession) -> NotarialTemplateLibraryItem:
        item = (
            self.db.query(NotarialTemplateLibraryItem)
            .filter(NotarialTemplateLibraryItem.notary_id == self.notary_id, NotarialTemplateLibraryItem.metadata_json.contains(f'"review_session_id": {session.id}'))
            .order_by(NotarialTemplateLibraryItem.id.desc())
            .first()
        )
        if item is None:
            raise ValueError("La sesion no tiene plantilla asociada.")
        return item

    def _field_rows(self, session_id: int) -> list[NotarialHumanFieldReview]:
        return (
            self.db.query(NotarialHumanFieldReview)
            .filter(NotarialHumanFieldReview.session_id == session_id, NotarialHumanFieldReview.notary_id == self.notary_id)
            .order_by(NotarialHumanFieldReview.id.asc())
            .all()
        )

    def _next_template_version(self, item_id: int) -> int:
        current = self.db.query(func.coalesce(func.max(NotarialTemplateVersion.version_number), 0)).filter(NotarialTemplateVersion.library_item_id == item_id).scalar()
        return int(current or 0) + 1

    def _visual_review(self, session: NotarialHumanReviewSession, fields: list[NotarialHumanFieldReview]) -> list[dict[str, Any]]:
        by_block: dict[int, list[NotarialHumanFieldReview]] = {}
        for field in fields:
            if field.block_id is not None:
                by_block.setdefault(field.block_id, []).append(field)
        blocks = (
            self.db.query(NotarialDocumentBlock)
            .filter(NotarialDocumentBlock.id.in_(list(by_block) or [0]))
            .order_by(NotarialDocumentBlock.block_index.asc())
            .all()
        )
        return [
            {
                "block_id": block.id,
                "location_key": block.location_key,
                "original_text": block.text,
                "proposed_text": _replace_with_markers(block.text, by_block.get(block.id, [])),
                "fields": [field.id for field in by_block.get(block.id, [])],
            }
            for block in blocks
        ]

    def _audit(self, entity_type: str, entity_id: int, event_type: str, idempotency_key: str | None, payload: dict[str, Any]) -> None:
        if idempotency_key and self._audit_seen(idempotency_key):
            return
        self.db.add(
            NotarialHumanReviewAudit(
                notary_id=self.notary_id,
                entity_type=entity_type,
                entity_id=entity_id,
                event_type=event_type,
                actor_user_id=self.actor_user_id,
                idempotency_key=idempotency_key,
                payload_json=json.dumps(payload, ensure_ascii=False, sort_keys=True),
            )
        )

    def _audit_seen(self, idempotency_key: str | None) -> bool:
        if not idempotency_key:
            return False
        return (
            self.db.query(NotarialHumanReviewAudit.id)
            .filter(NotarialHumanReviewAudit.notary_id == self.notary_id, NotarialHumanReviewAudit.idempotency_key == idempotency_key)
            .first()
            is not None
        )


def _template_paragraphs(blocks: list[NotarialDocumentBlock], fields: list[NotarialHumanFieldReview]) -> tuple[list[dict[str, Any]], dict[str, Any]]:
    by_block: dict[int, list[NotarialHumanFieldReview]] = {}
    for field in fields:
        if field.block_id is not None and field.status == "approved":
            by_block.setdefault(field.block_id, []).append(field)
    placeholders: dict[str, Any] = {}
    paragraphs = []
    for block in blocks:
        text = _replace_with_markers(block.text, by_block.get(block.id, []), placeholders)
        paragraphs.append({"block_id": block.id, "location_key": block.location_key, "text": text})
    return paragraphs, placeholders


def _replace_with_markers(text: str, fields: list[NotarialHumanFieldReview], placeholders: dict[str, Any] | None = None) -> str:
    output = text or ""
    for field in sorted(fields, key=lambda item: len(item.original_value or ""), reverse=True):
        if not field.original_value or not field.proposed_field_code:
            continue
        marker = "{{" + _normalize_code(field.proposed_field_code) + "}}"
        output = re.sub(re.escape(field.original_value), marker, output, count=1)
        if placeholders is not None:
            placeholders[marker] = {
                "field_review_id": field.id,
                "field_code": _normalize_code(field.proposed_field_code),
                "original_value": field.original_value,
                "fixed_variable_label": field.fixed_variable_label,
            }
    return output


def _session_payload(session: NotarialHumanReviewSession) -> dict[str, Any]:
    return {
        "id": session.id,
        "decision_id": session.decision_id,
        "document_id": session.document_id,
        "parse_run_id": session.parse_run_id,
        "version": session.version,
        "status": session.status,
        "reviewer_user_id": session.reviewer_user_id,
        "locked_by_user_id": session.locked_by_user_id,
        "approved_at": session.approved_at.isoformat() if session.approved_at else None,
        "metadata": _json_object(session.metadata_json),
    }


def _field_payload(field: NotarialHumanFieldReview) -> dict[str, Any]:
    return {
        "id": field.id,
        "session_id": field.session_id,
        "block_id": field.block_id,
        "entity_id": field.entity_id,
        "occurrence_key": field.occurrence_key,
        "field_code": field.field_code,
        "proposed_field_code": field.proposed_field_code,
        "original_value": field.original_value,
        "proposed_value": field.proposed_value,
        "action": field.action,
        "apply_scope": field.apply_scope,
        "fixed_variable_label": field.fixed_variable_label,
        "is_new_field_proposal": field.is_new_field_proposal,
        "status": field.status,
        "metadata": _json_object(field.metadata_json),
    }


def _library_payload(item: NotarialTemplateLibraryItem) -> dict[str, Any]:
    return {
        "id": item.id,
        "library_key": item.library_key,
        "name": item.name,
        "template_kind": item.template_kind,
        "status": item.status,
        "act_code": item.act_code,
        "document_type": item.document_type,
        "family_id": item.family_id,
        "bank_name": item.bank_name,
        "legal_entity_id": item.legal_entity_id,
        "project_name": item.project_name,
        "source_document_id": item.source_document_id,
        "latest_version_id": item.latest_version_id,
        "approved_version_id": item.approved_version_id,
        "metadata": _json_object(item.metadata_json),
    }


def _version_payload(version: NotarialTemplateVersion) -> dict[str, Any]:
    return {
        "id": version.id,
        "library_item_id": version.library_item_id,
        "version_number": version.version_number,
        "status": version.status,
        "source_decision_id": version.source_decision_id,
        "review_session_id": version.review_session_id,
        "source_document_id": version.source_document_id,
        "placeholder_map": _json_object(version.placeholder_map_json),
        "provenance": _json_object(version.provenance_json),
        "storage_path": version.storage_path,
        "approved_at": version.approved_at.isoformat() if version.approved_at else None,
        "rollback_of_version_id": version.rollback_of_version_id,
    }


def _first_act(decision: NotarialDocumentDecision) -> str | None:
    hybrid = _json_object(decision.hybrid_decision_json)
    acts = (((hybrid.get("selected") or {}).get("classification") or {}).get("acts") or [])
    return str(acts[0]) if acts else None


def _lock_token(decision_id: int, version: int, actor_user_id: int | None) -> str:
    raw = f"{decision_id}:{version}:{actor_user_id}:{_now().isoformat()}"
    return hashlib.sha1(raw.encode("utf-8")).hexdigest()


def _normalize_code(value: str | None) -> str | None:
    if not value:
        return None
    normalized = re.sub(r"[^A-Za-z0-9]+", "_", value).strip("_").upper()
    return normalized or None


def _json_object(raw: str | None) -> dict[str, Any]:
    try:
        payload = json.loads(raw or "{}")
    except json.JSONDecodeError:
        return {}
    return payload if isinstance(payload, dict) else {}


def _merge_json(raw: str | None, patch: dict[str, Any]) -> str:
    payload = _json_object(raw)
    payload.update({key: value for key, value in patch.items() if value is not None})
    return json.dumps(payload, ensure_ascii=False, sort_keys=True)


def _now() -> datetime:
    return datetime.now(timezone.utc)
