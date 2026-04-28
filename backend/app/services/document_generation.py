from __future__ import annotations

import json
import re
import zipfile
from bisect import bisect_right
from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX

PLACEHOLDER_PATTERN = re.compile(r"\{\{.*?\}\}", re.DOTALL)
DASH_PLACEHOLDER = "[[--]]"
TAB_NOTARIAL = "\t"


def normalize_placeholder(raw: str) -> str:
    plain = raw.replace("\n", "").replace("\r", "").replace("\t", "")
    plain = re.sub(r"\s+", "", plain)
    return plain


def build_notarial_dash_fill(text: str | None = None) -> str:
    """
    EasyPro_1 no construye una cadena de guiones a mano.
    Inserta un tab literal y deja que Word/LibreOffice aplique
    el tab stop con líder de dashes definido en la plantilla.
    """
    return TAB_NOTARIAL


def _iter_paragraphs(container):
    paragraphs = getattr(container, "paragraphs", [])
    for paragraph in paragraphs:
        yield paragraph

    for table in getattr(container, "tables", []):
        for row in table.rows:
            for cell in row.cells:
                yield from _iter_paragraphs(cell)


def _replace_tokens_in_paragraph(paragraph, replacements: dict[str, str]) -> bool:
    runs = list(paragraph.runs)
    if not runs:
        return False

    original_texts = [run.text or "" for run in runs]
    full_text = "".join(original_texts)
    if not full_text:
        return False

    matches: list[tuple[int, int, str]] = []
    for match in PLACEHOLDER_PATTERN.finditer(full_text):
        token = normalize_placeholder(match.group(0))
        if token in replacements:
            matches.append((match.start(), match.end(), str(replacements[token] or "")))

    if not matches and DASH_PLACEHOLDER not in full_text:
        return False

    run_starts: list[int] = []
    cursor = 0
    for text in original_texts:
        run_starts.append(cursor)
        cursor += len(text)

    mutated_texts = list(original_texts)

    for start, end, replacement in reversed(matches):
        start_run = bisect_right(run_starts, start) - 1
        end_run = bisect_right(run_starts, end - 1) - 1
        if start_run < 0 or end_run < 0:
            continue

        start_offset = start - run_starts[start_run]
        end_offset = end - run_starts[end_run]

        if start_run == end_run:
            mutated_texts[start_run] = (
                mutated_texts[start_run][:start_offset]
                + replacement
                + mutated_texts[start_run][end_offset:]
            )
            continue

        mutated_texts[start_run] = mutated_texts[start_run][:start_offset] + replacement
        for index in range(start_run + 1, end_run):
            mutated_texts[index] = ""
        mutated_texts[end_run] = mutated_texts[end_run][end_offset:]

    paragraph_text_after = "".join(mutated_texts)
    if DASH_PLACEHOLDER in paragraph_text_after:
        dash_fill = build_notarial_dash_fill(paragraph_text_after)
        dash_run_starts: list[int] = []
        cursor = 0
        for text in mutated_texts:
            dash_run_starts.append(cursor)
            cursor += len(text)

        dash_matches = list(re.finditer(re.escape(DASH_PLACEHOLDER), paragraph_text_after))
        for start, end in reversed([(m.start(), m.end()) for m in dash_matches]):
            start_run = bisect_right(dash_run_starts, start) - 1
            end_run = bisect_right(dash_run_starts, end - 1) - 1
            if start_run < 0 or end_run < 0:
                continue

            start_offset = start - dash_run_starts[start_run]
            end_offset = end - dash_run_starts[end_run]

            if start_run == end_run:
                mutated_texts[start_run] = (
                    mutated_texts[start_run][:start_offset]
                    + dash_fill
                    + mutated_texts[start_run][end_offset:]
                )
                continue

            mutated_texts[start_run] = mutated_texts[start_run][:start_offset] + dash_fill
            for index in range(start_run + 1, end_run):
                mutated_texts[index] = ""
            mutated_texts[end_run] = mutated_texts[end_run][end_offset:]

    changed = False
    for run, new_text in zip(runs, mutated_texts):
        if run.text != new_text:
            run.text = new_text
            changed = True
    return changed


def _replace_tokens_in_container(container, replacements: dict[str, str]) -> bool:
    changed = False
    for paragraph in _iter_paragraphs(container):
        changed = _replace_tokens_in_paragraph(paragraph, replacements) or changed
    return changed


def render_docx_template(source_path: str | Path, destination_path: str | Path, replacements: dict[str, str]) -> None:
    source = Path(source_path)
    destination = Path(destination_path)
    destination.parent.mkdir(parents=True, exist_ok=True)
    document = Document(str(source))
    _replace_tokens_in_container(document, replacements)
    for section in document.sections:
        _replace_tokens_in_container(section.header, replacements)
        _replace_tokens_in_container(section.footer, replacements)
    document.save(str(destination))


def recalculate_dash_fills(docx_path: str | Path) -> None:
    """
    Después de render_docx_template(), recorre cada párrafo del docx
    y ajusta los runs de guiones '- - -' al final para que llenen
    exactamente hasta el margen derecho (seguridad antifraude notarial).

    Lógica:
    - Cada párrafo notarial termina con un run de guiones "- - - - -"
    - Detectar runs que son SOLO guiones y espacios al final del párrafo
    - Calcular el texto real del párrafo sin los guiones finales
    - Calcular cuántos guiones caben para llegar a LONGITUD_OBJETIVO = 90 chars
    - Reemplazar el run de guiones con la cantidad correcta
    """
    doc = Document(str(docx_path))

    for paragraph in _iter_paragraphs(doc):
        runs = list(paragraph.runs)
        if not runs:
            continue
        full_text = "".join(run.text or "" for run in runs)
        if DASH_PLACEHOLDER not in full_text:
            continue
        _replace_tokens_in_paragraph(paragraph, {DASH_PLACEHOLDER: build_notarial_dash_fill(full_text)})

    for section in doc.sections:
        for paragraph in _iter_paragraphs(section.header):
            if DASH_PLACEHOLDER in "".join(run.text or "" for run in paragraph.runs):
                _replace_tokens_in_paragraph(paragraph, {DASH_PLACEHOLDER: build_notarial_dash_fill("".join(run.text or "" for run in paragraph.runs))})
        for paragraph in _iter_paragraphs(section.footer):
            if DASH_PLACEHOLDER in "".join(run.text or "" for run in paragraph.runs):
                _replace_tokens_in_paragraph(paragraph, {DASH_PLACEHOLDER: build_notarial_dash_fill("".join(run.text or "" for run in paragraph.runs))})

    doc.save(str(docx_path))


def extract_text_from_docx(source_path: str | Path) -> str:
    """
    Extrae el texto plano del Word para usarlo como referencia para Gari.
    """
    source = Path(source_path)
    text_parts: list[str] = []
    with zipfile.ZipFile(source, "r") as zf:
        if "word/document.xml" in zf.namelist():
            xml_content = zf.read("word/document.xml").decode("utf-8", errors="ignore")
            clean = re.sub(r"<[^>]+>", " ", xml_content)
            clean = re.sub(r"\s+", " ", clean).strip()
            text_parts.append(clean)
    return "\n".join(text_parts)


def generate_plain_pdf(destination_path: str | Path, title: str, lines: list[str]) -> None:
    escaped_lines: list[str] = []
    for raw in [title, *lines]:
        text = (raw or "").replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
        escaped_lines.append(text)

    content_lines = ["BT", "/F1 12 Tf", "50 800 Td"]
    first = True
    for line in escaped_lines:
        if first:
            content_lines.append(f"({line}) Tj")
            first = False
        else:
            content_lines.append("0 -16 Td")
            content_lines.append(f"({line}) Tj")
    content_lines.append("ET")
    stream = "\n".join(content_lines).encode("latin-1", errors="replace")

    objects = [
        b"1 0 obj<< /Type /Catalog /Pages 2 0 R >>endobj\n",
        b"2 0 obj<< /Type /Pages /Kids [3 0 R] /Count 1 >>endobj\n",
        b"3 0 obj<< /Type /Page /Parent 2 0 R /MediaBox [0 0 595 842] /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>endobj\n",
        f"4 0 obj<< /Length {len(stream)} >>stream\n".encode("ascii") + stream + b"\nendstream endobj\n",
        b"5 0 obj<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>endobj\n",
    ]

    destination = Path(destination_path)
    destination.parent.mkdir(parents=True, exist_ok=True)
    with destination.open("wb") as handle:
        handle.write(b"%PDF-1.4\n")
        offsets = [0]
        for obj in objects:
            offsets.append(handle.tell())
            handle.write(obj)
        xref_start = handle.tell()
        handle.write(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
        handle.write(b"0000000000 65535 f \n")
        for offset in offsets[1:]:
            handle.write(f"{offset:010d} 00000 n \n".encode("ascii"))
        handle.write(f"trailer<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_start}\n%%EOF".encode("ascii"))


def build_case_text_snapshot(case_number: str, act_type: str, participants: list[dict], act_data: dict) -> list[str]:
    lines = [
        f"Caso: {case_number}",
        f"Acto: {act_type}",
        "Intervinientes:",
    ]
    for item in participants:
        lines.append(f"- {item.get('role_label')}: {item.get('full_name')} ({item.get('document_type')} {item.get('document_number')})")
    lines.append("Datos del acto:")
    for key, value in act_data.items():
        lines.append(f"- {key}: {value}")
    return lines


def serialize_placeholder_snapshot(replacements: dict[str, str]) -> str:
    return json.dumps(replacements, ensure_ascii=False, indent=2)


def extract_highlighted_fields_from_docx(source_path: str | Path) -> list[dict]:
    from docx import Document
    from docx.enum.text import WD_COLOR_INDEX
    import re

    source = Path(source_path)
    doc = Document(str(source))

    VALID_HIGHLIGHTS = {WD_COLOR_INDEX.TURQUOISE, WD_COLOR_INDEX.YELLOW}

    # PASO 1: buscar etiquetas {{CAMPO}} - prioridad alta
    seen_codes: dict[str, int] = {}
    etiqueta_fields: list[dict] = []
    order = 1

    for para in doc.paragraphs:
        matches = re.findall(r"\{\{([^}]+)\}\}", para.text)
        for match in matches:
            raw = match.strip()
            if not raw or len(raw) < 2:
                continue
            if re.match(r"^[\s\.\-\_\,\;\:]+$", raw):
                continue
            field_code = re.sub(r"[^a-z0-9]+", "_", raw.lower()).strip("_")
            if not field_code or field_code in seen_codes:
                continue
            label = raw.replace("_", " ").title().strip()
            if len(label) < 2:
                continue
            seen_codes[field_code] = order
            etiqueta_fields.append({
                "field_code": field_code,
                "label": label,
                "display_order": order,
                "highlight_color": "none",
            })
            order += 1

    if etiqueta_fields:
        return etiqueta_fields

    # PASO 2: fallback - buscar highlights turquesa o amarillo
    seen_codes = {}
    highlight_fields: list[dict] = []
    order = 1

    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.highlight_color not in VALID_HIGHLIGHTS:
                continue
            raw = (run.text or "").strip()
            if not raw or len(raw) < 2:
                continue
            if re.match(r"^[\s\.\-\_\,\;\:]+$", raw):
                continue
            field_code = re.sub(r"[^a-z0-9]+", "_", raw.lower()).strip("_")
            if not field_code or field_code in seen_codes:
                continue
            label = raw.replace("_", " ").strip()
            if len(label) < 2:
                continue
            color_name = "turquoise" if run.font.highlight_color == WD_COLOR_INDEX.TURQUOISE else "yellow"
            seen_codes[field_code] = order
            highlight_fields.append({
                "field_code": field_code,
                "label": label,
                "display_order": order,
                "highlight_color": color_name,
            })
            order += 1

    return highlight_fields
