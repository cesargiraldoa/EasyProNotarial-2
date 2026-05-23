"""
EasyPro 2 - Motor de reemplazo de datos en .docx
==================================================

Aplica reemplazos en el documento preservando:
- Formato (negrita, italica, mayusculas)
- Multiples ocurrencias del mismo dato
- Variantes de mayusculas/minusculas

Estrategia:
1. Para cada dato a reemplazar, buscar TODAS las ocurrencias
2. Manejar variantes: "Juan Perez" vs "JUAN PEREZ" vs "Juan Pérez"
3. Reemplazar a nivel de "run" para preservar formato
4. Devolver un nuevo .docx + lista de reemplazos efectuados (trazabilidad)
"""
import re
from copy import deepcopy
from pathlib import Path
from typing import Optional

from docx import Document
from docx.document import Document as DocumentType


def normalizar(texto: str) -> str:
    """Normaliza para matching: quita tildes, pasa a mayusculas, comprime espacios."""
    if not texto:
        return ""
    # Quitar tildes
    reemplazos_tildes = str.maketrans("áéíóúÁÉÍÓÚñÑ", "aeiouAEIOUnN")
    t = texto.translate(reemplazos_tildes)
    # Mayusculas + comprimir espacios
    t = re.sub(r"\s+", " ", t.upper().strip())
    return t


def aplicar_caso_de(plantilla_original: str, valor_nuevo: str) -> str:
    """
    Aplica el caso (mayusculas/minusculas) del original al nuevo valor.
    Si el original estaba en MAYUSCULAS -> el nuevo va en MAYUSCULAS.
    Si estaba en Capitalizado -> Capitalizado.
    """
    if plantilla_original.isupper():
        return valor_nuevo.upper()
    elif plantilla_original.istitle():
        return valor_nuevo.title()
    elif plantilla_original.islower():
        return valor_nuevo.lower()
    else:
        # Mezcla rara, devolver tal cual
        return valor_nuevo


def buscar_todas_variantes(texto_completo: str, valor_buscado: str) -> list[tuple[int, int, str]]:
    """
    Busca todas las apariciones de un valor en el texto, considerando variantes
    de mayusculas/minusculas y tildes.

    Retorna lista de (inicio, fin, texto_encontrado_exacto).
    """
    if not valor_buscado or len(valor_buscado.strip()) < 2:
        return []

    apariciones = []

    # Buscar tanto el valor original como variantes case-insensitive
    # Usamos regex case-insensitive con escape de caracteres especiales
    patron_escapado = re.escape(valor_buscado)

    for match in re.finditer(patron_escapado, texto_completo, re.IGNORECASE):
        apariciones.append((match.start(), match.end(), match.group(0)))

    return apariciones


def construir_patron_flexible(texto: str) -> str:
    """
    Construye un patron regex que tolera espacios y caracteres invisibles
    entre las palabras del texto. Util para nombres que pueden estar
    fragmentados en multiples runs con caracteres especiales entre medio.

    Ejemplo:
        "JUAN PEREZ" -> r"JUAN[\\s\\u00a0\\u200b]*PEREZ"

    Acepta como separador:
    - Espacios normales \\s
    - Espacios no-rompibles \\u00a0
    - Zero-width space \\u200b
    - Multiples de cualquier combinacion
    """
    palabras = texto.split()
    if len(palabras) <= 1:
        return re.escape(texto)

    # Cada palabra escapada, separadas por un patron flexible
    separador = r"[\s ​]+"
    return separador.join(re.escape(p) for p in palabras)


def reemplazar_en_runs(paragraph, valor_viejo: str, valor_nuevo: str, palabra_completa: bool = False) -> int:
    """
    Reemplaza dentro de los runs de un parrafo preservando formato.
    Maneja correctamente matches que cruzan multiples runs.

    Args:
        paragraph: parrafo python-docx
        valor_viejo: texto a buscar
        valor_nuevo: texto que lo reemplaza
        palabra_completa: si True, exige limites de palabra (\\b...\\b)
                         - usar True para concordancia (comprador != compradores)
                         - usar False para datos (cedulas, valores, nombres completos)

    Estrategia:
    1. Construir mapa de posicion absoluta -> (run_idx, offset_local)
    2. Buscar todas las ocurrencias en el texto del parrafo
    3. Para cada match, identificar runs afectados
    4. Si el match esta en 1 run: reemplazo simple
    5. Si el match cruza varios runs:
       - Poner el valor nuevo completo en el primer run afectado
       - Vaciar el resto de runs que esten dentro del match

    Retorna numero de reemplazos hechos en este parrafo.
    """
    if not valor_viejo or not paragraph.text:
        return 0

    # Texto completo del parrafo
    texto_completo = paragraph.text

    # Decidir patron:
    # - Si es palabra_completa: usar \b...\b para limites estrictos
    # - Si valor tiene multiples palabras: usar patron flexible (tolera separadores raros)
    # - Si es valor simple: re.escape normal
    if palabra_completa:
        patron = r"\b" + re.escape(valor_viejo) + r"\b"
    elif " " in valor_viejo.strip():
        # Nombre compuesto - usar patron flexible que tolera caracteres invisibles
        patron = construir_patron_flexible(valor_viejo)
    else:
        patron = re.escape(valor_viejo)

    matches = list(re.finditer(patron, texto_completo, re.IGNORECASE))

    if not matches:
        return 0

    reemplazos = 0

    # Procesar matches en ORDEN INVERSO para que las posiciones no se invaliden
    for match in reversed(matches):
        encontrado_exacto = match.group(0)
        valor_final = aplicar_caso_de(encontrado_exacto, valor_nuevo)
        match_inicio = match.start()
        match_fin = match.end()

        # Mapear cada run a su rango de posiciones
        # NOTA: regenerar el mapa en cada iteracion porque los reemplazos previos
        # cambian las longitudes
        runs_info = []
        pos = 0
        for run in paragraph.runs:
            runs_info.append({
                "run": run,
                "inicio": pos,
                "fin": pos + len(run.text),
            })
            pos += len(run.text)

        # Identificar runs afectados por este match
        runs_afectados = []
        for info in runs_info:
            if info["fin"] <= match_inicio:
                continue  # run antes del match
            if info["inicio"] >= match_fin:
                break  # run despues del match
            # Hay solape
            runs_afectados.append(info)

        if not runs_afectados:
            continue

        if len(runs_afectados) == 1:
            # Caso simple: match en un solo run
            info = runs_afectados[0]
            run = info["run"]
            inicio_local = match_inicio - info["inicio"]
            fin_local = match_fin - info["inicio"]
            run.text = run.text[:inicio_local] + valor_final + run.text[fin_local:]
            reemplazos += 1
        else:
            # Caso complejo: match cruza varios runs
            # Estrategia:
            # - Primer run afectado: mantener prefijo antes del match + valor_nuevo
            # - Runs intermedios: vaciar
            # - Ultimo run afectado: mantener sufijo despues del match

            primer = runs_afectados[0]
            ultimo = runs_afectados[-1]

            prefijo = primer["run"].text[: match_inicio - primer["inicio"]]
            sufijo = ultimo["run"].text[match_fin - ultimo["inicio"]:]

            # Poner el valor nuevo en el primer run (hereda su formato)
            primer["run"].text = prefijo + valor_final

            # Vaciar runs intermedios
            for info in runs_afectados[1:-1]:
                info["run"].text = ""

            # El ultimo run mantiene solo el sufijo
            if primer is not ultimo:
                ultimo["run"].text = sufijo

            reemplazos += 1

    return reemplazos


def _normalizar_guiones(doc):
    """
    Une el texto de todos los runs de cada párrafo,
    limpia guiones dobles, y redistribuye el texto
    preservando el formato del primer run.
    """
    for paragraph in doc.paragraphs:
        if not paragraph.runs:
            continue

        texto_completo = ''.join(run.text for run in paragraph.runs)

        texto_limpio = re.sub(r'-{2,}', '-', texto_completo)
        texto_limpio = re.sub(r'- {2,}-', '- -', texto_limpio)
        texto_limpio = re.sub(r' {2,}', ' ', texto_limpio)

        if texto_limpio == texto_completo:
            continue

        # Todo el texto limpio va al primer run; el resto se vacía.
        # Evita el problema de redistribuir con longitudes originales cuando
        # texto_limpio es más corto que texto_completo.
        paragraph.runs[0].text = texto_limpio
        for run in paragraph.runs[1:]:
            run.text = ""

    return doc


def aplicar_reemplazos_docx(
    docx_path_origen: str,
    reemplazos: list[dict],
    docx_path_destino: str
) -> dict:
    """
    Aplica una lista de reemplazos al .docx y guarda el resultado.

    Args:
        docx_path_origen: ruta del .docx original
        reemplazos: lista de dicts {viejo, nuevo, etiqueta, palabra_completa?}
                    Si palabra_completa=True, exige limites \\b...\\b
                    (necesario para concordancia: "comprador" != "compradores")
        docx_path_destino: ruta donde guardar el .docx modificado

    Returns:
        dict con estadisticas: { total_reemplazos, por_etiqueta, no_encontrados }
    """
    doc = Document(docx_path_origen)

    estadisticas = {
        "total_reemplazos": 0,
        "por_etiqueta": {},
        "no_encontrados": [],
    }

    for rep in reemplazos:
        viejo = rep.get("viejo", "").strip()
        nuevo = rep.get("nuevo", "").strip()
        etiqueta = rep.get("etiqueta", "sin_etiqueta")
        palabra_completa = rep.get("palabra_completa", False)

        if not viejo or not nuevo or viejo == nuevo:
            continue

        count_para_este = 0

        # Reemplazar en parrafos
        for paragraph in doc.paragraphs:
            count_para_este += reemplazar_en_runs(paragraph, viejo, nuevo, palabra_completa)

        # Reemplazar en tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        count_para_este += reemplazar_en_runs(paragraph, viejo, nuevo, palabra_completa)

        if count_para_este == 0:
            estadisticas["no_encontrados"].append({
                "etiqueta": etiqueta,
                "viejo": viejo,
                "nuevo": nuevo,
            })
        else:
            estadisticas["por_etiqueta"][etiqueta] = count_para_este
            estadisticas["total_reemplazos"] += count_para_este

    # Limpiar guiones dobles que puedan quedar tras los reemplazos
    _normalizar_guiones(doc)

    # Guardar
    doc.save(docx_path_destino)

    return estadisticas


def aplicar_reemplazos_por_contexto(
    doc: DocumentType,
    nombre_persona: str,
    reemplazos: list[dict],
    nombres_excluir: Optional[list[str]] = None,
) -> int:
    """
    Aplica reemplazos solo en párrafos que contienen nombre_persona
    Y NO contienen ningún nombre de nombres_excluir.

    El filtro de exclusión evita contaminar párrafos que listan varias
    personas juntas (ej. "DANIELA CAMPO... y SEBASTIÁN NIETO GIRALDO...").
    Retorna el número de reemplazos efectuados.
    """
    nombre_upper = nombre_persona.upper()
    excluir_upper = [n.upper() for n in (nombres_excluir or []) if n.strip()]

    def _procesar(paragraph) -> int:
        texto_upper = paragraph.text.upper()
        if nombre_upper not in texto_upper:
            return 0
        if any(excl in texto_upper for excl in excluir_upper):
            return 0
        total = 0
        for rep in reemplazos:
            viejo = rep.get("viejo", "").strip()
            nuevo = rep.get("nuevo", "").strip()
            if not viejo or not nuevo or viejo == nuevo:
                continue
            total += reemplazar_en_runs(paragraph, viejo, nuevo, rep.get("palabra_completa", False))
        return total

    total = 0
    for paragraph in doc.paragraphs:
        total += _procesar(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    total += _procesar(paragraph)
    return total


def aplicar_genero_contextual_docx(
    docx_path: str,
    pares: list[dict],
    todos_nombres: Optional[list[str]] = None,
) -> dict:
    """
    Aplica reemplazos de género solo en párrafos que mencionan a cada persona
    y no mencionan a ninguna otra persona del documento.
    Modifica el archivo in-place.

    Args:
        docx_path: ruta del .docx de salida (ya procesado por aplicar_reemplazos_docx)
        pares: lista de {"nombre": str, "reemplazos": list[dict]}
        todos_nombres: lista completa de nombres de personas en el documento;
                       un párrafo se salta si menciona a cualquier otro nombre

    Returns:
        dict nombre → cantidad de reemplazos efectuados
    """
    doc = Document(docx_path)
    nombres_ref = [n.strip() for n in (todos_nombres or []) if n.strip()]
    stats: dict = {}
    for par in pares:
        nombre = (par.get("nombre") or "").strip()
        reps = par.get("reemplazos") or []
        if nombre and reps:
            otros = [n for n in nombres_ref if n.upper() != nombre.upper()]
            stats[nombre] = aplicar_reemplazos_por_contexto(doc, nombre, reps, nombres_excluir=otros)
    doc.save(docx_path)
    return stats


def construir_lista_reemplazos(datos_anteriores: dict, datos_nuevos: dict) -> list[dict]:
    """
    Construye la lista de reemplazos a partir de los datos viejos y nuevos.

    Args:
        datos_anteriores: estructura con personas, valores, inmueble del caso anterior
        datos_nuevos: misma estructura con los datos nuevos a aplicar

    Returns:
        lista de dicts {viejo, nuevo, etiqueta}
    """
    reemplazos = []

    # PERSONAS
    personas_old = {(p.get("rol") or p.get("ROL") or ""): p for p in datos_anteriores.get("personas", [])}
    personas_new = {(p.get("rol") or p.get("ROL") or ""): p for p in datos_nuevos.get("personas", [])}

    # Para evitar duplicar reemplazos cuando misma persona aparece en varios roles,
    # usar un set de (viejo, nuevo) ya agregados
    pares_ya_agregados = set()

    def agregar_reemplazo(viejo, nuevo, etiqueta):
        if not viejo or not nuevo:
            return
        viejo = viejo.strip()
        nuevo = nuevo.strip()
        if not viejo or not nuevo or viejo == nuevo:
            return
        clave = (viejo.upper(), nuevo.upper())
        if clave in pares_ya_agregados:
            return
        pares_ya_agregados.add(clave)
        reemplazos.append({"viejo": viejo, "nuevo": nuevo, "etiqueta": etiqueta})

    for rol, persona_new in personas_new.items():
        persona_old = personas_old.get(rol, {})

        # Datos basicos
        agregar_reemplazo(persona_old.get("nombre_completo"), persona_new.get("nombre_completo"), f"{rol}.nombre")
        agregar_reemplazo(persona_old.get("numero_documento"), persona_new.get("numero_documento"), f"{rol}.documento")

        # Datos de contacto y firma (v2.2+)
        agregar_reemplazo(persona_old.get("direccion"), persona_new.get("direccion"), f"{rol}.direccion")
        agregar_reemplazo(persona_old.get("telefono"), persona_new.get("telefono"), f"{rol}.telefono")
        agregar_reemplazo(persona_old.get("email"), persona_new.get("email"), f"{rol}.email")
        agregar_reemplazo(persona_old.get("profesion"), persona_new.get("profesion"), f"{rol}.profesion")
        agregar_reemplazo(persona_old.get("nacionalidad"), persona_new.get("nacionalidad"), f"{rol}.nacionalidad")
        agregar_reemplazo(persona_old.get("actividad_economica"), persona_new.get("actividad_economica"), f"{rol}.actividad_economica")

    # VALORES
    valores_old = datos_anteriores.get("valores", [])
    valores_new = datos_nuevos.get("valores", [])

    for v_new in valores_new:
        tipo = v_new.get("tipo")
        acto = v_new.get("acto_relacionado")
        v_old = next(
            (v for v in valores_old
             if v.get("tipo") == tipo and v.get("acto_relacionado") == acto),
            None
        )
        texto_viejo = v_old.get("texto_en_documento") if v_old else None
        texto_nuevo = v_new.get("texto_en_documento")
        if texto_viejo and texto_nuevo and texto_viejo != texto_nuevo:
            reemplazos.append({
                "viejo": texto_viejo,
                "nuevo": texto_nuevo,
                "etiqueta": f"valor.{tipo}.acto_{acto}",
            })

    # INMUEBLE
    inm_old = datos_anteriores.get("inmueble", {}) or {}
    inm_new = datos_nuevos.get("inmueble", {}) or {}

    for campo in ["numero", "matricula_inmobiliaria", "conjunto_o_edificio", "municipio", "departamento"]:
        val_old = inm_old.get(campo)
        val_new = inm_new.get(campo)
        if val_old and val_new and val_old != val_new:
            reemplazos.append({
                "viejo": val_old,
                "nuevo": val_new,
                "etiqueta": f"inmueble.{campo}",
            })

    # NOTARIA
    not_old = datos_anteriores.get("notaria", {}) or {}
    not_new = datos_nuevos.get("notaria", {}) or {}

    for campo in ["nombre_notaria", "municipio_notaria", "numero_escritura"]:
        val_old = not_old.get(campo)
        val_new = not_new.get(campo)
        if val_old and val_new and val_old != val_new:
            reemplazos.append({
                "viejo": val_old,
                "nuevo": val_new,
                "etiqueta": f"notaria.{campo}",
            })

    # FECHAS
    fec_old = datos_anteriores.get("fechas", {}) or {}
    fec_new = datos_nuevos.get("fechas", {}) or {}

    val_old = fec_old.get("fecha_otorgamiento")
    val_new = fec_new.get("fecha_otorgamiento")
    if val_old and val_new and val_old != val_new:
        reemplazos.append({
            "viejo": val_old,
            "nuevo": val_new,
            "etiqueta": "fechas.fecha_otorgamiento",
        })

    return reemplazos

# deploy 2026-05-23

# deploy 2026-05-23
