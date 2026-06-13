from __future__ import annotations

import re
import unicodedata
from pathlib import Path
from typing import Iterable

from docx import Document


def _iter_paragraphs_in_table(table) -> Iterable:
    seen_cells: set[int] = set()
    for row in table.rows:
        for cell in row.cells:
            cell_id = id(cell._tc)
            if cell_id in seen_cells:
                continue
            seen_cells.add(cell_id)
            for paragraph in cell.paragraphs:
                yield paragraph
            for nested_table in cell.tables:
                yield from _iter_paragraphs_in_table(nested_table)


def _iter_paragraphs_in_container(container) -> Iterable:
    for paragraph in container.paragraphs:
        yield paragraph
    for table in container.tables:
        yield from _iter_paragraphs_in_table(table)


def _iter_document_paragraphs(document: Document) -> Iterable:
    yield from _iter_paragraphs_in_container(document)
    for section in document.sections:
        yield from _iter_paragraphs_in_container(section.header)
        yield from _iter_paragraphs_in_container(section.footer)


def _replace_literal_in_paragraph(paragraph, old: str, new: str) -> int:
    if not old or paragraph is None or not paragraph.text:
        return 0

    matches = list(re.finditer(re.escape(old), paragraph.text))
    if not matches:
        return 0

    replacements = 0
    for match in reversed(matches):
        match_start = match.start()
        match_end = match.end()

        runs_info = []
        position = 0
        for run in paragraph.runs:
            runs_info.append({
                "run": run,
                "start": position,
                "end": position + len(run.text),
            })
            position += len(run.text)

        affected = []
        for info in runs_info:
            if info["end"] <= match_start:
                continue
            if info["start"] >= match_end:
                break
            affected.append(info)

        if not affected:
            continue

        if len(affected) == 1:
            info = affected[0]
            run = info["run"]
            start_local = match_start - info["start"]
            end_local = match_end - info["start"]
            run.text = run.text[:start_local] + new + run.text[end_local:]
            replacements += 1
            continue

        first = affected[0]
        last = affected[-1]
        prefix = first["run"].text[: match_start - first["start"]]
        suffix = last["run"].text[match_end - last["start"]:]

        first["run"].text = prefix + new
        for info in affected[1:-1]:
            info["run"].text = ""
        if first is not last:
            last["run"].text = suffix

        replacements += 1

    return replacements


def _normalize_value(value) -> str:
    if value is None:
        return ""
    return str(value)


def _normalize_key(value: str) -> str:
    normalized = unicodedata.normalize("NFD", str(value or ""))
    normalized = "".join(char for char in normalized if unicodedata.category(char) != "Mn")
    normalized = re.sub(r"[^a-zA-Z0-9]+", "_", normalized.lower()).strip("_")
    return normalized


SECOND_BUYER_KEYS = {
    "nombre_comprador_2",
    "tipo_documento_comprador_2",
    "numero_documento_comprador_2",
    "comprador_2_es_hombre_o_mujer",
    "nacionalidad_comprador_2",
    "municipio_domicilio_comprador_2",
    "estado_civil_comprador_2",
    "direccion_comprador_2",
    "celular_comprador_2",
    "email_comprador_2",
    "profesion_u_oficio_comprador_2",
    "actividad_economica_comprador_2",
}


def is_second_buyer_empty(normalized_values: dict[str, str]) -> bool:
    return all(not _normalize_value(normalized_values.get(key, "")).strip() for key in SECOND_BUYER_KEYS)


def _extract_digits(value: str) -> str:
    return re.sub(r"\D+", "", _normalize_value(value))


MONEY_FIELD_KEYWORDS = (
    "valor",
    "precio",
    "cuota",
    "cuantia",
    "cuantia",
    "avaluo",
    "avaluo",
    "retencion",
    "retencion",
    "derechos",
    "iva",
    "aporte",
    "fondo",
    "impuesto",
    "hipoteca",
    "notariado",
    "superintendencia",
)

MONEY_FIELD_EXCLUDED_KEYWORDS = (
    "documento",
    "cedula",
    "telefono",
    "celular",
    "matricula",
    "escritura",
    "consecutivo",
    "anio",
    "ano",
    "year",
)


def _is_money_field(field: dict) -> bool:
    haystack = " ".join(
        [
            _normalize_key(field.get("key") or ""),
            _normalize_key(field.get("label") or ""),
            _normalize_key(field.get("section") or ""),
        ]
    )
    if not haystack:
        return False
    if any(keyword in haystack for keyword in MONEY_FIELD_EXCLUDED_KEYWORDS):
        return False
    return any(keyword in haystack for keyword in MONEY_FIELD_KEYWORDS)


def _format_money_value(value: str) -> str:
    raw = _normalize_value(value).strip()
    if not raw:
        return ""
    digits = _extract_digits(raw)
    if not digits:
        return raw
    normalized_raw = re.sub(r"[\s$]", "", raw)
    if re.fullmatch(r"[\d.,]+", normalized_raw) is None and not raw.isdigit():
        return raw
    return f"{int(digits):,}".replace(",", ".")


UNIDADES = [
    "",
    "UNO",
    "DOS",
    "TRES",
    "CUATRO",
    "CINCO",
    "SEIS",
    "SIETE",
    "OCHO",
    "NUEVE",
    "DIEZ",
    "ONCE",
    "DOCE",
    "TRECE",
    "CATORCE",
    "QUINCE",
    "DIECISÉIS",
    "DIECISIETE",
    "DIECIOCHO",
    "DIECINUEVE",
]
DECENAS = ["", "", "VEINTE", "TREINTA", "CUARENTA", "CINCUENTA", "SESENTA", "SETENTA", "OCHENTA", "NOVENTA"]
CENTENAS = [
    "",
    "CIENTO",
    "DOSCIENTOS",
    "TRESCIENTOS",
    "CUATROCIENTOS",
    "QUINIENTOS",
    "SEISCIENTOS",
    "SETECIENTOS",
    "OCHOCIENTOS",
    "NOVECIENTOS",
]


def _centenas_a_letras(number: int) -> str:
    if number == 0:
        return ""
    if number == 100:
        return "CIEN"
    hundreds = number // 100
    remainder = number % 100
    parts: list[str] = []
    if hundreds > 0:
        parts.append(CENTENAS[hundreds])
    if remainder > 0:
        if remainder < 20:
            parts.append(UNIDADES[remainder])
        elif remainder < 30:
            units = remainder % 10
            if units == 0:
                parts.append("VEINTE")
            else:
                parts.append(f"VEINTI{UNIDADES[units].lower()}".upper())
        else:
            tens = remainder // 10
            units = remainder % 10
            if units == 0:
                parts.append(DECENAS[tens])
            else:
                parts.append(f"{DECENAS[tens]} Y {UNIDADES[units]}")
    return " ".join(part for part in parts if part).strip()


def _numero_a_letras(number: int, suffix: str) -> str:
    if number <= 0:
        return ""
    if number == 0:
        return f"CERO {suffix}"

    parts: list[str] = []
    billions = number // 1_000_000_000
    millions = (number % 1_000_000_000) // 1_000_000
    thousands = (number % 1_000_000) // 1_000
    remainder = number % 1_000

    if billions > 0:
        text = _centenas_a_letras(billions)
        parts.append("MIL MILLONES" if billions == 1 else f"{text} MIL MILLONES")
    if millions > 0:
        text = _centenas_a_letras(millions)
        parts.append("UN MILLÓN" if millions == 1 else f"{text} MILLONES")
    if thousands > 0:
        text = _centenas_a_letras(thousands)
        parts.append("MIL" if thousands == 1 else f"{text} MIL")
    if remainder > 0:
        parts.append(_centenas_a_letras(remainder))
    return f"{' '.join(part for part in parts if part).strip()} {suffix}".strip()


DERIVED_MONEY_RULES = (
    ("valor_de_la_venta_en_numeros", "valor_apartamento_en_letras", "PESOS MONEDA LEGAL"),
    ("en_numeros_cuota_inicial", "en_letras_cuota_inicial", "PESOS"),
    ("valor_del_acto_de_la_hipoteca", "valor_del_acto_de_la_hipoteca_en_letras", "PESOS"),
)


def _derive_missing_marked_values(normalized_values: dict[str, str], normalized_field_keys: set[str]) -> dict[str, str]:
    derived = dict(normalized_values)
    for money_key, letter_key, suffix in DERIVED_MONEY_RULES:
        if money_key not in normalized_field_keys or letter_key not in normalized_field_keys:
            continue
        if derived.get(letter_key, "").strip():
            continue
        digits = _extract_digits(derived.get(money_key, ""))
        if not digits:
            continue
        derived[letter_key] = _numero_a_letras(int(digits), suffix)

    if "origen_cuota_inicial" in normalized_field_keys and not derived.get("origen_cuota_inicial", "").strip():
        derived["origen_cuota_inicial"] = "recursos propios"

    return derived


def _normalize_money_values_for_fields(normalized_values: dict[str, str], fields: list[dict]) -> dict[str, str]:
    normalized = dict(normalized_values)
    for field in fields:
        if not isinstance(field, dict) or not _is_money_field(field):
            continue
        key = _normalize_key(field.get("key") or "")
        if not key or key not in normalized:
            continue
        normalized[key] = _format_money_value(normalized[key])
    return normalized


EMPTY_SECOND_BUYER_LABEL_PATTERN = re.compile(
    r"^(?:DIRECCI[ÓO]N|CELULAR|CORREO|PROFESI[ÓO]N\s+U\s+OFICIO|ACTIVIDAD\s+ECON[ÓO]MICA|ESTADO\s+CIVIL)\s*:\s*\.?\s*$",
    re.IGNORECASE,
)

SIGNATURE_LINE_PATTERN = re.compile(r"^[\s._-]{6,}$")


def cleanup_empty_second_buyer_text(text: str) -> str:
    if not text:
        return text

    normalized_text = text.replace("\xa0", " ")
    lines: list[str] = []
    for raw_line in normalized_text.splitlines():
        stripped = raw_line.strip()
        if not stripped:
            lines.append("")
            continue
        if EMPTY_SECOND_BUYER_LABEL_PATTERN.match(stripped):
            continue

        line = raw_line
        replacements = [
            (r"\s+y\s*-\s*-\s*", ""),
            (r"\s+y\s*,\s*de las condiciones", " de las condiciones"),
            (r"\s+y\s*,\s*identificado", ", identificado"),
            (r"\s+y\s*,\s*quien", " quien"),
            (r"\s+y\s*,\s*respectivamente", ""),
            (r",\s*respectivamente", ""),
            (r"\s+y\s+respectivamente", ""),
            (r"\s+y\s*,\s*", " "),
        ]
        for pattern, replacement in replacements:
            line = re.sub(pattern, replacement, line, flags=re.IGNORECASE)

        line = re.sub(r"\s{2,}", " ", line)
        line = re.sub(r"\s+([,.;:])", r"\1", line)
        line = line.strip()
        if EMPTY_SECOND_BUYER_LABEL_PATTERN.match(line):
            continue
        lines.append(line)

    return "\n".join(lines).strip()


def _set_paragraph_text(paragraph, new_text: str) -> None:
    runs = list(paragraph.runs)
    if not runs:
        if new_text:
            paragraph.add_run(new_text)
        return

    runs[0].text = new_text
    for run in runs[1:]:
        run.text = ""


def _cleanup_empty_second_buyer_signature_block(paragraphs: list) -> None:
    for index, paragraph in enumerate(paragraphs):
        text = (paragraph.text or "").strip()
        if not SIGNATURE_LINE_PATTERN.match(text):
            continue

        window = paragraphs[index + 1 : index + 7]
        if not window:
            continue

        cleaned_window = [cleanup_empty_second_buyer_text(candidate.text or "") for candidate in window]
        if not cleaned_window:
            continue

        if all(not entry.strip() for entry in cleaned_window):
            _set_paragraph_text(paragraph, "")


def _cleanup_residual_markers(document: Document) -> None:
    all_paragraphs = list(_iter_document_paragraphs(document))
    for paragraph in all_paragraphs:
        cleaned_text = (paragraph.text or "").replace("[[--]]", "")
        cleaned_text = re.sub(r"\s{2,}", " ", cleaned_text)
        cleaned_text = re.sub(r"\s+([,.;:])", r"\1", cleaned_text).strip()
        if cleaned_text != (paragraph.text or ""):
            _set_paragraph_text(paragraph, cleaned_text)


def apply_marked_template_replacements(
    docx_path_origen: str | Path,
    docx_path_destino: str | Path,
    values: dict,
    fields: list[dict] | None = None,
) -> dict:
    """
    Reemplaza marcadores exactos de plantilla marcada usando la metadata detectada.
    No usa IA ni concordancia.
    """

    document = Document(str(docx_path_origen))
    values_map = {str(key): _normalize_value(value) for key, value in (values or {}).items()}
    fields_list = fields or []
    normalized_field_keys = {
        _normalize_key(str(field.get("key") or ""))
        for field in fields_list
        if isinstance(field, dict) and str(field.get("key") or "").strip()
    }
    normalized_values = {_normalize_key(key): value for key, value in values_map.items()}
    normalized_values = _derive_missing_marked_values(normalized_values, normalized_field_keys)
    normalized_values = _normalize_money_values_for_fields(normalized_values, fields_list)

    print("### MARKED TEMPLATE REAL GENERATOR HIT ###")
    print("### SECOND BUYER EMPTY:", is_second_buyer_empty(normalized_values))
    print("### VALUE valor_apartamento_en_letras:", normalized_values.get("valor_apartamento_en_letras"))
    print("### VALUE en_letras_cuota_inicial:", normalized_values.get("en_letras_cuota_inicial"))
    print(
        "### VALUE valor_del_acto_de_la_hipoteca_en_letras:",
        normalized_values.get("valor_del_acto_de_la_hipoteca_en_letras"),
    )
    print("### VALUE origen_cuota_inicial:", normalized_values.get("origen_cuota_inicial"))

    replacement_targets: list[tuple[str, str, str]] = []
    for field in fields_list:
        if not isinstance(field, dict):
            continue
        key = str(field.get("key") or "").strip()
        if not key:
            continue
        normalized_key = _normalize_key(key)
        replacement_value = normalized_values.get(normalized_key, values_map.get(key, ""))
        if _is_money_field(field):
            replacement_value = _format_money_value(replacement_value)
        raw_markers = field.get("raw_markers") or []
        if not isinstance(raw_markers, list):
            continue
        for raw_marker in raw_markers:
            marker = str(raw_marker or "").strip()
            if not marker:
                continue
            replacement_targets.append((marker, key, replacement_value))

    replacement_targets.sort(key=lambda item: len(item[0]), reverse=True)

    stats = {
        "total_replacements": 0,
        "by_key": {},
        "not_found": [],
    }

    seen_markers: set[str] = set()
    for marker, key, replacement_value in replacement_targets:
        if marker in seen_markers:
            continue
        seen_markers.add(marker)

        count = 0
        for paragraph in _iter_document_paragraphs(document):
            count += _replace_literal_in_paragraph(paragraph, marker, replacement_value)

        if count == 0:
            stats["not_found"].append({
                "key": key,
                "old": marker,
                "new": replacement_value,
            })
        else:
            stats["by_key"][key] = stats["by_key"].get(key, 0) + count
            stats["total_replacements"] += count

    if is_second_buyer_empty(normalized_values):
        all_paragraphs = list(_iter_document_paragraphs(document))
        for paragraph in all_paragraphs:
            cleaned_text = cleanup_empty_second_buyer_text(paragraph.text or "")
            if cleaned_text != (paragraph.text or ""):
                _set_paragraph_text(paragraph, cleaned_text)
        _cleanup_empty_second_buyer_signature_block(all_paragraphs)

    _cleanup_residual_markers(document)

    document.save(str(docx_path_destino))
    return stats
