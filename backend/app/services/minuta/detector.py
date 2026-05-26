"""
EasyPro 2 - Detector inteligente de minutas
=============================================

Funcion principal: analizar_documento(docx_path, api_key)
  -> Detecta automaticamente si es B1 (en blanco) o B2 (ya diligenciada)
  -> Extrae los datos del caso anterior (B2) o los campos vacios (B1)
  -> Retorna estructura unificada para el panel lateral

Logica de deteccion automatica:
  - Si encuentra >5 nombres propios + >3 cedulas/NITs reales -> B2
  - Si encuentra placeholders (XXX, _____, [campo]) o muy pocos datos -> B1

Adaptado del detector V4 de ayer pero con:
  - Modo dual (B1/B2)
  - Posiciones de cada dato (para reemplazo posterior)
  - Genero detectado por persona
"""
import json
import os
import re
from pathlib import Path
from typing import Optional

import httpx
from docx import Document
from openai import OpenAI


def _make_openai_client(api_key: str) -> OpenAI:
    """Crea el cliente OpenAI. En desarrollo con SSL roto, usa verify=False."""
    if os.environ.get("OPENAI_DISABLE_SSL_VERIFY", "").lower() == "true":
        return OpenAI(api_key=api_key, http_client=httpx.Client(verify=False))
    return OpenAI(api_key=api_key)


MODEL = "gpt-4o-mini"


# ============================================================
# UTILIDADES DE EXTRACCION
# ============================================================

def extraer_texto_estructurado(path: str) -> str:
    """Extrae texto del .docx preservando estructura (parrafos y tablas)."""
    doc = Document(str(path))
    partes = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            partes.append(text)
    for tbl_idx, table in enumerate(doc.tables, start=1):
        partes.append(f"\n[TABLA {tbl_idx}]")
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            partes.append(" | ".join(cells))
        partes.append(f"[/TABLA {tbl_idx}]\n")
    return "\n".join(partes)


# ============================================================
# CLASIFICACION B1 vs B2 (deteccion automatica de modo)
# ============================================================

def clasificar_modo(texto: str) -> dict:
    """
    Decide si el documento es B1 (en blanco) o B2 (diligenciada) por heuristica.

    Senales B2 (minuta diligenciada):
      - Muchas cedulas con formato real (X.XXX.XXX)
      - Nombres propios en MAYUSCULAS de mas de 2 palabras
      - Valores monetarios con cifras grandes ($XXX.XXX.XXX)

    Senales B1 (en blanco):
      - Placeholders: XXX, _____, [campo], {{var}}, ___________
      - Pocos datos reales detectados
    """
    # Conteo de senales B2
    cedulas = len(re.findall(r"\b\d{1,3}(?:\.\d{3}){2,3}\b", texto))
    nombres_mayus = len(re.findall(r"\b[A-ZÁÉÍÓÚÑ]{3,}(?:\s+[A-ZÁÉÍÓÚÑ]{3,}){2,4}\b", texto))
    valores = len(re.findall(r"\$\s*\d{1,3}(?:\.\d{3}){1,3}\b", texto))

    # Conteo de senales B1
    placeholders_xxx = len(re.findall(r"\bXXX+\b", texto))
    placeholders_lineas = len(re.findall(r"_{5,}", texto))
    placeholders_brackets = len(re.findall(r"\[[A-Z][^\]]{2,30}\]", texto))
    placeholders_llaves = len(re.findall(r"\{\{[^}]+\}\}", texto))

    senales_b2 = cedulas + nombres_mayus + valores
    senales_b1 = placeholders_xxx + placeholders_lineas + placeholders_brackets + placeholders_llaves

    # Decision
    if senales_b2 >= 10 and senales_b2 > senales_b1 * 2:
        modo = "B2"
        confianza = min(1.0, senales_b2 / 30)
    elif senales_b1 >= 5 and senales_b1 > senales_b2:
        modo = "B1"
        confianza = min(1.0, senales_b1 / 15)
    elif senales_b2 > senales_b1:
        modo = "B2"
        confianza = 0.5
    else:
        modo = "B1"
        confianza = 0.5

    return {
        "modo": modo,
        "confianza": round(confianza, 2),
        "senales_b2": {
            "cedulas": cedulas,
            "nombres_mayus": nombres_mayus,
            "valores_monetarios": valores,
            "total": senales_b2,
        },
        "senales_b1": {
            "placeholders_xxx": placeholders_xxx,
            "lineas_subrayadas": placeholders_lineas,
            "brackets": placeholders_brackets,
            "llaves": placeholders_llaves,
            "total": senales_b1,
        },
    }


# ============================================================
# PROMPT B2: extraer datos del caso anterior
# ============================================================

PROMPT_B2 = """Eres experto en escrituras publicas notariales colombianas.

TU TAREA: Te paso una minuta NOTARIAL YA DILIGENCIADA con datos de un caso REAL pasado. Debes extraer TODOS los datos variables del caso anterior, organizados por rol.

INSTRUCCIONES:

1. DEVUELVES SOLO JSON VALIDO. Sin explicaciones.

2. Para CADA persona detectada (cada comprador, vendedor, deudor, apoderado, conyuge, otorgante, notario, protocolista, representante legal):
   - rol: "comprador_1", "vendedor", "deudor_hipotecante", "apoderado", "conyuge_comprador_1", "notario", "protocolista", etc.
   - nombre_completo: tal cual aparece (preservar mayusculas)
   - tipo_documento: usar exactamente uno de estos valores: "C.C", "C.E", "T.I", "PAS", "P.P.T", "R.C", "NIT"
   - numero_documento: con el formato original (con puntos o sin puntos como aparece)
   - genero: "M" (masculino), "F" (femenino), "J" (juridica empresa/banco/fideicomiso)
   - estado_civil: si se menciona ("CASADO", "SOLTERO", "UNION_MARITAL", etc.)
   - nacionalidad: nacionalidad de la persona si se menciona (ej: "colombiano", "venezolana")
   - domicilio: ciudad de domicilio si se menciona
   - direccion: direccion completa de domicilio si se menciona (calle, numero, complemento)
   - telefono: numero de telefono/celular si se menciona
   - email: correo electronico si se menciona (tal cual aparece, preservando mayusculas)
   - profesion: profesion u ocupacion si se menciona
   - actividad_economica: actividad economica si se menciona (puede ser distinta a profesion)

3. Para CADA valor monetario detectado:
   - tipo: "valor_compraventa", "valor_hipoteca", "valor_liberacion", "cuota_inicial", "derechos_notariales", "iva", "aporte_superintendencia", "fondo_notariado", "otro"
   - monto_numerico: solo el numero entero (sin $ ni puntos)
   - texto_en_documento: como aparece literalmente ($212.600.000)
   - acto_relacionado: a que acto pertenece (1, 2, 3...) si la escritura tiene multiples actos

4. Para el INMUEBLE (extraer con maxima precision):
   - tipo: tipo en minusculas — "apartamento", "casa", "local", "lote", "parqueadero", "bodega", "oficina", "deposito"
   - numero: numero del apartamento/casa/unidad
   - matricula_inmobiliaria: numero de matricula (formato XXX-XXXXXXX)
   - conjunto_o_edificio: nombre del edificio o conjunto
   - municipio: ciudad donde esta el inmueble
   - departamento: departamento donde esta el inmueble
   - coeficiente_copropiedad: si se menciona
   - cedula_catastral: numero completo tal como aparece (ej: "0881001060001000020000100002"). Buscar etiquetas "CEDULA CATASTRAL:", "Cedula Catastral:", "CODIGO CATASTRAL SEGUN CERTIFICADO"
   - codigo_catastral: codigo alfanumerico CORTO (ej: "AAX0009PUYC", "ABF0003NRDC") que aparece despues de "CODIGO CATASTRAL SEGUN CERTIFICADO DE TRADICION:" o "Codigo catastral:". Es distinto a la cedula catastral (que es numerica larga de 28 digitos). Buscar especificamente la etiqueta "CODIGO CATASTRAL SEGUN CERTIFICADO DE TRADICION" en el documento.
   - area_construida: con unidades exactas (ej: "59.66 M2"). Buscar "Area construida", "area construida interior", "AREA CONSTRUIDA"
   - area_privada: con unidades (ej: "38.39 m2"). Buscar "Area privada", "area privada cubierta"
   - area_total: con unidades. Buscar "area total", "area total construida"
   - piso: planta o piso donde esta (ej: "SEGUNDA PLANTA", "piso catorce", "decimo noveno piso")
   - barrio: nombre del barrio si aparece
   - direccion: direccion completa con calle/carrera y numero
   - nota_linderos: si hay nota introductoria de linderos (ej: "segun planos aprobados")
   - linderos: texto COMPLETO de linderos tal como aparece — NO resumir ni truncar

5. Para la NOTARIA:
   - nombre_notaria: nombre completo
   - municipio_notaria
   - numero_escritura: si esta diligenciado, "PENDIENTE" si esta en blanco

6. Para FECHAS:
   - fecha_otorgamiento: dia, mes, ano

7. Para los ACTOS contenidos en la escritura (puede haber multiples):
   - El PRIMER elemento debe ser el acto PRINCIPAL: el que da nombre a la escritura y tiene cuantia.
   - Actos principales (van primero): COMPRAVENTA, COMPRAVENTA_VIS, HIPOTECA, LIBERACION_HIPOTECA, AFECTACION_VIVIENDA_FAMILIAR, PODER_GENERAL, CAPITULACIONES_MATRIMONIALES, CORRECCION_REGISTRO_CIVIL
   - Actos secundarios (van despues): RENUNCIA_CONDICION_RESOLUTORIA, PROTOCOLIZACION_CTO, PODER_ESPECIAL, CANCELACION_COMODATO, PATRIMONIO_FAMILIA
   - Orden correcto: [acto_principal, ...actos_secundarios]

8. Para DECISIONES / RESPUESTAS del cliente:
   - vivienda_familiar: true/false/null si se menciona
   - patrimonio_familia: true/false/null
   - notificacion_electronica: true/false/null

ESQUEMA JSON:
{
  "personas": [
    {
      "rol": "comprador_1",
      "nombre_completo": "JUAN CAMILO VASQUEZ MIRA",
      "tipo_documento": "C.C",
      "numero_documento": "1.037.657.164",
      "genero": "M",
      "estado_civil": null,
      "nacionalidad": "colombiano",
      "domicilio": "Envigado, Antioquia",
      "direccion": "Calle 45 # 32-12",
      "telefono": "300-1234567",
      "email": "jcamilovasquezm@gmail.com",
      "profesion": "Ingeniero",
      "actividad_economica": "Servicios profesionales"
    }
  ],
  "valores": [
    {
      "tipo": "valor_compraventa",
      "monto_numerico": 212600000,
      "texto_en_documento": "$212.600.000",
      "acto_relacionado": 3
    }
  ],
  "inmueble": {
    "tipo": "apartamento",
    "numero": "804",
    "matricula_inmobiliaria": "001-1528731",
    "conjunto_o_edificio": "CONJUNTO RESIDENCIAL JAGGUA P.H.",
    "municipio": "Caldas",
    "departamento": "Antioquia",
    "coeficiente_copropiedad": null,
    "cedula_catastral": "0881001060001000020000100002",
    "codigo_catastral": "AAX0009PUYC",
    "area_construida": "59.66 M2",
    "area_privada": "38.39 m2",
    "area_total": null,
    "piso": "OCTAVO PISO",
    "barrio": null,
    "direccion": "Carrera 45 # 12-30",
    "nota_linderos": null,
    "linderos": null
  },
  "notaria": {
    "nombre_notaria": "Notaria Unica del Circulo de Caldas",
    "municipio_notaria": "Caldas",
    "numero_escritura": "PENDIENTE"
  },
  "fechas": {
    "fecha_otorgamiento": null
  },
  "actos": ["COMPRAVENTA_VIS", "HIPOTECA", "LIBERACION_HIPOTECA", "PATRIMONIO_FAMILIA", "RENUNCIA_CONDICION_RESOLUTORIA", "PROTOCOLIZACION_CTO", "PODER_ESPECIAL"],
  "decisiones": {
    "vivienda_familiar": null,
    "patrimonio_familia": null,
    "notificacion_electronica": null
  }
}

SE EXHAUSTIVO. Si dudas si algo es dato, INCLUYELO.
NO INVENTES datos: si no aparece en el texto, devuelve null o lista vacia.
"""


# ============================================================
# PROMPT B1: detectar campos vacios en plantilla
# ============================================================

PROMPT_B1 = """Eres experto en escrituras publicas notariales colombianas.

TU TAREA: Te paso una PLANTILLA notarial SIN DILIGENCIAR (con espacios vacios, placeholders como XXX, _____, [campo], o {{variable}}). Debes detectar TODOS los campos que el protocolista debe llenar.

DEVUELVE el mismo esquema JSON que para minutas diligenciadas, pero con valores null en los campos a llenar:

{
  "personas": [
    {
      "rol": "comprador_1",
      "nombre_completo": null,
      "tipo_documento": null,
      "numero_documento": null,
      "genero": null,
      "estado_civil": null,
      "nacionalidad": null,
      "domicilio": null
    }
  ],
  ...
}

REGLAS:
- Si la plantilla tiene 2 compradores marcados (aunque vacios), devuelve 2 personas con rol comprador_1 y comprador_2
- Si solo se menciona 1 vendedor, devuelve 1 persona con rol vendedor
- Para cada placeholder que veas (XXX, ____), identifica de que campo es por el contexto

SOLO JSON VALIDO. Sin explicaciones.
"""


# ============================================================
# FUNCION PRINCIPAL
# ============================================================

def llamar_gpt(client: OpenAI, system_prompt: str, texto: str, modo: str) -> tuple[dict, dict]:
    """Llama a GPT-4o-mini y devuelve el JSON parseado mas info de uso."""
    user_prompt = (
        f"Analiza este texto de minuta notarial colombiana (modo {modo}).\n\n"
        f"TEXTO:\n\"\"\"\n{texto}\n\"\"\"\n\n"
        "Devuelve SOLO el JSON con los datos detectados."
    )
    response = client.chat.completions.create(
        model=MODEL,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        temperature=0.1,
        max_tokens=8000,
        response_format={"type": "json_object"},
    )
    raw = response.choices[0].message.content or "{}"
    usage = response.usage
    return json.loads(raw), {
        "prompt_tokens": usage.prompt_tokens,
        "completion_tokens": usage.completion_tokens,
        "total_tokens": usage.total_tokens,
    }


def calcular_costo(uso: dict) -> float:
    """Calcula costo USD basado en tarifas GPT-4o-mini.
    Tarifas: $0.15 USD por millón IN / $0.60 USD por millón OUT
    """
    return round(
        (uso["prompt_tokens"] / 1_000_000) * 0.15 +
        (uso["completion_tokens"] / 1_000_000) * 0.60,
        4
    )


def analizar_documento(docx_path: str, api_key: str) -> dict:
    """
    Funcion principal: analiza un .docx y devuelve datos estructurados.

    Returns:
        {
            "archivo": str,
            "modo_detectado": "B1" | "B2",
            "confianza_modo": float,
            "datos": { personas, valores, inmueble, notaria, fechas, actos, decisiones },
            "costo_usd": float,
            "tokens": {...},
            "texto_original": str,
        }
    """
    docx_path = Path(docx_path)
    if not docx_path.exists():
        raise FileNotFoundError(f"No existe: {docx_path}")

    # 1. Extraer texto
    texto = extraer_texto_estructurado(str(docx_path))

    # 2. Clasificar modo (heuristica local, sin IA)
    clasificacion = clasificar_modo(texto)
    modo = clasificacion["modo"]

    # 3. Llamar IA con el prompt correcto
    client = _make_openai_client(api_key)
    prompt = PROMPT_B2 if modo == "B2" else PROMPT_B1
    datos, uso = llamar_gpt(client, prompt, texto, modo)

    return {
        "archivo": docx_path.name,
        "modo_detectado": modo,
        "confianza_modo": clasificacion["confianza"],
        "senales_clasificacion": {
            "b2": clasificacion["senales_b2"],
            "b1": clasificacion["senales_b1"],
        },
        "datos": datos,
        "costo_usd": calcular_costo(uso),
        "tokens": uso,
        "texto_original": texto,
        "chars": len(texto),
    }
