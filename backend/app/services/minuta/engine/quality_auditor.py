from __future__ import annotations

from pathlib import Path
import re

from docx import Document

from app.services.minuta.engine.models import RenderAudit, RenderIssue, RenderSeverity
from app.services.minuta.engine.template_analyzer import RESIDUAL_TOKEN_PATTERN, TemplateAnalyzer, iter_document_paragraphs

INCOMPLETE_PHRASE_PATTERNS = (
    re.compile(r"\bproveniente de\s*\.", re.IGNORECASE),
    re.compile(r"\bcon\s*\.", re.IGNORECASE),
    re.compile(r"\bmediante\s*\.", re.IGNORECASE),
    re.compile(r"\bpor valor de\s*\.", re.IGNORECASE),
    re.compile(r"\bidentificad[oa]s?\s+con\s*\.", re.IGNORECASE),
    re.compile(r"\bcon c[ée]dula de ciudadan[íi]a n[úu]mero\s*\.", re.IGNORECASE),
    re.compile(r"^\s*C\.C\.?\s*\.?\s*$", re.IGNORECASE),
)
EMPTY_SIGNATURE_SEQUENCE = re.compile(
    r"(?:_{6,}|-{6,}|\s{6,})\s*\n\s*C\.?C\.?\s*\n\s*Celular\.?\s*\n\s*Direcci[óo]n:?\s*\n\s*Correo:?",
    re.IGNORECASE,
)


class QualityAuditor:
    def audit_post_render(self, output_path: str | Path, audit: RenderAudit) -> RenderAudit:
        analysis = TemplateAnalyzer().analyze(output_path)
        audit.unresolved_placeholders = analysis.unresolved_placeholders
        audit.residual_tokens = sorted(set(analysis.residual_tokens))

        for placeholder in audit.unresolved_placeholders:
            audit.issues.append(
                RenderIssue(
                    code="unresolved_placeholder",
                    message=f"Quedó placeholder sin resolver: {placeholder}",
                    severity=RenderSeverity.BLOCKER,
                    placeholder=placeholder,
                )
            )

        for token in audit.residual_tokens:
            audit.issues.append(
                RenderIssue(
                    code="residual_token",
                    message=f"Quedó token residual: {token}",
                    severity=RenderSeverity.BLOCKER,
                    details={"token": token},
                )
            )

        self._verify_red_replacements(output_path, audit)
        self._verify_structure(audit)
        self._verify_incomplete_phrases(output_path, audit)
        self._verify_empty_signature_blocks(output_path, audit)
        return audit

    def _verify_red_replacements(self, output_path: str | Path, audit: RenderAudit) -> None:
        document = Document(str(output_path))
        red_values = {
            item.inserted_value
            for item in audit.replaced
            if item.inserted_value
        }
        if not red_values:
            return

        found_red: set[str] = set()
        red_runs_detected = 0
        for paragraph, _location in iter_document_paragraphs(document):
            for run in paragraph.runs:
                if run.font.color.rgb and str(run.font.color.rgb).upper() == "FF0000" and run.text.strip():
                    red_runs_detected += 1
                    if run.text in red_values:
                        found_red.add(run.text)
        audit.red_runs_detected = red_runs_detected

        for item in audit.replaced:
            if not item.inserted_value:
                continue
            if item.inserted_value in found_red:
                continue
            audit.issues.append(
                RenderIssue(
                    code="replacement_not_red",
                    message=f"El valor insertado no quedó en rojo: {item.placeholder}",
                    severity=RenderSeverity.ERROR,
                    placeholder=item.placeholder,
                    field_key=item.field_key,
                    location=item.location,
                )
            )

    def _verify_incomplete_phrases(self, output_path: str | Path, audit: RenderAudit) -> None:
        document = Document(str(output_path))
        for paragraph, location in iter_document_paragraphs(document):
            text = paragraph.text or ""
            for pattern in INCOMPLETE_PHRASE_PATTERNS:
                match = pattern.search(text)
                if not match:
                    continue
                audit.issues.append(
                    RenderIssue(
                        code="incomplete_required_phrase",
                        message="Frase incompleta por campo requerido vacío o no mapeado.",
                        severity=RenderSeverity.BLOCKER,
                        location=location,
                        details={"text": match.group(0)},
                    )
                )

    def _verify_empty_signature_blocks(self, output_path: str | Path, audit: RenderAudit) -> None:
        document = Document(str(output_path))
        text = "\n".join(paragraph.text for paragraph, _location in iter_document_paragraphs(document))
        matches = EMPTY_SIGNATURE_SEQUENCE.findall(text)
        audit.empty_signature_blocks_detected = len(matches)
        if not matches:
            return
        audit.issues.append(
            RenderIssue(
                code="empty_signature_block",
                message="Quedó bloque de firma vacío sin persona asociada.",
                severity=RenderSeverity.BLOCKER,
                details={"count": len(matches)},
            )
        )

    def _verify_structure(self, audit: RenderAudit) -> None:
        before = audit.structure_before
        after = audit.structure_after
        if not before or not after:
            return
        if after.get("paragraphs", 0) < max(1, before.get("paragraphs", 0) // 2):
            audit.issues.append(
                RenderIssue(
                    code="structure_changed_too_much",
                    message="El documento generado perdió demasiados párrafos frente al modelo.",
                    severity=RenderSeverity.WARNING,
                    details={"before": before, "after": after},
                )
            )
        if after.get("dash_sequences", 0) < before.get("dash_sequences", 0):
            audit.issues.append(
                RenderIssue(
                    code="notarial_dash_sequences_lost",
                    message="El documento generado perdiÃ³ secuencias de guiones notariales frente al modelo.",
                    severity=RenderSeverity.WARNING,
                    details={"before": before, "after": after},
                )
            )
