from __future__ import annotations

import re
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import RGBColor
from docx.text.run import Run

from app.services.minuta.engine.models import PlaceholderAudit, RenderAudit, RenderIssue, RenderSeverity
from app.services.minuta.engine.template_analyzer import iter_document_paragraphs
from app.services.minuta.rules.common_rules import normalize_key
from app.services.minuta.rules.conditional_blocks import should_remove_si_aplica_block
from app.services.minuta.rules.gender_rules import normalize_gender
from app.services.minuta.rules.person_rules import collective_adjective, collective_label, grammar_for_gender
from app.services.minuta.rules.phrase_rules import normalize_notarial_phrase

RED_COLOR = RGBColor(0xFF, 0x00, 0x00)
TECHNICAL_TAB_TOKEN = "[[--]]"
AUXILIARY_TOKEN_PATTERN = re.compile(r"\[SI APLICA[^\]]*\]|\[NO APLICA[^\]]*\]|\[N[úu]mero\]|\[Informaci[óo]n\]", re.IGNORECASE)
DASH_SEQUENCE_PATTERN = re.compile(r"(?:-\s*){3,}|-{3,}")
SIGNATURE_LINE_PATTERN = re.compile(r"^[\s._-]{6,}$")
EMPTY_SIGNATURE_LABEL_PATTERN = re.compile(
    r"^(?:C\.?C\.?|CELULAR\.?|DIRECCI[ÓO]N:?\s*|CORREO:?\s*|PROFESI[ÓO]N\s+U\s+OFICIO:?\s*|ACTIVIDAD\s+ECON[ÓO]MICA:?\s*|ESTADO\s+CIVIL:?\s*)\.?$",
    re.IGNORECASE,
)
OPTIONAL_EMPTY_SEGMENTS = (
    {
        "field_key": "origen_cuota_inicial",
        "message": "Se omitió el segmento asociado a Origen cuota inicial porque no fue diligenciado.",
        "patterns": (
            re.compile(r",\s*proveniente de\s*\.", re.IGNORECASE),
            re.compile(r"\s+proveniente de\s*\.", re.IGNORECASE),
        ),
        "replacement": ".",
    },
    {
        "field_key": "optional_payment_detail",
        "message": "Se omitió un segmento descriptivo de forma de pago porque no fue diligenciado.",
        "forbidden_prefix": re.compile(r"(identificad|c[eé]dula|ciudadan[ií]a|documento|c\.?c\.?)", re.IGNORECASE),
        "patterns": (
            re.compile(r"\s+con\s*\.", re.IGNORECASE),
        ),
        "replacement": ".",
    },
)


def copy_run_format(source_run, target_run) -> None:
    if source_run is None or target_run is None:
        return
    if source_run._r.rPr is not None:
        target_run._r.insert(0, deepcopy(source_run._r.rPr))


def insert_run_after(paragraph, anchor_run, text: str, template_run=None, color: RGBColor | None = None):
    new_r = OxmlElement("w:r")
    anchor_run._r.addnext(new_r)
    new_run = Run(new_r, paragraph)
    if template_run is not None:
        copy_run_format(template_run, new_run)
    new_run.text = text
    if color is not None and text:
        new_run.font.color.rgb = color
    return new_run


def set_paragraph_text_preserving_first_run(paragraph, new_text: str) -> None:
    runs = list(paragraph.runs)
    if not runs:
        if new_text:
            paragraph.add_run(new_text)
        return
    runs[0].text = new_text
    for run in runs[1:]:
        run.text = ""


class DocxRenderer:
    def render(
        self,
        source_path: str | Path,
        destination_path: str | Path,
        marker_values: dict[str, str],
        marker_keys: dict[str, str | None],
        normalized_values: dict[str, str],
    ) -> RenderAudit:
        document = Document(str(source_path))
        audit = RenderAudit(structure_before=self._structure(document))

        for marker in sorted(marker_values, key=len, reverse=True):
            value = marker_values[marker]
            field_key = marker_keys.get(marker)
            marker_seen = False
            for paragraph, location in iter_document_paragraphs(document):
                audit_items = self._replace_marker_in_paragraph(paragraph, marker, value, field_key, location)
                if audit_items:
                    marker_seen = True
                for item in audit_items:
                    if not item.inserted_value:
                        audit.empty.append(item)
                    elif item.replaced:
                        audit.replaced.append(item)
                    else:
                        audit.missing.append(item)
            if not marker_seen:
                audit.missing.append(
                    PlaceholderAudit(
                        placeholder=marker,
                        field_key=field_key,
                        inserted_value=value,
                        location="document",
                        color_applied=None,
                        style_preserved=False,
                        replaced=False,
                    )
                )

        self._remove_conditional_blocks(document, normalized_values)
        audit.technical_tabs_resolved = self._resolve_technical_tabs(document)
        self._remove_auxiliary_tokens(document)
        self._cleanup_optional_empty_segments(document, audit)
        self._normalize_known_gender_literals(document, normalized_values)
        self._normalize_composed_phrases(document)
        audit.empty_signature_blocks_removed = self._remove_empty_signature_blocks(document)
        audit.structure_after = self._structure(document)
        audit.notarial_dash_sequences_preserved = min(
            audit.structure_before.get("dash_sequences", 0),
            audit.structure_after.get("dash_sequences", 0),
        )
        document.save(str(destination_path))
        return audit

    def _replace_marker_in_paragraph(
        self,
        paragraph,
        marker: str,
        value: str,
        field_key: str | None,
        location: str,
        color: RGBColor | None = RED_COLOR,
    ) -> list[PlaceholderAudit]:
        text = paragraph.text or ""
        matches = list(re.finditer(re.escape(marker), text))
        if not matches:
            return []
        audits: list[PlaceholderAudit] = []
        for match in reversed(matches):
            affected = self._affected_runs(paragraph, match.start(), match.end())
            if not affected:
                continue
            self._replace_affected_runs(paragraph, affected, match.start(), match.end(), value, color)
            audits.append(
                PlaceholderAudit(
                    placeholder=marker,
                    field_key=field_key,
                    inserted_value=value,
                    location=location,
                    color_applied="FF0000" if value and color is not None else None,
                    style_preserved=True,
                    replaced=True,
                )
            )
        return audits

    def _replace_first_marker_in_paragraph(self, paragraph, marker: str, value: str) -> bool:
        text = paragraph.text or ""
        match = re.search(re.escape(marker), text)
        if not match:
            return False
        affected = self._affected_runs(paragraph, match.start(), match.end())
        if not affected:
            return False
        self._replace_affected_runs(paragraph, affected, match.start(), match.end(), value, None)
        return True

    def _affected_runs(self, paragraph, start: int, end: int) -> list[dict]:
        runs_info: list[dict] = []
        position = 0
        for run in paragraph.runs:
            run_end = position + len(run.text)
            runs_info.append({"run": run, "start": position, "end": run_end})
            position = run_end
        affected: list[dict] = []
        for info in runs_info:
            if info["end"] <= start:
                continue
            if info["start"] >= end:
                break
            affected.append(info)
        return affected

    def _replace_affected_runs(
        self,
        paragraph,
        affected: list[dict],
        start: int,
        end: int,
        value: str,
        color: RGBColor | None = RED_COLOR,
    ) -> None:
        first = affected[0]
        last = affected[-1]
        first_run = first["run"]
        last_run = last["run"]
        prefix = first_run.text[: start - first["start"]]
        suffix = last_run.text[end - last["start"] :]
        first_run.text = prefix
        for info in affected[1:]:
            info["run"].text = ""
        inserted = insert_run_after(paragraph, first_run, value, template_run=first_run, color=color)
        if suffix:
            insert_run_after(paragraph, inserted, suffix, template_run=last_run)

    def _remove_conditional_blocks(self, document: Document, normalized_values: dict[str, str]) -> None:
        for paragraph, _location in iter_document_paragraphs(document):
            if should_remove_si_aplica_block(normalized_values, paragraph.text or ""):
                set_paragraph_text_preserving_first_run(paragraph, "")

    def _resolve_technical_tabs(self, document: Document) -> int:
        resolved = 0
        for paragraph, _location in iter_document_paragraphs(document):
            audit_items = self._replace_marker_in_paragraph(paragraph, TECHNICAL_TAB_TOKEN, "\t", None, "technical-tab", color=None)
            resolved += len(audit_items)
        return resolved

    def _remove_auxiliary_tokens(self, document: Document) -> None:
        for paragraph, _location in iter_document_paragraphs(document):
            tokens = sorted(set(AUXILIARY_TOKEN_PATTERN.findall(paragraph.text or "")), key=len, reverse=True)
            for token in tokens:
                self._replace_marker_in_paragraph(paragraph, token, "", None, "cleanup", color=None)

    def _cleanup_optional_empty_segments(self, document: Document, audit: RenderAudit) -> None:
        omitted_keys: set[str] = set()
        for paragraph, _location in iter_document_paragraphs(document):
            for config in OPTIONAL_EMPTY_SEGMENTS:
                for pattern in config["patterns"]:
                    if self._replace_optional_empty_pattern(
                        paragraph,
                        pattern,
                        str(config["replacement"]),
                        config.get("forbidden_prefix"),
                    ):
                        omitted_keys.add(str(config["field_key"]))
        for config in OPTIONAL_EMPTY_SEGMENTS:
            field_key = str(config["field_key"])
            if field_key not in omitted_keys:
                continue
            audit.optional_segments_omitted += 1
            audit.optional_segments_omitted_keys.append(field_key)
            audit.issues.append(
                RenderIssue(
                    code="optional_field_omitted",
                    message=str(config["message"]),
                    severity=RenderSeverity.WARNING,
                    field_key=field_key,
                )
            )

    def _replace_optional_empty_pattern(self, paragraph, pattern: re.Pattern, replacement: str, forbidden_prefix: object = None) -> bool:
        text = paragraph.text or ""
        if not text:
            return False
        forbidden = forbidden_prefix if isinstance(forbidden_prefix, re.Pattern) else None
        replaced = False
        for match in reversed(list(pattern.finditer(text))):
            if forbidden and forbidden.search(text[max(0, match.start() - 40) : match.start()]):
                continue
            affected = self._affected_runs(paragraph, match.start(), match.end())
            if not affected:
                continue
            self._replace_affected_runs(paragraph, affected, match.start(), match.end(), replacement, color=None)
            replaced = True
        return replaced

    def _normalize_known_gender_literals(self, document: Document, normalized_values: dict[str, str]) -> None:
        buyer_1 = bool(normalized_values.get("nombre_comprador_1") or normalized_values.get("nombre_comprador"))
        buyer_2 = bool(normalized_values.get("nombre_comprador_2"))
        gender_1 = normalize_gender(normalized_values.get("comprador_es_hombre_o_mujer") or normalized_values.get("comprador_1_es_hombre_o_mujer"))
        gender_2 = normalize_gender(normalized_values.get("comprador_2_es_hombre_o_mujer"))
        genders = [gender for present, gender in ((buyer_1, gender_1), (buyer_2, gender_2)) if present]
        individual_1 = grammar_for_gender(normalized_values.get("comprador_es_hombre_o_mujer") or normalized_values.get("comprador_1_es_hombre_o_mujer"))
        individual_2 = grammar_for_gender(normalized_values.get("comprador_2_es_hombre_o_mujer"))

        for paragraph, _location in iter_document_paragraphs(document):
            text = paragraph.text or ""
            if buyer_1 and buyer_2 and normalized_values.get("nombre_comprador_1", "") in text and normalized_values.get("nombre_comprador_2", "") in text:
                for marker, first, second in (
                    ("domiciliado(a)", individual_1.domiciled, individual_2.domiciled),
                    ("identificado(a/s)", individual_1.identified, individual_2.identified),
                    ("identificado(a)", individual_1.identified, individual_2.identified),
                ):
                    if self._replace_first_marker_in_paragraph(paragraph, marker, first):
                        self._replace_first_marker_in_paragraph(paragraph, marker, second)
            else:
                self._replace_marker_in_paragraph(paragraph, "domiciliado(a)", collective_adjective("domiciliado", genders), None, "gender-literal-cleanup", color=None)
                self._replace_marker_in_paragraph(paragraph, "identificado(a/s)", collective_adjective("identificado", genders), None, "gender-literal-cleanup", color=None)
                self._replace_marker_in_paragraph(paragraph, "identificado(a)", collective_adjective("identificado", genders), None, "gender-literal-cleanup", color=None)

            replacements = {
                "COMPRADOR(A/ES)": collective_label("comprador", genders).upper(),
                "comprador(a)(es)": collective_label("comprador", genders).replace("el ", "").replace("la ", "").replace("los ", "").replace("las ", ""),
                "EL COMPRADOR": collective_label("comprador", genders).upper(),
                "EL(LOS) HIPOTECANTE(S)": collective_label("hipotecante", genders).upper(),
                "EL(LOS) DEUDOR(ES)": collective_label("deudor", genders).upper(),
                "de nacionalidad colombiana": "colombiana",
                "de nacionalidad colombiano": "colombiano",
                "de nacionalidad colombianas": "colombianas",
                "de nacionalidad colombianos": "colombianos",
                " quien en adelante se denominará(n)": " quienes en adelante se denominarán",
            }
            for marker, replacement in replacements.items():
                self._replace_marker_in_paragraph(paragraph, marker, replacement, None, "collective-literal-cleanup", color=None)
            while ",  " in (paragraph.text or ""):
                self._replace_first_marker_in_paragraph(paragraph, ",  ", ", ")

    def _normalize_composed_phrases(self, document: Document) -> None:
        for paragraph, _location in iter_document_paragraphs(document):
            text = paragraph.text or ""
            normalized = normalize_notarial_phrase(text)
            if normalized != text:
                set_paragraph_text_preserving_first_run(paragraph, normalized)

    def _remove_empty_signature_blocks(self, document: Document) -> int:
        paragraphs = [paragraph for paragraph, _location in iter_document_paragraphs(document)]
        removed = 0
        for index, paragraph in enumerate(paragraphs):
            if not SIGNATURE_LINE_PATTERN.match((paragraph.text or "").strip()):
                continue
            block = paragraphs[index + 1 : index + 8]
            raw_block = [(item.text or "").strip() for item in block]
            has_name_or_document = any(re.search(r"[A-ZÁÉÍÓÚÑ]{3,}.*\d{4,}", raw, re.IGNORECASE) for raw in raw_block[:2])
            has_empty_labels = sum(1 for raw in raw_block if EMPTY_SIGNATURE_LABEL_PATTERN.match(raw)) >= 3
            if has_name_or_document or not has_empty_labels:
                continue
            set_paragraph_text_preserving_first_run(paragraph, "")
            for candidate in block:
                raw = (candidate.text or "").strip()
                if not raw or EMPTY_SIGNATURE_LABEL_PATTERN.match(raw):
                    set_paragraph_text_preserving_first_run(candidate, "")
            removed += 1
        return removed

    def _structure(self, document: Document) -> dict[str, int]:
        dash_sequences = 0
        for paragraph, _location in iter_document_paragraphs(document):
            dash_sequences += len(DASH_SEQUENCE_PATTERN.findall(paragraph.text or ""))
        return {
            "paragraphs": len(document.paragraphs),
            "tables": len(document.tables),
            "sections": len(document.sections),
            "dash_sequences": dash_sequences,
        }
