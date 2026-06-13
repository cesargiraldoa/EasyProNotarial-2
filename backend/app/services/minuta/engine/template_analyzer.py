from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable

from docx import Document

from app.services.minuta.engine.models import TemplateAnalysis, TemplatePlaceholder
from app.services.minuta.rules.common_rules import normalize_key

PLACEHOLDER_PATTERN = re.compile(r"(\{\{[^{}]+\}\}|\[[^\[\]]+\])")
RESIDUAL_TOKEN_PATTERN = re.compile(r"(\[\[--\]\]|\[N[úu]mero\]|\[Informaci[óo]n\]|\[SI APLICA[^\]]*\])", re.IGNORECASE)


def iter_paragraphs_in_table(table) -> Iterable:
    seen_cells: set[int] = set()
    for row_index, row in enumerate(table.rows):
        for cell_index, cell in enumerate(row.cells):
            cell_id = id(cell._tc)
            if cell_id in seen_cells:
                continue
            seen_cells.add(cell_id)
            for paragraph_index, paragraph in enumerate(cell.paragraphs):
                yield paragraph, f"table:r{row_index}:c{cell_index}:p{paragraph_index}"
            for nested_index, nested_table in enumerate(cell.tables):
                for paragraph, location in iter_paragraphs_in_table(nested_table):
                    yield paragraph, f"table:r{row_index}:c{cell_index}:nested{nested_index}:{location}"


def iter_paragraphs_in_container(container, prefix: str) -> Iterable:
    for index, paragraph in enumerate(container.paragraphs):
        yield paragraph, f"{prefix}:p{index}"
    for table_index, table in enumerate(container.tables):
        for paragraph, location in iter_paragraphs_in_table(table):
            yield paragraph, f"{prefix}:t{table_index}:{location}"


def iter_document_paragraphs(document: Document) -> Iterable:
    yield from iter_paragraphs_in_container(document, "body")
    for section_index, section in enumerate(document.sections):
        yield from iter_paragraphs_in_container(section.header, f"header:{section_index}")
        yield from iter_paragraphs_in_container(section.footer, f"footer:{section_index}")


class TemplateAnalyzer:
    def analyze(self, docx_path: str | Path, fields: list[dict] | None = None) -> TemplateAnalysis:
        document = Document(str(docx_path))
        marker_to_key = self._marker_to_key(fields or [])
        analysis = TemplateAnalysis(structure=self._structure(document))

        for paragraph, location in iter_document_paragraphs(document):
            text = paragraph.text or ""
            for match in PLACEHOLDER_PATTERN.finditer(text):
                marker = match.group(1)
                field_key = marker_to_key.get(marker)
                analysis.placeholders.append(TemplatePlaceholder(marker=marker, field_key=field_key, location=location))
                normalized_key = normalize_key(field_key or marker)
                if any(role in normalized_key for role in ("comprador_2", "vendedor_2")):
                    role = "comprador_2" if "comprador_2" in normalized_key else "vendedor_2"
                    analysis.required_actor_roles.add(role)
            analysis.residual_tokens.extend(match.group(1) for match in RESIDUAL_TOKEN_PATTERN.finditer(text))

        analysis.unresolved_placeholders = sorted({item.marker for item in analysis.placeholders})
        return analysis

    def _marker_to_key(self, fields: list[dict]) -> dict[str, str]:
        mapping: dict[str, str] = {}
        for field in fields:
            if not isinstance(field, dict):
                continue
            key = str(field.get("key") or "").strip()
            raw_markers = field.get("raw_markers") or []
            if not key or not isinstance(raw_markers, list):
                continue
            for marker in raw_markers:
                marker_text = str(marker or "").strip()
                if marker_text:
                    mapping[marker_text] = key
        return mapping

    def _structure(self, document: Document) -> dict[str, int]:
        table_count = len(document.tables)
        header_paragraphs = sum(len(section.header.paragraphs) for section in document.sections)
        footer_paragraphs = sum(len(section.footer.paragraphs) for section in document.sections)
        return {
            "paragraphs": len(document.paragraphs),
            "tables": table_count,
            "sections": len(document.sections),
            "header_paragraphs": header_paragraphs,
            "footer_paragraphs": footer_paragraphs,
        }
