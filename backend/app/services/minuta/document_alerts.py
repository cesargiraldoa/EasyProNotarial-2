from __future__ import annotations

import re
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Iterable

from docx import Document

from app.services.minuta.rules.common_rules import normalize_key, normalize_value
from app.services.minuta.rules.phrase_rules import has_duplicate_notarial_phrase


@dataclass(frozen=True)
class DocumentAlert:
    code: str
    message: str
    severity: str = "warning"
    field_key: str | None = None
    location: str | None = None
    value: str | None = None
    details: dict | None = None

    def to_dict(self) -> dict:
        return asdict(self)


LIVE_TEXT_PATTERNS: tuple[tuple[str, re.Pattern[str]], ...] = (
    ("pending_literal", re.compile(r"\bPendiente\b", re.IGNORECASE)),
    ("placeholder_xxxx", re.compile(r"\bX{4,}\b", re.IGNORECASE)),
    ("partner_name_placeholder", re.compile(r"\bNOMBRE\s+PAREJA\b", re.IGNORECASE)),
    ("spouse_name_placeholder", re.compile(r"\bNOMBRE\s+C[OÓ]NYUGE\b", re.IGNORECASE)),
    ("document_number_placeholder", re.compile(r"\bN[UÚ]MERO\s+DE\s+DOCUMENTO\b", re.IGNORECASE)),
    ("buyer_literal_placeholder", re.compile(r"\bCOMPRADOR\(A\)\b", re.IGNORECASE)),
    ("seller_literal_placeholder", re.compile(r"\bVENDEDOR\(A\)\b", re.IGNORECASE)),
    ("article_literal_placeholder", re.compile(r"\b(?:EL\(LOS\)|LA\(S\))\b", re.IGNORECASE)),
)

MATRICULA_PATTERN = re.compile(r"\b\d{2,3}[A-Z]?-\d{4,}\b", re.IGNORECASE)


def document_text(docx_path: str | Path) -> list[tuple[str, str]]:
    document = Document(str(docx_path))
    paragraphs: list[tuple[str, str]] = []
    for index, paragraph in enumerate(document.paragraphs):
        paragraphs.append((paragraph.text or "", f"body:p{index}"))
    for table_index, table in enumerate(document.tables):
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                for paragraph_index, paragraph in enumerate(cell.paragraphs):
                    paragraphs.append((paragraph.text or "", f"body:t{table_index}:r{row_index}:c{cell_index}:p{paragraph_index}"))
    for section_index, section in enumerate(document.sections):
        for prefix, container in (("header", section.header), ("footer", section.footer)):
            for paragraph_index, paragraph in enumerate(container.paragraphs):
                paragraphs.append((paragraph.text or "", f"{prefix}:{section_index}:p{paragraph_index}"))
    return paragraphs


def detect_live_text_alerts(paragraphs: Iterable[tuple[str, str]]) -> list[DocumentAlert]:
    alerts: list[DocumentAlert] = []
    seen: set[tuple[str, str, str]] = set()
    for text, location in paragraphs:
        for code, pattern in LIVE_TEXT_PATTERNS:
            for match in pattern.finditer(text or ""):
                key = (code, match.group(0), location)
                if key in seen:
                    continue
                seen.add(key)
                alerts.append(
                    DocumentAlert(
                        code=code,
                        message=f"Quedó texto pendiente por revisar: {match.group(0)}",
                        location=location,
                        value=match.group(0),
                    )
                )
        if has_duplicate_notarial_phrase(text):
            alerts.append(
                DocumentAlert(
                    code="duplicated_phrase",
                    message="Se detectó una frase duplicada o mal compuesta en el documento.",
                    location=location,
                    value=text[:220],
                )
            )
    return alerts


def canonical_matricula(values: dict[str, object]) -> str:
    preferred_keys = (
        "matricula_inmobiliaria",
        "numero_matricula",
        "matricula",
        "matricula_principal",
        "folio_matricula",
    )
    normalized = {normalize_key(key): normalize_value(value) for key, value in (values or {}).items()}
    for key in preferred_keys:
        value = normalized.get(key, "")
        if value:
            return value.upper()
    for key, value in normalized.items():
        if "matricula" in key and value:
            return value.upper()
    return ""


def is_composite_case(values: dict[str, object], case_acts: list[dict] | None = None) -> bool:
    acts = case_acts or []
    if len(acts) > 1:
        return True
    joined = " ".join(
        [
            normalize_value(values.get("variante_id")),
            normalize_value(values.get("proyecto")),
            normalize_value(values.get("actos")),
            normalize_value(values.get("subactos")),
        ]
    )
    normalized = normalize_key(joined)
    return any(token in normalized for token in ("jaggua", "aragua", "hipoteca", "liberacion", "patrimonio_familia", "comodato", "cto"))


def detect_matricula_alerts(docx_path: str | Path, values: dict[str, object], case_acts: list[dict] | None = None) -> list[DocumentAlert]:
    canonical = canonical_matricula(values)
    if not canonical:
        return [
            DocumentAlert(
                code="canonical_matricula_missing",
                message="No hay matrícula inmobiliaria canónica para validar el documento.",
                field_key="matricula_inmobiliaria",
            )
        ]
    found: set[str] = set()
    for text, _location in document_text(docx_path):
        found.update(match.group(0).upper() for match in MATRICULA_PATTERN.finditer(text or ""))
    mismatches = sorted(item for item in found if item != canonical)
    if not mismatches or is_composite_case(values, case_acts):
        return []
    return [
        DocumentAlert(
            code="matricula_mismatch",
            message="La matrícula del documento no coincide con la matrícula canónica del inmueble principal.",
            field_key="matricula_inmobiliaria",
            value=canonical,
            details={"canonical": canonical, "found": mismatches},
        )
    ]


def alerts_to_dicts(alerts: Iterable[DocumentAlert]) -> list[dict]:
    return [alert.to_dict() for alert in alerts]
