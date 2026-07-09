from __future__ import annotations

from pathlib import Path

from docx import Document

from app.services.minuta.assisted_tagging.models import DocxStructure, DocxTextBlock
from app.services.minuta.engine.template_analyzer import iter_document_paragraphs


class DocxStructureExtractor:
    def extract(self, docx_path: str | Path) -> DocxStructure:
        path = Path(docx_path)
        document = Document(str(path))
        blocks: list[DocxTextBlock] = []

        for index, (paragraph, location) in enumerate(iter_document_paragraphs(document), start=1):
            text = " ".join((paragraph.text or "").split())
            if not text:
                continue
            blocks.append(
                DocxTextBlock(
                    id=f"b{index}",
                    location=location,
                    text=text,
                    char_count=len(text),
                )
            )

        return DocxStructure(
            blocks=blocks,
            filename=path.name,
            paragraph_count=len(document.paragraphs),
            table_count=len(document.tables),
        )
