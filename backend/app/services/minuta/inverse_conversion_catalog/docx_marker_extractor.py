from __future__ import annotations

import re
from pathlib import Path

from docx import Document

from app.services.minuta.inverse_conversion_catalog.models import MarkerLocation, MarkerOccurrence


MARKER_RE = re.compile(r"\{\{\s*([^{}]+?)\s*\}\}")


class DocxMarkerExtractor:
    """Read-only extractor for human-authored {{FIELD}} markers in DOCX files."""

    def __init__(self, context_chars: int = 120) -> None:
        self.context_chars = context_chars

    def extract_from_docx(self, path: str | Path) -> list[MarkerOccurrence]:
        document = Document(str(path))
        occurrences: list[MarkerOccurrence] = []

        for paragraph_index, paragraph in enumerate(document.paragraphs):
            occurrences.extend(
                self.extract_from_text(
                    paragraph.text,
                    MarkerLocation(paragraph_index=paragraph_index),
                )
            )

        for table_index, table in enumerate(document.tables):
            for row_index, row in enumerate(table.rows):
                for cell_index, cell in enumerate(row.cells):
                    for paragraph_index, paragraph in enumerate(cell.paragraphs):
                        occurrences.extend(
                            self.extract_from_text(
                                paragraph.text,
                                MarkerLocation(
                                    paragraph_index=paragraph_index,
                                    table_index=table_index,
                                    row_index=row_index,
                                    cell_index=cell_index,
                                ),
                            )
                        )

        return occurrences

    def extract_from_text(self, text: str, location: MarkerLocation | None = None) -> list[MarkerOccurrence]:
        if location is None:
            location = MarkerLocation()

        occurrences: list[MarkerOccurrence] = []
        for match in MARKER_RE.finditer(text or ""):
            raw_marker = match.group(0)
            raw_field_code = match.group(1).strip()
            before_start = max(0, match.start() - self.context_chars)
            after_end = min(len(text), match.end() + self.context_chars)
            occurrences.append(
                MarkerOccurrence(
                    raw_marker=raw_marker,
                    raw_field_code=raw_field_code,
                    text=text,
                    text_before=self._clean_context(text[before_start : match.start()]),
                    text_after=self._clean_context(text[match.end() : after_end]),
                    location=location,
                    start_index=match.start(),
                    end_index=match.end(),
                )
            )
        return occurrences

    @staticmethod
    def _clean_context(value: str) -> str:
        return re.sub(r"\s+", " ", value or "").strip()
