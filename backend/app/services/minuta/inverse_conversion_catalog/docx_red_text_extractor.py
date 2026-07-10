from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import RGBColor

from app.services.minuta.inverse_conversion_catalog.models import MarkerLocation, RedTextOccurrence


class DocxRedTextExtractor:
    """Extract red run text without modifying the source DOCX."""

    def extract_from_docx(self, path: str | Path) -> list[RedTextOccurrence]:
        document = Document(str(path))
        occurrences: list[RedTextOccurrence] = []

        for paragraph_index, paragraph in enumerate(document.paragraphs):
            occurrences.extend(self._extract_from_paragraph(paragraph, MarkerLocation(paragraph_index=paragraph_index)))

        for table_index, table in enumerate(document.tables):
            for row_index, row in enumerate(table.rows):
                for cell_index, cell in enumerate(row.cells):
                    for paragraph_index, paragraph in enumerate(cell.paragraphs):
                        occurrences.extend(
                            self._extract_from_paragraph(
                                paragraph,
                                MarkerLocation(
                                    paragraph_index=paragraph_index,
                                    table_index=table_index,
                                    row_index=row_index,
                                    cell_index=cell_index,
                                ),
                            )
                        )

        return [occurrence for occurrence in occurrences if occurrence.text.strip()]

    def _extract_from_paragraph(self, paragraph, location: MarkerLocation) -> list[RedTextOccurrence]:
        return [
            RedTextOccurrence(text=run.text, location=location)
            for run in paragraph.runs
            if run.text and self._is_red(run.font.color.rgb)
        ]

    @staticmethod
    def _is_red(color: RGBColor | None) -> bool:
        if color is None:
            return False
        return color[0] >= 0xC0 and color[1] <= 0x50 and color[2] <= 0x50
