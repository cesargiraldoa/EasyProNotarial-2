from __future__ import annotations

import io
import os
from pathlib import Path

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from openai import OpenAI
from supabase import Client, create_client

from app.core.config import get_settings


SYSTEM_PROMPT_GARI = """Eres Gari, motor de redacción notarial colombiano de EasyPro.

IDENTIDAD Y ROL:
- Redactas instrumentos públicos notariales en Colombia con precisión jurídica total.
- Sigues el Decreto 960 de 1970, Decreto 2148 de 1983 y demás normas notariales colombianas.
- Solo produces el texto del documento. Nunca incluyes explicaciones, comentarios ni metadata.

REGLAS DE FORMATO NOTARIAL OBLIGATORIAS:
- Usa guiones para llenar espacios: - - - - - - - - - - - - - - - - - - - - - - - - - - -
- Números siempre en texto seguido del numeral: "diecinueve (19)", "dos mil veintiséis (2026)"
- Valores monetarios: "$6.000.000" Y "seis millones de pesos colombianos ($6.000.000)"
- Negrilla para títulos de actos: **PRIMER ACTO: LIBERACIÓN PARCIAL DE HIPOTECA**
- Cada acto inicia en nueva sección con su número ordinal en negrilla
- Los intervinientes se presentan con su calidad completa: "quien obra en calidad de apoderado especial de [ENTIDAD] identificada con NIT [NIT]"

ESTRUCTURA OBLIGATORIA DEL DOCUMENTO:
1. Encabezado: NOTARÍA + NÚMERO ESCRITURA + CLASE Y CUANTÍA DE ACTOS + PERSONAS QUE INTERVIENEN
2. Apertura: ciudad, fecha en texto, notario titular
3. Un bloque por cada acto en el orden exacto indicado en el prompt
4. Cada acto: comparecencia del interviniente  declaraciones  aceptación
5. Liquidación de derechos notariales
6. Constancias legales (Ley 1581/2012, Art. 102 Decreto 960/1970)
7. Firmas de todos los comparecientes + Notario

REGLAS DE INTERVINIENTES:
- Cada apoderado SIEMPRE menciona: "quien obra como apoderado especial de [ENTIDAD, NIT]"
- La personería se acredita con: "escritura pública número X de la Notaría Y de [ciudad], la cual se protocoliza"
- El género gramatical SIEMPRE debe concordar con el sexo del interviniente
- Si el interviniente está de tránsito: "domiciliado en [municipio], de tránsito por Caldas"

MODO CORRECCIÓN:
- Cuando recibes un borrador anterior + instrucción de corrección
- Aplica SOLO el cambio solicitado
- Reproduce el resto del documento sin alteraciones
- No resumas ni acortes el documento

PROHIBICIONES ABSOLUTAS:
- Nunca inventes datos que no estén en el prompt
- Nunca uses placeholders como [DATO] o {{VARIABLE}}
- Nunca agregues comentarios fuera del texto notarial
- Nunca omitas actos que estén en la lista de actos requeridos
"""


def get_supabase_client() -> Client:
    url = os.environ.get("SUPABASE_URL", "")
    key = os.environ.get("SUPABASE_SERVICE_KEY", "")
    if not url or not key:
        raise ValueError("SUPABASE_URL y SUPABASE_SERVICE_KEY requeridos")
    return create_client(url, key)


def get_openai_client() -> OpenAI:
    settings = get_settings()
    api_key = getattr(settings, "openai_api_key", "") or os.environ.get("OPENAI_API_KEY")
    if not api_key:
        raise ValueError("OPENAI_API_KEY no configurada")
    return OpenAI(api_key=api_key)


def build_gari_prompt(
    act_type: str,
    notary_label: str,
    notary_name: str,
    participants: list[dict],
    act_data: dict,
    template_reference_text: str | None = None,
    variante_id: str | None = None,
    correction_note: str | None = None,
    previous_draft: str | None = None,
) -> str:
    """Construye el prompt para Gari con todos los datos del acto notarial."""
    ORDEN_ACTOS_POR_VARIANTE = {
        "aragua_apto_1c": [
            "LIBERACIÓN PARCIAL DE HIPOTECA",
            "PROTOCOLIZACIÓN CERTIFICADO TÉCNICO DE OCUPACIÓN",
            "COMPRAVENTA VIS",
            "RENUNCIA A CONDICIÓN RESOLUTORIA",
            "CANCELACIÓN DE COMODATO",
            "CONSTITUCIÓN DE PATRIMONIO DE FAMILIA INEMBARGABLE",
            "PODER ESPECIAL",
        ],
        "aragua_apto_2c": [
            "LIBERACIÓN PARCIAL DE HIPOTECA",
            "PROTOCOLIZACIÓN CERTIFICADO TÉCNICO DE OCUPACIÓN",
            "COMPRAVENTA VIS",
            "RENUNCIA A CONDICIÓN RESOLUTORIA",
            "CANCELACIÓN DE COMODATO",
            "CONSTITUCIÓN DE PATRIMONIO DE FAMILIA INEMBARGABLE",
            "PODER ESPECIAL",
        ],
        "aragua_parq_2c": [
            "LIBERACIÓN PARCIAL DE HIPOTECA",
            "PROTOCOLIZACIÓN CERTIFICADO TÉCNICO DE OCUPACIÓN",
            "COMPRAVENTA VIS",
            "RENUNCIA A CONDICIÓN RESOLUTORIA",
            "CANCELACIÓN DE COMODATO",
            "PODER ESPECIAL",
        ],
        "aragua_parq_3c": [
            "LIBERACIÓN PARCIAL DE HIPOTECA",
            "PROTOCOLIZACIÓN CERTIFICADO TÉCNICO DE OCUPACIÓN",
            "COMPRAVENTA VIS",
            "RENUNCIA A CONDICIÓN RESOLUTORIA",
            "CANCELACIÓN DE COMODATO",
            "PODER ESPECIAL",
        ],
        "jaggua_fna_1c": [
            "LIBERACIÓN PARCIAL DE HIPOTECA",
            "PROTOCOLIZACIÓN CERTIFICADO TÉCNICO DE OCUPACIÓN",
            "COMPRAVENTA VIS",
            "RENUNCIA A CONDICIÓN RESOLUTORIA",
            "CANCELACIÓN DE COMODATO",
            "CONSTITUCIÓN DE HIPOTECA ABIERTA DE PRIMER GRADO",
            "CONSTITUCIÓN DE PATRIMONIO DE FAMILIA INEMBARGABLE",
            "PODER ESPECIAL",
        ],
        "jaggua_fna_2c": [
            "LIBERACIÓN PARCIAL DE HIPOTECA",
            "PROTOCOLIZACIÓN CERTIFICADO TÉCNICO DE OCUPACIÓN",
            "COMPRAVENTA VIS",
            "RENUNCIA A CONDICIÓN RESOLUTORIA",
            "CANCELACIÓN DE COMODATO",
            "CONSTITUCIÓN DE HIPOTECA ABIERTA DE PRIMER GRADO",
            "CONSTITUCIÓN DE PATRIMONIO DE FAMILIA INEMBARGABLE",
            "PODER ESPECIAL",
        ],
        "jaggua_bogota_1c": [
            "LIBERACIÓN PARCIAL DE HIPOTECA",
            "PROTOCOLIZACIÓN CERTIFICADO TÉCNICO DE OCUPACIÓN",
            "COMPRAVENTA VIS",
            "RENUNCIA A CONDICIÓN RESOLUTORIA",
            "CONSTITUCIÓN DE HIPOTECA ABIERTA DE PRIMER GRADO",
            "CONSTITUCIÓN DE PATRIMONIO DE FAMILIA INEMBARGABLE",
            "PODER ESPECIAL",
        ],
        "jaggua_bogota_2c": [
            "LIBERACIÓN PARCIAL DE HIPOTECA",
            "PROTOCOLIZACIÓN CERTIFICADO TÉCNICO DE OCUPACIÓN",
            "COMPRAVENTA VIS",
            "RENUNCIA A CONDICIÓN RESOLUTORIA",
            "CONSTITUCIÓN DE HIPOTECA ABIERTA DE PRIMER GRADO",
            "CONSTITUCIÓN DE PATRIMONIO DE FAMILIA INEMBARGABLE",
            "PODER ESPECIAL",
        ],
    }

    grouped_participants: dict[str, list[dict]] = {}
    for participant in participants:
        role_code = participant.get("role_code", "") or ""
        grouped_participants.setdefault(role_code, []).append(participant)

    participants_text = ""
    ordered_role_codes = list(grouped_participants.keys())
    for role_code in ordered_role_codes:
        for p in grouped_participants.get(role_code, []):
            role_label = p.get("role_label", "") or p.get("role_code", "")
            participants_text += (
                f"\n- ROL: {str(p.get('role_label', role_code)).upper()}"
            )
            participants_text += f"\n  Nombre completo: {p.get('full_name', '')}"
            participants_text += f"\n  Tipo documento: {p.get('document_type', '')}"
            participants_text += f"\n  Número documento: {p.get('document_number', '')}"
            participants_text += f"\n  Sexo: {p.get('sex', '')}"
            participants_text += f"\n  Nacionalidad: {p.get('nationality', '')}"
            participants_text += f"\n  Estado civil: {p.get('marital_status', '')}"
            participants_text += f"\n  Profesión u oficio: {p.get('profession', '')}"
            participants_text += f"\n  Municipio de domicilio: {p.get('municipality', '')}"
            participants_text += (
                f"\n  ¿Está de tránsito?: {'Sí' if p.get('is_transient') else 'No'}"
            )
            participants_text += f"\n  Teléfono: {p.get('phone', '')}"
            participants_text += f"\n  Dirección: {p.get('address', '')}"
            participants_text += f"\n  Email: {p.get('email', '')}"

    entidades_section = """
ENTIDADES JURÃDICAS PREEXISTENTES (usar cuando el rol sea apoderado):
- Fideicomiso P.A. Aragua de Primavera / vocera: Fiduciaria Bancolombia S.A.  NIT 830.054.539-0
- Constructora Contex S.A.S. BIC  NIT 900.082.107-5
- Bancolombia S.A.  NIT 890.903.938-8
- Banco Davivienda S.A.  NIT 860.034.313-7
- Fondo Nacional del Ahorro S.A.  NIT 899.999.284-4
- Banco de Bogotá S.A.  NIT 860.002.964-4
- Fideicomiso P.A. Jaggua / vocera: Fiduciaria Bancolombia S.A.  NIT 830.054.539-0

El rol del interviniente indica quÃ© entidad representa:
- apoderado_banco_libera  representa al banco acreedor que libera la hipoteca
- apoderado_fideicomiso  representa al Fideicomiso (vendedor)
- apoderado_fideicomitente  representa a la Constructora (fideicomitente)
- apoderado_banco_hipoteca  representa al banco que otorga el nuevo crÃ©dito hipotecario
"""

    act_data_text = ""
    for key, value in act_data.items():
        act_data_text += f"\n- {key}: {value}"

    reference_section = ""
    if template_reference_text:
        _REF_MAX_CHARS = 14_000
        _ref_text = template_reference_text[:_REF_MAX_CHARS]
        _truncated = len(template_reference_text) > _REF_MAX_CHARS
        reference_section = f"""
PLANTILLA DE REFERENCIA DEL ACTO (usar como guía de estructura y redacción notarial colombiana):
{"[NOTA: texto de referencia truncado por longitud — mantener estructura y estilo del fragmento mostrado]" if _truncated else ""}
---
{_ref_text}
---
"""

    variante_section = ""
    if variante_id and variante_id in ORDEN_ACTOS_POR_VARIANTE:
        actos_ordenados = ORDEN_ACTOS_POR_VARIANTE[variante_id]
        actos_texto = "\n".join(
            f"{idx}. {acto}" for idx, acto in enumerate(actos_ordenados, start=1)
        )
        variante_section = f"""
VARIANTE DOCUMENTAL: {variante_id}
ACTOS QUE DEBE CONTENER EN ESTE ORDEN EXACTO:
{actos_texto}
"""

    correction_section = ""
    if correction_note and previous_draft:
        correction_section = f"""
MODO CORRECCIÓN ACTIVO:
El siguiente es el borrador actual que debes corregir:
---BORRADOR ACTUAL---
{previous_draft[:8000]}
---FIN BORRADOR---

Instrucción de corrección: {correction_note}

INSTRUCCIÓN CRÍTICA: Aplica ÚNICAMENTE el cambio solicitado en la instrucción anterior.
No reescribas párrafos que no fueron mencionados. No cambies datos que no se pidan.
Devuelve el documento completo con solo esa corrección aplicada.
"""
    elif correction_note:
        correction_section = f"""
INSTRUCCIÓN DE CORRECCIÓN: {correction_note}
Aplica únicamente este cambio. No modifiques nada más del documento.
"""
    prompt = f"""Eres un experto en derecho notarial colombiano. Tu tarea es redactar un instrumento público notarial completo y correcto jurídicamente.

NOTARÍA: {notary_label}
NOTARIO TITULAR: {notary_name}
TIPO DE ACTO: {act_type}

INTERVINIENTES:{participants_text}

{entidades_section}

DATOS DEL ACTO:{act_data_text}

{variante_section}

{reference_section}

{correction_section}

INSTRUCCIONES DE REDACCIÓN:
1. Redacta el documento completo en español formal notarial colombiano
2. Incluye todas las secciones requeridas según la ley colombiana
3. Inserta los datos de los intervinientes exactamente como se proporcionan — nunca inventar datos
5. Usa el género gramatical correcto según el sexo de cada interviniente
6. Incluye constancias notariales obligatorias: Ley 1581/2012, Art. 102 Decreto 960/1970
7. Incluye liquidación de derechos notariales con los valores proporcionados
8. Incluye sección de firmas con datos de comparecientes
9. Documento listo para firmar, sin placeholders ni etiquetas
10. Si hay VARIANTE DOCUMENTAL, genera los actos en ese orden exacto

FORMATO DE SALIDA:
- Solo el texto del documento
- Sin explicaciones ni etiquetas
- Listo para exportar a Word
"""

    return prompt


def generate_notarial_document(
    act_type: str,
    notary_label: str,
    notary_name: str,
    participants: list[dict],
    act_data: dict,
    template_reference_text: str | None = None,
    max_tokens: int = 4000,
    variante_id: str | None = None,
    correction_note: str | None = None,
    previous_draft: str | None = None,
) -> str:
    """Genera el documento notarial completo usando GPT-4o."""
    client = get_openai_client()
    prompt = build_gari_prompt(
        act_type=act_type,
        notary_label=notary_label,
        notary_name=notary_name,
        participants=participants,
        act_data=act_data,
        template_reference_text=template_reference_text,
        variante_id=variante_id,
        correction_note=correction_note,
        previous_draft=previous_draft,
    )

    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {
                "role": "system",
                "content": SYSTEM_PROMPT_GARI,
            },
            {
                "role": "user",
                "content": prompt,
            },
        ],
        temperature=0.2,
        max_tokens=max_tokens,
    )

    return response.choices[0].message.content or ""


def build_gari_docx_buffer(text: str) -> io.BytesIO:
    """Construye un documento .docx en memoria con formato Gari."""
    doc = DocxDocument()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)

    for line in text.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        texto = line
        # Solo agregar guiones si la línea parece un campo para rellenar
        # (termina con ":" o es corta y no es un título/cláusula completa)
        es_encabezado = any(line.upper().startswith(kw) for kw in [
            "PRIMERO", "SEGUNDO", "TERCERO", "CUARTO", "QUINTO", "SEXTO",
            "SÉPTIMO", "OCTAVO", "NOVENO", "DÉCIMO", "ESCRITURA", "ACTO:",
            "OTORGAMIENTO", "CONSTANCIA", "AUTORIZACIÓN", "PARÁGRAFO",
            "PARAGRAFO", "NOTA", "DERECHOS", "SUPERFONDO", "LIQUIDACIÓN",
        ])
        es_campo_vacio = line.endswith(":") or (len(line) < 60 and not es_encabezado and not line.endswith("."))
        if es_campo_vacio and not line.strip().startswith("- -"):
            texto = texto + " - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -"
        p = doc.add_paragraph(texto)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if line.startswith("- - -") or line.startswith("---"):
            run = p.runs[0] if p.runs else p.add_run(line)
            run.bold = True

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def save_gari_document_as_docx(text: str, output_path: str | Path) -> str:
    """Genera el .docx en memoria y lo sube a Supabase Storage. Retorna signed URL."""
    buffer = build_gari_docx_buffer(text)
    file_bytes = buffer.getvalue()

    supabase = get_supabase_client()
    storage_path = str(output_path).replace("\\", "/")
    supabase.storage.from_("documentos").upload(
        path=storage_path,
        file=file_bytes,
        file_options={
            "content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "upsert": "true",
        },
    )

    signed = supabase.storage.from_("documentos").create_signed_url(storage_path, 3600)
    url = signed.get("signedUrl") or signed.get("signedURL") or ""
    if not url:
        raise ValueError(f"Supabase no retornó signed URL para {storage_path}. Respuesta: {signed}")
    return url

def resolver_escritura(
    proyecto: str,
    tipo_inmueble: str,
    num_compradores: int,
    banco_hipotecante: str | None,
    campos_caso: dict,
) -> dict:
    """Resuelve la variante de escritura y valida campos requeridos para el caso."""

    def _esta_vacio(valor: object) -> bool:
        if valor is None:
            return True
        if isinstance(valor, str):
            return valor.strip() == ""
        if isinstance(valor, (list, dict, tuple, set)):
            return len(valor) == 0
        return False

    actos_aragua_apto = [
        "liberacion_hipoteca",
        "cto",
        "compraventa_vis",
        "renuncia_resolutoria",
        "cancelacion_comodato",
        "patrimonio_familia",
        "poder_especial",
    ]
    campos_aragua_apto = [
        "numero_apartamento",
        "matricula_inmobiliaria",
        "cedula_catastral",
        "linderos",
        "numero_piso",
        "area_privada",
        "area_total",
        "altura",
        "coeficiente_copropiedad",
        "avaluo_catastral",
        "valor_venta",
        "fecha_promesa_compraventa",
        "paz_salvo_predial_numero",
        "paz_salvo_predial_fecha",
        "paz_salvo_predial_vigencia",
    ]

    actos_aragua_parq = [
        "liberacion_hipoteca",
        "cto",
        "compraventa_vis",
        "renuncia_resolutoria",
        "cancelacion_comodato",
        "poder_especial",
    ]
    campos_aragua_parq = [
        "numero_parqueadero",
        "matricula_inmobiliaria",
        "cedula_catastral",
        "linderos",
        "area_privada",
        "altura",
        "coeficiente_copropiedad",
        "avaluo_catastral",
        "valor_venta",
        "fecha_promesa_compraventa",
        "paz_salvo_predial_numero",
        "paz_salvo_predial_fecha",
        "paz_salvo_predial_vigencia",
    ]

    actos_jaggua_fna = [
        "liberacion_hipoteca",
        "cto",
        "compraventa_vis",
        "renuncia_resolutoria",
        "cancelacion_comodato",
        "hipoteca_primer_grado",
        "patrimonio_familia",
        "poder_especial",
    ]
    campos_jaggua_fna = [
        "numero_apartamento",
        "matricula_inmobiliaria",
        "linderos",
        "numero_piso",
        "area_privada",
        "area_total",
        "altura",
        "coeficiente_copropiedad",
        "valor_venta",
        "cuota_inicial",
        "valor_hipoteca",
        "origen_cuota_inicial",
        "fecha_promesa_compraventa",
        "inmueble_sera_casa_habitacion",
        "tiene_bien_afectado",
    ]

    actos_jaggua_bogota = [
        "liberacion_hipoteca",
        "cto",
        "compraventa_vis",
        "renuncia_resolutoria",
        "hipoteca_primer_grado",
        "patrimonio_familia",
        "poder_especial",
    ]
    campos_jaggua_bogota = [
        "numero_apartamento",
        "matricula_inmobiliaria",
        "linderos",
        "numero_piso",
        "area_privada",
        "area_total",
        "altura",
        "valor_venta",
        "cuota_inicial",
        "valor_hipoteca",
        "origen_cuota_inicial",
        "fecha_promesa_compraventa",
        "inmueble_sera_casa_habitacion",
        "tiene_bien_afectado",
    ]

    variantes = {
        ("aragua", "apartamento", 1, None): {
            "variante_id": "aragua_apto_1c",
            "plantilla_id": "aragua_apto_1c",
            "actos_requeridos": actos_aragua_apto,
            "campos_requeridos": campos_aragua_apto,
            "banco_nit": "890.903.938-8",
            "max_tokens_estimado": 5500,
        },
        ("aragua", "apartamento", 2, None): {
            "variante_id": "aragua_apto_2c",
            "plantilla_id": "aragua_apto_2c",
            "actos_requeridos": actos_aragua_apto,
            "campos_requeridos": campos_aragua_apto,
            "banco_nit": "890.903.938-8",
            "max_tokens_estimado": 5800,
        },
        ("aragua", "parqueadero", 2, None): {
            "variante_id": "aragua_parq_2c",
            "plantilla_id": "aragua_parq_2c",
            "actos_requeridos": actos_aragua_parq,
            "campos_requeridos": campos_aragua_parq,
            "banco_nit": "890.903.938-8",
            "max_tokens_estimado": 5200,
        },
        ("aragua", "parqueadero", 3, None): {
            "variante_id": "aragua_parq_3c",
            "plantilla_id": "aragua_parq_3c",
            "actos_requeridos": actos_aragua_parq,
            "campos_requeridos": campos_aragua_parq,
            "banco_nit": "890.903.938-8",
            "max_tokens_estimado": 5400,
        },
        ("jaggua", "apartamento", 1, "fna"): {
            "variante_id": "jaggua_fna_1c",
            "plantilla_id": "jaggua_fna_1c",
            "actos_requeridos": actos_jaggua_fna,
            "campos_requeridos": campos_jaggua_fna,
            "banco_nit": "899.999.284-4",
            "max_tokens_estimado": 6500,
        },
        ("jaggua", "apartamento", 2, "fna"): {
            "variante_id": "jaggua_fna_2c",
            "plantilla_id": "jaggua_fna_2c",
            "actos_requeridos": actos_jaggua_fna,
            "campos_requeridos": campos_jaggua_fna,
            "banco_nit": "899.999.284-4",
            "max_tokens_estimado": 6800,
        },
        ("jaggua", "apartamento", 1, "bogota"): {
            "variante_id": "jaggua_bogota_1c",
            "plantilla_id": "jaggua_bogota_1c",
            "actos_requeridos": actos_jaggua_bogota,
            "campos_requeridos": campos_jaggua_bogota,
            "banco_nit": "860.002.964-4",
            "max_tokens_estimado": 6800,
        },
        ("jaggua", "apartamento", 2, "bogota"): {
            "variante_id": "jaggua_bogota_2c",
            "plantilla_id": "jaggua_bogota_2c",
            "actos_requeridos": actos_jaggua_bogota,
            "campos_requeridos": campos_jaggua_bogota,
            "banco_nit": "860.002.964-4",
            "max_tokens_estimado": 7200,
        },
    }

    proyecto_norm = (proyecto or "").strip().lower()
    tipo_inmueble_norm = (tipo_inmueble or "").strip().lower()
    banco_norm = (banco_hipotecante or "").strip().lower() or None

    key = (proyecto_norm, tipo_inmueble_norm, num_compradores, banco_norm)
    if key not in variantes:
        raise ValueError(
            "No existe una variante de escritura para la combinación "
            f"proyecto={proyecto_norm}, tipo_inmueble={tipo_inmueble_norm}, "
            f"num_compradores={num_compradores}, banco_hipotecante={banco_norm}."
        )

    variante = variantes[key]
    campos_requeridos = variante["campos_requeridos"]
    campos_faltantes = [
        campo for campo in campos_requeridos if _esta_vacio(campos_caso.get(campo))
    ]

    return {
        "variante_id": variante["variante_id"],
        "plantilla_id": variante["plantilla_id"],
        "actos_requeridos": list(variante["actos_requeridos"]),
        "campos_requeridos": list(campos_requeridos),
        "campos_faltantes": campos_faltantes,
        "banco_nit": variante["banco_nit"],
        "max_tokens_estimado": variante["max_tokens_estimado"],
    }


def resolver_escritura_desde_template(template) -> dict:
    """
    Reemplaza resolver_escritura(). Recibe el objeto DocumentTemplate de BD.
    Retorna variante_id, campos_requeridos y max_tokens sin hardcodear nada.
    """
    campos_requeridos = [
        f.field_code for f in (template.fields or []) if f.is_required
    ]
    max_tokens = min(4000 + len(campos_requeridos) * 500, 8000)
    return {
        "variante_id": template.slug,
        "plantilla_id": template.slug,
        "campos_requeridos": campos_requeridos,
        "campos_faltantes": [],   # se valida en el endpoint
        "max_tokens_estimado": max_tokens,
    }


if __name__ == "__main__":
    ejemplo_aragua = resolver_escritura(
        proyecto="aragua",
        tipo_inmueble="apartamento",
        num_compradores=1,
        banco_hipotecante=None,
        campos_caso={
            "numero_apartamento": "1201",
            "matricula_inmobiliaria": "50N-123456",
            "cedula_catastral": "AA-001",
            "linderos": "Norte con...",
            "numero_piso": "12",
            "area_privada": "80",
            "area_total": "90",
            "altura": "2.40",
            "coeficiente_copropiedad": "1.25%",
            "avaluo_catastral": "180000000",
            "valor_venta": "250000000",
            "fecha_promesa_compraventa": "2026-04-01",
            "paz_salvo_predial_numero": "PSP-55",
            "paz_salvo_predial_fecha": "2026-03-20",
            "paz_salvo_predial_vigencia": "2026",
        },
    )
    print("Ejemplo aragua apartamento 1c:")
    print(ejemplo_aragua)

    ejemplo_jaggua = resolver_escritura(
        proyecto="jaggua",
        tipo_inmueble="apartamento",
        num_compradores=2,
        banco_hipotecante="bogota",
        campos_caso={
            "numero_apartamento": "905",
            "matricula_inmobiliaria": "50C-765432",
            "linderos": "Sur con...",
            "numero_piso": "9",
            "area_privada": "70",
            "area_total": "78",
            "altura": "2.35",
            "valor_venta": "320000000",
            "cuota_inicial": "80000000",
            "valor_hipoteca": "240000000",
            "origen_cuota_inicial": "Ahorros",
            "fecha_promesa_compraventa": "2026-04-10",
            "inmueble_sera_casa_habitacion": True,
            "tiene_bien_afectado": False,
        },
    )
    print("Ejemplo jaggua bogota 2c:")
    print(ejemplo_jaggua)
