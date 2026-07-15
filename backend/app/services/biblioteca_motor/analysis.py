from __future__ import annotations

import hashlib
import io
import json
import re
import time
from dataclasses import dataclass
from typing import Any, Protocol

from docx import Document


CONTEXT_CHARS = 90
ANALYSIS_TOTAL_BUDGET_SECONDS = 25.0
AI_BUDGET_SECONDS = 10.0
MAX_AI_CANDIDATES = 40
MAX_OPERATIONAL_SUGGESTIONS = 120
PROVISIONAL_FIELD_PREFIX = "PENDING_FIELD_"

ROLE_KEYWORDS = (
    "COMPRADOR",
    "COMPRADORES",
    "VENDEDOR",
    "VENDEDORES",
    "OTORGANTE",
    "OTORGANTES",
    "COMPARECIENTE",
    "COMPARECIENTES",
    "APODERADO",
    "APODERADA",
    "ACREEDOR",
    "DEUDOR",
)

LEGAL_PERSON_STOPWORDS = (
    "ACTO",
    "ACTOS",
    "BANCO",
    "CLAUSULA",
    "CODIGO",
    "COMPRADORES",
    "CONTRATO",
    "CORPORACION",
    "DECLARACION",
    "DOCUMENTO",
    "ESCRITURA",
    "FIDUCIARIA",
    "INMOBILIARIA",
    "INSTRUMENTO",
    "JURIDICA",
    "LEGAL",
    "LTDA",
    "MATRICULA",
    "NOTARIA",
    "NUMERO",
    "PUBLICO",
    "SAS",
    "SOCIEDAD",
    "VALOR",
    "VENDEDORES",
)

STRUCTURAL_PHRASES = (
    "LOS COMPRADORES",
    "LOS VENDEDORES",
    "VALOR DEL ACTO",
    "INSTRUMENTO PUBLICO",
    "ESCRITURA PUBLICA",
    "MATRICULA INMOBILIARIA",
    "CEDULA CATASTRAL",
    "DATOS DEL INMUEBLE",
    "RESUMEN DE ACTOS",
)

TYPE_PRIORITY = {
    "email": 100,
    "nit": 95,
    "document_number": 90,
    "matricula_inmobiliaria": 88,
    "codigo_catastral": 86,
    "money": 82,
    "date": 80,
    "deed_number": 78,
    "phone": 76,
    "area": 74,
    "address": 70,
    "bank": 65,
    "person_name": 50,
}

CANDIDATE_CATEGORIES = {
    "person_name": "persona",
    "document_number": "persona",
    "nit": "persona_juridica",
    "matricula_inmobiliaria": "inmueble",
    "codigo_catastral": "inmueble",
    "area": "inmueble",
    "address": "ubicacion",
    "money": "valor",
    "date": "fecha",
    "deed_number": "escritura",
    "bank": "entidad",
    "email": "contacto",
    "phone": "contacto",
}

AI_USEFUL_TYPES = {
    "person_name",
    "document_number",
    "nit",
    "matricula_inmobiliaria",
    "codigo_catastral",
    "money",
    "date",
    "deed_number",
    "bank",
    "email",
    "phone",
    "address",
    "area",
}

AI_CONTEXT_KEYWORDS = ROLE_KEYWORDS + (
    "INMUEBLE",
    "PRECIO",
    "VALOR",
    "ESCRITURA",
    "IDENTIFICACION",
    "CEDULA",
    "MATRICULA",
)


@dataclass(frozen=True)
class DocumentBlock:
    block_type: str
    text: str
    block_index: int
    paragraph_index: int | None = None
    table_index: int | None = None
    row_index: int | None = None
    cell_index: int | None = None


@dataclass(frozen=True)
class Candidate:
    candidate_id: str
    original_text: str
    candidate_type: str
    context_before: str
    context_after: str
    location: dict[str, Any]


@dataclass(frozen=True)
class FieldDefinition:
    code: str
    label: str
    category: str


@dataclass(frozen=True)
class _RawCandidate:
    candidate_type: str
    original_text: str
    start: int
    end: int
    score: int


class CandidateClassifier(Protocol):
    def classify(
        self,
        candidates: list[Candidate],
        fields: list[FieldDefinition],
        timeout_seconds: float | None = None,
    ) -> list[dict[str, Any]]:
        ...


class AIClassifierTimeoutError(RuntimeError):
    pass


class AIClassifierProviderError(RuntimeError):
    pass


class AIClassifierInvalidResponseError(RuntimeError):
    pass


PATTERNS: tuple[tuple[str, re.Pattern[str]], ...] = (
    ("email", re.compile(r"\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b", re.IGNORECASE)),
    ("nit", re.compile(r"\b(?:NIT\.?\s*)?\d{3}[.\s]?\d{3}[.\s]?\d{3}-?\d\b", re.IGNORECASE)),
    ("matricula_inmobiliaria", re.compile(r"\b\d{3}-\d{5,10}\b")),
    ("codigo_catastral", re.compile(r"\b(?:c[ée]dula|c[oó]digo)\s+catastral\s*(?:n[úu]mero\s*)?[:#-]?\s*([A-Z0-9-]{10,35})\b", re.IGNORECASE)),
    ("money", re.compile(r"(?:\$|COP\s*)\s?\d{1,3}(?:[.,]\d{3})+(?:,\d{2})?|\b\d{1,3}(?:[.,]\d{3})+\s+PESOS\b", re.IGNORECASE)),
    ("date", re.compile(r"\b\d{1,2}\s+de\s+(?:enero|febrero|marzo|abril|mayo|junio|julio|agosto|septiembre|setiembre|octubre|noviembre|diciembre)\s+de\s+\d{4}\b|\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b", re.IGNORECASE)),
    ("deed_number", re.compile(r"\b(?:escritura(?:\s+p[úu]blica)?\s*(?:n[úu]mero|no\.?)?\s*)[:#-]?\s*(\d{2,8})\b", re.IGNORECASE)),
    ("bank", re.compile(r"\bBanco\s+[A-ZÁÉÍÓÚÑ][A-Za-zÁÉÍÓÚÜÑáéíóúüñ]+(?:\s+[A-ZÁÉÍÓÚÑ][A-Za-zÁÉÍÓÚÜÑáéíóúüñ]+){0,3}\b")),
    ("phone", re.compile(r"\b(?:\+57\s*)?(?:3\d{2}|60\d)\s?\d{3}\s?\d{4}\b")),
    ("address", re.compile(r"\b(?:Calle|Carrera|Avenida|Diagonal|Transversal|Autopista)\s+\d+[A-Za-z]?(?:\s*(?:#|No\.?|Nro\.?)\s*\d+[A-Za-z]?(?:-\d+)?)?(?:\s+[A-Za-zÁÉÍÓÚÜÑáéíóúüñ0-9#.-]+){0,8}", re.IGNORECASE)),
    ("area", re.compile(r"\b\d+(?:[.,]\d+)?\s*(?:m2|m²|metros\s+cuadrados|hect[áa]reas?)\b", re.IGNORECASE)),
    ("document_number", re.compile(r"\b(?:c(?:é|e)dula(?:\s+de\s+ciudadan(?:í|i)a)?|CC|C\.C\.)\s*(?:n[úu]mero|no\.?)?\s*[:#-]?\s*([\d.]{6,15})\b", re.IGNORECASE)),
    ("person_name", re.compile(r"\b[A-ZÁÉÍÓÚÜÑ]{2,}(?:\s+[A-ZÁÉÍÓÚÜÑ]{2,}){1,5}\b")),
    ("person_name", re.compile(r"\b[A-ZÁÉÍÓÚÜÑ][a-záéíóúüñ]+(?:\s+[A-ZÁÉÍÓÚÜÑ][a-záéíóúüñ]+){1,4}\b")),
)


def extract_docx_blocks(docx_bytes: bytes) -> list[DocumentBlock]:
    doc = Document(io.BytesIO(docx_bytes))
    blocks: list[DocumentBlock] = []
    block_index = 0

    for paragraph_index, paragraph in enumerate(doc.paragraphs, start=1):
        text = paragraph.text
        if text:
            block_index += 1
            blocks.append(DocumentBlock("paragraph", text, block_index, paragraph_index=paragraph_index))

    for table_index, table in enumerate(doc.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            for cell_index, cell in enumerate(row.cells, start=1):
                text = "\n".join(paragraph.text for paragraph in cell.paragraphs if paragraph.text)
                if text:
                    block_index += 1
                    blocks.append(
                        DocumentBlock(
                            "table_cell",
                            text,
                            block_index,
                            table_index=table_index,
                            row_index=row_index,
                            cell_index=cell_index,
                        ),
                    )
    return blocks


def extract_candidates(docx_bytes: bytes) -> list[Candidate]:
    candidates: list[Candidate] = []
    for block in extract_docx_blocks(docx_bytes):
        raw_candidates: list[_RawCandidate] = []
        occurrence_counts: dict[str, int] = {}
        for candidate_type, pattern in PATTERNS:
            for match in pattern.finditer(block.text):
                start, end = _match_span(match)
                original = block.text[start:end]
                if not original.strip():
                    continue
                if _is_structural_or_low_evidence(candidate_type, original, block.text, start, end):
                    continue
                raw_candidates.append(
                    _RawCandidate(
                        candidate_type=candidate_type,
                        original_text=original,
                        start=start,
                        end=end,
                        score=_candidate_score(candidate_type, original, block.text, start, end),
                    ),
                )

        for raw in _dedupe_raw_candidates(raw_candidates):
            occurrence_counts[raw.original_text] = occurrence_counts.get(raw.original_text, 0) + 1
            occurrence_index = occurrence_counts[raw.original_text]
            location = _build_location(block, raw.start, raw.end, occurrence_index)
            if block.text[raw.start:raw.end] != raw.original_text:
                raise ValueError("La ubicacion calculada no coincide con el texto original.")
            candidate_id = "cand_" + _short_hash(f"{raw.candidate_type}:{raw.original_text}:{location['location_key']}")[:16]
            candidates.append(
                Candidate(
                    candidate_id=candidate_id,
                    original_text=raw.original_text,
                    candidate_type=raw.candidate_type,
                    context_before=_context_before(block.text, raw.start),
                    context_after=_context_after(block.text, raw.end),
                    location=location,
                ),
            )
    candidates.sort(key=lambda item: (item.location["block_index"], item.location["char_start"], item.candidate_id))
    return candidates


def analyze_biblioteca_document(
    docx_bytes: bytes,
    fields: list[Any],
    *,
    ai_classifier: CandidateClassifier | None = None,
    max_suggestions: int = MAX_OPERATIONAL_SUGGESTIONS,
    total_budget_seconds: float = ANALYSIS_TOTAL_BUDGET_SECONDS,
    ai_budget_seconds: float = AI_BUDGET_SECONDS,
    max_ai_candidates: int = MAX_AI_CANDIDATES,
) -> dict[str, Any]:
    started = time.perf_counter()
    extraction_started = time.perf_counter()
    field_defs = [_field_definition(item) for item in fields]
    field_by_code = {field.code: field for field in field_defs}
    candidates = extract_candidates(docx_bytes)
    extraction_ms = _elapsed_ms(extraction_started)

    deterministic_started = time.perf_counter()
    deterministic = {candidate.candidate_id: _deterministic_field_code(candidate, field_by_code) for candidate in candidates}
    operational = _build_operational_suggestions(candidates, deterministic, {}, field_by_code)
    deterministic_ms = _elapsed_ms(deterministic_started)

    eligible_ai, selected_ai = _select_ai_candidates(candidates, deterministic, max_ai_candidates)
    elapsed_before_ai = time.perf_counter() - started
    remaining_budget = max(0.0, total_budget_seconds - elapsed_before_ai)
    effective_ai_budget = min(ai_budget_seconds, remaining_budget)
    ai_results, ai_diag = _classify_with_ai_once(
        ai_classifier,
        selected_ai,
        field_defs,
        field_by_code,
        eligible=len(eligible_ai),
        omitted=max(0, len(eligible_ai) - len(selected_ai)),
        timeout_seconds=effective_ai_budget,
    )

    if ai_results:
        operational = _build_operational_suggestions(candidates, deterministic, ai_results, field_by_code)

    operational.sort(key=lambda item: (-item["_priority"], item["location"]["block_index"], item["location"]["char_start"]))
    omitted = max(0, len(operational) - max_suggestions)
    suggestions = [{key: value for key, value in item.items() if key != "_priority"} for item in operational[:max_suggestions]]

    catalogued_total = sum(1 for item in operational if item["catalog_status"] == "matched")
    provisional_total = sum(1 for item in operational if item["catalog_status"] == "unmapped")
    mode = "hybrid" if ai_results else "deterministic"
    status = _analysis_status(ai_diag, bool(ai_results))
    total_ms = _elapsed_ms(started)
    return {
        "analysis_id": "analysis_" + _short_hash(hashlib.sha256(docx_bytes).hexdigest())[:16],
        "mode": mode,
        "status": status,
        "suggestions": suggestions,
        "stats": {
            "deterministic_candidates": len(candidates),
            "classified_candidates": catalogued_total,
            "unclassified_candidates": provisional_total,
            "catalogued_suggestions": catalogued_total,
            "provisional_suggestions": provisional_total,
            "suggestions": len(suggestions),
            "omitted_suggestions": omitted,
            "excluded_low_priority": omitted,
            "ai_candidates_eligible": len(eligible_ai),
            "ai_candidates_sent": len(selected_ai),
            "ai_candidates_omitted": max(0, len(eligible_ai) - len(selected_ai)),
        },
        "diagnostics": {
            "ai": ai_diag,
        },
        "timing": {
            "download_ms": 0,
            "extraction_ms": extraction_ms,
            "deterministic_ms": deterministic_ms,
            "ai_ms": ai_diag["duration_ms"],
            "total_ms": total_ms,
        },
    }


class OpenAIBibliotecaClassifier:
    def __init__(self, api_key: str, model: str = "gpt-4o-mini") -> None:
        self.api_key = api_key
        self.model = model

    def classify(
        self,
        candidates: list[Candidate],
        fields: list[FieldDefinition],
        timeout_seconds: float | None = None,
    ) -> list[dict[str, Any]]:
        try:
            from openai import APIConnectionError, APIError, APITimeoutError, OpenAI
        except Exception:  # pragma: no cover
            from openai import OpenAI
            APIConnectionError = APIError = APITimeoutError = Exception

        timeout = max(1.0, min(float(timeout_seconds or AI_BUDGET_SECONDS), AI_BUDGET_SECONDS))
        client = OpenAI(api_key=self.api_key, timeout=timeout)
        payload = {
            "candidates": [
                {
                    "candidate_id": item.candidate_id,
                    "original_text": item.original_text,
                    "candidate_type": item.candidate_type,
                    "context_before": item.context_before,
                    "context_after": item.context_after,
                }
                for item in candidates
            ],
            "allowed_fields": [{"code": item.code, "label": item.label, "category": item.category} for item in fields],
        }
        try:
            response = client.chat.completions.create(
                model=self.model,
                temperature=0,
                response_format={"type": "json_object"},
                messages=[
                    {
                        "role": "system",
                        "content": (
                            "Clasifica candidatos notariales. Responde JSON con key classifications. "
                            "Cada item debe usar candidate_id existente, field_code permitido o null, "
                            "confidence 0..1 y reason breve. No calcules ubicaciones."
                        ),
                    },
                    {"role": "user", "content": json.dumps(payload, ensure_ascii=False)},
                ],
            )
            content = response.choices[0].message.content or "{}"
            parsed = json.loads(content)
            classifications = parsed.get("classifications", [])
            if not isinstance(classifications, list):
                raise AIClassifierInvalidResponseError("invalid classifications")
            return classifications
        except (APITimeoutError, TimeoutError) as exc:
            raise AIClassifierTimeoutError("timeout") from exc
        except json.JSONDecodeError as exc:
            raise AIClassifierInvalidResponseError("invalid_json") from exc
        except (APIConnectionError, APIError) as exc:
            raise AIClassifierProviderError("provider_error") from exc
        except Exception as exc:
            if "timeout" in exc.__class__.__name__.lower():
                raise AIClassifierTimeoutError("timeout") from exc
            raise


def _classify_with_ai_once(
    ai_classifier: CandidateClassifier | None,
    candidates: list[Candidate],
    fields: list[FieldDefinition],
    field_by_code: dict[str, FieldDefinition],
    *,
    eligible: int,
    omitted: int,
    timeout_seconds: float,
) -> tuple[dict[str, dict[str, Any]], dict[str, Any]]:
    diag: dict[str, Any] = {
        "attempted": False,
        "status": "not_configured" if ai_classifier is None else "not_attempted",
        "eligible": eligible,
        "sent": len(candidates),
        "omitted": omitted,
        "classified": 0,
        "duration_ms": 0,
    }
    if ai_classifier is None or not candidates:
        return {}, diag
    if timeout_seconds <= 0:
        diag["status"] = "budget_exhausted"
        return {}, diag

    started = time.perf_counter()
    diag["attempted"] = True
    ai_results: dict[str, dict[str, Any]] = {}
    allowed_ids = {candidate.candidate_id for candidate in candidates}
    try:
        raw_results = _invoke_classifier(ai_classifier, candidates, fields, timeout_seconds)
    except (AIClassifierTimeoutError, TimeoutError):
        diag["status"] = "timeout"
        diag["duration_ms"] = min(_elapsed_ms(started), int(timeout_seconds * 1000))
        return {}, diag
    except (AIClassifierInvalidResponseError, json.JSONDecodeError, ValueError):
        diag["status"] = "invalid_json"
        diag["duration_ms"] = _elapsed_ms(started)
        return {}, diag
    except Exception:
        diag["status"] = "provider_error"
        diag["duration_ms"] = _elapsed_ms(started)
        return {}, diag

    if not isinstance(raw_results, list):
        diag["status"] = "invalid_json"
        diag["duration_ms"] = _elapsed_ms(started)
        return {}, diag
    for raw in raw_results:
        if not isinstance(raw, dict):
            continue
        candidate_id = raw.get("candidate_id")
        if not isinstance(candidate_id, str) or candidate_id not in allowed_ids or candidate_id in ai_results:
            continue
        if _validated_ai_code(raw, field_by_code):
            ai_results[candidate_id] = raw
    diag["classified"] = len(ai_results)
    diag["status"] = "completed"
    diag["duration_ms"] = _elapsed_ms(started)
    return ai_results, diag


def _invoke_classifier(
    ai_classifier: CandidateClassifier,
    candidates: list[Candidate],
    fields: list[FieldDefinition],
    timeout_seconds: float,
) -> list[dict[str, Any]]:
    try:
        return ai_classifier.classify(candidates, fields, timeout_seconds=timeout_seconds)
    except TypeError as exc:
        if "timeout" not in str(exc):
            raise
        return ai_classifier.classify(candidates, fields)


def _build_operational_suggestions(
    candidates: list[Candidate],
    deterministic: dict[str, str | None],
    ai_results: dict[str, dict[str, Any]],
    field_by_code: dict[str, FieldDefinition],
) -> list[dict[str, Any]]:
    operational: list[dict[str, Any]] = []
    for candidate in candidates:
        if not candidate.original_text.strip():
            continue
        ai_result = ai_results.get(candidate.candidate_id)
        ai_code = _validated_ai_code(ai_result, field_by_code)
        det_code = deterministic.get(candidate.candidate_id)
        field_code = det_code or ai_code
        field = field_by_code.get(field_code or "")
        is_provisional = not field_code or field is None

        if is_provisional:
            field_code = _provisional_field_code(candidate)
            field_label = "Campo nuevo por definir"
            category = CANDIDATE_CATEGORIES.get(candidate.candidate_type, "otro")
            source = "provisional"
            reason = _clean_reason(ai_result) or "Dato variable detectado sin equivalencia en la biblioteca de campos."
        else:
            field_label = field.label
            category = field.category
            if det_code and ai_code and det_code == ai_code:
                source = "hybrid"
            elif ai_code and not det_code:
                source = "ai"
            else:
                source = "deterministic"
            reason = _clean_reason(ai_result)

        operational.append(
            {
                "suggestion_id": "sug_" + _short_hash(candidate.candidate_id + ":" + field_code)[:16],
                "candidate_id": candidate.candidate_id,
                "candidate_type": candidate.candidate_type,
                "original_text": candidate.original_text,
                "field_code": field_code,
                "field_label": field_label,
                "category": category,
                "catalog_status": "unmapped" if is_provisional else "matched",
                "requires_field_assignment": is_provisional,
                "confidence": _confidence(source, not is_provisional),
                "source": source,
                "needs_human_review": True if is_provisional else source != "hybrid",
                "location": candidate.location,
                "context_before": candidate.context_before,
                "context_after": candidate.context_after,
                "reason": reason,
                "_priority": _suggestion_priority(candidate, source),
            },
        )
    return operational


def _select_ai_candidates(
    candidates: list[Candidate],
    deterministic: dict[str, str | None],
    max_ai_candidates: int,
) -> tuple[list[Candidate], list[Candidate]]:
    eligible = [
        candidate
        for candidate in candidates
        if deterministic.get(candidate.candidate_id) is None and _is_ai_candidate_eligible(candidate)
    ]
    eligible.sort(key=_ai_candidate_sort_key)
    return eligible, eligible[:max_ai_candidates]


def _match_span(match: re.Match[str]) -> tuple[int, int]:
    if match.lastindex:
        for index in range(1, match.lastindex + 1):
            if match.group(index):
                return match.span(index)
    return match.span(0)


def _build_location(block: DocumentBlock, start: int, end: int, occurrence_index: int) -> dict[str, Any]:
    if block.block_type == "paragraph":
        location_key = f"paragraph:{block.paragraph_index}:{start}:{end}:{occurrence_index}"
    else:
        location_key = f"table_cell:{block.table_index}:{block.row_index}:{block.cell_index}:{start}:{end}:{occurrence_index}"
    return {
        "block_type": block.block_type,
        "block_index": block.block_index,
        "paragraph_index": block.paragraph_index,
        "table_index": block.table_index,
        "row_index": block.row_index,
        "cell_index": block.cell_index,
        "char_start": start,
        "char_end": end,
        "occurrence_index": occurrence_index,
        "location_key": location_key,
        "block_hash": hashlib.sha256(block.text.encode("utf-8")).hexdigest(),
    }


def _context_before(text: str, start: int) -> str:
    return text[max(0, start - CONTEXT_CHARS) : start].strip()


def _context_after(text: str, end: int) -> str:
    return text[end : end + CONTEXT_CHARS].strip()


def _short_hash(value: str) -> str:
    return hashlib.sha256(value.encode("utf-8")).hexdigest()


def _field_definition(item: Any) -> FieldDefinition:
    if isinstance(item, dict):
        raw_code = item.get("code", "")
        raw_label = item.get("label", "")
        raw_category = item.get("category", "otro")
    else:
        raw_code = getattr(item, "code", "")
        raw_label = getattr(item, "label", "")
        raw_category = getattr(item, "category", "otro")
    return FieldDefinition(
        code=str(raw_code or "").strip(),
        label=str(raw_label or raw_code or "").strip(),
        category=str(raw_category or "otro").strip() or "otro",
    )


def _elapsed_ms(started: float) -> int:
    return max(0, int((time.perf_counter() - started) * 1000))


def _analysis_status(ai_diag: dict[str, Any], has_ai_results: bool) -> str:
    if has_ai_results:
        return "completed_hybrid"
    if not ai_diag.get("attempted"):
        return "completed_deterministic"
    status = ai_diag.get("status")
    if status == "timeout":
        return "completed_with_ai_timeout"
    if status == "invalid_json":
        return "completed_with_ai_invalid_response"
    if status == "provider_error":
        return "completed_with_ai_provider_fallback"
    return "completed_deterministic"


def _is_structural_or_low_evidence(candidate_type: str, original: str, block_text: str, start: int, end: int) -> bool:
    normalized = _normalize_ascii(original)
    if candidate_type != "person_name":
        return False
    if any(phrase == normalized or phrase in normalized for phrase in STRUCTURAL_PHRASES):
        return True
    words = normalized.split()
    if len(words) < 2 or len(words) > 5:
        return True
    if any(word in LEGAL_PERSON_STOPWORDS for word in words):
        return True
    if not _has_person_role_context(block_text, start, end):
        return True
    return False


def _is_ai_candidate_eligible(candidate: Candidate) -> bool:
    if candidate.candidate_type not in AI_USEFUL_TYPES:
        return False
    if not candidate.original_text.strip():
        return False
    context = _normalize_ascii(f"{candidate.context_before} {candidate.original_text} {candidate.context_after}")
    if any(keyword in context for keyword in AI_CONTEXT_KEYWORDS):
        return True
    return candidate.candidate_type in {"matricula_inmobiliaria", "codigo_catastral", "money", "deed_number", "nit"}


def _ai_candidate_sort_key(candidate: Candidate) -> tuple[int, int, int, str]:
    context = _normalize_ascii(f"{candidate.context_before} {candidate.original_text} {candidate.context_after}")
    context_score = sum(10 for keyword in AI_CONTEXT_KEYWORDS if keyword in context)
    type_score = TYPE_PRIORITY.get(candidate.candidate_type, 0)
    return (-(context_score + type_score), candidate.location["block_index"], candidate.location["char_start"], candidate.candidate_id)


def _candidate_score(candidate_type: str, original: str, block_text: str, start: int, end: int) -> int:
    score = TYPE_PRIORITY.get(candidate_type, 0)
    if candidate_type == "person_name" and _has_person_role_context(block_text, start, end):
        score += 25
    if candidate_type != "person_name":
        score += min(20, len(original))
    return score


def _dedupe_raw_candidates(raw_candidates: list[_RawCandidate]) -> list[_RawCandidate]:
    selected: list[_RawCandidate] = []
    for candidate in sorted(raw_candidates, key=lambda item: (-item.score, item.start, -(item.end - item.start))):
        if any(_spans_overlap(candidate.start, candidate.end, existing.start, existing.end) for existing in selected):
            continue
        selected.append(candidate)
    selected.sort(key=lambda item: (item.start, item.end, item.candidate_type))
    return selected


def _spans_overlap(start: int, end: int, other_start: int, other_end: int) -> bool:
    return start < other_end and other_start < end


def _has_person_role_context(block_text: str, start: int, end: int) -> bool:
    window = _normalize_ascii(block_text[max(0, start - 80) : min(len(block_text), end + 80)])
    return any(keyword in window for keyword in ROLE_KEYWORDS)


def _normalize_ascii(value: str) -> str:
    replacements = {
        "Á": "A",
        "É": "E",
        "Í": "I",
        "Ó": "O",
        "Ú": "U",
        "Ü": "U",
        "Ñ": "N",
        "á": "A",
        "é": "E",
        "í": "I",
        "ó": "O",
        "ú": "U",
        "ü": "U",
        "ñ": "N",
    }
    normalized = "".join(replacements.get(char, char.upper()) for char in value)
    return re.sub(r"\s+", " ", normalized).strip()


def _deterministic_field_code(candidate: Candidate, field_by_code: dict[str, FieldDefinition]) -> str | None:
    text = _normalize_ascii(f"{candidate.context_before} {candidate.original_text} {candidate.context_after}")
    candidate_map = {
        "matricula_inmobiliaria": ["MATRICULA_INMOBILIARIA"],
        "codigo_catastral": ["CEDULA_CATASTRAL", "CODIGO_CATASTRAL"],
        "deed_number": ["NUMERO_ESCRITURA"],
        "money": ["VALOR_VENTA", "PRECIO", "VALOR_ACTO"],
        "date": ["FECHA_ESCRITURA", "FECHA_OTORGAMIENTO"],
        "email": ["EMAIL", "CORREO"],
        "phone": ["TELEFONO", "CELULAR"],
        "address": ["DIRECCION_INMUEBLE", "DIRECCION"],
        "area": ["AREA_CONSTRUIDA", "AREA"],
        "nit": ["NIT"],
        "bank": ["BANCO"],
    }
    if candidate.candidate_type == "person_name" and "COMPRADOR" in text:
        return _first_existing(["COMPRADOR_1", "NOMBRE_COMPRADOR_1", "COMPRADOR"], field_by_code)
    if candidate.candidate_type == "document_number" and "COMPRADOR" in text:
        return _first_existing(["CEDULA_COMPRADOR_1", "DOCUMENTO_COMPRADOR_1"], field_by_code)
    return _first_existing(candidate_map.get(candidate.candidate_type, []), field_by_code)


def _first_existing(codes: list[str], field_by_code: dict[str, FieldDefinition]) -> str | None:
    for code in codes:
        if code in field_by_code:
            return code
    return None


def _validated_ai_code(result: dict[str, Any] | None, field_by_code: dict[str, FieldDefinition]) -> str | None:
    if not result:
        return None
    code = result.get("field_code")
    if isinstance(code, str) and code.strip() in field_by_code:
        return code.strip()
    return None


def _provisional_field_code(candidate: Candidate) -> str:
    normalized_value = _normalize_ascii(candidate.original_text)
    digest = _short_hash(f"{candidate.candidate_type}:{normalized_value}")[:12].upper()
    candidate_type = re.sub(r"[^A-Z0-9]+", "_", candidate.candidate_type.upper()).strip("_") or "OTRO"
    return f"{PROVISIONAL_FIELD_PREFIX}{candidate_type}_{digest}"


def _suggestion_priority(candidate: Candidate, source: str) -> int:
    score = TYPE_PRIORITY.get(candidate.candidate_type, 0)
    if source == "hybrid":
        score += 40
    elif source == "ai":
        score += 25
    elif source == "provisional":
        score -= 10
    if candidate.candidate_type == "person_name":
        score += 10
    return score


def _confidence(source: str, has_field: bool) -> float:
    if source == "provisional":
        return 0.55
    if not has_field:
        return 0.0
    if source == "hybrid":
        return 0.97
    if source == "ai":
        return 0.9
    return 0.78


def _clean_reason(result: dict[str, Any] | None) -> str | None:
    if not result:
        return None
    reason = result.get("reason")
    return reason.strip()[:240] if isinstance(reason, str) and reason.strip() else None
