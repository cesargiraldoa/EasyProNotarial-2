from __future__ import annotations

import hashlib
import os
import re
import unittest
from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document

from app.services.minuta.inverse_conversion_catalog.corpus_importer import CorpusImporter
from app.services.minuta.inverse_conversion_catalog.corpus_reporter import CorpusReporter
from app.services.minuta.inverse_conversion_catalog.docx_marker_extractor import DocxMarkerExtractor
from app.services.minuta.inverse_conversion_catalog.field_alias_builder import FieldAliasBuilder
from app.services.minuta.inverse_conversion_catalog.field_code_normalizer import FieldCodeNormalizer
from app.services.minuta.inverse_conversion_catalog.models import ImportResult, ImportedDocument


MODULE_DIR = Path(__file__).resolve().parents[1] / "services" / "minuta" / "inverse_conversion_catalog"


class InverseConversionCatalogTests(unittest.TestCase):
    def _docx(self, path: Path, paragraphs: list[str] | None = None, table_text: str | None = None) -> Path:
        document = Document()
        for text in paragraphs or []:
            document.add_paragraph(text)
        if table_text is not None:
            table = document.add_table(rows=1, cols=1)
            table.cell(0, 0).text = table_text
        document.save(path)
        return path

    def test_extracts_single_marker(self):
        with TemporaryDirectory() as tmp_dir:
            path = self._docx(Path(tmp_dir) / "single.docx", ["Precio {{VALOR_VENTA}} pactado."])

            markers = DocxMarkerExtractor().extract_from_docx(path)

            self.assertEqual(len(markers), 1)
            self.assertEqual(markers[0].raw_field_code, "VALOR_VENTA")
            self.assertEqual(markers[0].location.paragraph_index, 0)

    def test_extracts_multiple_markers_in_same_paragraph(self):
        with TemporaryDirectory() as tmp_dir:
            path = self._docx(Path(tmp_dir) / "multi.docx", ["{{NOMBRE_COMPRADOR}} compra por {{PRECIO_VENTA}}."])

            markers = DocxMarkerExtractor().extract_from_docx(path)

            self.assertEqual([marker.raw_field_code for marker in markers], ["NOMBRE_COMPRADOR", "PRECIO_VENTA"])

    def test_extracts_marker_in_table(self):
        with TemporaryDirectory() as tmp_dir:
            path = self._docx(Path(tmp_dir) / "table.docx", table_text="Documento {{CEDULA_COMPRADOR}}")

            marker = DocxMarkerExtractor().extract_from_docx(path)[0]

            self.assertEqual(marker.raw_field_code, "CEDULA_COMPRADOR")
            self.assertEqual(marker.location.table_index, 0)
            self.assertEqual(marker.location.row_index, 0)
            self.assertEqual(marker.location.cell_index, 0)

    def test_extracts_text_before_and_text_after(self):
        markers = DocxMarkerExtractor().extract_from_text("El precio es {{VALOR_VENTA}} moneda corriente.")

        self.assertEqual(markers[0].text_before, "El precio es")
        self.assertEqual(markers[0].text_after, "moneda corriente.")

    def test_normalizes_uppercase_snake_case(self):
        self.assertEqual(FieldCodeNormalizer().normalize(" valor acto "), "VALOR_ACTO")

    def test_cleans_accents_and_weird_characters(self):
        self.assertEqual(FieldCodeNormalizer().normalize("{{ Nùmero  documento--comprador!! }}"), "NUMERO_DOCUMENTO_COMPRADOR")

    def test_value_aliases_are_suggested_towards_valor_venta(self):
        aliases = FieldAliasBuilder().build(
            {
                "VALOR_ACTO": 2,
                "PRECIO_VENTA": 4,
                "EN_NUMEROS_VALOR_DE_LA_VENTA": 1,
                "VALOR_VENTA": 5,
            }
        )

        by_raw = {alias.raw_field_code: alias for alias in aliases}
        self.assertEqual(by_raw["PRECIO_VENTA"].canonical_field_code, "VALOR_VENTA")
        self.assertEqual(by_raw["EN_NUMEROS_VALOR_DE_LA_VENTA"].canonical_field_code, "VALOR_VENTA")
        self.assertNotEqual(by_raw["VALOR_ACTO"].canonical_field_code, "VALOR_VENTA")

    def test_tipo_documento_comprador_is_not_grouped_with_numero_documento_comprador(self):
        by_raw = self._aliases_by_raw(
            {
                "TIPO_DOCUMENTO_COMPRADOR": 8,
                "NUMERO_DOCUMENTO_COMPRADOR": 14,
            }
        )

        self.assertNotEqual(by_raw["TIPO_DOCUMENTO_COMPRADOR"].canonical_field_code, "NUMERO_DOCUMENTO_COMPRADOR")

    def test_tipo_de_documento_otorgante_is_not_grouped_with_numero_documento_otorgante(self):
        by_raw = self._aliases_by_raw(
            {
                "TIPO_DE_DOCUMENTO_OTORGANTE": 16,
                "NUMERO_DOCUMENTO_OTORGANTE": 19,
            }
        )

        self.assertNotEqual(by_raw["TIPO_DE_DOCUMENTO_OTORGANTE"].canonical_field_code, "NUMERO_DOCUMENTO_OTORGANTE")

    def test_municipio_expedicion_documento_is_not_grouped_with_numero_documento(self):
        by_raw = self._aliases_by_raw(
            {
                "MUNICIPIO_DE_EXPEDICION_DOCUMENTO_COMPRADOR": 2,
                "NUMERO_DOCUMENTO_COMPRADOR": 14,
            }
        )

        self.assertNotEqual(
            by_raw["MUNICIPIO_DE_EXPEDICION_DOCUMENTO_COMPRADOR"].canonical_field_code,
            "NUMERO_DOCUMENTO_COMPRADOR",
        )

    def test_fecha_nacimiento_is_not_grouped_with_municipio_nacimiento(self):
        by_raw = self._aliases_by_raw(
            {
                "FECHA_NACIMIENTO_DEL_MENOR": 1,
                "MUNICIPIO_NACIMIENTO_DEL_MENOR": 1,
            }
        )

        self.assertNotEqual(by_raw["FECHA_NACIMIENTO_DEL_MENOR"].canonical_field_code, "MUNICIPIO_NACIMIENTO_DEL_MENOR")
        self.assertNotEqual(by_raw["MUNICIPIO_NACIMIENTO_DEL_MENOR"].canonical_field_code, "FECHA_NACIMIENTO_DEL_MENOR")

    def test_date_parts_are_not_grouped_together(self):
        by_raw = self._aliases_by_raw(
            {
                "DIA_ELABORACION_ESCRITURA": 10,
                "MES_ELABORACION_ESCRITURA": 11,
                "ANO_ELABORACION_ESCRITURA": 6,
            }
        )

        self.assertNotEqual(by_raw["DIA_ELABORACION_ESCRITURA"].canonical_field_code, "MES_ELABORACION_ESCRITURA")
        self.assertNotEqual(by_raw["DIA_ELABORACION_ESCRITURA"].canonical_field_code, "ANO_ELABORACION_ESCRITURA")
        self.assertNotEqual(by_raw["MES_ELABORACION_ESCRITURA"].canonical_field_code, "ANO_ELABORACION_ESCRITURA")

    def test_numero_escritura_letters_and_numbers_are_not_direct_aliases(self):
        by_raw = self._aliases_by_raw(
            {
                "NUMERO_ESCRITURA_EN_LETRAS": 7,
                "NUMERO_ESCRITURA_EN_NUMEROS": 7,
            }
        )

        self.assertNotEqual(by_raw["NUMERO_ESCRITURA_EN_LETRAS"].canonical_field_code, "NUMERO_ESCRITURA_EN_NUMEROS")
        self.assertNotEqual(by_raw["NUMERO_ESCRITURA_EN_NUMEROS"].canonical_field_code, "NUMERO_ESCRITURA_EN_LETRAS")

    def test_buyer_ordinals_are_not_collapsed_automatically(self):
        by_raw = self._aliases_by_raw(
            {
                "NOMBRE_COMPRADOR": 22,
                "NOMBRE_COMPRADOR_1": 34,
                "NOMBRE_COMPRADOR_2": 23,
            }
        )

        self.assertEqual(by_raw["NOMBRE_COMPRADOR_1"].canonical_field_code, "NOMBRE_COMPRADOR_1")
        self.assertEqual(by_raw["NOMBRE_COMPRADOR_2"].canonical_field_code, "NOMBRE_COMPRADOR_2")
        self.assertNotEqual(by_raw["NOMBRE_COMPRADOR_1"].canonical_field_code, "NOMBRE_COMPRADOR")
        self.assertNotEqual(by_raw["NOMBRE_COMPRADOR_2"].canonical_field_code, "NOMBRE_COMPRADOR")

    def test_safe_sale_value_aliases_still_point_to_valor_venta(self):
        by_raw = self._aliases_by_raw(
            {
                "VALOR_DE_LA_VENTA_EN_NUMEROS": 10,
                "EN_NUMEROS_VALOR_DE_LA_VENTA": 9,
                "VALOR_VENTA": 4,
            }
        )

        self.assertEqual(by_raw["VALOR_DE_LA_VENTA_EN_NUMEROS"].canonical_field_code, "VALOR_VENTA")
        self.assertEqual(by_raw["EN_NUMEROS_VALOR_DE_LA_VENTA"].canonical_field_code, "VALOR_VENTA")
        self.assertNotEqual(by_raw["VALOR_DE_LA_VENTA_EN_NUMEROS"].status, "approved")
        self.assertNotEqual(by_raw["EN_NUMEROS_VALOR_DE_LA_VENTA"].status, "approved")

    def test_importer_tolerates_bad_document_and_continues(self):
        with TemporaryDirectory() as tmp_dir:
            root = Path(tmp_dir)
            self._docx(root / "good.docx", ["Valor {{VALOR_VENTA}}"])
            (root / "bad.docx").write_text("not a real docx", encoding="utf-8")

            result = CorpusImporter().import_path(root, dry_run=True)

            self.assertEqual(len(result.documents), 2)
            self.assertEqual(len(result.errors), 1)
            self.assertIn("VALOR_VENTA", result.field_frequency)

    def test_reporter_generates_json_and_csv(self):
        with TemporaryDirectory() as tmp_dir:
            result = ImportResult(documents=[ImportedDocument(filename="a.docx", source_zip=None, source_path="a.docx")])
            result.field_frequency = {"VALOR_VENTA": 1}

            written = CorpusReporter().write_import_reports(result, tmp_dir)

            self.assertTrue((Path(tmp_dir) / "import_summary.json").exists())
            self.assertTrue((Path(tmp_dir) / "import_summary.csv").exists())
            self.assertIn("field_frequency.json", written)

    def test_no_generated_field_is_approved_automatically(self):
        with TemporaryDirectory() as tmp_dir:
            root = Path(tmp_dir)
            self._docx(root / "good.docx", ["Valor {{VALOR_ACTO}}"])

            result = CorpusImporter().import_path(root, dry_run=True)

            self.assertNotIn("approved", {field.status for document in result.documents for field in document.fields})
            self.assertNotIn("approved", {alias.status for alias in result.aliases})

    def test_module_works_without_openai_api_key(self):
        old_value = os.environ.pop("OPENAI_API_KEY", None)
        try:
            self.assertEqual(FieldCodeNormalizer().normalize("precio venta"), "PRECIO_VENTA")
        finally:
            if old_value is not None:
                os.environ["OPENAI_API_KEY"] = old_value

    def test_module_does_not_depend_on_rag(self):
        self.assertFalse(self._module_matches(r"\b(import|from)\s+[\w.]*rag\b|\bRAG\b"))

    def test_module_does_not_touch_document_templates(self):
        self.assertFalse(self._module_contains("document_templates"))

    def test_module_does_not_touch_marked_template(self):
        self.assertFalse(self._module_contains("marked_template"))

    def test_module_does_not_modify_input_docx(self):
        with TemporaryDirectory() as tmp_dir:
            path = self._docx(Path(tmp_dir) / "source.docx", ["Precio {{VALOR_VENTA}}"])
            before = hashlib.sha256(path.read_bytes()).hexdigest()

            DocxMarkerExtractor().extract_from_docx(path)
            CorpusImporter().import_path(path, dry_run=True)

            after = hashlib.sha256(path.read_bytes()).hexdigest()
            self.assertEqual(before, after)

    def _module_contains(self, needle: str) -> bool:
        for path in MODULE_DIR.glob("*.py"):
            if needle in path.read_text(encoding="utf-8").lower():
                return True
        return False

    def _module_matches(self, pattern: str) -> bool:
        compiled = re.compile(pattern)
        for path in MODULE_DIR.glob("*.py"):
            if compiled.search(path.read_text(encoding="utf-8")):
                return True
        return False

    def _aliases_by_raw(self, frequency: dict[str, int]):
        return {alias.raw_field_code: alias for alias in FieldAliasBuilder().build(frequency)}


if __name__ == "__main__":
    unittest.main()
