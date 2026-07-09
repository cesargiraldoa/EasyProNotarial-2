from __future__ import annotations

import sys
import os
import json
import unittest
from dataclasses import asdict
from pathlib import Path
from tempfile import TemporaryDirectory
from unittest.mock import patch

from docx import Document

REPO_ROOT = Path(__file__).resolve().parents[3]
if str(REPO_ROOT) not in sys.path:
    sys.path.insert(0, str(REPO_ROOT))

from tools.notarial_template_lab.docx_structural_extractor import DocxStructuralExtractor
from tools.notarial_template_lab.document_profiler_agent import DocumentProfilerAgent
from tools.notarial_template_lab.field_proposal_agent import FieldProposalAgent
from tools.notarial_template_lab.field_proposer import FieldProposer
from tools.notarial_template_lab.human_review import apply_human_review, load_review_decisions
from tools.notarial_template_lab.models import FieldProposal, ProposalOccurrence
from tools.notarial_template_lab.prompt_contracts import compact_document_map
from tools.notarial_template_lab.roundtrip_validator import RoundtripValidator, iter_all_text
from tools.notarial_template_lab.run_lab import LabLLMExecutionError, mark_human_review_applied, run_lab
from tools.notarial_template_lab.report_writer import safe_document_name
from tools.notarial_template_lab.template_draft_writer import TemplateDraftWriter


class FakeLLMClient:
    def __init__(self):
        self.calls = []
        self.payload_sizes = []

    def complete_json(self, system_prompt: str, payload: dict) -> dict:
        self.calls.append((system_prompt, payload))
        self.payload_sizes.append(len(json.dumps(payload, ensure_ascii=False)))
        if "Contrato JSON" in system_prompt and "document_type" in system_prompt:
            return {
                "document_type": "escritura_publica",
                "recommended_mode": "documento_individual",
                "acts_detected": ["COMPRAVENTA"],
                "structural_sections": [{"title": "Comparecencia", "block_id": "block_x", "location": "body/paragraph 1"}],
                "parties_summary": [],
                "property_summary": {"items": []},
                "money_summary": [],
                "risk_notes": ["Revision humana requerida."],
                "confidence": 0.7,
                "evidence": [{"block_id": "block_x", "location": "body/paragraph 1", "reason": "Bloque sintetico."}],
            }
        target_block = next((block for block in payload["blocks"] if "matricula inmobiliaria" in block["text"]), None)
        if target_block is None:
            return {"proposals": []}
        value = "123-4567890"
        start = target_block["text"].index(value)
        return {
            "proposals": [
                {
                    "field_key": "matricula_inmobiliaria",
                    "label": "Matricula inmobiliaria",
                    "marker": "{{matricula_inmobiliaria}}",
                    "value": value,
                    "proposal_type": "project_field",
                    "role": None,
                    "scope": "inmueble",
                    "confidence": 0.91,
                    "occurrences": [
                        {
                            "block_id": target_block["block_id"],
                            "location": target_block["location"],
                            "start": start,
                            "end": start + len(value),
                        }
                    ],
                    "evidence": ["Patron de matricula con rotulo cercano."],
                    "reason": "Ubicacion exacta en DocumentMap compacto.",
                    "apply_strategy": "selected_occurrences",
                },
                {
                    "field_key": "sin_ubicacion",
                    "label": "Sin ubicacion",
                    "marker": "{{sin_ubicacion}}",
                    "value": "NO EXISTE",
                    "proposal_type": "document_field",
                    "role": None,
                    "scope": None,
                    "confidence": 0.95,
                    "occurrences": [],
                    "evidence": [],
                    "reason": "Debe normalizarse como review_required.",
                    "apply_strategy": "all_occurrences",
                },
            ]
        }


class DuplicateBatchLLMClient:
    def __init__(self):
        self.calls = []

    def complete_json(self, system_prompt: str, payload: dict) -> dict:
        self.calls.append((system_prompt, payload))
        if "Contrato JSON" in system_prompt and "document_type" in system_prompt:
            return {
                "document_type": "escritura_publica",
                "recommended_mode": "documento_individual",
                "acts_detected": [],
                "structural_sections": [],
                "parties_summary": [],
                "property_summary": {},
                "money_summary": [],
                "risk_notes": [],
                "confidence": 0.5,
                "evidence": [],
            }
        blocks = payload.get("blocks") or []
        if not blocks:
            return {"proposals": []}
        block = blocks[0]
        value = block["text"][: min(12, len(block["text"]))]
        if not value:
            return {"proposals": []}
        return {
            "proposals": [
                {
                    "field_key": "campo_repetido",
                    "label": "Campo repetido",
                    "marker": "{{campo_repetido}}",
                    "value": value,
                    "proposal_type": "document_field",
                    "role": None,
                    "scope": None,
                    "confidence": 0.9,
                    "occurrences": [
                        {
                            "block_id": block["block_id"],
                            "location": block["location"],
                            "start": 0,
                            "end": len(value),
                        }
                    ],
                    "evidence": ["Duplicado sintetico."],
                    "reason": "Test de deduplicacion.",
                    "apply_strategy": "selected_occurrences",
                },
                {
                    "field_key": "campo_repetido",
                    "label": "Campo repetido",
                    "marker": "{{campo_repetido}}",
                    "value": value,
                    "proposal_type": "document_field",
                    "role": None,
                    "scope": None,
                    "confidence": 0.9,
                    "occurrences": [
                        {
                            "block_id": block["block_id"],
                            "location": block["location"],
                            "start": 0,
                            "end": len(value),
                        }
                    ],
                    "evidence": ["Duplicado sintetico."],
                    "reason": "Test de deduplicacion.",
                    "apply_strategy": "selected_occurrences",
                },
            ]
        }


class NotarialTemplateLabTests(unittest.TestCase):
    def _build_docx(self, path: Path) -> None:
        doc = Document()
        doc.sections[0].header.add_paragraph("NOTARIA DE PRUEBA")
        doc.sections[0].footer.add_paragraph("PROTOCOLO 12/04/2026")

        paragraph = doc.add_paragraph()
        paragraph.add_run("EL COMPRADOR ")
        bold_run = paragraph.add_run("NOMBRE GENERICO UNO")
        bold_run.bold = True
        paragraph.add_run(", identificado con C.C. 123.456.789, celular 3012345678 y correo usuario@example.org.")

        doc.add_paragraph(
            "El inmueble corresponde al APARTAMENTO 1201 con matricula inmobiliaria 123-4567890 "
            "y codigo predial 123456789012345678901234567890."
        )
        doc.add_paragraph("El precio de venta es $98.765.432 y el saldo es $12.345.678.")

        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Rotulo"
        table.cell(0, 1).text = "Valor"
        table.cell(1, 0).text = "NIT otorgante"
        table.cell(1, 1).text = "800.111.222-3"
        nested = table.cell(1, 1).add_table(rows=1, cols=1)
        nested.cell(0, 0).text = "Fecha escritura 05/06/2026"
        doc.save(path)

    def _proposal_for_value(
        self,
        document_map,
        field_key: str,
        value: str,
        proposal_type: str = "document_field",
        confidence: float = 0.95,
    ) -> FieldProposal:
        occurrences = []
        for block in document_map.blocks:
            start = 0
            while True:
                index = block.raw_text.find(value, start)
                if index < 0:
                    break
                occurrences.append(
                    ProposalOccurrence(
                        occurrence_id=f"occ_{field_key}_{len(occurrences) + 1}",
                        block_id=block.block_id,
                        location=block.location,
                        text=value,
                        start=index,
                        end=index + len(value),
                        before=block.raw_text[max(0, index - 20):index],
                        after=block.raw_text[index + len(value):index + len(value) + 20],
                    )
                )
                start = index + len(value)
        return FieldProposal(
            field_key=field_key,
            label=field_key.replace("_", " ").capitalize(),
            marker=f"{{{{{field_key}}}}}",
            value=value,
            confidence=confidence,
            proposal_type=proposal_type,  # type: ignore[arg-type]
            occurrences=occurrences,
            apply_strategy="all_occurrences",
            reason="Synthetic proposal for human review tests.",
        )

    def _build_same_block_docx(self, path: Path) -> None:
        doc = Document()
        paragraph = doc.add_paragraph()
        paragraph.add_run("Datos confirmados: ")
        paragraph.add_run("123-4567890")
        paragraph.add_run(", ")
        paragraph.add_run("usuario@example.org")
        paragraph.add_run(", ")
        paragraph.add_run("$98.765.432")
        paragraph.add_run(".")
        doc.save(path)

    def test_extracts_paragraphs_tables_headers_footers_and_runs(self):
        with TemporaryDirectory() as tmp:
            source = Path(tmp) / "synthetic.docx"
            self._build_docx(source)

            document_map = DocxStructuralExtractor().extract(source)

            self.assertGreaterEqual(document_map.quality.total_blocks, 8)
            self.assertGreater(document_map.quality.total_runs, 0)
            self.assertTrue(any("header section" in block.location for block in document_map.blocks))
            self.assertTrue(any("footer section" in block.location for block in document_map.blocks))
            self.assertTrue(any("/table 1/row 2/cell 2" in block.location for block in document_map.blocks))
            self.assertTrue(any("/table 1/row 2/cell 2/table 1" in block.location for block in document_map.blocks))

            buyer_block = next(block for block in document_map.blocks if "NOMBRE GENERICO UNO" in block.raw_text)
            self.assertEqual(buyer_block.runs[0].start, 0)
            self.assertEqual(buyer_block.runs[0].end, len("EL COMPRADOR "))
            self.assertEqual(buyer_block.runs[1].start, len("EL COMPRADOR "))
            self.assertTrue(buyer_block.runs[1].style.bold)

    def test_detects_general_occurrences_and_proposes_generic_fields(self):
        with TemporaryDirectory() as tmp:
            source = Path(tmp) / "synthetic.docx"
            self._build_docx(source)

            document_map = DocxStructuralExtractor().extract(source)
            occurrence_types = set(document_map.occurrences_index)

            self.assertIn("real_estate_registration", occurrence_types)
            self.assertIn("dotted_document", occurrence_types)
            self.assertIn("nit", occurrence_types)
            self.assertIn("money", occurrence_types)
            self.assertIn("email", occurrence_types)
            self.assertIn("phone", occurrence_types)
            self.assertIn("long_property_code", occurrence_types)
            self.assertIn("date", occurrence_types)
            self.assertIn("uppercase_relevant_phrase", occurrence_types)

            proposals = FieldProposer().propose(document_map)
            keys = {proposal.field_key for proposal in proposals}

            self.assertIn("matricula_inmobiliaria", keys)
            self.assertIn("valor_precio", keys)
            self.assertIn("correo_comprador", keys)
            self.assertTrue(any(proposal.proposal_type == "review_required" for proposal in proposals))

    def test_writes_experimental_docx_and_roundtrips_with_marked_detector(self):
        with TemporaryDirectory() as tmp:
            source = Path(tmp) / "synthetic.docx"
            draft = Path(tmp) / "draft.docx"
            self._build_docx(source)

            document_map = DocxStructuralExtractor().extract(source)
            proposals = FieldProposer().propose(document_map)
            draft_result = TemplateDraftWriter().write(source, draft, document_map, proposals)
            validation = RoundtripValidator().validate(draft)

            self.assertTrue(draft.exists())
            self.assertGreater(len(draft_result.replacements), 0)
            self.assertTrue(validation.opens)
            self.assertTrue(validation.contains_markers)
            self.assertTrue(validation.passed, validation.errors)
            self.assertGreater(validation.marked_fields_count, 0)
            detected_keys = {field["key"] for field in validation.marked_fields}
            self.assertIn("matricula_inmobiliaria", detected_keys)

    def test_writer_applies_multiple_replacements_in_same_paragraph(self):
        with TemporaryDirectory() as tmp:
            source = Path(tmp) / "same_block.docx"
            draft = Path(tmp) / "draft.docx"
            self._build_same_block_docx(source)
            document_map = DocxStructuralExtractor().extract(source)
            proposals = [
                self._proposal_for_value(document_map, "matricula", "123-4567890"),
                self._proposal_for_value(document_map, "correo", "usuario@example.org"),
                self._proposal_for_value(document_map, "precio", "$98.765.432"),
            ]

            result = TemplateDraftWriter().write(source, draft, document_map, proposals)
            text = docx_text(draft)

            self.assertEqual(len(result.replacements), 3)
            self.assertIn("{{matricula}}", text)
            self.assertIn("{{correo}}", text)
            self.assertIn("{{precio}}", text)

    def test_writer_uses_descending_offsets_so_later_replacement_does_not_shift_earlier_one(self):
        with TemporaryDirectory() as tmp:
            source = Path(tmp) / "offsets.docx"
            draft = Path(tmp) / "draft.docx"
            doc = Document()
            doc.add_paragraph("Campos: 111111 y 222222.")
            doc.save(source)
            document_map = DocxStructuralExtractor().extract(source)
            proposals = [
                self._proposal_for_value(document_map, "campo_inicio", "111111"),
                self._proposal_for_value(document_map, "campo_final", "222222"),
            ]
            proposals[0].marker = "{{MARCADOR_MUCHO_MAS_LARGO_QUE_VALOR_ORIGINAL}}"

            result = TemplateDraftWriter().write(source, draft, document_map, proposals)
            text = docx_text(draft)

            self.assertEqual(len(result.replacements), 2)
            self.assertIn("{{MARCADOR_MUCHO_MAS_LARGO_QUE_VALOR_ORIGINAL}}", text)
            self.assertIn("{{campo_final}}", text)
            self.assertNotIn("111111", text)
            self.assertNotIn("222222", text)

    def test_writer_replaces_value_in_long_paragraph(self):
        with TemporaryDirectory() as tmp:
            source = Path(tmp) / "long.docx"
            draft = Path(tmp) / "draft.docx"
            doc = Document()
            doc.add_paragraph(f"{'texto ' * 250} FECHA 12/04/2026 {'contexto ' * 250}")
            doc.save(source)
            document_map = DocxStructuralExtractor().extract(source)
            proposal = self._proposal_for_value(document_map, "fecha_larga", "12/04/2026")

            result = TemplateDraftWriter().write(source, draft, document_map, [proposal])
            text = docx_text(draft)

            self.assertEqual(len(result.replacements), 1)
            self.assertIn("{{fecha_larga}}", text)
            self.assertNotIn("12/04/2026", text)

    def test_writer_replaces_value_split_across_runs_by_rebuilding_paragraph(self):
        with TemporaryDirectory() as tmp:
            source = Path(tmp) / "runs.docx"
            draft = Path(tmp) / "draft.docx"
            doc = Document()
            paragraph = doc.add_paragraph("Codigo: ")
            paragraph.add_run("ABC").bold = True
            paragraph.add_run("DEF").italic = True
            paragraph.add_run(" confirmado.")
            doc.save(source)
            document_map = DocxStructuralExtractor().extract(source)
            proposal = self._proposal_for_value(document_map, "codigo", "ABCDEF")

            result = TemplateDraftWriter().write(source, draft, document_map, [proposal])
            text = docx_text(draft)

            self.assertEqual(len(result.replacements), 1)
            self.assertIn("{{codigo}}", text)
            self.assertNotIn("ABCDEF", text)

    def test_writer_reports_block_id_not_found(self):
        with TemporaryDirectory() as tmp:
            source = Path(tmp) / "synthetic.docx"
            draft = Path(tmp) / "draft.docx"
            self._build_docx(source)
            document_map = DocxStructuralExtractor().extract(source)
            proposal = self._proposal_for_value(document_map, "correo", "usuario@example.org")
            proposal.occurrences[0].block_id = "block_missing"

            result = TemplateDraftWriter().write(source, draft, document_map, [proposal])

            self.assertEqual(len(result.replacements), 0)
            self.assertIn("block_id not found", result.skipped[0]["reason"])
            self.assertEqual(result.skipped[0]["start"], proposal.occurrences[0].start)

    def test_writer_uses_unique_text_fallback_on_offset_mismatch(self):
        with TemporaryDirectory() as tmp:
            source = Path(tmp) / "fallback.docx"
            draft = Path(tmp) / "draft.docx"
            doc = Document()
            doc.add_paragraph("Valor unico: ABC123.")
            doc.save(source)
            document_map = DocxStructuralExtractor().extract(source)
            proposal = self._proposal_for_value(document_map, "codigo", "ABC123")
            proposal.occurrences[0].start = 0
            proposal.occurrences[0].end = 6

            result = TemplateDraftWriter().write(source, draft, document_map, [proposal])
            text = docx_text(draft)

            self.assertEqual(len(result.replacements), 1)
            self.assertIn("unique text fallback", result.replacements[0].reason)
            self.assertIn("{{codigo}}", text)

    def test_writer_reports_ambiguous_text_fallback(self):
        with TemporaryDirectory() as tmp:
            source = Path(tmp) / "ambiguous.docx"
            draft = Path(tmp) / "draft.docx"
            doc = Document()
            doc.add_paragraph("Codigo ABC123 y otro ABC123.")
            doc.save(source)
            document_map = DocxStructuralExtractor().extract(source)
            proposal = self._proposal_for_value(document_map, "codigo", "ABC123")
            proposal.occurrences = [proposal.occurrences[0]]
            proposal.occurrences[0].start = 0
            proposal.occurrences[0].end = 6

            result = TemplateDraftWriter().write(source, draft, document_map, [proposal])

            self.assertEqual(len(result.replacements), 0)
            self.assertIn("expected text appears multiple times", result.skipped[0]["failure_reason"])

    def test_compacts_document_map_for_llm_with_auditable_offsets(self):
        with TemporaryDirectory() as tmp:
            source = Path(tmp) / "synthetic.docx"
            self._build_docx(source)

            document_map = DocxStructuralExtractor().extract(source)
            compact = compact_document_map(document_map)

            self.assertEqual(compact["document_id"], document_map.document_id)
            self.assertLessEqual(len(compact["blocks"]), len(document_map.blocks))
            self.assertIn("occurrences_index", compact)
            money = compact["occurrences_index"]["money"][0]
            self.assertIn("block_id", money)
            self.assertIn("location", money)
            self.assertIsInstance(money["start"], int)
            self.assertIsInstance(money["end"], int)
            self.assertIn("reduction_stats", compact)
            self.assertNotIn("runs", compact["blocks"][0])

    def test_compact_document_map_reduces_large_simulated_document(self):
        with TemporaryDirectory() as tmp:
            source = Path(tmp) / "large.docx"
            doc = Document()
            irrelevant = "Texto operativo sin senales notariales " * 60
            for _ in range(40):
                doc.add_paragraph(irrelevant)
            doc.add_paragraph("El inmueble tiene matricula inmobiliaria 123-4567890 y precio $98.765.432.")
            doc.save(source)

            document_map = DocxStructuralExtractor().extract(source)
            compact = compact_document_map(document_map, max_blocks=5, max_block_chars=90)
            compact_text = "\n".join(block["text"] for block in compact["blocks"])

            self.assertLess(compact["reduction_stats"]["selected_chars"], compact["reduction_stats"]["original_chars"])
            self.assertLessEqual(len(compact["blocks"]), 5)
            self.assertIn("123-4567890", compact_text)
            self.assertNotIn(irrelevant.strip(), compact_text)

    def test_llm_agents_parse_and_normalize_json(self):
        with TemporaryDirectory() as tmp:
            source = Path(tmp) / "synthetic.docx"
            self._build_docx(source)
            document_map = DocxStructuralExtractor().extract(source)
            client = FakeLLMClient()

            profile = DocumentProfilerAgent(client).run(document_map)
            proposals = FieldProposalAgent(client).run(document_map, profile)

            self.assertEqual(profile.recommended_mode, "documento_individual")
            self.assertEqual(proposals[0].field_key, "matricula_inmobiliaria")
            self.assertEqual(proposals[0].apply_strategy, "selected_occurrences")
            self.assertEqual(proposals[1].proposal_type, "review_required")
            self.assertEqual(proposals[1].apply_strategy, "review_required")

    def test_field_proposal_agent_processes_batches_dedupes_and_preserves_locations(self):
        with TemporaryDirectory() as tmp:
            source = Path(tmp) / "multi_batch.docx"
            doc = Document()
            for index in range(6):
                doc.add_paragraph(f"COMPRADOR GENERICO {index} correo usuario{index}@example.org.")
            doc.save(source)
            document_map = DocxStructuralExtractor().extract(source)
            client = DuplicateBatchLLMClient()

            profile = DocumentProfilerAgent(client).run(document_map, max_blocks=6, max_block_chars=120)
            proposals = FieldProposalAgent(client).run(
                document_map,
                profile,
                max_blocks_per_batch=2,
                max_block_chars=120,
            )
            field_calls = [call for call in client.calls if "propuesta de campos" in call[0]]

            self.assertGreater(len(field_calls), 1)
            self.assertEqual(len(proposals), len(field_calls))
            self.assertTrue(all(item.occurrences[0].block_id for item in proposals))
            self.assertTrue(all(item.occurrences[0].location for item in proposals))

    def test_proposal_without_block_id_is_not_replaceable(self):
        with TemporaryDirectory() as tmp:
            source = Path(tmp) / "synthetic.docx"
            draft = Path(tmp) / "draft.docx"
            self._build_docx(source)
            document_map = DocxStructuralExtractor().extract(source)

            proposal = FieldProposal(
                field_key="correo_comprador",
                label="Correo comprador",
                marker="{{correo_comprador}}",
                value="usuario@example.org",
                confidence=0.99,
                proposal_type="document_field",
                occurrences=[
                    ProposalOccurrence(
                        occurrence_id="occ_fake",
                        block_id="",
                        location="",
                        text="usuario@example.org",
                        start=0,
                        end=19,
                        before="",
                        after="",
                    )
                ],
                apply_strategy="selected_occurrences",
                reason="Sin ubicacion exacta.",
            )

            result = TemplateDraftWriter().write(source, draft, document_map, [proposal])

            self.assertEqual(len(result.replacements), 0)
            self.assertTrue(any("block_id" in item["reason"] for item in result.skipped))

    def test_run_lab_with_llm_mock_writes_profile_and_llm_proposals(self):
        with TemporaryDirectory() as tmp:
            source = Path(tmp) / "synthetic.docx"
            artifacts_root = Path(tmp) / "artifacts"
            self._build_docx(source)

            client = FakeLLMClient()
            result = run_lab(
                source,
                artifacts_root,
                use_llm=True,
                llm_client=client,
                llm_max_blocks_per_batch=2,
                llm_max_block_chars=160,
                llm_debug_payloads=True,
            )
            artifacts_dir = Path(result["artifacts_dir"])

            self.assertTrue((artifacts_dir / "02_document_profile.json").exists())
            self.assertTrue((artifacts_dir / "03_field_proposals_llm.json").exists())
            self.assertTrue((artifacts_dir / "02_llm_profile_payload.json").exists())
            self.assertTrue((artifacts_dir / "03_llm_field_payload_batch_001.json").exists())
            self.assertGreaterEqual(result["total_llm_proposals"], 1)
            self.assertLess(max(client.payload_sizes), 20000)

    def test_run_lab_with_llm_aborts_before_api_when_payload_is_too_large(self):
        with TemporaryDirectory() as tmp:
            source = Path(tmp) / "synthetic.docx"
            artifacts_root = Path(tmp) / "artifacts"
            self._build_docx(source)
            client = FakeLLMClient()

            with self.assertRaises(LabLLMExecutionError) as raised:
                run_lab(
                    source,
                    artifacts_root,
                    use_llm=True,
                    llm_client=client,
                    llm_max_blocks_per_batch=2,
                    llm_max_block_chars=160,
                    llm_max_estimated_tokens=1,
                )

            self.assertIn("payload too large, reduce batch size", str(raised.exception))
            self.assertEqual(client.calls, [])

    def test_reuse_llm_artifacts_loads_existing_proposals_without_openai_key(self):
        with TemporaryDirectory() as tmp:
            source = Path(tmp) / "synthetic.docx"
            artifacts_root = Path(tmp) / "artifacts"
            self._build_docx(source)
            document_map = DocxStructuralExtractor().extract(source)
            llm_proposal = self._proposal_for_value(document_map, "correo_llm", "usuario@example.org")
            write_llm_artifacts(artifacts_root / safe_document_name(source.stem), [llm_proposal])

            with patch.dict(os.environ, {"OPENAI_API_KEY": ""}):
                result = run_lab(source, artifacts_root, reuse_llm_artifacts=True)

            self.assertTrue(result["reused_llm_artifacts"])
            self.assertEqual(result["total_llm_proposals"], 1)

    def test_reuse_llm_artifacts_with_review_file_generates_confirmed_outputs(self):
        with TemporaryDirectory() as tmp:
            source = Path(tmp) / "synthetic.docx"
            artifacts_root = Path(tmp) / "artifacts"
            self._build_docx(source)
            document_map = DocxStructuralExtractor().extract(source)
            llm_proposal = self._proposal_for_value(document_map, "correo_llm", "usuario@example.org")
            write_llm_artifacts(artifacts_root / safe_document_name(source.stem), [llm_proposal])
            review_path = write_review_file(
                Path(tmp) / "review.json",
                [
                    {
                        "field_key": "correo_llm",
                        "value": "usuario@example.org",
                        "decision": "confirm",
                        "marker": "{{correo_desde_llm}}",
                        "apply_strategy": "all_occurrences",
                    }
                ],
            )

            result = run_lab(source, artifacts_root, reuse_llm_artifacts=True, review_file=review_path)
            artifacts_dir = Path(result["artifacts_dir"])

            self.assertTrue((artifacts_dir / "08_review_decisions_applied.json").exists())
            self.assertTrue((artifacts_dir / "09_template_confirmed.docx").exists())
            self.assertTrue((artifacts_dir / "10_confirmed_validation_report.json").exists())
            self.assertGreaterEqual(result["confirmed_replacements"], 1)
            self.assertTrue(result["confirmed_validation_passed"])

    def test_human_review_three_confirmed_fields_in_same_block_leave_three_markers(self):
        with TemporaryDirectory() as tmp:
            source = Path(tmp) / "same_block.docx"
            confirmed = Path(tmp) / "confirmed.docx"
            self._build_same_block_docx(source)
            document_map = DocxStructuralExtractor().extract(source)
            proposals = [
                self._proposal_for_value(document_map, "matricula", "123-4567890"),
                self._proposal_for_value(document_map, "correo", "usuario@example.org"),
                self._proposal_for_value(document_map, "precio", "$98.765.432"),
            ]
            decisions = load_review_decisions(
                write_review_file(
                    Path(tmp) / "review.json",
                    [
                        {"field_key": "matricula", "decision": "confirm", "apply_strategy": "all_occurrences"},
                        {"field_key": "correo", "decision": "confirm", "apply_strategy": "all_occurrences"},
                        {"field_key": "precio", "decision": "confirm", "apply_strategy": "all_occurrences"},
                    ],
                )
            )

            review = apply_human_review(proposals, decisions)
            draft_result = TemplateDraftWriter().write(source, confirmed, document_map, review.confirmed_proposals)
            mark_human_review_applied(review, draft_result)
            text = docx_text(confirmed)

            self.assertEqual(len(draft_result.replacements), 3)
            self.assertTrue(all(item.applied for item in review.applied_decisions if item.replaceable))
            self.assertEqual(text.count("{{"), 3)

    def test_human_review_failed_occurrences_include_writer_reason(self):
        with TemporaryDirectory() as tmp:
            source = Path(tmp) / "ambiguous.docx"
            confirmed = Path(tmp) / "confirmed.docx"
            doc = Document()
            doc.add_paragraph("Codigo ABC123 y otro ABC123.")
            doc.save(source)
            document_map = DocxStructuralExtractor().extract(source)
            proposal = self._proposal_for_value(document_map, "codigo", "ABC123")
            proposal.occurrences = [proposal.occurrences[0]]
            proposal.occurrences[0].start = 0
            proposal.occurrences[0].end = 6
            decisions = load_review_decisions(
                write_review_file(
                    Path(tmp) / "review.json",
                    [{"field_key": "codigo", "decision": "confirm", "apply_strategy": "all_occurrences"}],
                )
            )

            review = apply_human_review([proposal], decisions)
            draft_result = TemplateDraftWriter().write(source, confirmed, document_map, review.confirmed_proposals)
            mark_human_review_applied(review, draft_result)
            applied = review.applied_decisions[0]

            self.assertFalse(applied.applied)
            self.assertEqual(applied.applied_count, 0)
            self.assertEqual(applied.expected_count, 1)
            self.assertIn("expected text appears multiple times", applied.failed_occurrences[0]["failure_reason"])

    def test_regression_seventeen_confirmed_fields_produce_seventeen_markers(self):
        with TemporaryDirectory() as tmp:
            source = Path(tmp) / "seventeen.docx"
            artifacts_root = Path(tmp) / "artifacts"
            values = [f"dato-{index:02d}" for index in range(17)]
            doc = Document()
            doc.add_paragraph(" ".join(values))
            doc.save(source)
            document_map = DocxStructuralExtractor().extract(source)
            proposals = [self._proposal_for_value(document_map, f"campo_{index:02d}", value) for index, value in enumerate(values)]
            write_llm_artifacts(artifacts_root / safe_document_name(source.stem), proposals)
            review_path = write_review_file(
                Path(tmp) / "review.json",
                [
                    {"field_key": f"campo_{index:02d}", "decision": "confirm", "apply_strategy": "all_occurrences"}
                    for index in range(17)
                ],
            )

            result = run_lab(source, artifacts_root, reuse_llm_artifacts=True, review_file=review_path)
            text = docx_text(Path(result["artifacts_dir"]) / "09_template_confirmed.docx")
            review_payload = json.loads((Path(result["artifacts_dir"]) / "08_review_decisions_applied.json").read_text(encoding="utf-8"))

            self.assertEqual(result["confirmed_replacements"], 17)
            self.assertEqual(text.count("{{"), 17)
            self.assertTrue(all(item["applied"] for item in review_payload["applied_decisions"] if item["replaceable"]))

    def test_regression_confirmed_replacements_match_expected_markers_with_reused_llm_artifacts(self):
        with TemporaryDirectory() as tmp:
            source = Path(tmp) / "same_block.docx"
            artifacts_root = Path(tmp) / "artifacts"
            self._build_same_block_docx(source)
            document_map = DocxStructuralExtractor().extract(source)
            proposals = [
                self._proposal_for_value(document_map, "matricula", "123-4567890"),
                self._proposal_for_value(document_map, "correo", "usuario@example.org"),
                self._proposal_for_value(document_map, "precio", "$98.765.432"),
            ]
            write_llm_artifacts(artifacts_root / safe_document_name(source.stem), proposals)
            review_path = write_review_file(
                Path(tmp) / "review.json",
                [
                    {"field_key": "matricula", "decision": "confirm", "apply_strategy": "all_occurrences"},
                    {"field_key": "correo", "decision": "confirm", "apply_strategy": "all_occurrences"},
                    {"field_key": "precio", "decision": "confirm", "apply_strategy": "all_occurrences"},
                ],
            )

            result = run_lab(source, artifacts_root, reuse_llm_artifacts=True, review_file=review_path)
            text = docx_text(Path(result["artifacts_dir"]) / "09_template_confirmed.docx")

            self.assertEqual(result["confirmed_replacements"], 3)
            self.assertIn("{{matricula}}", text)
            self.assertIn("{{correo}}", text)
            self.assertIn("{{precio}}", text)
            self.assertEqual(text.count("{{"), 3)

    def test_load_review_decisions_accepts_utf8_bom(self):
        with TemporaryDirectory() as tmp:
            review_path = Path(tmp) / "review.json"
            review_path.write_text(
                json.dumps({"decisions": [{"field_key": "campo", "decision": "confirm"}]}),
                encoding="utf-8-sig",
            )

            decisions = load_review_decisions(review_path)

            self.assertEqual(decisions[0].field_key, "campo")

    def test_reuse_llm_artifacts_fails_cleanly_when_proposals_are_missing(self):
        with TemporaryDirectory() as tmp:
            source = Path(tmp) / "synthetic.docx"
            artifacts_root = Path(tmp) / "artifacts"
            self._build_docx(source)
            artifacts_dir = artifacts_root / safe_document_name(source.stem)
            artifacts_dir.mkdir(parents=True, exist_ok=True)
            (artifacts_dir / "02_document_profile.json").write_text(json.dumps(document_profile_payload()), encoding="utf-8")

            with self.assertRaises(LabLLMExecutionError) as raised:
                run_lab(source, artifacts_root, reuse_llm_artifacts=True)

            self.assertIn("no existing LLM artifacts found", str(raised.exception))
            self.assertIn("03_field_proposals_llm.json", str(raised.exception))

    def test_human_review_confirmed_exact_location_replaces_and_renames_marker(self):
        with TemporaryDirectory() as tmp:
            source = Path(tmp) / "synthetic.docx"
            confirmed = Path(tmp) / "confirmed.docx"
            self._build_docx(source)
            document_map = DocxStructuralExtractor().extract(source)
            proposal = self._proposal_for_value(document_map, "correo_comprador", "usuario@example.org")
            decisions = load_review_decisions(
                write_review_file(
                    Path(tmp) / "review.json",
                    [
                        {
                            "field_key": "correo_comprador",
                            "decision": "confirm",
                            "marker": "{{email_cliente}}",
                            "apply_strategy": "all_occurrences",
                        }
                    ],
                )
            )

            review = apply_human_review([proposal], decisions)
            draft_result = TemplateDraftWriter().write(source, confirmed, document_map, review.confirmed_proposals)
            text = "\n".join(paragraph.text for paragraph in Document(str(confirmed)).paragraphs)

            self.assertEqual(len(draft_result.replacements), 1)
            self.assertIn("{{email_cliente}}", text)
            self.assertNotIn("usuario@example.org", text)

    def test_human_review_selected_occurrences_replaces_only_selected(self):
        with TemporaryDirectory() as tmp:
            source = Path(tmp) / "repeated.docx"
            confirmed = Path(tmp) / "confirmed.docx"
            doc = Document()
            doc.add_paragraph("Correos: usuario@example.org y usuario@example.org.")
            doc.save(source)
            document_map = DocxStructuralExtractor().extract(source)
            proposal = self._proposal_for_value(document_map, "correo", "usuario@example.org")
            selected_id = proposal.occurrences[1].occurrence_id
            decisions = load_review_decisions(
                write_review_file(
                    Path(tmp) / "review.json",
                    [
                        {
                            "field_key": "correo",
                            "decision": "confirm",
                            "apply_strategy": "selected_occurrences",
                            "occurrence_ids": [selected_id],
                        }
                    ],
                )
            )

            review = apply_human_review([proposal], decisions)
            draft_result = TemplateDraftWriter().write(source, confirmed, document_map, review.confirmed_proposals)
            text = Document(str(confirmed)).paragraphs[0].text

            self.assertEqual(len(draft_result.replacements), 1)
            self.assertEqual(text.count("usuario@example.org"), 1)
            self.assertEqual(text.count("{{correo}}"), 1)
            self.assertTrue(text.startswith("Correos: usuario@example.org y {{correo}}"))

    def test_human_review_blocks_missing_location_review_required_and_fixed_text(self):
        with TemporaryDirectory() as tmp:
            source = Path(tmp) / "synthetic.docx"
            self._build_docx(source)
            document_map = DocxStructuralExtractor().extract(source)
            missing_location = FieldProposal(
                field_key="sin_ubicacion",
                label="Sin ubicacion",
                marker="{{sin_ubicacion}}",
                value="usuario@example.org",
                confidence=0.99,
                proposal_type="document_field",
                occurrences=[
                    ProposalOccurrence(
                        occurrence_id="occ_missing",
                        block_id="",
                        location="",
                        text="usuario@example.org",
                        start=0,
                        end=19,
                        before="",
                        after="",
                    )
                ],
                apply_strategy="all_occurrences",
                reason="Missing location.",
            )
            review_required = self._proposal_for_value(document_map, "nombre_revision", "NOMBRE GENERICO UNO", "review_required", 0.99)
            fixed_text = self._proposal_for_value(document_map, "texto_fijo", "EL COMPRADOR", "fixed_text", 0.99)
            low_confidence = self._proposal_for_value(document_map, "correo_bajo", "usuario@example.org", "document_field", 0.5)
            decisions = [
                {"field_key": item.field_key, "decision": "confirm", "apply_strategy": "all_occurrences"}
                for item in [missing_location, review_required, fixed_text, low_confidence]
            ]
            review = apply_human_review(
                [missing_location, review_required, fixed_text, low_confidence],
                [decision for decision in load_review_decisions(write_review_file(Path(tmp) / "review.json", decisions))],
            )

            self.assertEqual(len(review.confirmed_proposals), 0)
            reasons = " ".join(item.block_reason or "" for item in review.applied_decisions)
            self.assertIn("block_id", reasons)
            self.assertIn("review_required", reasons)
            self.assertIn("fixed_text", reasons)
            self.assertIn("Confianza insuficiente", reasons)

    def test_run_lab_without_review_file_does_not_generate_confirmed_template(self):
        with TemporaryDirectory() as tmp:
            source = Path(tmp) / "synthetic.docx"
            artifacts_root = Path(tmp) / "artifacts"
            self._build_docx(source)

            result = run_lab(source, artifacts_root)
            artifacts_dir = Path(result["artifacts_dir"])

            self.assertTrue((artifacts_dir / "06_template_draft.docx").exists())
            self.assertFalse((artifacts_dir / "09_template_confirmed.docx").exists())

    def test_run_lab_with_review_file_generates_confirmed_template_and_validation(self):
        with TemporaryDirectory() as tmp:
            source = Path(tmp) / "synthetic.docx"
            artifacts_root = Path(tmp) / "artifacts"
            self._build_docx(source)
            document_map = DocxStructuralExtractor().extract(source)
            proposal = self._proposal_for_value(document_map, "correo_comprador", "usuario@example.org")
            review_path = write_review_file(
                Path(tmp) / "review.json",
                [
                    {
                        "field_key": "correo_comprador",
                        "value": proposal.value,
                        "decision": "confirm",
                        "marker": "{{email_cliente}}",
                        "apply_strategy": "all_occurrences",
                    }
                ],
            )

            result = run_lab(source, artifacts_root, review_file=review_path)
            artifacts_dir = Path(result["artifacts_dir"])

            self.assertTrue((artifacts_dir / "08_review_decisions_applied.json").exists())
            self.assertTrue((artifacts_dir / "09_template_confirmed.docx").exists())
            self.assertTrue((artifacts_dir / "10_confirmed_validation_report.json").exists())
            self.assertGreaterEqual(result["confirmed_replacements"], 1)
            self.assertTrue(result["confirmed_validation_passed"])


def write_review_file(path: Path, decisions: list[dict]) -> Path:
    path.write_text(json.dumps({"decisions": decisions}, ensure_ascii=False), encoding="utf-8")
    return path


def write_llm_artifacts(artifacts_dir: Path, proposals: list[FieldProposal]) -> None:
    artifacts_dir.mkdir(parents=True, exist_ok=True)
    (artifacts_dir / "02_document_profile.json").write_text(
        json.dumps(document_profile_payload(), ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    (artifacts_dir / "03_field_proposals_llm.json").write_text(
        json.dumps([asdict(proposal) for proposal in proposals], ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


def document_profile_payload() -> dict:
    return {
        "document_type": "escritura_publica",
        "recommended_mode": "documento_individual",
        "acts_detected": ["COMPRAVENTA"],
        "structural_sections": [],
        "parties_summary": [],
        "property_summary": {},
        "money_summary": [],
        "risk_notes": [],
        "confidence": 0.7,
        "evidence": [],
    }


def docx_text(path: Path) -> str:
    return "\n".join(iter_all_text(Document(str(path))))


if __name__ == "__main__":
    unittest.main()
