import unittest
from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document
from docx.shared import RGBColor

from app.services.minuta.marked_template_generator import (
    _derive_missing_marked_values,
    _format_money_value,
    _is_money_field,
    _normalize_money_values_for_fields,
    apply_marked_template_replacements,
    cleanup_empty_second_buyer_text,
    is_second_buyer_empty,
    validate_gender_concordance,
)
from app.services.minuta.engine.notarial_document_engine import NotarialDocumentEngine, NotarialRenderBlockedError
from app.services.minuta.rules.date_rules import contractual_date_text, notarial_date_text
from app.services.minuta.rules.money_rules import number_to_words
from app.services.minuta.rules.person_rules import collective_adjective, collective_label


class MarkedTemplateGeneratorCleanupTests(unittest.TestCase):
    def test_is_second_buyer_empty_returns_true_when_all_known_fields_are_blank(self):
        values = {
            "nombre_comprador_2": "",
            "tipo_documento_comprador_2": "",
            "numero_documento_comprador_2": "",
            "direccion_comprador_2": "   ",
            "celular_comprador_2": "",
            "email_comprador_2": "",
        }

        self.assertTrue(is_second_buyer_empty(values))

    def test_is_second_buyer_empty_returns_false_when_any_known_field_has_value(self):
        values = {
            "nombre_comprador_2": "MARIA PEREZ",
            "tipo_documento_comprador_2": "",
        }

        self.assertFalse(is_second_buyer_empty(values))

    def test_cleanup_empty_second_buyer_text_removes_inline_residue_patterns(self):
        self.assertEqual(
            cleanup_empty_second_buyer_text("JUAN CAMILO VASQUEZ MIRA - C.C - 1.037.657.164 y  -  -"),
            "JUAN CAMILO VASQUEZ MIRA - C.C - 1.037.657.164",
        )
        self.assertEqual(
            cleanup_empty_second_buyer_text("JUAN CAMILO VASQUEZ MIRA y ,"),
            "JUAN CAMILO VASQUEZ MIRA",
        )
        self.assertEqual(
            cleanup_empty_second_buyer_text("JUAN CAMILO VASQUEZ MIRA y , de las condiciones del contrato"),
            "JUAN CAMILO VASQUEZ MIRA de las condiciones del contrato",
        )
        self.assertEqual(
            cleanup_empty_second_buyer_text("JUAN CAMILO VASQUEZ MIRA y , identificado con cédula"),
            "JUAN CAMILO VASQUEZ MIRA, identificado con cédula",
        )
        self.assertEqual(
            cleanup_empty_second_buyer_text("JUAN CAMILO VASQUEZ MIRA y , quien comparece"),
            "JUAN CAMILO VASQUEZ MIRA quien comparece",
        )
        self.assertEqual(
            cleanup_empty_second_buyer_text("JUAN CAMILO VASQUEZ MIRA y , respectivamente"),
            "JUAN CAMILO VASQUEZ MIRA",
        )

    def test_cleanup_empty_second_buyer_text_removes_empty_profile_lines(self):
        self.assertEqual(cleanup_empty_second_buyer_text("DIRECCIÓN: ."), "")
        self.assertEqual(cleanup_empty_second_buyer_text("CELULAR: ."), "")
        self.assertEqual(cleanup_empty_second_buyer_text("CORREO: ."), "")
        self.assertEqual(cleanup_empty_second_buyer_text("Profesión u Oficio:"), "")
        self.assertEqual(cleanup_empty_second_buyer_text("Actividad Económica:"), "")
        self.assertEqual(cleanup_empty_second_buyer_text("Estado Civil:"), "")

    def test_derive_missing_marked_values_fills_letters_without_forcing_origin(self):
        values = {
            "valor_de_la_venta_en_numeros": "212.600.000",
            "en_numeros_cuota_inicial": "63.040.480",
            "valor_del_acto_de_la_hipoteca": "149.559.520",
            "origen_cuota_inicial": "",
        }
        field_keys = {
            "valor_de_la_venta_en_numeros",
            "valor_apartamento_en_letras",
            "en_numeros_cuota_inicial",
            "en_letras_cuota_inicial",
            "valor_del_acto_de_la_hipoteca",
            "valor_del_acto_de_la_hipoteca_en_letras",
            "origen_cuota_inicial",
        }

        derived = _derive_missing_marked_values(values, field_keys)

        self.assertEqual(
            derived["valor_apartamento_en_letras"],
            "DOSCIENTOS DOCE MILLONES SEISCIENTOS MIL PESOS MONEDA CORRIENTE",
        )
        self.assertEqual(
            derived["en_letras_cuota_inicial"],
            "SESENTA Y TRES MILLONES CUARENTA MIL CUATROCIENTOS OCHENTA PESOS MONEDA CORRIENTE",
        )
        self.assertEqual(
            derived["valor_del_acto_de_la_hipoteca_en_letras"],
            "CIENTO CUARENTA Y NUEVE MILLONES QUINIENTOS CINCUENTA Y NUEVE MIL QUINIENTOS VEINTE PESOS MONEDA CORRIENTE",
        )
        self.assertEqual(derived["origen_cuota_inicial"], "")

    def test_derive_missing_marked_values_includes_cents_in_money_letters(self):
        values = {"en_numeros_cuota_inicial": "63.040.480,76"}
        field_keys = {"en_numeros_cuota_inicial", "en_letras_cuota_inicial"}

        derived = _derive_missing_marked_values(values, field_keys)

        self.assertEqual(derived["en_numeros_cuota_inicial"], "63.040.480,76")
        self.assertEqual(
            derived["en_letras_cuota_inicial"],
            "SESENTA Y TRES MILLONES CUARENTA MIL CUATROCIENTOS OCHENTA PESOS MONEDA CORRIENTE CON SETENTA Y SEIS CENTAVOS",
        )

    def test_derive_sale_quota_and_mortgage_letters_use_current_currency_and_cents(self):
        values = {
            "valor_de_la_venta_en_numeros": "680785432,12",
            "en_numeros_cuota_inicial": "456789021,45",
            "valor_del_acto_de_la_hipoteca": "345876234,12",
        }
        field_keys = {
            "valor_de_la_venta_en_numeros",
            "valor_apartamento_en_letras",
            "en_numeros_cuota_inicial",
            "en_letras_cuota_inicial",
            "valor_del_acto_de_la_hipoteca",
            "valor_del_acto_de_la_hipoteca_en_letras",
        }

        derived = _derive_missing_marked_values(values, field_keys)

        self.assertEqual(derived["valor_de_la_venta_en_numeros"], "680.785.432,12")
        self.assertEqual(derived["en_numeros_cuota_inicial"], "456.789.021,45")
        self.assertEqual(derived["valor_del_acto_de_la_hipoteca"], "345.876.234,12")
        self.assertEqual(
            derived["valor_apartamento_en_letras"],
            "SEISCIENTOS OCHENTA MILLONES SETECIENTOS OCHENTA Y CINCO MIL CUATROCIENTOS TREINTA Y DOS PESOS MONEDA CORRIENTE CON DOCE CENTAVOS",
        )
        self.assertEqual(
            derived["en_letras_cuota_inicial"],
            "CUATROCIENTOS CINCUENTA Y SEIS MILLONES SETECIENTOS OCHENTA Y NUEVE MIL VEINTE Y UN PESOS MONEDA CORRIENTE CON CUARENTA Y CINCO CENTAVOS",
        )
        self.assertEqual(
            derived["valor_del_acto_de_la_hipoteca_en_letras"],
            "TRESCIENTOS CUARENTA Y CINCO MILLONES OCHOCIENTOS SETENTA Y SEIS MIL DOSCIENTOS TREINTA Y CUATRO PESOS MONEDA CORRIENTE CON DOCE CENTAVOS",
        )

    def test_money_normalization_formats_general_monetary_fields_without_touching_non_monetary_numbers(self):
        fields = [
            {"key": "avaluo_total_catastral", "label": "Avalúo total catastral", "section": "values"},
            {"key": "retencion_en_la_fuente", "label": "Retención en la fuente", "section": "liquidation"},
            {"key": "numero_documento_comprador_1", "label": "Número documento comprador 1", "section": "buyers"},
        ]
        values = {
            "avaluo_total_catastral": "120000000",
            "retencion_en_la_fuente": "119000",
            "numero_documento_comprador_1": "1037657164",
        }

        normalized = _normalize_money_values_for_fields(values, fields)

        self.assertTrue(_is_money_field(fields[0]))
        self.assertTrue(_is_money_field(fields[1]))
        self.assertFalse(_is_money_field(fields[2]))
        self.assertEqual(normalized["avaluo_total_catastral"], "120.000.000")
        self.assertEqual(normalized["retencion_en_la_fuente"], "119.000")
        self.assertEqual(normalized["numero_documento_comprador_1"], "1037657164")
        self.assertEqual(_format_money_value("212.600.000"), "212.600.000")

    def test_money_words_use_de_pesos_for_exact_millions(self):
        self.assertEqual(number_to_words(3_000_000, "PESOS"), "TRES MILLONES DE PESOS")
        self.assertEqual(number_to_words(4_000_000, "PESOS"), "CUATRO MILLONES DE PESOS")
        self.assertEqual(number_to_words(600_000_000, "PESOS"), "SEISCIENTOS MILLONES DE PESOS")
        self.assertEqual(
            number_to_words(678_943_221, "PESOS"),
            "SEISCIENTOS SETENTA Y OCHO MILLONES NOVECIENTOS CUARENTA Y TRES MIL DOSCIENTOS VEINTE Y UN PESOS",
        )

    def test_contractual_and_deed_dates_have_different_notarial_formats(self):
        self.assertEqual(
            notarial_date_text("2025-06-12"),
            "doce (12) días del mes de junio del año dos mil veinticinco (2025)",
        )
        self.assertEqual(
            contractual_date_text("2026-06-12"),
            "doce (12) de junio de dos mil veintiséis (2026)",
        )
        self.assertNotIn("veintiseis", contractual_date_text("2026-06-12"))
        self.assertNotIn("días del mes", contractual_date_text("2026-06-12"))

    def test_collective_person_concordance_labels(self):
        self.assertEqual(collective_label("comprador", ["F"]), "la compradora")
        self.assertEqual(collective_adjective("identificado", ["F"]), "identificada")
        self.assertEqual(collective_adjective("domiciliado", ["F"]), "domiciliada")
        self.assertEqual(collective_adjective("colombiano", ["F"]), "colombiana")
        self.assertEqual(collective_label("comprador", ["M"]), "el comprador")
        self.assertEqual(collective_adjective("identificado", ["M"]), "identificado")
        self.assertEqual(collective_label("comprador", ["M", "F"]), "los compradores")
        self.assertEqual(collective_adjective("identificado", ["M", "F"]), "identificados")
        self.assertEqual(collective_label("comprador", ["F", "F"]), "las compradoras")
        self.assertEqual(collective_adjective("domiciliado", ["F", "F"]), "domiciliadas")
        self.assertEqual(collective_label("deudor", ["M", "M"]), "los deudores")

    def test_validate_gender_concordance_returns_structured_warnings(self):
        values = {
            "comprador_es_hombre_o_mujer": "MUJER",
            "nacionalidad_comprador": "colombiano",
            "estado_civil_comprador": "casado",
        }

        warnings = validate_gender_concordance(values)

        self.assertGreaterEqual(len(warnings), 2)
        self.assertTrue(any(item["field_key"] == "nacionalidad_comprador" for item in warnings))
        self.assertTrue(any(item["field_key"] == "estado_civil_comprador" for item in warnings))

    def test_apply_marked_template_replacements_preserves_runs_and_paints_inserted_value_red(self):
        with TemporaryDirectory() as tmp_dir:
            source = Path(tmp_dir) / "source.docx"
            target = Path(tmp_dir) / "target.docx"

            doc = Document()
            paragraph = doc.add_paragraph()
            prefix_run = paragraph.add_run("CLÁUSULA CUARTA: ")
            prefix_run.bold = True
            paragraph.add_run("es la suma de ")
            placeholder_run_start = paragraph.add_run("{{VALOR")
            placeholder_run_start.bold = True
            placeholder_run_end = paragraph.add_run("_VENTA}}")
            placeholder_run_end.bold = True
            doc.save(source)

            stats = apply_marked_template_replacements(
                source,
                target,
                {"valor_venta": "212600000"},
                [
                    {
                        "key": "valor_venta",
                        "label": "Valor venta",
                        "section": "values",
                        "raw_markers": ["{{VALOR_VENTA}}"],
                    }
                ],
            )

            generated = Document(str(target))
            generated_paragraph = generated.paragraphs[0]
            self.assertIn("CLÁUSULA CUARTA: es la suma de 212.600.000", generated_paragraph.text)
            self.assertGreaterEqual(stats["total_replacements"], 1)

            red_runs = [run for run in generated_paragraph.runs if run.text == "212.600.000"]
            self.assertEqual(len(red_runs), 1)
            self.assertEqual(red_runs[0].font.color.rgb, RGBColor(0xFF, 0x00, 0x00))
            self.assertTrue(red_runs[0].bold)
            self.assertTrue(generated_paragraph.runs[0].bold)

    def test_notarial_dashes_are_preserved_and_technical_token_becomes_tab(self):
        with TemporaryDirectory() as tmp_dir:
            source = Path(tmp_dir) / "source.docx"
            target = Path(tmp_dir) / "target.docx"

            doc = Document()
            paragraph = doc.add_paragraph()
            paragraph.add_run("CLASES DE ACTOS: - - - - - - ")
            paragraph.add_run("{{acto}}")
            paragraph.add_run(" - - - - ")
            paragraph.add_run("[[--]]")
            paragraph.add_run(" FINAL")
            doc.save(source)

            stats = apply_marked_template_replacements(
                source,
                target,
                {"acto": "COMPRAVENTA"},
                [{"key": "acto", "label": "Acto", "section": "actos", "raw_markers": ["{{acto}}"]}],
            )

            generated_paragraph = Document(str(target)).paragraphs[0]
            text = generated_paragraph.text

            self.assertIn("CLASES DE ACTOS: - - - - - - COMPRAVENTA - - - - \t FINAL", text)
            self.assertNotIn("[[--]]", text)
            self.assertEqual(stats["technical_tabs_resolved"], 1)
            self.assertGreaterEqual(stats["notarial_dash_sequences_preserved"], 2)

            red_runs = [run for run in generated_paragraph.runs if run.text == "COMPRAVENTA"]
            self.assertEqual(len(red_runs), 1)
            self.assertEqual(red_runs[0].font.color.rgb, RGBColor(0xFF, 0x00, 0x00))
            self.assertIn("<w:tab", generated_paragraph._p.xml)
            self.assertEqual(generated_paragraph.runs[0].text, "CLASES DE ACTOS: - - - - - - ")
            self.assertEqual(generated_paragraph.runs[-1].text, " FINAL")

    def test_engine_generates_with_two_complete_buyers(self):
        with TemporaryDirectory() as tmp_dir:
            source = Path(tmp_dir) / "source.docx"
            target = Path(tmp_dir) / "target.docx"
            doc = Document()
            doc.add_paragraph("Compradores: {{NOMBRE_COMPRADOR_1}} y {{NOMBRE_COMPRADOR_2}}")
            doc.save(source)

            result = NotarialDocumentEngine().render_marked_template(
                source,
                target,
                {
                    "nombre_comprador_1": "JUAN CAMILO",
                    "nombre_comprador_2": "MARIA PEREZ",
                    "numero_documento_comprador_2": "123",
                },
                [
                    {"key": "nombre_comprador_1", "label": "Comprador 1", "raw_markers": ["{{NOMBRE_COMPRADOR_1}}"]},
                    {"key": "nombre_comprador_2", "label": "Comprador 2", "raw_markers": ["{{NOMBRE_COMPRADOR_2}}"]},
                ],
            )

            self.assertEqual(result.statistics["blockers_count"], 0)
            self.assertIn("JUAN CAMILO y MARIA PEREZ", Document(str(target)).paragraphs[0].text)

    def test_engine_warns_two_buyer_template_with_one_buyer_data(self):
        with TemporaryDirectory() as tmp_dir:
            source = Path(tmp_dir) / "source.docx"
            target = Path(tmp_dir) / "target.docx"
            doc = Document()
            doc.add_paragraph("Compradores: {{NOMBRE_COMPRADOR_1}} y {{NOMBRE_COMPRADOR_2}}")
            doc.save(source)

            result = NotarialDocumentEngine().render_marked_template(
                source,
                target,
                {"nombre_comprador_1": "JUAN CAMILO"},
                [
                    {"key": "nombre_comprador_1", "label": "Comprador 1", "raw_markers": ["{{NOMBRE_COMPRADOR_1}}"]},
                    {"key": "nombre_comprador_2", "label": "Comprador 2", "raw_markers": ["{{NOMBRE_COMPRADOR_2}}"]},
                ],
            )

            self.assertEqual(result.statistics["blockers_count"], 0)
            self.assertTrue(any(issue.code in {"template_requires_missing_actor", "empty_replacement_value"} for issue in result.warnings))

    def test_engine_formats_money_letters_date_residuals_and_conditional_block(self):
        with TemporaryDirectory() as tmp_dir:
            source = Path(tmp_dir) / "source.docx"
            target = Path(tmp_dir) / "target.docx"
            doc = Document()
            doc.add_paragraph("Precio {{VALOR_VENTA}} letras {{VALOR_LETRAS}} fecha {{FECHA_OTORGAMIENTO}} [[--]] [Número]")
            doc.add_paragraph("[SI APLICA] Afectación a vivienda familiar {{QUEDA_AFECTADO}}")
            doc.save(source)

            result = NotarialDocumentEngine().render_marked_template(
                source,
                target,
                {
                    "valor_de_la_venta_en_numeros": "212600000",
                    "fecha_otorgamiento": "2025-06-12",
                    "queda_afectado": "NO",
                },
                [
                    {"key": "valor_de_la_venta_en_numeros", "label": "Valor de la venta", "section": "values", "raw_markers": ["{{VALOR_VENTA}}"]},
                    {"key": "valor_apartamento_en_letras", "label": "Valor apartamento en letras", "section": "values", "raw_markers": ["{{VALOR_LETRAS}}"]},
                    {"key": "fecha_otorgamiento", "label": "Fecha otorgamiento", "section": "protocol", "raw_markers": ["{{FECHA_OTORGAMIENTO}}"]},
                    {"key": "queda_afectado", "label": "Queda afectado", "section": "decisions", "raw_markers": ["{{QUEDA_AFECTADO}}"]},
                ],
            )

            text = "\n".join(paragraph.text for paragraph in Document(str(target)).paragraphs)
            self.assertIn("212.600.000", text)
            self.assertIn("DOSCIENTOS DOCE MILLONES SEISCIENTOS MIL PESOS MONEDA CORRIENTE", text)
            self.assertIn("doce (12) días del mes de junio del año DOS MIL VEINTICINCO (2025)".lower(), text.lower())
            self.assertNotIn("[[--]]", text)
            self.assertIn("\t", text)
            self.assertNotIn("[Número]", text)
            self.assertNotIn("[SI APLICA]", text)
            self.assertGreaterEqual(result.statistics["total_replacements"], 4)
            self.assertEqual(result.statistics["technical_tabs_resolved"], 1)
            self.assertEqual(result.audit.residual_tokens, [])

    def test_optional_origin_empty_generates_warning_and_cleans_payment_clause(self):
        with TemporaryDirectory() as tmp_dir:
            source = Path(tmp_dir) / "source.docx"
            target = Path(tmp_dir) / "target.docx"
            doc = Document()
            doc.add_paragraph("A) CUOTA INICIAL: proveniente de {{ORIGEN_CUOTA_INICIAL}}.")
            doc.save(source)

            result = NotarialDocumentEngine().render_marked_template(
                source,
                target,
                {"origen_cuota_inicial": ""},
                [{"key": "origen_cuota_inicial", "label": "Origen cuota inicial", "raw_markers": ["{{ORIGEN_CUOTA_INICIAL}}"]}],
            )

            text = Document(str(target)).paragraphs[0].text
            self.assertEqual(result.statistics["blockers_count"], 0)
            self.assertEqual(result.statistics["optional_segments_omitted"], 1)
            self.assertTrue(any(issue.code == "optional_field_omitted" and issue.field_key == "origen_cuota_inicial" for issue in result.warnings))
            self.assertNotIn("proveniente de .", text)
            self.assertNotIn("recursos propios", text)

    def test_required_origin_fills_payment_clause_when_present(self):
        with TemporaryDirectory() as tmp_dir:
            source = Path(tmp_dir) / "source.docx"
            target = Path(tmp_dir) / "target.docx"
            doc = Document()
            doc.add_paragraph("A) CUOTA INICIAL: proveniente de {{ORIGEN_CUOTA_INICIAL}}.")
            doc.save(source)

            result = NotarialDocumentEngine().render_marked_template(
                source,
                target,
                {"origen_cuota_inicial": "recursos propios"},
                [{"key": "origen_cuota_inicial", "label": "Origen cuota inicial", "raw_markers": ["{{ORIGEN_CUOTA_INICIAL}}"]}],
            )

            text = Document(str(target)).paragraphs[0].text
            self.assertIn("proveniente de recursos propios.", text)
            self.assertNotIn("proveniente de .", text)
            self.assertFalse(any(issue.field_key == "origen_cuota_inicial" for issue in result.warnings))

    def test_optional_saldo_detail_empty_generates_warning_and_cleans_phrase(self):
        with TemporaryDirectory() as tmp_dir:
            source = Path(tmp_dir) / "source.docx"
            target = Path(tmp_dir) / "target.docx"
            doc = Document()
            doc.add_paragraph("El saldo será pagado con {{FUENTE_SALDO}}.")
            doc.save(source)

            result = NotarialDocumentEngine().render_marked_template(
                source,
                target,
                {"fuente_saldo": ""},
                [{"key": "fuente_saldo", "label": "Fuente saldo", "raw_markers": ["{{FUENTE_SALDO}}"]}],
            )

            text = Document(str(target)).paragraphs[0].text
            self.assertEqual(result.statistics["blockers_count"], 0)
            self.assertTrue(any(issue.code == "optional_field_omitted" for issue in result.warnings))
            self.assertNotIn("con .", text)

    def test_critical_field_empty_generates_warning_without_blocking_generation(self):
        with TemporaryDirectory() as tmp_dir:
            source = Path(tmp_dir) / "source.docx"
            target = Path(tmp_dir) / "target.docx"
            doc = Document()
            doc.add_paragraph("Matrícula {{NUMERO_MATRICULA}}.")
            doc.save(source)

            result = NotarialDocumentEngine().render_marked_template(
                source,
                target,
                {"numero_matricula": ""},
                [{"key": "numero_matricula", "label": "Matrícula", "raw_markers": ["{{NUMERO_MATRICULA}}"]}],
            )

            self.assertEqual(result.statistics["blockers_count"], 0)
            self.assertTrue(any(issue.code == "required_field_missing" and issue.field_key == "numero_matricula" for issue in result.warnings))

    def test_person_concordance_literals_are_resolved_for_mixed_buyers(self):
        with TemporaryDirectory() as tmp_dir:
            source = Path(tmp_dir) / "source.docx"
            target = Path(tmp_dir) / "target.docx"
            doc = Document()
            doc.add_paragraph(
                "{{NOMBRE_COMPRADOR_1}}, {{COMPRADOR_ES_HOMBRE_O_MUJER}}, mayor de edad, "
                "de nacionalidad {{NACIONALIDAD_COMPRADOR}}, domiciliado(a) y residente en {{MUNICIPIO_DOMICILIO_COMPRADOR}}, "
                "{{SELECCIONE_SI_COMPRADOR_ESTA_DE_TRANSITO}} identificado(a/s) con {{TIPO_DOCUMENTO_COMPRADOR_1}} {{NUMERO_DOCUMENTO_COMPRADOR_1}}, "
                "de estado civil {{ESTADO_CIVIL_COMPRADOR}} y {{NOMBRE_COMPRADOR_2}}, {{COMPRADOR_2_ES_HOMBRE_O_MUJER}}, mayor de edad, "
                "de nacionalidad {{NACIONALIDAD_COMPRADOR_2}}, domiciliado(a) y residente en {{MUNICIPIO_DOMICILIO_COMPRADOR_2}}, "
                "{{SELECCIONE_SI_COMPRADOR_2_ESTA_DE_TRANSITO}} identificado(a/s) con {{TIPO_DOCUMENTO_COMPRADOR_2}} {{NUMERO_DOCUMENTO_COMPRADOR_2}}, "
                "de estado civil {{ESTADO_CIVIL_COMPRADOR_2}}."
            )
            doc.save(source)

            fields = [
                {"key": key.lower(), "label": key, "raw_markers": [f"{{{{{key}}}}}"]}
                for key in (
                    "NOMBRE_COMPRADOR_1",
                    "COMPRADOR_ES_HOMBRE_O_MUJER",
                    "NACIONALIDAD_COMPRADOR",
                    "MUNICIPIO_DOMICILIO_COMPRADOR",
                    "SELECCIONE_SI_COMPRADOR_ESTA_DE_TRANSITO",
                    "TIPO_DOCUMENTO_COMPRADOR_1",
                    "NUMERO_DOCUMENTO_COMPRADOR_1",
                    "ESTADO_CIVIL_COMPRADOR",
                    "NOMBRE_COMPRADOR_2",
                    "COMPRADOR_2_ES_HOMBRE_O_MUJER",
                    "NACIONALIDAD_COMPRADOR_2",
                    "MUNICIPIO_DOMICILIO_COMPRADOR_2",
                    "SELECCIONE_SI_COMPRADOR_2_ESTA_DE_TRANSITO",
                    "TIPO_DOCUMENTO_COMPRADOR_2",
                    "NUMERO_DOCUMENTO_COMPRADOR_2",
                    "ESTADO_CIVIL_COMPRADOR_2",
                )
            ]
            values = {
                "nombre_comprador_1": "DANIELA CAMPO",
                "comprador_es_hombre_o_mujer": "MUJER",
                "nacionalidad_comprador": "Colombiana",
                "municipio_domicilio_comprador": "Bogotá D.C.",
                "seleccione_si_comprador_esta_de_transito": "No está de transito",
                "tipo_documento_comprador_1": "C.C.",
                "numero_documento_comprador_1": "10203040",
                "estado_civil_comprador": "Casada con sociedad conyugal disuelta y liquidada",
                "nombre_comprador_2": "GERÓNIMO GIRALDO",
                "comprador_2_es_hombre_o_mujer": "HOMBRE",
                "nacionalidad_comprador_2": "Colombiano",
                "municipio_domicilio_comprador_2": "Bogotá D.C.",
                "seleccione_si_comprador_2_esta_de_transito": "No está de transito",
                "tipo_documento_comprador_2": "C.C.",
                "numero_documento_comprador_2": "20406080",
                "estado_civil_comprador_2": "Casado sin sociedad conyugal por efecto de capitulaciones",
            }

            NotarialDocumentEngine().render_marked_template(source, target, values, fields)
            text = Document(str(target)).paragraphs[0].text

            self.assertIn("DANIELA CAMPO, mujer", text)
            self.assertIn("colombiana, domiciliada", text)
            self.assertIn("identificada con cédula de ciudadanía número 10203040", text)
            self.assertIn("GERÓNIMO GIRALDO, varón", text)
            self.assertIn("colombiano, domiciliado", text)
            self.assertIn("identificado con cédula de ciudadanía número 20406080", text)
            self.assertNotIn("identificado(a/s)", text)
            self.assertNotIn("domiciliado(a)", text)
            self.assertNotIn("transito", text)

    def test_empty_signature_block_is_removed_but_valid_signature_is_preserved(self):
        with TemporaryDirectory() as tmp_dir:
            source = Path(tmp_dir) / "source.docx"
            target = Path(tmp_dir) / "target.docx"
            doc = Document()
            doc.add_paragraph("____________________________")
            doc.add_paragraph("C.C.")
            doc.add_paragraph("Celular.")
            doc.add_paragraph("Dirección:")
            doc.add_paragraph("Correo:")
            doc.add_paragraph("Profesión u Oficio:")
            doc.add_paragraph("Actividad Económica:")
            doc.add_paragraph("Estado Civil:")
            doc.add_paragraph("____________________________")
            doc.add_paragraph("DANIELA CAMPO")
            doc.add_paragraph("C.C. 10203040")
            doc.save(source)

            result = NotarialDocumentEngine().render_marked_template(source, target, {}, [])
            text = "\n".join(paragraph.text for paragraph in Document(str(target)).paragraphs)

            self.assertEqual(result.statistics["empty_signature_blocks_removed"], 1)
            self.assertNotIn("C.C.\nCelular.\nDirección:", text)
            self.assertIn("DANIELA CAMPO", text)
            self.assertIn("C.C. 10203040", text)

    def test_engine_reports_replaced_empty_missing_and_unresolved_placeholders(self):
        with TemporaryDirectory() as tmp_dir:
            source = Path(tmp_dir) / "source.docx"
            target = Path(tmp_dir) / "target.docx"
            doc = Document()
            doc.add_paragraph("A {{A}} B {{B}} vivo {{C}}")
            doc.save(source)

            result = NotarialDocumentEngine().render_marked_template(
                source,
                target,
                {"a": "uno", "b": ""},
                [
                    {"key": "a", "label": "A", "raw_markers": ["{{A}}"]},
                    {"key": "b", "label": "B", "raw_markers": ["{{B}}"]},
                    {"key": "d", "label": "D", "raw_markers": ["{{D}}"]},
                ],
            )

            self.assertEqual(result.statistics["blockers_count"], 0)
            self.assertTrue(any(issue.code == "unresolved_placeholder" for issue in result.warnings))

    def test_composed_phrase_normalizer_cleans_common_notarial_duplicates(self):
        with TemporaryDirectory() as tmp_dir:
            source = Path(tmp_dir) / "source.docx"
            target = Path(tmp_dir) / "target.docx"
            doc = Document()
            doc.add_paragraph("número número de escritura por cantidad de de pesos y P.H. P.H. varón(fideicomisario)")
            doc.save(source)

            result = NotarialDocumentEngine().render_marked_template(source, target, {}, [])
            text = Document(str(target)).paragraphs[0].text

            self.assertEqual(result.statistics["blockers_count"], 0)
            self.assertIn("número de escritura", text)
            self.assertIn("cantidad de pesos", text)
            self.assertIn("P.H.", text)
            self.assertNotIn("P.H. P.H.", text)
            self.assertNotIn("varón(fideicomisario)", text)


if __name__ == "__main__":
    unittest.main()
