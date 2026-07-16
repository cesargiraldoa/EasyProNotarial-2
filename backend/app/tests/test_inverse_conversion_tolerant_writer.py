from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document

from app.services.minuta.inverse_conversion_writer.models import MarkedCandidate
from app.services.minuta.inverse_conversion_writer.service import InverseConversionMarkedDocxService


class TolerantInverseWriterTests(unittest.TestCase):
    def test_unmapped_candidate_does_not_block_valid_marked_docx(self):
        with tempfile.TemporaryDirectory() as tmp:
            source = Path(tmp) / "source.docx"
            output = Path(tmp) / "marked.docx"
            document = Document()
            document.add_paragraph("APARTAMENTO 804 - DATO NUEVO")
            document.save(source)

            candidates = [
                MarkedCandidate(
                    id="cand_valid",
                    text="804",
                    suggested_key="NUMERO_APARTAMENTO",
                    canonical_field_code="NUMERO_APARTAMENTO",
                    status="accepted",
                    occurrences=1,
                ),
                MarkedCandidate(
                    id="cand_unmapped",
                    text="DATO NUEVO",
                    suggested_key="CAMPO_QUE_NO_EXISTE",
                    status="accepted",
                    occurrences=1,
                ),
            ]

            result = InverseConversionMarkedDocxService().generate(
                source,
                output,
                source.name,
                candidates,
                allowed_field_codes={"NUMERO_APARTAMENTO"},
            )

            marked = Document(output)
            visible = "\n".join(paragraph.text for paragraph in marked.paragraphs)
            self.assertIn("{{NUMERO_APARTAMENTO}}", visible)
            self.assertIn("DATO NUEVO", visible)
            self.assertEqual(result.marked_occurrences, 1)


if __name__ == "__main__":
    unittest.main()
