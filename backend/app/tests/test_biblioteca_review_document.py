from __future__ import annotations

import hashlib
import io
import zipfile
import unittest

from docx import Document
from lxml import etree

from app.services.biblioteca_motor.review_document import (
    FIELD_TAG_PREFIX,
    SUGGESTION_TAG_PREFIX,
    accept_suggestions_in_docx,
    apply_review_decision_in_docx,
    cascade_field_controls_in_docx,
    prepare_review_document,
    reject_suggestions_in_docx,
    visible_text_from_docx,
)


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}
W = f"{{{W_NS}}}"


def _docx_with_split_runs() -> bytes:
    doc = Document()
    paragraph = doc.add_paragraph()
    paragraph.add_run("LOS COMPRADORES: ")
    bold = paragraph.add_run("Daniela")
    bold.bold = True
    italic = paragraph.add_run(" Campo")
    italic.italic = True
    paragraph.add_run(" firma.")
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _docx_with_table() -> bytes:
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "MATRICULA INMOBILIARIA 050-123456"
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _docx_with_paragraph(text: str) -> bytes:
    doc = Document()
    doc.add_paragraph(text)
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _document_xml(docx_bytes: bytes) -> etree._Element:
    with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as archive:
        return etree.fromstring(archive.read("word/document.xml"))


def _tags(docx_bytes: bytes) -> list[str]:
    root = _document_xml(docx_bytes)
    return [item.get(f"{W}val") for item in root.findall(".//w:sdtPr/w:tag", NS)]


def _suggestion(
    original: str,
    start: int,
    end: int,
    text: str,
    *,
    suggestion_id: str = "sug_1",
    candidate_id: str = "cand_1",
    **location_overrides,
):
    location = {
        "block_type": "paragraph",
        "block_index": 1,
        "paragraph_index": 1,
        "table_index": None,
        "row_index": None,
        "cell_index": None,
        "char_start": start,
        "char_end": end,
        "occurrence_index": 1,
        "location_key": f"paragraph:1:{start}:{end}:1",
        "block_hash": hashlib.sha256(text.encode("utf-8")).hexdigest(),
    }
    location.update(location_overrides)
    return {
        "suggestion_id": suggestion_id,
        "candidate_id": candidate_id,
        "original_text": original,
        "field_code": "COMPRADOR_1",
        "visible_code": "COMPRADOR_1",
        "field_instance_id": "fi_comprador_1",
        "field_label": "Comprador 1",
        "category": "persona",
        "confidence": 0.97,
        "source": "deterministic",
        "location": location,
        "context_before": "",
        "context_after": "",
    }


class BibliotecaReviewDocumentTests(unittest.TestCase):
    def test_wraps_exact_paragraph_span_in_content_control_preserving_visible_text(self):
        docx = _docx_with_split_runs()
        text = "LOS COMPRADORES: Daniela Campo firma."
        start = text.index("Daniela Campo")
        end = start + len("Daniela Campo")
        result = prepare_review_document(docx, [_suggestion("Daniela Campo", start, end, text)], analysis_id="analysis_1")
        root = _document_xml(result.docx_bytes)

        self.assertEqual(visible_text_from_docx(result.docx_bytes), visible_text_from_docx(docx))
        self.assertEqual(result.wrapped_count, 1)
        self.assertEqual(len(root.findall(".//w:sdt", NS)), 1)
        self.assertTrue(_tags(result.docx_bytes)[0].startswith(SUGGESTION_TAG_PREFIX))
        self.assertEqual(len(root.findall(".//w:highlight[@w:val='yellow']", NS)), 2)
        self.assertTrue(root.findall(".//w:sdtContent//w:b", NS))
        self.assertTrue(root.findall(".//w:sdtContent//w:i", NS))

    def test_wraps_multiple_non_overlapping_controls_in_same_paragraph(self):
        text = "LOS COMPRADORES: Daniela Campo y Carlos Ruiz firman."
        docx = _docx_with_paragraph(text)
        first_start = text.index("Daniela Campo")
        second_start = text.index("Carlos Ruiz")
        first = _suggestion("Daniela Campo", first_start, first_start + len("Daniela Campo"), text)
        second = _suggestion(
            "Carlos Ruiz",
            second_start,
            second_start + len("Carlos Ruiz"),
            text,
            suggestion_id="sug_2",
            candidate_id="cand_2",
        )
        second["field_code"] = "COMPRADOR_2"
        second["visible_code"] = "COMPRADOR_2"
        second["field_instance_id"] = "fi_comprador_2"

        result = prepare_review_document(docx, [first, second], analysis_id="analysis_1")
        root = _document_xml(result.docx_bytes)

        self.assertEqual(result.wrapped_count, 2)
        self.assertEqual(visible_text_from_docx(result.docx_bytes), visible_text_from_docx(docx))
        self.assertEqual(len(root.findall(".//w:sdt", NS)), 2)

    def test_wraps_table_cell_span(self):
        docx = _docx_with_table()
        text = "MATRICULA INMOBILIARIA 050-123456"
        start = text.index("050-123456")
        suggestion = _suggestion(
            "050-123456",
            start,
            start + len("050-123456"),
            text,
            block_type="table_cell",
            paragraph_index=None,
            table_index=1,
            row_index=1,
            cell_index=1,
            location_key=f"table_cell:1:1:1:{start}:{start + len('050-123456')}:1",
        )
        suggestion["field_code"] = "MATRICULA_INMOBILIARIA"

        result = prepare_review_document(docx, [suggestion], analysis_id="analysis_1")

        self.assertEqual(result.wrapped_count, 1)
        self.assertIn("MATRICULA INMOBILIARIA 050-123456", visible_text_from_docx(result.docx_bytes))
        self.assertTrue(_tags(result.docx_bytes)[0].startswith(SUGGESTION_TAG_PREFIX))

    def test_accept_transforms_suggestion_to_field_control_and_replaces_visible_value(self):
        docx = _docx_with_split_runs()
        text = "LOS COMPRADORES: Daniela Campo firma."
        start = text.index("Daniela Campo")
        result = prepare_review_document(docx, [_suggestion("Daniela Campo", start, start + 13, text)], analysis_id="analysis_1")
        tag = _tags(result.docx_bytes)[0]

        accepted = accept_suggestions_in_docx(
            result.docx_bytes,
            [{"suggestion_tag": tag, "field_instance_id": "COMPRADOR_1", "occurrence_id": "occ_001"}],
        )

        self.assertIn("{{COMPRADOR_1}}", visible_text_from_docx(accepted))
        self.assertFalse([item for item in _tags(accepted) if item.startswith(SUGGESTION_TAG_PREFIX)])
        self.assertTrue([item for item in _tags(accepted) if item.startswith(FIELD_TAG_PREFIX)])
        self.assertFalse(_document_xml(accepted).findall(".//w:highlight", NS))

    def test_reject_unwraps_suggestion_and_keeps_original_text(self):
        docx = _docx_with_split_runs()
        text = "LOS COMPRADORES: Daniela Campo firma."
        start = text.index("Daniela Campo")
        result = prepare_review_document(docx, [_suggestion("Daniela Campo", start, start + 13, text)], analysis_id="analysis_1")
        rejected = reject_suggestions_in_docx(result.docx_bytes, [_tags(result.docx_bytes)[0]])

        self.assertEqual(visible_text_from_docx(rejected), visible_text_from_docx(docx))
        self.assertFalse(_document_xml(rejected).findall(".//w:sdt", NS))

    def test_apply_review_decision_accepts_exact_suggestion_and_audits(self):
        docx = _docx_with_split_runs()
        text = "LOS COMPRADORES: Daniela Campo firma."
        start = text.index("Daniela Campo")
        result = prepare_review_document(docx, [_suggestion("Daniela Campo", start, start + 13, text)], analysis_id="analysis_1")
        occurrence = result.groups[0]["occurrences"][0]

        applied = apply_review_decision_in_docx(
            result.docx_bytes,
            {
                "action": "accept",
                "suggestion_tag": occurrence["tag"],
                "occurrence_id": occurrence["occurrence_id"],
                "field_instance_id": "fi_comprador_1",
                "field_code": "COMPRADOR_1",
                "visible_code": "COMPRADOR_1",
                "original_text": occurrence["original_text"],
                "location": occurrence["location"],
            },
        )

        self.assertIn("{{COMPRADOR_1}}", visible_text_from_docx(applied.docx_bytes))
        self.assertFalse([item for item in _tags(applied.docx_bytes) if item.startswith(SUGGESTION_TAG_PREFIX)])
        self.assertTrue(any(item.startswith(f"{FIELD_TAG_PREFIX}fi_comprador_1:{occurrence['occurrence_id']}:COMPRADOR_1:COMPRADOR_1") for item in _tags(applied.docx_bytes)))
        self.assertEqual(applied.audit["status"], "accepted")
        self.assertTrue(applied.audit["validated_text"])
        self.assertTrue(applied.audit["validated_location"])

    def test_apply_review_decision_rejects_and_keeps_original_text(self):
        docx = _docx_with_split_runs()
        text = "LOS COMPRADORES: Daniela Campo firma."
        start = text.index("Daniela Campo")
        result = prepare_review_document(docx, [_suggestion("Daniela Campo", start, start + 13, text)], analysis_id="analysis_1")
        occurrence = result.groups[0]["occurrences"][0]

        applied = apply_review_decision_in_docx(
            result.docx_bytes,
            {
                "action": "reject",
                "suggestion_tag": occurrence["tag"],
                "occurrence_id": occurrence["occurrence_id"],
                "original_text": occurrence["original_text"],
                "location": occurrence["location"],
            },
        )

        self.assertEqual(visible_text_from_docx(applied.docx_bytes), visible_text_from_docx(docx))
        self.assertFalse(_document_xml(applied.docx_bytes).findall(".//w:sdt", NS))
        self.assertEqual(applied.audit["status"], "rejected")

    def test_apply_review_decision_change_assigns_different_field(self):
        docx = _docx_with_split_runs()
        text = "LOS COMPRADORES: Daniela Campo firma."
        start = text.index("Daniela Campo")
        result = prepare_review_document(docx, [_suggestion("Daniela Campo", start, start + 13, text)], analysis_id="analysis_1")
        occurrence = result.groups[0]["occurrences"][0]

        applied = apply_review_decision_in_docx(
            result.docx_bytes,
            {
                "action": "change",
                "suggestion_tag": occurrence["tag"],
                "occurrence_id": occurrence["occurrence_id"],
                "field_instance_id": "fi_vendedor_1",
                "field_code": "VENDEDOR_1",
                "visible_code": "VENDEDOR_1",
                "original_text": occurrence["original_text"],
                "location": occurrence["location"],
            },
        )

        self.assertIn("{{VENDEDOR_1}}", visible_text_from_docx(applied.docx_bytes))
        self.assertEqual(applied.audit["status"], "changed")
        self.assertTrue(any(item.startswith(f"{FIELD_TAG_PREFIX}fi_vendedor_1:{occurrence['occurrence_id']}:VENDEDOR_1:VENDEDOR_1") for item in _tags(applied.docx_bytes)))

    def test_apply_review_decision_does_not_accept_provisional_field(self):
        docx = _docx_with_split_runs()
        text = "LOS COMPRADORES: Daniela Campo firma."
        start = text.index("Daniela Campo")
        result = prepare_review_document(docx, [_suggestion("Daniela Campo", start, start + 13, text)], analysis_id="analysis_1")
        occurrence = result.groups[0]["occurrences"][0]

        with self.assertRaisesRegex(ValueError, "provisional_field_requires_assignment"):
            apply_review_decision_in_docx(
                result.docx_bytes,
                {
                    "action": "accept",
                    "suggestion_tag": occurrence["tag"],
                    "occurrence_id": occurrence["occurrence_id"],
                    "field_instance_id": "prov_1",
                    "field_code": "PENDING_FIELD_PERSON_NAME_DEMO",
                    "visible_code": "PENDING_FIELD_PERSON_NAME_DEMO",
                    "original_text": occurrence["original_text"],
                    "location": occurrence["location"],
                },
            )

    def test_apply_review_decision_allows_sequential_same_paragraph_when_control_text_matches(self):
        docx = _docx_with_split_runs()
        text = "LOS COMPRADORES: Daniela Campo firma."
        start = text.index("Daniela Campo")
        result = prepare_review_document(docx, [_suggestion("Daniela Campo", start, start + 13, text)], analysis_id="analysis_1")
        occurrence = result.groups[0]["occurrences"][0]
        stale_location = {**occurrence["location"], "block_hash": "stale"}

        applied = apply_review_decision_in_docx(
            result.docx_bytes,
            {
                "action": "accept",
                "suggestion_tag": occurrence["tag"],
                "occurrence_id": occurrence["occurrence_id"],
                "field_instance_id": "fi_comprador_1",
                "field_code": "COMPRADOR_1",
                "visible_code": "COMPRADOR_1",
                "original_text": occurrence["original_text"],
                "location": stale_location,
            },
        )

        self.assertEqual(applied.audit["location_validation"], "content_control_current:block_hash_mismatch")

    def test_accepts_two_suggestions_in_same_paragraph_across_consecutive_versions(self):
        text = "LOS COMPRADORES: Daniela Campo y Carlos Ruiz firman."
        docx = _docx_with_paragraph(text)
        first_start = text.index("Daniela Campo")
        second_start = text.index("Carlos Ruiz")
        first = _suggestion("Daniela Campo", first_start, first_start + len("Daniela Campo"), text)
        second = _suggestion(
            "Carlos Ruiz",
            second_start,
            second_start + len("Carlos Ruiz"),
            text,
            suggestion_id="sug_2",
            candidate_id="cand_2",
        )
        second["field_code"] = "COMPRADOR_2"
        second["visible_code"] = "COMPRADOR_2"
        second["field_instance_id"] = "fi_comprador_2"
        prepared = prepare_review_document(docx, [first, second], analysis_id="analysis_1")
        first_occurrence = next(item for group in prepared.groups for item in group["occurrences"] if item["original_text"] == "Daniela Campo")
        second_occurrence = next(item for group in prepared.groups for item in group["occurrences"] if item["original_text"] == "Carlos Ruiz")

        first_applied = apply_review_decision_in_docx(
            prepared.docx_bytes,
            {
                "action": "accept",
                "suggestion_tag": first_occurrence["tag"],
                "occurrence_id": first_occurrence["occurrence_id"],
                "field_instance_id": "fi_comprador_1",
                "field_code": "COMPRADOR_1",
                "visible_code": "COMPRADOR_1",
                "original_text": first_occurrence["original_text"],
                "location": first_occurrence["location"],
            },
        )
        second_applied = apply_review_decision_in_docx(
            first_applied.docx_bytes,
            {
                "action": "accept",
                "suggestion_tag": second_occurrence["tag"],
                "occurrence_id": second_occurrence["occurrence_id"],
                "field_instance_id": "fi_comprador_2",
                "field_code": "COMPRADOR_2",
                "visible_code": "COMPRADOR_2",
                "original_text": second_occurrence["original_text"],
                "location": second_occurrence["location"],
            },
        )

        self.assertIn("{{COMPRADOR_1}}", visible_text_from_docx(second_applied.docx_bytes))
        self.assertIn("{{COMPRADOR_2}}", visible_text_from_docx(second_applied.docx_bytes))
        self.assertTrue(second_applied.audit["location_validation"].startswith("content_control_current:"))

    def test_apply_review_decision_detects_control_text_changed_after_analysis(self):
        docx = _docx_with_split_runs()
        text = "LOS COMPRADORES: Daniela Campo firma."
        start = text.index("Daniela Campo")
        result = prepare_review_document(docx, [_suggestion("Daniela Campo", start, start + 13, text)], analysis_id="analysis_1")
        occurrence = result.groups[0]["occurrences"][0]

        with self.assertRaisesRegex(ValueError, "suggestion_text_mismatch"):
            apply_review_decision_in_docx(
                result.docx_bytes,
                {
                    "action": "accept",
                    "suggestion_tag": occurrence["tag"],
                    "occurrence_id": occurrence["occurrence_id"],
                    "field_instance_id": "fi_comprador_1",
                    "field_code": "COMPRADOR_1",
                    "visible_code": "COMPRADOR_1",
                    "original_text": "Daniela Campos",
                    "location": occurrence["location"],
                },
            )

    def test_backend_cascade_updates_only_matching_field_instance(self):
        docx = _docx_with_split_runs()
        text = "LOS COMPRADORES: Daniela Campo firma."
        start = text.index("Daniela Campo")
        result = prepare_review_document(docx, [_suggestion("Daniela Campo", start, start + 13, text)], analysis_id="analysis_1")
        field_docx = accept_suggestions_in_docx(
            result.docx_bytes,
            [{"suggestion_tag": _tags(result.docx_bytes)[0], "field_instance_id": "COMPRADOR_1", "occurrence_id": "occ_001"}],
        )

        cascaded, updated = cascade_field_controls_in_docx(field_docx, "COMPRADOR_1", "Maria Perez")

        self.assertEqual(updated, 1)
        self.assertIn("Maria Perez", visible_text_from_docx(cascaded))


if __name__ == "__main__":
    unittest.main()
