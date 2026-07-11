from __future__ import annotations

import json
import os
import tempfile
import time
from io import BytesIO
from pathlib import Path

from docx import Document
from sqlalchemy import text

from app.db.session import SessionLocal
from app.models.notarial_document_intelligence import (
    NotarialDocument,
    NotarialDocumentBatch,
    NotarialDocumentBatchItem,
    NotarialDocumentBlock,
    NotarialDocumentDecision,
    NotarialDocumentEmbedding,
    NotarialDocumentEvidence,
    NotarialIntelligenceRun,
    NotarialDocumentParseRun,
    NotarialEmbeddingVersion,
)
from app.models.notary import Notary
from app.services.notarial_document_intelligence.contracts import DocumentUpload, IngestBatchRequest
from app.services.notarial_document_intelligence.engine import NotarialIntelligenceEngine
from app.services.notarial_document_intelligence.ingestion import NotarialDocumentIngestionService, TASK_PROCESS_INTELLIGENCE_RUN
from app.services.notarial_document_intelligence.storage import NotarialIntelligenceStorage
from app.services.notarial_document_intelligence.vector_store import NotarialVectorStore
from app.workers.document_worker import (
    process_queued_document_batch_task,
    reparse_document_task,
    transient_retry_probe_task,
)
from app.workers.intelligence_worker import (
    process_intelligence_run_task,
)


def build_docx_bytes(apartment: str) -> bytes:
    buffer = BytesIO()
    document = Document()
    document.add_paragraph("ESCRITURA DE COMPRAVENTA INMUEBLE")
    document.add_paragraph(f"APARTAMENTO: {apartment}")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "MATRICULA"
    table.cell(0, 1).text = f"001-REAL-{apartment}"
    document.save(buffer)
    return buffer.getvalue()


def wait_for_batch(batch_id: int, expected_status: str = "completed", timeout_seconds: int = 90) -> None:
    deadline = time.time() + timeout_seconds
    while time.time() < deadline:
        db = SessionLocal()
        try:
            batch = db.get(NotarialDocumentBatch, batch_id)
            if batch is not None and batch.status == expected_status:
                return
            if batch is not None and batch.status in {"error", "partial_error"}:
                raise AssertionError(f"Batch {batch_id} finished with {batch.status}: {batch.error_message}")
        finally:
            db.close()
        time.sleep(1)
    raise AssertionError(f"Batch {batch_id} did not reach {expected_status} within {timeout_seconds}s")


def wait_for_parse_run(parse_run_id: int, expected_status: str = "completed", timeout_seconds: int = 90) -> None:
    deadline = time.time() + timeout_seconds
    while time.time() < deadline:
        db = SessionLocal()
        try:
            run = db.get(NotarialDocumentParseRun, parse_run_id)
            if run is not None and run.status == expected_status:
                return
            if run is not None and run.status == "error":
                raise AssertionError(f"Parse run {parse_run_id} failed: {run.error_message}")
        finally:
            db.close()
        time.sleep(1)
    raise AssertionError(f"Parse run {parse_run_id} did not reach {expected_status} within {timeout_seconds}s")


def seed_notaries() -> None:
    db = SessionLocal()
    try:
        for notary_id, slug in [(1, "real-notaria-uno"), (2, "real-notaria-dos")]:
            if db.get(Notary, notary_id) is None:
                db.add(
                    Notary(
                        id=notary_id,
                        slug=slug,
                        catalog_identity_key=f"{slug}-identity",
                        commercial_name=slug,
                        legal_name=slug,
                        municipality="Medellin",
                        notary_label=slug,
                        city="Medellin",
                    )
                )
        db.commit()
    finally:
        db.close()


def validate_pgvector() -> None:
    db = SessionLocal()
    try:
        extension = db.execute(text("select extname from pg_extension where extname = 'vector'")).scalar_one()
        assert extension == "vector"
        index = db.execute(
            text(
                """
                select indexname
                from pg_indexes
                where tablename = 'notarial_document_embeddings'
                  and indexname = 'ix_notarial_embeddings_embedding_hnsw'
                  and indexdef ilike '%hnsw%'
                  and indexdef ilike '%vector_cosine_ops%'
                """
            )
        ).first()
        assert index is not None

        version = NotarialEmbeddingVersion(version_key="real-ci-3d", provider="ci", model_name="ci-mini", dimensions=384)
        db.add(version)
        db.flush()
        store = NotarialVectorStore(db, dimensions=384)
        vector_a = [1.0] + [0.0] * 383
        vector_b = [0.0, 1.0] + [0.0] * 382
        db.add_all(
            [
                NotarialDocumentEmbedding(
                    notary_id=1,
                    embedding_version_id=version.id,
                    source_type="block",
                    source_id=101,
                    content_hash="a" * 64,
                    embedding_dimensions=384,
                    embedding=store.encode_vector(vector_a),
                ),
                NotarialDocumentEmbedding(
                    notary_id=1,
                    embedding_version_id=version.id,
                    source_type="block",
                    source_id=102,
                    content_hash="b" * 64,
                    embedding_dimensions=384,
                    embedding=store.encode_vector(vector_b),
                ),
                NotarialDocumentEmbedding(
                    notary_id=2,
                    embedding_version_id=version.id,
                    source_type="block",
                    source_id=201,
                    content_hash="c" * 64,
                    embedding_dimensions=384,
                    embedding=store.encode_vector(vector_a),
                ),
            ]
        )
        db.commit()

        results = store.search(1, vector_a, embedding_version_id=version.id, limit=5)
        assert [item.source_id for item in results] == [101, 102]
    finally:
        db.close()


def validate_worker_and_reconciliation(storage: NotarialIntelligenceStorage) -> None:
    db = SessionLocal()
    try:
        service = NotarialDocumentIngestionService(db, notary_id=1, storage=storage)
        queued = service.queue_batch(
            IngestBatchRequest(
                name="Real Celery queued",
                source_type="ci",
                documents=[DocumentUpload(filename="real-celery.docx", content=build_docx_bytes("3001"))],
            )
        )
        publication = service.get_publication_for_target("batch", queued.batch_id)
        task_id = service.publish_task_publication(publication.id, process_queued_document_batch_task.delay)
        assert task_id
    finally:
        db.close()
    wait_for_batch(queued.batch_id)

    db = SessionLocal()
    try:
        assert db.query(NotarialDocument).filter(NotarialDocument.notary_id == 1).count() >= 1
        assert db.query(NotarialDocumentBatchItem).filter(NotarialDocumentBatchItem.batch_id == queued.batch_id).count() == 1
        assert db.query(NotarialDocumentBlock).join(NotarialDocumentParseRun).filter(NotarialDocumentParseRun.is_active.is_(True)).count() > 0
    finally:
        db.close()

    db = SessionLocal()
    try:
        service = NotarialDocumentIngestionService(db, notary_id=1, storage=storage)
        failed = service.queue_batch(
            IngestBatchRequest(
                name="Real failed publication",
                source_type="ci",
                documents=[DocumentUpload(filename="real-reconcile.docx", content=build_docx_bytes("3002"))],
            )
        )
        publication = service.get_publication_for_target("batch", failed.batch_id)

        def fail_delay(*args):
            raise ConnectionError("ci broker publication simulation")

        try:
            service.publish_task_publication(publication.id, fail_delay)
        except Exception:
            pass
        db.refresh(publication)
        assert publication.status == "publication_failed"
        results = service.dispatch_recoverable_publications(
            {"notarial_document_intelligence.document.process_queued_batch": process_queued_document_batch_task.delay}
        )
        assert results and results[0]["status"] == "published"
    finally:
        db.close()
    wait_for_batch(failed.batch_id)


def validate_reparse(storage: NotarialIntelligenceStorage) -> None:
    db = SessionLocal()
    try:
        document = db.query(NotarialDocument).filter(NotarialDocument.notary_id == 1).order_by(NotarialDocument.id.asc()).first()
        assert document is not None
        Path(document.storage_path).write_bytes(build_docx_bytes("3999"))
        service = NotarialDocumentIngestionService(db, notary_id=1, storage=storage)
        parse_run = service.queue_reparse_document(document.id)
        parse_run_id = parse_run.id
        publication = service.get_publication_for_target("reparse", parse_run.id)
        task_id = service.publish_task_publication(publication.id, reparse_document_task.delay)
        assert task_id
    finally:
        db.close()
    wait_for_parse_run(parse_run_id)

    db = SessionLocal()
    try:
        run = db.get(NotarialDocumentParseRun, parse_run_id)
        assert run is not None and run.is_active is True
        metadata = json.loads(run.metadata_json or "{}")
        assert "unstructured" in metadata
    finally:
        db.close()


def validate_celery_retry_probe() -> None:
    result = transient_retry_probe_task.delay()
    payload = result.get(timeout=90)
    assert payload["status"] == "completed"
    assert payload["retries"] >= 1


def validate_intelligence_worker() -> None:
    db = SessionLocal()
    try:
        document = db.query(NotarialDocument).filter(NotarialDocument.notary_id == 1).order_by(NotarialDocument.id.asc()).first()
        assert document is not None
        document_id = document.id
        run = NotarialIntelligenceEngine(db).create_document_run(1, document_id)
        service = NotarialDocumentIngestionService(db, notary_id=1)
        publication = service.ensure_task_publication(
            request_key=f"real:intelligence_run:{run.id}",
            target_type="intelligence_run",
            target_id=run.id,
            task_name=TASK_PROCESS_INTELLIGENCE_RUN,
            task_args=[run.id],
            document_id=document_id,
            parse_run_id=run.parse_run_id,
            metadata={"source": "real_service_check"},
        )
        task_id = service.publish_task_publication(publication.id, process_intelligence_run_task.delay)
        assert task_id
        run_id = run.id
    finally:
        db.close()

    result = process_intelligence_run_task.AsyncResult(task_id).get(timeout=180)
    assert result["document_id"] == document_id
    assert result["classification"]["document_type"] in {"escritura_compraventa", "escritura_hipoteca"}
    assert result["embedding"]["indexed"] > 0 or result["embedding"]["skipped"] > 0
    assert result["decision_id"] > 0

    db = SessionLocal()
    try:
        reindex_run = NotarialIntelligenceEngine(db).create_reindex_run(1)
        service = NotarialDocumentIngestionService(db, notary_id=1)
        publication = service.ensure_task_publication(
            request_key=f"real:intelligence_run:{reindex_run.id}",
            target_type="intelligence_run",
            target_id=reindex_run.id,
            task_name=TASK_PROCESS_INTELLIGENCE_RUN,
            task_args=[reindex_run.id],
            metadata={"source": "real_service_reindex_check"},
        )
        reindex_task_id = service.publish_task_publication(publication.id, process_intelligence_run_task.delay)
        assert reindex_task_id
    finally:
        db.close()
    reindex = process_intelligence_run_task.AsyncResult(reindex_task_id).get(timeout=180)
    assert reindex["documents"] >= 1
    assert reindex["version_key"].startswith("sentence-transformers:")

    db = SessionLocal()
    try:
        version = (
            db.query(NotarialEmbeddingVersion)
            .filter(NotarialEmbeddingVersion.provider == "sentence-transformers")
            .order_by(NotarialEmbeddingVersion.id.desc())
            .first()
        )
        assert version is not None
        assert version.dimensions == 384
        assert db.query(NotarialDocumentEmbedding).filter(NotarialDocumentEmbedding.notary_id == 1, NotarialDocumentEmbedding.embedding_version_id == version.id).count() > 0
        assert db.query(NotarialDocumentDecision).filter(NotarialDocumentDecision.notary_id == 1, NotarialDocumentDecision.document_id == document_id).count() > 0
        assert db.get(NotarialIntelligenceRun, run_id).status == "completed"
        evidence_types = {
            row[0]
            for row in db.query(NotarialDocumentEvidence.evidence_type)
            .filter(NotarialDocumentEvidence.notary_id == 1, NotarialDocumentEvidence.document_id == document_id)
            .all()
        }
        assert "document_classification" in evidence_types
        assert "hybrid_rag" in evidence_types
        refreshed = db.get(NotarialDocument, document_id)
        assert refreshed is not None
        assert refreshed.family_id is not None
        assert refreshed.document_type is not None
    finally:
        db.close()


def main() -> None:
    assert os.environ.get("DATABASE_URL", "").startswith("postgresql"), "real check requires PostgreSQL"
    assert os.environ.get("REDIS_URL", "").startswith("redis://"), "real check requires Redis"
    with tempfile.TemporaryDirectory(prefix="easypro2-real-ci-storage-") as tmp_dir:
        storage = NotarialIntelligenceStorage(local_root=Path(tmp_dir) / "storage", use_supabase=False)
        seed_notaries()
        validate_pgvector()
        validate_worker_and_reconciliation(storage)
        validate_reparse(storage)
        validate_celery_retry_probe()
        validate_intelligence_worker()


if __name__ == "__main__":
    main()
