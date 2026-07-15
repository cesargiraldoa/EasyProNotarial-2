from __future__ import annotations

import hashlib
import io
import unittest

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.services.biblioteca_motor.review_document import (
    prepare_review_document,
    visible_text_from_docx,
)
from app.services.biblioteca_motor.review_document_safe import (
    VISIBLE_TEXT_ERROR,
    VISIBLE_TEXT_SKIP_REASON,
    prepare_review_document_safe,
)


def _docx_with_inline_container_between_runs() -> bytes:
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("A")

    hyperlink = OxmlElement("w:hyperlink")
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "X"
    run.append(text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)

    paragraph.add_run("B")
    document.add_paragraph("SAFE 123")

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _suggestion(
    *,
    suggestion_id: str,
    candidate_id: str,
    paragraph_index: int,
    paragraph_text: str,
    original_text: str,
    start: int,
    field_code: str,
) -> dict:
    end = start + len(original_text)
    return {
        "suggestion_id": suggestion_id,
        "candidate_id": candidate_id,
        "occurrence_id": f"occ_{suggestion_id}",
        "original_text": original_text,
        "field_code": field_code,
        "visible_code": field_code,
        "base_field_code": field_code,
        "field_instance_id": f"fi_{field_code.lower()}",
        "field_label": field_code,
        "category": "otro",
        "confidence": 0.95,
        "source": "llm",
        "location": {
            "block_type": "paragraph",
            "block_index": paragraph_index,
            "paragraph_index": paragraph_index,
            "table_index": None,
            "row_index": None,
            "cell_index": None,
            "char_start": start,
            "char_end": end,
            "occurrence_index": 1,
            "location_key": f"paragraph:{paragraph_index}:{start}:{end}:1",
            "block_hash": hashlib.sha256(paragraph_text.encode("utf-8")).hexdigest(),
        },
        "context_before": "",
        "context_after": "",
    }


class BibliotecaReviewVisibleGuardTests(unittest.TestCase):
    def test_original_mutator_detects_inline_container_reordering(self):
        docx = _docx_with_inline_container_between_runs()
        unsafe = _suggestion(
            suggestion_id="unsafe",
            candidate_id="cand_unsafe",
            paragraph_index=1,
            paragraph_text="AXB",
            original_text="AXB",
            start=0,
            field_code="UNSAFE",
        )

        with self.assertRaisesRegex(ValueError, VISIBLE_TEXT_ERROR):
            prepare_review_document(docx, [unsafe], analysis_id="analysis_guard")

    def test_guard_skips_only_incompatible_suggestion_and_preserves_safe_one(self):
        docx = _docx_with_inline_container_between_runs()
        unsafe = _suggestion(
            suggestion_id="unsafe",
            candidate_id="cand_unsafe",
            paragraph_index=1,
            paragraph_text="AXB",
            original_text="AXB",
            start=0,
            field_code="UNSAFE",
        )
        safe = _suggestion(
            suggestion_id="safe",
            candidate_id="cand_safe",
            paragraph_index=2,
            paragraph_text="SAFE 123",
            original_text="123",
            start=5,
            field_code="SAFE_NUMBER",
        )

        result = prepare_review_document_safe(
            docx,
            [unsafe, safe],
            analysis_id="analysis_guard",
        )

        self.assertEqual(visible_text_from_docx(result.docx_bytes), visible_text_from_docx(docx))
        self.assertEqual(result.wrapped_count, 1)
        self.assertEqual(len(result.groups), 1)
        self.assertEqual(result.groups[0]["field_code"], "SAFE_NUMBER")
        self.assertIn(
            {"suggestion_id": "unsafe", "reason": VISIBLE_TEXT_SKIP_REASON},
            result.skipped,
        )

    def test_guard_keeps_fast_path_when_all_suggestions_are_safe(self):
        document = Document()
        document.add_paragraph("ESCRITURA 123")
        output = io.BytesIO()
        document.save(output)
        docx = output.getvalue()
        safe = _suggestion(
            suggestion_id="safe",
            candidate_id="cand_safe",
            paragraph_index=1,
            paragraph_text="ESCRITURA 123",
            original_text="123",
            start=10,
            field_code="NUMERO_ESCRITURA",
        )

        result = prepare_review_document_safe(docx, [safe], analysis_id="analysis_guard")

        self.assertEqual(result.wrapped_count, 1)
        self.assertEqual(result.skipped, [])
        self.assertEqual(visible_text_from_docx(result.docx_bytes), visible_text_from_docx(docx))


if __name__ == "__main__":
    unittest.main()
