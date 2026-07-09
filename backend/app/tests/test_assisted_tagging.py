from __future__ import annotations

import os
import json
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document
from docx.shared import RGBColor
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker

from app.db.base import Base
from app.models.document_template import DocumentTemplate
from app.models.notary import Notary
from app.models.template_field import TemplateField
from app.models.user import User
from app.services.minuta.assisted_tagging.approved_template_parser import ApprovedTemplateParser
from app.services.minuta.assisted_tagging.assisted_tagging_service import AssistedTaggingService
from app.services.minuta.assisted_tagging.docx_red_writer import DocxRedWriter
from app.services.minuta.assisted_tagging.docx_structure_extractor import DocxStructureExtractor
from app.services.minuta.assisted_tagging.llm_tagging_client import LlmTaggingClient
from app.services.minuta.assisted_tagging.tagging_response_validator import TaggingResponseValidator
from app.services.storage import download_storage_bytes


class AssistedTaggingTest(unittest.TestCase):
    def _sample_docx(self, path: Path) -> None:
        doc = Document()
        doc.add_paragraph("ESCRITURA PUBLICA NUMERO 1234")
        doc.add_paragraph("Comparece JUAN CAMILO VASQUEZ MIRA identificado con C.C. 1.037.657.164.")
        doc.add_paragraph("El precio de venta es $212.600.000 y la fecha es 9 de marzo de 2024.")
        doc.save(str(path))

    def test_pretagging_writes_red_and_parser_builds_placeholders(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            source = Path(tmp_dir) / "sample.docx"
            red = Path(tmp_dir) / "red.docx"
            self._sample_docx(source)

            structure = DocxStructureExtractor().extract(source)
            raw = LlmTaggingClient()._deterministic_proposal(structure)
            validated = TaggingResponseValidator().validate(raw, structure)
            self.assertGreaterEqual(len(validated.fields), 2)

            audit = DocxRedWriter().write(source, red, validated.fields)
            self.assertGreater(audit["red_runs_written"], 0)

            parsed = ApprovedTemplateParser().parse(red)
            self.assertGreater(parsed.variable_count, 0)
            technical = Path(tmp_dir) / "technical.docx"
            technical.write_bytes(parsed.technical_docx)
            doc = Document(str(technical))
            text = "\n".join(p.text for p in doc.paragraphs)
            self.assertIn("{{", text)
            self.assertIn("}}", text)
            self.assertIn("{{CAMPO_HUMANO_1}}", text)
            self.assertNotIn("{{JUAN_CAMILO_VASQUEZ_MIRA}}", text)

    def test_validator_rejects_text_not_present_in_docx(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            source = Path(tmp_dir) / "sample.docx"
            self._sample_docx(source)
            structure = DocxStructureExtractor().extract(source)
            result = TaggingResponseValidator().validate(
                {"fields": [{"field_code": "inventado", "label": "Inventado", "text": "NO EXISTE"}]},
                structure,
            )
            self.assertEqual(result.fields, [])
            self.assertTrue(result.warnings)

    def test_llm_fallback_requires_explicit_flag(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            source = Path(tmp_dir) / "sample.docx"
            self._sample_docx(source)
            structure = DocxStructureExtractor().extract(source)
            with patch.dict(os.environ, {"OPENAI_API_KEY": "", "ASSISTED_TAGGING_ALLOW_FALLBACK": ""}, clear=False):
                with self.assertRaisesRegex(RuntimeError, "ASSISTED_TAGGING_ALLOW_FALLBACK=true"):
                    LlmTaggingClient().propose(structure, "Compraventa")

    def test_integral_service_flow_uses_uploaded_approved_docx_and_safe_placeholders(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            tmp = Path(tmp_dir)
            engine = create_engine("sqlite:///:memory:")
            Base.metadata.create_all(engine)
            SessionLocal = sessionmaker(bind=engine)
            db = SessionLocal()
            try:
                notary = Notary(
                    slug="notaria-test",
                    catalog_identity_key="notaria-test",
                    commercial_name="Notaria Test",
                    legal_name="Notaria Test",
                    department="Antioquia",
                    municipality="Medellin",
                    notary_label="Notaria Test",
                    city="Medellin",
                    institutional_data="",
                )
                db.add(notary)
                db.flush()
                user = User(
                    email="protocolista@example.com",
                    full_name="Protocolista Test",
                    password_hash="x",
                    default_notary_id=notary.id,
                )
                db.add(user)
                db.flush()

                source = tmp / "sample.docx"
                self._sample_docx(source)
                service = AssistedTaggingService()

                def save_bytes(notary_id: int, job_uuid: str, filename: str, content: bytes) -> str:
                    target_dir = tmp / "storage" / str(notary_id) / job_uuid
                    target_dir.mkdir(parents=True, exist_ok=True)
                    target = target_dir / filename
                    target.write_bytes(content)
                    return str(target)

                service._save_bytes = save_bytes  # type: ignore[method-assign]

                with patch.dict(os.environ, {"OPENAI_API_KEY": "", "ASSISTED_TAGGING_ALLOW_FALLBACK": "true"}, clear=False):
                    job = service.create_job(
                        db,
                        filename="minuta-origen.docx",
                        content=source.read_bytes(),
                        document_type="Compraventa con hipoteca",
                        notary_id=notary.id,
                        user_id=user.id,
                    )
                    service.run_job(db, job)
                self.assertEqual(job.status, "human_review")
                self.assertTrue(job.pretagged_docx_storage_path)
                warnings = json.loads(job.warnings_json)
                audit = json.loads(job.audit_json)
                self.assertTrue(any("fallback local" in warning for warning in warnings))
                self.assertEqual(audit["llm"]["llm_mode"], "fallback")
                self.assertIn("model_not_used", audit["llm"])

                approved_path = tmp / "approved.docx"
                approved_path.write_bytes(download_storage_bytes(job.pretagged_docx_storage_path))
                approved_doc = Document(str(approved_path))
                run = approved_doc.add_paragraph("Dato agregado: ").add_run("TEXTO HUMANO NUEVO")
                run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                approved_doc.save(str(approved_path))

                service.upload_approved_docx(db, job, filename="aprobada.docx", content=approved_path.read_bytes())
                self.assertTrue(job.approved_docx_storage_path)
                service.approve_job(db, job, user_id=user.id)

                self.assertEqual(job.status, "saved_to_library")
                template = db.query(DocumentTemplate).filter(DocumentTemplate.id == job.template_id).one()
                fields = db.query(TemplateField).filter(TemplateField.template_id == template.id).all()
                self.assertGreaterEqual(len(fields), 2)

                field_codes = {field.field_code for field in fields}
                placeholder_keys = {field.placeholder_key for field in fields}
                self.assertIn("nombre_comprador_1", field_codes)
                self.assertIn("campo_humano_1", field_codes)
                self.assertIn("NOMBRE_COMPRADOR_1", placeholder_keys)
                self.assertIn("CAMPO_HUMANO_1", placeholder_keys)
                self.assertNotIn("juan", template.slug)
                self.assertNotIn("1037657164", template.slug)
                for field in fields:
                    raw = f"{field.field_code} {field.placeholder_key or ''}".lower()
                    self.assertNotIn("juan", raw)
                    self.assertNotIn("vasquez", raw)
                    self.assertNotIn("1037657164", raw)
                    self.assertNotIn("212600000", raw)

                technical = Document(str(job.technical_template_storage_path))
                technical_text = "\n".join(paragraph.text for paragraph in technical.paragraphs)
                self.assertIn("{{NOMBRE_COMPRADOR_1}}", technical_text)
                self.assertIn("{{CAMPO_HUMANO_1}}", technical_text)
                self.assertNotIn("{{JUAN_CAMILO_VASQUEZ_MIRA}}", technical_text)
                self.assertNotIn("{{1037657164}}", technical_text)
                red_placeholders = [
                    run.text
                    for paragraph in technical.paragraphs
                    for run in paragraph.runs
                    if run.font.color.rgb and str(run.font.color.rgb).upper() == "FF0000"
                ]
                self.assertTrue(any("{{NOMBRE_COMPRADOR_1}}" in text for text in red_placeholders))
            finally:
                db.close()


if __name__ == "__main__":
    unittest.main()
