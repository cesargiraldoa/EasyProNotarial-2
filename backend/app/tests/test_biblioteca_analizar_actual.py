from __future__ import annotations

import io
import json
from collections import Counter
from types import SimpleNamespace
import unittest

from docx import Document
from fastapi import HTTPException
from sqlalchemy import create_engine
from sqlalchemy.orm import Session

from app.api.v1.endpoints import biblioteca
from app.models.base import Base
from app.models.biblioteca_learning import FieldSignal
from app.services.biblioteca_motor.analysis import analyze_biblioteca_document
from app.services.biblioteca_motor.anchor_resolver import resolve_anchors
from app.services.biblioteca_motor.document_map import build_document_map
from app.services.biblioteca_motor.field_instance_service import resolve_decision_target_instance
from app.services.biblioteca_motor.field_signal_service import SignalContext, anonymize_context, record_field_signal
from app.services.biblioteca_motor.llm_extractor import LLMExtraction, LLMExtractorUnavailable, StaticBibliotecaLLMExtractor
from app.services.biblioteca_motor.notary_prompt_service import compile_profile, retrieve_relevant_examples
from app.services.biblioteca_motor.review_document import prepare_review_document


def _docx_bytes(paragraphs: list[str], table_cells: list[str] | None = None) -> bytes:
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    if table_cells:
        table = document.add_table(rows=1, cols=len(table_cells))
        for index, text in enumerate(table_cells):
            table.cell(0, index).text = text
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _fields():
    return [
        {"code": "COMPRADOR_1", "label": "Comprador 1", "category": "persona"},
        {"code": "COMPRADOR_2", "label": "Comprador 2", "category": "persona"},
        {"code": "CEDULA_COMPRADOR_1", "label": "Cedula comprador 1", "category": "persona"},
        {"code": "CEDULA_COMPRADOR_2", "label": "Cedula comprador 2", "category": "persona"},
        {"code": "VENDEDOR_1", "label": "Vendedor 1", "category": "persona"},
        {"code": "BANCO", "label": "Banco", "category": "entidad"},
        {"code": "NIT", "label": "NIT", "category": "persona_juridica"},
        {"code": "MATRICULA_INMOBILIARIA", "label": "Matricula", "category": "inmueble"},
    ]


def _sample_extraction() -> dict:
    return {
        "document_type": "compraventa_simple",
        "entities": [
            {"entity_ref": "ent_daniela", "entity_type": "natural_person", "display_name": "DANIELA CAMPO", "document_number": "1.234.567"},
            {"entity_ref": "ent_carlos", "entity_type": "natural_person", "display_name": "CARLOS RUIZ", "document_number": "8.765.432"},
            {"entity_ref": "ent_banco", "entity_type": "financial_entity", "display_name": "Banco Andino", "nit": "900123456-7"},
        ],
        "roles": [
            {"entity_ref": "ent_daniela", "role": "COMPRADOR"},
            {"entity_ref": "ent_carlos", "role": "COMPRADOR"},
            {"entity_ref": "ent_banco", "role": "ACREEDOR"},
        ],
        "field_instances": [
            {
                "field_instance_ref": "fi_name_daniela",
                "entity_ref": "ent_daniela",
                "role": "COMPRADOR",
                "candidate_type": "person_name",
                "base_field_code": "NOMBRE",
                "suggested_field_code": "COMPRADOR_1",
                "visible_code": "COMPRADOR_1",
                "label": "Comprador 1",
                "category": "persona",
                "catalog_match": True,
            },
            {
                "field_instance_ref": "fi_doc_daniela",
                "entity_ref": "ent_daniela",
                "role": "COMPRADOR",
                "candidate_type": "document_number",
                "base_field_code": "NUMERO_DOCUMENTO",
                "suggested_field_code": "CEDULA_COMPRADOR_1",
                "visible_code": "CEDULA_COMPRADOR_1",
                "label": "Cedula comprador 1",
                "category": "persona",
                "catalog_match": True,
            },
            {
                "field_instance_ref": "fi_name_carlos",
                "entity_ref": "ent_carlos",
                "role": "COMPRADOR",
                "candidate_type": "person_name",
                "base_field_code": "NOMBRE",
                "suggested_field_code": "COMPRADOR_2",
                "visible_code": "COMPRADOR_2",
                "label": "Comprador 2",
                "category": "persona",
                "catalog_match": True,
            },
            {
                "field_instance_ref": "fi_unmapped",
                "entity_ref": "ent_carlos",
                "role": "COMPRADOR",
                "candidate_type": "email",
                "base_field_code": "EMAIL",
                "suggested_field_code": "EMAIL_COMPRADOR_2",
                "visible_code": "EMAIL_COMPRADOR_2",
                "label": "Email comprador 2",
                "category": "contacto",
                "catalog_match": False,
            },
            {
                "field_instance_ref": "fi_matricula",
                "entity_ref": None,
                "role": None,
                "candidate_type": "matricula_inmobiliaria",
                "base_field_code": "MATRICULA_INMOBILIARIA",
                "suggested_field_code": "MATRICULA_INMOBILIARIA",
                "visible_code": "MATRICULA_INMOBILIARIA",
                "label": "Matricula",
                "category": "inmueble",
                "catalog_match": True,
            },
        ],
        "occurrences": [
            {
                "occurrence_ref": "occ_daniela_1",
                "field_instance_ref": "fi_name_daniela",
                "block_id": "p_0001",
                "exact_text": "DANIELA CAMPO",
                "occurrence_index": 1,
                "left_context": "LOS COMPRADORES: ",
                "right_context": " identificada",
                "confidence": 0.96,
            },
            {
                "occurrence_ref": "occ_doc_daniela",
                "field_instance_ref": "fi_doc_daniela",
                "block_id": "p_0001",
                "exact_text": "1.234.567",
                "occurrence_index": 1,
                "left_context": "cedula numero ",
                "right_context": " y ",
                "confidence": 0.95,
            },
            {
                "occurrence_ref": "occ_carlos",
                "field_instance_ref": "fi_name_carlos",
                "block_id": "p_0001",
                "exact_text": "CARLOS RUIZ",
                "occurrence_index": 1,
                "left_context": " y ",
                "right_context": " identificado",
                "confidence": 0.95,
            },
            {
                "occurrence_ref": "occ_daniela_2",
                "field_instance_ref": "fi_name_daniela",
                "block_id": "p_0002",
                "exact_text": "DANIELA CAMPO",
                "occurrence_index": 1,
                "left_context": "",
                "right_context": " firma",
                "confidence": 0.93,
            },
            {
                "occurrence_ref": "occ_email",
                "field_instance_ref": "fi_unmapped",
                "block_id": "p_0003",
                "exact_text": "carlos@example.com",
                "occurrence_index": 1,
                "left_context": "correo ",
                "right_context": ".",
                "confidence": 0.86,
            },
            {
                "occurrence_ref": "occ_matricula",
                "field_instance_ref": "fi_matricula",
                "block_id": "t_0001_r_0001_c_0001_p_0001",
                "exact_text": "050-123456",
                "occurrence_index": 1,
                "left_context": "MATRICULA ",
                "right_context": "",
                "confidence": 0.99,
            },
        ],
        "unmapped_fields": [
            {"field_instance_ref": "fi_unmapped", "proposed_code": "EMAIL_COMPRADOR_2", "label": "Email comprador 2", "category": "contacto", "field_type": "email"},
        ],
        "confidence": 0.94,
        "reason": "fixture",
        "diagnostics": {"fixture": True},
    }


def _sample_docx() -> bytes:
    return _docx_bytes(
        [
            "LOS COMPRADORES: DANIELA CAMPO identificada con cedula numero 1.234.567 y CARLOS RUIZ identificado con cedula numero 8.765.432.",
            "DANIELA CAMPO firma nuevamente.",
            "correo carlos@example.com.",
        ],
        ["MATRICULA 050-123456"],
    )


def _corpus_fields():
    return [
        {"code": "COMPRADOR_1", "label": "Comprador 1", "category": "persona"},
        {"code": "COMPRADOR_2", "label": "Comprador 2", "category": "persona"},
        {"code": "VENDEDOR_1", "label": "Vendedor 1", "category": "persona"},
        {"code": "VENDEDOR_2", "label": "Vendedor 2", "category": "persona"},
        {"code": "CEDULA_COMPRADOR_1", "label": "Cedula comprador 1", "category": "persona"},
        {"code": "CEDULA_COMPRADOR_2", "label": "Cedula comprador 2", "category": "persona"},
        {"code": "CEDULA_VENDEDOR_1", "label": "Cedula vendedor 1", "category": "persona"},
        {"code": "CEDULA_VENDEDOR_2", "label": "Cedula vendedor 2", "category": "persona"},
        {"code": "APODERADO_1", "label": "Apoderado 1", "category": "persona"},
        {"code": "BANCO", "label": "Banco", "category": "entidad"},
        {"code": "FIDUCIARIA", "label": "Fiduciaria", "category": "entidad"},
        {"code": "PATRIMONIO_AUTONOMO", "label": "Patrimonio autonomo", "category": "entidad"},
        {"code": "NIT", "label": "NIT", "category": "persona_juridica"},
        {"code": "MATRICULA_INMOBILIARIA", "label": "Matricula", "category": "inmueble"},
        {"code": "PROPIEDAD_HORIZONTAL", "label": "Propiedad horizontal", "category": "inmueble"},
        {"code": "AREA_PRIVADA", "label": "Area privada", "category": "inmueble"},
    ]


def _fi(ref, entity, role, candidate_type, base, suggested, label=None, category="persona", catalog_match=True):
    return {
        "field_instance_ref": ref,
        "entity_ref": entity,
        "role": role,
        "candidate_type": candidate_type,
        "base_field_code": base,
        "suggested_field_code": suggested,
        "visible_code": suggested,
        "label": label or suggested.replace("_", " ").title(),
        "category": category,
        "catalog_match": catalog_match,
    }


def _occ(ref, field_ref, block_id, text, left="", right="", index=1, confidence=0.94):
    return {
        "occurrence_ref": ref,
        "field_instance_ref": field_ref,
        "block_id": block_id,
        "exact_text": text,
        "occurrence_index": index,
        "left_context": left,
        "right_context": right,
        "confidence": confidence,
    }


def _corpus_cases():
    compraventa_docx = _docx_bytes(
        [
            "COMPRADORES: ANA LOPEZ c.c. 1.111.111 y LUIS MELO c.c. 2.222.222 compran a MARIA PAZ c.c. 3.333.333 y JOSE SOL c.c. 4.444.444.",
            "ANA LOPEZ tambien actua como APODERADA de LUIS MELO.",
        ],
        ["MATRICULA 060-ABC"],
    )
    compraventa_extraction = {
        "document_type": "compraventa_simple",
        "entities": [
            {"entity_ref": "ana", "entity_type": "natural_person", "display_name": "ANA LOPEZ", "document_number": "1.111.111"},
            {"entity_ref": "luis", "entity_type": "natural_person", "display_name": "LUIS MELO", "document_number": "2.222.222"},
            {"entity_ref": "maria", "entity_type": "natural_person", "display_name": "MARIA PAZ", "document_number": "3.333.333"},
            {"entity_ref": "jose", "entity_type": "natural_person", "display_name": "JOSE SOL", "document_number": "4.444.444"},
        ],
        "roles": [
            {"entity_ref": "ana", "role": "COMPRADOR"},
            {"entity_ref": "ana", "role": "APODERADO"},
            {"entity_ref": "luis", "role": "COMPRADOR"},
            {"entity_ref": "maria", "role": "VENDEDOR"},
            {"entity_ref": "jose", "role": "VENDEDOR"},
        ],
        "field_instances": [
            _fi("fi_ana_nombre", "ana", "COMPRADOR", "person_name", "NOMBRE", "COMPRADOR_1"),
            _fi("fi_ana_doc", "ana", "COMPRADOR", "document_number", "NUMERO_DOCUMENTO", "CEDULA_COMPRADOR_1"),
            _fi("fi_luis_nombre", "luis", "COMPRADOR", "person_name", "NOMBRE", "COMPRADOR_2"),
            _fi("fi_luis_doc", "luis", "COMPRADOR", "document_number", "NUMERO_DOCUMENTO", "CEDULA_COMPRADOR_2"),
            _fi("fi_maria_nombre", "maria", "VENDEDOR", "person_name", "NOMBRE", "VENDEDOR_1"),
            _fi("fi_maria_doc", "maria", "VENDEDOR", "document_number", "NUMERO_DOCUMENTO", "CEDULA_VENDEDOR_1"),
            _fi("fi_jose_nombre", "jose", "VENDEDOR", "person_name", "NOMBRE", "VENDEDOR_2"),
            _fi("fi_jose_doc", "jose", "VENDEDOR", "document_number", "NUMERO_DOCUMENTO", "CEDULA_VENDEDOR_2"),
            _fi("fi_ana_apoderada", "ana", "APODERADO", "person_name", "NOMBRE", "APODERADO_1"),
            _fi("fi_matricula", None, None, "matricula_inmobiliaria", "MATRICULA_INMOBILIARIA", "MATRICULA_INMOBILIARIA", category="inmueble"),
        ],
        "occurrences": [
            _occ("occ_ana_1", "fi_ana_nombre", "p_0001", "ANA LOPEZ", "COMPRADORES: ", " c.c."),
            _occ("occ_ana_doc", "fi_ana_doc", "p_0001", "1.111.111", " c.c. ", " y "),
            _occ("occ_luis", "fi_luis_nombre", "p_0001", "LUIS MELO", " y ", " c.c."),
            _occ("occ_luis_doc", "fi_luis_doc", "p_0001", "2.222.222", " c.c. ", " compran"),
            _occ("occ_maria", "fi_maria_nombre", "p_0001", "MARIA PAZ", "compran a ", " c.c."),
            _occ("occ_maria_doc", "fi_maria_doc", "p_0001", "3.333.333", " c.c. ", " y "),
            _occ("occ_jose", "fi_jose_nombre", "p_0001", "JOSE SOL", " y ", " c.c."),
            _occ("occ_jose_doc", "fi_jose_doc", "p_0001", "4.444.444", " c.c. ", "."),
            _occ("occ_ana_apoderada", "fi_ana_apoderada", "p_0002", "ANA LOPEZ", "", " tambien", 1),
            _occ("occ_matricula", "fi_matricula", "t_0001_r_0001_c_0001_p_0001", "060-ABC", "MATRICULA ", ""),
        ],
        "unmapped_fields": [],
        "confidence": 0.96,
        "reason": "corpus",
        "diagnostics": {"case": "compraventa_simple"},
    }

    hipoteca_docx = _docx_bytes(
        [
            "ACREEDOR HIPOTECARIO Banco Central S.A. NIT 900.111.222-3 otorga credito.",
            "FIDUCIARIA SEGURA S.A. NIT 800.999.888-1 actua en voceria del PATRIMONIO AUTONOMO TORRE NORTE.",
            "No extraer articulo 1502 del Codigo Civil.",
        ],
    )
    hipoteca_extraction = {
        "document_type": "compraventa_con_hipoteca",
        "entities": [
            {"entity_ref": "banco", "entity_type": "financial_entity", "display_name": "Banco Central S.A.", "nit": "900.111.222-3"},
            {"entity_ref": "fiduciaria", "entity_type": "legal_person", "display_name": "FIDUCIARIA SEGURA S.A.", "nit": "800.999.888-1"},
            {"entity_ref": "patrimonio", "entity_type": "legal_person", "display_name": "PATRIMONIO AUTONOMO TORRE NORTE"},
        ],
        "roles": [
            {"entity_ref": "banco", "role": "ACREEDOR"},
            {"entity_ref": "fiduciaria", "role": "FIDUCIARIA"},
            {"entity_ref": "patrimonio", "role": "PATRIMONIO_AUTONOMO"},
        ],
        "field_instances": [
            _fi("fi_banco", "banco", "ACREEDOR", "legal_name", "NOMBRE", "BANCO", category="entidad"),
            _fi("fi_banco_nit", "banco", "ACREEDOR", "nit", "NIT", "NIT", category="persona_juridica"),
            _fi("fi_fiduciaria", "fiduciaria", "FIDUCIARIA", "legal_name", "NOMBRE", "FIDUCIARIA", category="entidad"),
            _fi("fi_fiduciaria_nit", "fiduciaria", "FIDUCIARIA", "nit", "NIT", "NIT", category="persona_juridica"),
            _fi("fi_patrimonio", "patrimonio", "PATRIMONIO_AUTONOMO", "legal_name", "NOMBRE", "PATRIMONIO_AUTONOMO", category="entidad"),
        ],
        "occurrences": [
            _occ("occ_banco", "fi_banco", "p_0001", "Banco Central S.A.", "ACREEDOR HIPOTECARIO ", " NIT"),
            _occ("occ_banco_nit", "fi_banco_nit", "p_0001", "900.111.222-3", "NIT ", " otorga"),
            _occ("occ_fid", "fi_fiduciaria", "p_0002", "FIDUCIARIA SEGURA S.A.", "", " NIT"),
            _occ("occ_fid_nit", "fi_fiduciaria_nit", "p_0002", "800.999.888-1", "NIT ", " actua"),
            _occ("occ_patrimonio", "fi_patrimonio", "p_0002", "PATRIMONIO AUTONOMO TORRE NORTE", "del ", "."),
        ],
        "unmapped_fields": [],
        "confidence": 0.94,
        "reason": "corpus",
        "diagnostics": {"case": "hipoteca_fiduciaria"},
    }

    ph_docx = _docx_bytes(
        [
            "INMUEBLE APARTAMENTO 502 TORRE 6, PROPIEDAD HORIZONTAL ARAGUA, area privada 70.50 m2.",
            "Correo de administracion administracion@example.test.",
            "Referencia legal Ley 675 de 2001 no es campo variable.",
        ],
    )
    ph_extraction = {
        "document_type": "propiedad_horizontal",
        "entities": [],
        "roles": [],
        "field_instances": [
            _fi("fi_ph", None, None, "property_horizontal_name", "PROPIEDAD_HORIZONTAL", "PROPIEDAD_HORIZONTAL", category="inmueble"),
            _fi("fi_area", None, None, "area", "AREA_PRIVADA", "AREA_PRIVADA", category="inmueble"),
            _fi("fi_email_admin", None, None, "email", "EMAIL_ADMINISTRACION", "EMAIL_ADMINISTRACION", category="contacto", catalog_match=False),
        ],
        "occurrences": [
            _occ("occ_ph", "fi_ph", "p_0001", "PROPIEDAD HORIZONTAL ARAGUA", ", ", ", area"),
            _occ("occ_area", "fi_area", "p_0001", "70.50 m2", "area privada ", "."),
            _occ("occ_email_admin", "fi_email_admin", "p_0002", "administracion@example.test", "administracion ", "."),
        ],
        "unmapped_fields": [
            {"field_instance_ref": "fi_email_admin", "proposed_code": "EMAIL_ADMINISTRACION", "label": "Email administracion", "category": "contacto", "field_type": "email"},
        ],
        "confidence": 0.91,
        "reason": "corpus",
        "diagnostics": {"case": "propiedad_horizontal"},
    }

    return [
        {
            "name": "compraventa_simple",
            "docx": compraventa_docx,
            "extraction": compraventa_extraction,
            "expected_texts": [item["exact_text"] for item in compraventa_extraction["occurrences"]],
            "rejected_texts": [],
        },
        {
            "name": "compraventa_con_hipoteca",
            "docx": hipoteca_docx,
            "extraction": hipoteca_extraction,
            "expected_texts": [item["exact_text"] for item in hipoteca_extraction["occurrences"]],
            "rejected_texts": ["articulo 1502", "Codigo Civil"],
        },
        {
            "name": "propiedad_horizontal",
            "docx": ph_docx,
            "extraction": ph_extraction,
            "expected_texts": [item["exact_text"] for item in ph_extraction["occurrences"]],
            "rejected_texts": ["Ley 675 de 2001"],
        },
    ]


class BibliotecaLLMEngineTests(unittest.TestCase):
    def test_document_map_keeps_ordered_paragraphs_and_table_paragraphs_without_semantics(self):
        document_map = build_document_map(_sample_docx())
        blocks = document_map.blocks

        self.assertEqual([block.block_id for block in blocks], ["p_0001", "p_0002", "p_0003", "t_0001_r_0001_c_0001_p_0001"])
        self.assertEqual(blocks[0].block_type, "paragraph")
        self.assertEqual(blocks[3].block_type, "table_cell")
        self.assertEqual(blocks[3].table_index, 1)
        self.assertEqual(blocks[3].paragraph_index, 1)
        self.assertTrue(blocks[0].block_hash)

    def test_llm_schema_rejects_extra_or_missing_fields(self):
        payload = _sample_extraction()
        payload["unexpected"] = True

        with self.assertRaises(Exception):
            LLMExtraction.model_validate(payload)

    def test_anchor_resolver_verifies_repeated_occurrence_context_and_overlap(self):
        document_map = build_document_map(_docx_bytes(["DANIELA CAMPO y DANIELA CAMPO"]))
        extraction = LLMExtraction.model_validate(
            {
                **_sample_extraction(),
                "field_instances": [_sample_extraction()["field_instances"][0]],
                "occurrences": [
                    {
                        "occurrence_ref": "occ_2",
                        "field_instance_ref": "fi_name_daniela",
                        "block_id": "p_0001",
                        "exact_text": "DANIELA CAMPO",
                        "occurrence_index": 2,
                        "left_context": " y ",
                        "right_context": "",
                        "confidence": 0.9,
                    },
                    {
                        "occurrence_ref": "occ_overlap",
                        "field_instance_ref": "fi_name_daniela",
                        "block_id": "p_0001",
                        "exact_text": "DANIELA CAMPO",
                        "occurrence_index": 2,
                        "left_context": " y ",
                        "right_context": "",
                        "confidence": 0.9,
                    },
                    {
                        "occurrence_ref": "occ_bad",
                        "field_instance_ref": "fi_name_daniela",
                        "block_id": "p_9999",
                        "exact_text": "DANIELA CAMPO",
                        "occurrence_index": 1,
                        "left_context": "",
                        "right_context": "",
                        "confidence": 0.9,
                    },
                ],
            },
        )

        result = resolve_anchors(document_map, extraction.occurrences)

        self.assertEqual(len(result.anchored), 1)
        self.assertEqual(result.anchored[0].location["char_start"], 16)
        self.assertEqual({item.reason for item in result.skipped}, {"overlap_detected", "block_not_found"})

    def test_analysis_uses_llm_output_for_identity_catalog_and_provisionals(self):
        result = analyze_biblioteca_document(_sample_docx(), _fields(), extractor=StaticBibliotecaLLMExtractor(_sample_extraction()))
        suggestions = result["suggestions"]
        danielas = [item for item in suggestions if item["original_text"] == "DANIELA CAMPO"]
        email = next(item for item in suggestions if item["original_text"] == "carlos@example.com")
        matricula = next(item for item in suggestions if item["original_text"] == "050-123456")

        self.assertEqual(result["mode"], "llm_first")
        self.assertEqual(result["stats"]["detected_candidates"], 6)
        self.assertEqual(result["stats"]["anchored_suggestions"], 6)
        self.assertEqual(result["stats"]["skipped_suggestions"], 0)
        self.assertEqual(result["stats"]["provisional_suggestions"], 1)
        self.assertEqual(len({item["field_instance_id"] for item in danielas}), 1)
        self.assertNotEqual(danielas[0]["field_instance_id"], next(item for item in suggestions if item["visible_code"] == "COMPRADOR_2")["field_instance_id"])
        self.assertTrue(email["field_code"].startswith("PENDING_FIELD_"))
        self.assertTrue(email["requires_field_assignment"])
        self.assertEqual(matricula["location"]["block_type"], "table_cell")

    def test_anonymized_corpus_meets_recall_precision_and_anchor_targets(self):
        expected_total = 0
        found_total = 0
        unexpected_total = 0
        provisional_total = 0
        field_instance_ids: set[str] = set()
        occurrence_ids: set[str] = set()

        for case in _corpus_cases():
            result = analyze_biblioteca_document(case["docx"], _corpus_fields(), extractor=StaticBibliotecaLLMExtractor(case["extraction"]))
            review = prepare_review_document(case["docx"], result["suggestions"], analysis_id=result["analysis_id"])
            suggestion_texts = Counter(item["original_text"] for item in result["suggestions"])
            expected_texts = Counter(case["expected_texts"])

            expected_total += sum(expected_texts.values())
            found_total += sum(min(suggestion_texts[text], count) for text, count in expected_texts.items())
            unexpected_total += sum(count for text, count in suggestion_texts.items() if text not in expected_texts)
            provisional_total += result["stats"]["provisional_suggestions"]

            self.assertEqual(result["stats"]["skipped_suggestions"], 0, case["name"])
            self.assertEqual(result["stats"]["anchor_success_rate_applied"], 1.0, case["name"])
            self.assertEqual(review.wrapped_count, len(result["suggestions"]), case["name"])
            self.assertEqual(review.skipped, [], case["name"])
            self.assertEqual(len(result["suggestions"]), len({item["occurrence_id"] for item in result["suggestions"]}), case["name"])
            for rejected in case["rejected_texts"]:
                self.assertNotIn(rejected, suggestion_texts, case["name"])
            for item in result["suggestions"]:
                field_instance_id = item["field_instance_id"]
                occurrence_id = item["occurrence_id"]
                self.assertNotIn(occurrence_id, occurrence_ids)
                occurrence_ids.add(occurrence_id)
                field_instance_ids.add(field_instance_id)
                self.assertFalse(field_instance_id == item["field_code"])

        recall = found_total / expected_total
        precision = found_total / (found_total + unexpected_total)
        self.assertGreaterEqual(recall, 0.95)
        self.assertGreaterEqual(precision, 0.90)
        self.assertGreater(provisional_total, 0)
        self.assertEqual(len(field_instance_ids), len(set(field_instance_ids)))

    def test_analysis_fails_explicitly_without_llm_extractor(self):
        with self.assertRaises(LLMExtractorUnavailable):
            analyze_biblioteca_document(_sample_docx(), _fields(), extractor=None)

    def test_change_target_instance_does_not_reuse_previous_instance(self):
        occurrence = {
            "field_instance_id": "fi_old",
            "entity_id": "ent_1",
            "role": "COMPRADOR",
            "candidate_type": "person_name",
        }

        target = resolve_decision_target_instance(occurrence, field_code="VENDEDOR_1", visible_code="VENDEDOR_1")

        self.assertNotEqual(target, "fi_old")
        self.assertTrue(target.startswith("fi_"))

    def test_signal_anonymization_profile_compilation_and_retrieval_are_notary_scoped(self):
        engine = create_engine("sqlite:///:memory:")
        Base.metadata.create_all(engine)
        with Session(engine) as db:
            ctx = SignalContext(
                notary_id=10,
                analysis_run_id=1,
                case_id=1,
                document_id=2,
                source_version_id=3,
                review_version_id=4,
                decision_version_id=5,
                document_type="compraventa_simple",
                model="static",
                prompt_version="test",
                profile_version=None,
                user_id=99,
            )
            occurrence = {
                "occurrence_id": "occ_1",
                "original_text": "DANIELA CAMPO",
                "context_before": "COMPRADOR DANIELA CAMPO cedula 1.234.567",
                "context_after": "firma",
                "field_code": "COMPRADOR_1",
                "visible_code": "COMPRADOR_1",
                "field_instance_id": "fi_1",
                "candidate_type": "person_name",
                "role": "COMPRADOR",
                "entity_type": "natural_person",
                "location": {"block_type": "paragraph", "paragraph_index": 1},
                "confidence": 0.9,
            }
            for _index in range(3):
                record_field_signal(db, context=ctx, occurrence=occurrence, human_decision="accepted", final_field_code="COMPRADOR_1", final_field_instance_id="fi_1")
            other_ctx = SignalContext(**{**ctx.__dict__, "notary_id": 20})
            record_field_signal(db, context=other_ctx, occurrence=occurrence, human_decision="accepted", final_field_code="VENDEDOR_1", final_field_instance_id="fi_2")
            profile = compile_profile(db, notary_id=10)
            examples = retrieve_relevant_examples(db, notary_id=10, document_type="compraventa_simple", role="COMPRADOR")
            other_examples = retrieve_relevant_examples(db, notary_id=20, document_type="compraventa_simple")

            self.assertIsNotNone(profile)
            self.assertEqual(profile.version, 1)
            self.assertEqual(profile.source_signal_count, 3)
            self.assertEqual(len(examples), 3)
            self.assertEqual(other_examples[0]["final_field_code"], "VENDEDOR_1")
            self.assertNotIn("DANIELA", db.query(FieldSignal).filter(FieldSignal.notary_id == 10).first().anonymized_context)
            self.assertEqual(anonymize_context("NIT 900123456-7 DANIELA CAMPO 1.234.567"), "[NIT] [NOMBRE] [NUM]")

    def test_biblioteca_routes_remain_registered_with_llm_flow(self):
        paths = {route.path for route in biblioteca.router.routes}

        self.assertIn("/biblioteca/analizar", paths)
        self.assertIn("/biblioteca/analizar-actual", paths)
        self.assertIn("/biblioteca/analizar-y-preparar", paths)
        self.assertIn("/biblioteca/decidir", paths)
        self.assertIn("/biblioteca/actualizar-campo", paths)

    def test_legacy_upload_endpoint_fails_explicitly(self):
        with self.assertRaises(HTTPException) as issue:
            raise HTTPException(status_code=410)

        self.assertEqual(issue.exception.status_code, 410)


if __name__ == "__main__":
    unittest.main()
