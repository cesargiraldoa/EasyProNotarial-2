from __future__ import annotations

import json
import tempfile
import unittest
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.shared import RGBColor
from fastapi import FastAPI
from fastapi.testclient import TestClient
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker
from sqlalchemy.pool import StaticPool

from app.api.routes.inverse_conversion_engine import router as inverse_conversion_router
from app.core.deps import get_db
from app.models.base import Base
from app.services.minuta.inverse_conversion_catalog.models import FieldAlias, FieldDefinition
from app.services.minuta.inverse_conversion_writer import DocxMarkedWriter, MarkedCandidate, MarkedCandidateContext
from app.services.minuta.marked_template_detector import detect_marked_template
from app.services.minuta.marked_template_generator import apply_marked_template_replacements


RED = RGBColor(0xFF, 0x00, 0x00)
ALLOWED_CODES = {
    "NUMERO_APARTAMENTO",
    "NUMERO_MATRICULA",
    "NUMERO_DOCUMENTO_COMPRADOR",
    "NIT_COMPRADOR",
    "NIT_VENDEDOR",
    "VALOR_VENTA",
    "EMAIL_COMPRADOR",
}


def _iter_paragraphs(document):
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        seen_cells: set[int] = set()
        for row in table.rows:
            for cell in row.cells:
                cell_id = id(cell._tc)
                if cell_id in seen_cells:
                    continue
                seen_cells.add(cell_id)
                yield from _iter_paragraphs(cell)


class InverseConversionWriterTests(unittest.TestCase):
    def _save_doc(self, document: Document, directory: str, filename: str = "source.docx") -> Path:
        source = Path(directory) / filename
        document.save(source)
        return source

    def _candidate(
        self,
        text: str,
        canonical: str = "NUMERO_APARTAMENTO",
        status: str = "accepted",
        location: str = "paragraph 1",
        before: str = "",
        after: str = "",
        suggested_key: str | None = None,
        occurrences: int = 1,
    ) -> MarkedCandidate:
        return MarkedCandidate(
            id=f"cand-{canonical}-{text}-{before}",
            text=text,
            suggested_key=suggested_key or canonical.lower(),
            canonical_field_code=canonical,
            label=canonical.replace("_", " ").title(),
            status=status,
            confidence=0.95,
            occurrences=occurrences,
            contexts=(MarkedCandidateContext(location=location, before=before, after=after),),
        )

    def _runs_for_text(self, path: Path, text: str):
        document = Document(str(path))
        return [run for paragraph in _iter_paragraphs(document) for run in paragraph.runs if run.text == text]

    def test_writer_replaces_simple_value_with_exact_red_marker(self):
        with tempfile.TemporaryDirectory() as tmp:
            doc = Document()
            doc.add_paragraph("APARTAMENTO: 804")
            source = self._save_doc(doc, tmp)
            output = Path(tmp) / "marked.docx"

            result = DocxMarkedWriter().write(
                source,
                output,
                [self._candidate("804", "NUMERO_APARTAMENTO", before="APARTAMENTO:")],
                allowed_field_codes=ALLOWED_CODES,
            )

            paragraph = Document(str(output)).paragraphs[0]
            self.assertEqual(paragraph.text, "APARTAMENTO: {{NUMERO_APARTAMENTO}}")
            self.assertEqual(result.marked_occurrences, 1)
            marker_runs = self._runs_for_text(output, "{{NUMERO_APARTAMENTO}}")
            self.assertEqual(len(marker_runs), 1)
            self.assertEqual(marker_runs[0].font.color.rgb, RED)

    def test_writer_replaces_text_split_between_runs(self):
        with tempfile.TemporaryDirectory() as tmp:
            doc = Document()
            paragraph = doc.add_paragraph()
            paragraph.add_run("MATRICULA: ")
            paragraph.add_run("001-")
            paragraph.add_run("1528731")
            source = self._save_doc(doc, tmp)
            output = Path(tmp) / "marked.docx"

            DocxMarkedWriter().write(
                source,
                output,
                [
                    self._candidate(
                        "001-1528731",
                        "NUMERO_MATRICULA",
                        before="MATRICULA:",
                        suggested_key="matricula_inmobiliaria",
                    )
                ],
                allowed_field_codes=ALLOWED_CODES,
            )

            paragraph = Document(str(output)).paragraphs[0]
            self.assertEqual(paragraph.text, "MATRICULA: {{NUMERO_MATRICULA}}")
            self.assertEqual(self._runs_for_text(output, "{{NUMERO_MATRICULA}}")[0].font.color.rgb, RED)

    def test_writer_replaces_values_inside_tables(self):
        with tempfile.TemporaryDirectory() as tmp:
            doc = Document()
            table = doc.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "Apartamento"
            table.cell(0, 1).text = "804"
            source = self._save_doc(doc, tmp)
            output = Path(tmp) / "marked.docx"

            DocxMarkedWriter().write(
                source,
                output,
                [self._candidate("804", "NUMERO_APARTAMENTO", location="table 1 row 1 cell 2 paragraph 1")],
                allowed_field_codes=ALLOWED_CODES,
            )

            generated = Document(str(output))
            self.assertEqual(generated.tables[0].cell(0, 1).text, "{{NUMERO_APARTAMENTO}}")
            self.assertEqual(self._runs_for_text(output, "{{NUMERO_APARTAMENTO}}")[0].font.color.rgb, RED)

    def test_writer_replaces_multiple_occurrences_from_candidate_contexts(self):
        with tempfile.TemporaryDirectory() as tmp:
            doc = Document()
            doc.add_paragraph("APARTAMENTO: 804")
            doc.add_paragraph("UNIDAD PRIVADA: 804")
            source = self._save_doc(doc, tmp)
            output = Path(tmp) / "marked.docx"
            candidate = MarkedCandidate(
                id="cand-apartment",
                text="804",
                suggested_key="numero_apartamento",
                canonical_field_code="NUMERO_APARTAMENTO",
                label="Numero apartamento",
                status="accepted",
                confidence=0.95,
                occurrences=2,
                contexts=(
                    MarkedCandidateContext(location="paragraph 1", before="APARTAMENTO:", after=""),
                    MarkedCandidateContext(location="paragraph 2", before="UNIDAD PRIVADA:", after=""),
                ),
            )

            result = DocxMarkedWriter().write(source, output, [candidate], allowed_field_codes=ALLOWED_CODES)

            generated = Document(str(output))
            self.assertEqual(generated.paragraphs[0].text, "APARTAMENTO: {{NUMERO_APARTAMENTO}}")
            self.assertEqual(generated.paragraphs[1].text, "UNIDAD PRIVADA: {{NUMERO_APARTAMENTO}}")
            self.assertEqual(result.marked_occurrences, 2)

    def test_writer_ignores_rejected_and_pending_candidates(self):
        with tempfile.TemporaryDirectory() as tmp:
            doc = Document()
            doc.add_paragraph("APARTAMENTO: 804 MATRICULA: 001-1528731 CORREO: juan@example.com")
            source = self._save_doc(doc, tmp)
            output = Path(tmp) / "marked.docx"

            DocxMarkedWriter().write(
                source,
                output,
                [
                    self._candidate("804", "NUMERO_APARTAMENTO", before="APARTAMENTO:"),
                    self._candidate("001-1528731", "NUMERO_MATRICULA", status="rejected", before="MATRICULA:"),
                    self._candidate("juan@example.com", "EMAIL_COMPRADOR", status="pending", before="CORREO:"),
                ],
                allowed_field_codes=ALLOWED_CODES,
            )

            text = Document(str(output)).paragraphs[0].text
            self.assertIn("{{NUMERO_APARTAMENTO}}", text)
            self.assertIn("001-1528731", text)
            self.assertIn("juan@example.com", text)
            self.assertNotIn("{{NUMERO_MATRICULA}}", text)
            self.assertNotIn("{{EMAIL_COMPRADOR}}", text)

    def test_writer_rejects_invalid_canonical_code(self):
        with tempfile.TemporaryDirectory() as tmp:
            doc = Document()
            doc.add_paragraph("APARTAMENTO: 804")
            source = self._save_doc(doc, tmp)
            output = Path(tmp) / "marked.docx"

            with self.assertRaisesRegex(ValueError, "field_definitions"):
                DocxMarkedWriter().write(
                    source,
                    output,
                    [self._candidate("804", "CAMPO_INVENTADO", before="APARTAMENTO:")],
                    allowed_field_codes=ALLOWED_CODES,
                )
            self.assertFalse(output.exists())

    def test_writer_deduplicates_identical_candidates(self):
        with tempfile.TemporaryDirectory() as tmp:
            doc = Document()
            doc.add_paragraph("APARTAMENTO: 804")
            source = self._save_doc(doc, tmp)
            output = Path(tmp) / "marked.docx"
            candidate = self._candidate("804", "NUMERO_APARTAMENTO", before="APARTAMENTO:")

            result = DocxMarkedWriter().write(source, output, [candidate, candidate], allowed_field_codes=ALLOWED_CODES)

            self.assertEqual(Document(str(output)).paragraphs[0].text.count("{{NUMERO_APARTAMENTO}}"), 1)
            self.assertEqual(result.marked_occurrences, 1)

    def test_writer_uses_context_for_same_text_with_different_roles(self):
        with tempfile.TemporaryDirectory() as tmp:
            doc = Document()
            doc.add_paragraph("COMPRADOR NIT 830.054.539-0 y VENDEDOR NIT 830.054.539-0.")
            source = self._save_doc(doc, tmp)
            output = Path(tmp) / "marked.docx"

            DocxMarkedWriter().write(
                source,
                output,
                [
                    self._candidate("830.054.539-0", "NIT_COMPRADOR", before="COMPRADOR NIT", after="y"),
                    self._candidate("830.054.539-0", "NIT_VENDEDOR", before="VENDEDOR NIT", after="."),
                ],
                allowed_field_codes=ALLOWED_CODES,
            )

            text = Document(str(output)).paragraphs[0].text
            self.assertEqual(text, "COMPRADOR NIT {{NIT_COMPRADOR}} y VENDEDOR NIT {{NIT_VENDEDOR}}.")

    def test_writer_does_not_modify_original_file(self):
        with tempfile.TemporaryDirectory() as tmp:
            doc = Document()
            doc.add_paragraph("APARTAMENTO: 804")
            source = self._save_doc(doc, tmp)
            original_bytes = source.read_bytes()
            output = Path(tmp) / "marked.docx"

            DocxMarkedWriter().write(
                source,
                output,
                [self._candidate("804", "NUMERO_APARTAMENTO", before="APARTAMENTO:")],
                allowed_field_codes=ALLOWED_CODES,
            )

            self.assertEqual(source.read_bytes(), original_bytes)

    def test_writer_rejects_without_accepted_candidates(self):
        with tempfile.TemporaryDirectory() as tmp:
            doc = Document()
            doc.add_paragraph("APARTAMENTO: 804")
            source = self._save_doc(doc, tmp)
            output = Path(tmp) / "marked.docx"

            with self.assertRaises(ValueError):
                DocxMarkedWriter().write(
                    source,
                    output,
                    [self._candidate("804", "NUMERO_APARTAMENTO", status="rejected")],
                    allowed_field_codes=ALLOWED_CODES,
                )

    def test_marked_docx_handoff_to_traditional_detector_and_generator(self):
        with tempfile.TemporaryDirectory() as tmp:
            doc = Document()
            doc.add_paragraph("APARTAMENTO: 804")
            source = self._save_doc(doc, tmp)
            marked = Path(tmp) / "marked.docx"
            final = Path(tmp) / "final.docx"

            DocxMarkedWriter().write(
                source,
                marked,
                [self._candidate("804", "NUMERO_APARTAMENTO", before="APARTAMENTO:")],
                allowed_field_codes=ALLOWED_CODES,
            )
            detection = detect_marked_template(marked)

            self.assertEqual(detection["marker_types"]["curly"], 1)
            self.assertTrue(any(field["key"] == "numero_apartamento" for field in detection["fields"]))

            stats = apply_marked_template_replacements(
                marked,
                final,
                {"numero_apartamento": "905"},
                detection["fields"],
            )

            self.assertGreaterEqual(stats["total_replacements"], 1)
            self.assertEqual(Document(str(final)).paragraphs[0].text, "APARTAMENTO: 905")

    def _build_endpoint_docx_bytes(self) -> bytes:
        buffer = BytesIO()
        document = Document()
        document.add_paragraph("APARTAMENTO: 804")
        document.save(buffer)
        return buffer.getvalue()

    def _build_endpoint_client(self) -> tuple[TestClient, object]:
        engine = create_engine(
            "sqlite:///:memory:",
            connect_args={"check_same_thread": False},
            poolclass=StaticPool,
        )
        Base.metadata.create_all(engine, tables=[FieldDefinition.__table__, FieldAlias.__table__])
        Session = sessionmaker(bind=engine)
        db = Session()
        db.add(FieldDefinition(field_code="NUMERO_APARTAMENTO", display_name="Numero Apartamento", status="suggested", confidence=0.9))
        db.add(FieldAlias(raw_field_code="numero_apartamento", canonical_field_code="NUMERO_APARTAMENTO", frequency=1, status="suggested"))
        db.commit()

        app = FastAPI()
        app.include_router(inverse_conversion_router, prefix="/api/v1")

        def override_get_db():
            session = Session()
            try:
                yield session
            finally:
                session.close()

        app.dependency_overrides[get_db] = override_get_db
        return TestClient(app), engine

    def test_endpoint_generates_downloadable_semantic_docx(self):
        client, engine = self._build_endpoint_client()
        try:
            candidates = [
                {
                    "id": "cand_001",
                    "text": "804",
                    "suggested_key": "numero_apartamento",
                    "canonical_field_code": "NUMERO_APARTAMENTO",
                    "label": "Numero apartamento",
                    "status": "accepted",
                    "confidence": 0.95,
                    "occurrences": 1,
                    "contexts": [{"location": "paragraph 1", "before": "APARTAMENTO:", "after": ""}],
                }
            ]

            response = client.post(
                "/api/v1/inverse-conversion/generate-marked-docx",
                files={
                    "archivo": (
                        "minuta.docx",
                        self._build_endpoint_docx_bytes(),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                },
                data={"candidates": json.dumps(candidates)},
            )

            self.assertEqual(response.status_code, 200)
            self.assertEqual(response.headers["content-type"], "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            self.assertIn("minuta%20-%20marcado.docx", response.headers["content-disposition"])
            with tempfile.TemporaryDirectory() as tmp:
                generated = Path(tmp) / "generated.docx"
                generated.write_bytes(response.content)
                self.assertEqual(Document(str(generated)).paragraphs[0].text, "APARTAMENTO: {{NUMERO_APARTAMENTO}}")
        finally:
            engine.dispose()

    def test_endpoint_returns_controlled_error_for_invalid_canonical_code(self):
        client, engine = self._build_endpoint_client()
        try:
            candidates = [
                {
                    "id": "cand_001",
                    "text": "804",
                    "suggested_key": "numero_apartamento",
                    "canonical_field_code": "CAMPO_INVENTADO",
                    "status": "accepted",
                    "contexts": [{"location": "paragraph 1", "before": "APARTAMENTO:", "after": ""}],
                }
            ]

            response = client.post(
                "/api/v1/inverse-conversion/generate-marked-docx",
                files={
                    "archivo": (
                        "minuta.docx",
                        self._build_endpoint_docx_bytes(),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                },
                data={"candidates": json.dumps(candidates)},
            )

            self.assertEqual(response.status_code, 422)
            self.assertIn("field_definitions", response.json()["detail"])
        finally:
            engine.dispose()

    def test_endpoint_returns_controlled_error_without_accepted_candidates(self):
        client, engine = self._build_endpoint_client()
        try:
            candidates = [{"id": "cand_001", "text": "804", "status": "rejected"}]

            response = client.post(
                "/api/v1/inverse-conversion/generate-marked-docx",
                files={
                    "archivo": (
                        "minuta.docx",
                        self._build_endpoint_docx_bytes(),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                },
                data={"candidates": json.dumps(candidates)},
            )

            self.assertEqual(response.status_code, 422)
            self.assertIn("aceptar al menos un candidato", response.json()["detail"])
        finally:
            engine.dispose()


if __name__ == "__main__":
    unittest.main()
