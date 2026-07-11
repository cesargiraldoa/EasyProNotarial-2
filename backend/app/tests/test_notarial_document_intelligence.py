from __future__ import annotations

import importlib.util
import json
import tempfile
import threading
import unittest
import zipfile
from io import BytesIO
from pathlib import Path
from types import SimpleNamespace

from docx import Document
from fastapi import FastAPI
from fastapi.testclient import TestClient
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker
from sqlalchemy.pool import StaticPool

from app.api.routes import notarial_document_intelligence as route_module
from app.api.routes.notarial_document_intelligence import get_ingestion_service, router as intelligence_router
from app.core.deps import get_current_user, get_db
from app.models.base import Base
from app.models.notarial_document_intelligence import (
    NotarialDocument,
    NotarialDocumentBatch,
    NotarialDocumentBatchItem,
    NotarialBenchmarkRun,
    NotarialBlockAlignment,
    NotarialDocumentBlock,
    NotarialDocumentCluster,
    NotarialDocumentClusterMember,
    NotarialDocumentDecision,
    NotarialDocumentEmbedding,
    NotarialDocumentEntity,
    NotarialDocumentEvidence,
    NotarialDocumentFamily,
    NotarialDocumentFamilyMember,
    NotarialHumanFieldReview,
    NotarialHumanReviewAudit,
    NotarialHumanReviewSession,
    NotarialIntelligenceRun,
    NotarialDocumentParseRun,
    NotarialDocumentSection,
    NotarialTaskPublication,
    NotarialEmbeddingVersion,
    NotarialTemplateLibraryItem,
    NotarialTemplateVersion,
)
from app.models.notary import Notary
from app.services.minuta.inverse_conversion_catalog.models import FieldDefinition
from app.services.notarial_document_intelligence.contracts import (
    ParsedDocument,
    BatchStatus,
    DocumentProcessingStatus,
    DocumentUpload,
    IngestBatchRequest,
)
from app.services.notarial_document_intelligence.ingestion import NotarialDocumentIngestionService, TransientDocumentProcessingError
from app.services.notarial_document_intelligence.parser import NotarialDocumentParser
from app.services.notarial_document_intelligence.storage import NotarialIntelligenceStorage
from app.services.notarial_document_intelligence import storage as storage_module
from app.services.notarial_document_intelligence.embedding_provider import HashEmbeddingProvider
from app.services.notarial_document_intelligence.engine import NotarialIntelligenceEngine
from app.services.notarial_document_intelligence.llm_provider import CircuitOpenError, IntelligenceMode, NotarialLLMService, RedisCircuitBreaker
from app.services.notarial_document_intelligence.vector_store import NotarialVectorStore


TABLES = [
    Notary.__table__,
    NotarialDocumentBatch.__table__,
    NotarialDocumentFamily.__table__,
    NotarialDocument.__table__,
    NotarialDocumentBatchItem.__table__,
    NotarialDocumentParseRun.__table__,
    NotarialTaskPublication.__table__,
    NotarialDocumentSection.__table__,
    NotarialDocumentBlock.__table__,
    NotarialDocumentEntity.__table__,
    NotarialDocumentFamilyMember.__table__,
    NotarialDocumentCluster.__table__,
    NotarialDocumentClusterMember.__table__,
    NotarialIntelligenceRun.__table__,
    NotarialEmbeddingVersion.__table__,
    NotarialDocumentEmbedding.__table__,
    NotarialDocumentEvidence.__table__,
    NotarialBlockAlignment.__table__,
    NotarialDocumentDecision.__table__,
    NotarialBenchmarkRun.__table__,
    FieldDefinition.__table__,
    NotarialHumanReviewSession.__table__,
    NotarialHumanFieldReview.__table__,
    NotarialTemplateLibraryItem.__table__,
    NotarialTemplateVersion.__table__,
    NotarialHumanReviewAudit.__table__,
]


def build_docx_bytes(title: str = "COMPRAVENTA", apartment: str = "804", include_table: bool = True, extra_lines: list[str] | None = None) -> bytes:
    buffer = BytesIO()
    document = Document()
    document.add_paragraph(f"ESCRITURA DE {title}")
    document.add_paragraph(f"APARTAMENTO: {apartment}")
    for line in extra_lines or []:
        document.add_paragraph(line)
    if include_table:
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "MATRICULA"
        table.cell(0, 1).text = f"001-1528{apartment}"
    document.save(buffer)
    return buffer.getvalue()


class NotarialDocumentIntelligenceTests(unittest.TestCase):
    def setUp(self):
        self.engine = create_engine(
            "sqlite:///:memory:",
            connect_args={"check_same_thread": False},
            poolclass=StaticPool,
        )
        Base.metadata.create_all(self.engine, tables=TABLES)
        self.Session = sessionmaker(bind=self.engine)
        self.tmp = tempfile.TemporaryDirectory()
        self.storage = NotarialIntelligenceStorage(local_root=Path(self.tmp.name) / "storage", use_supabase=False)
        self.notary_id = 1
        self._seed_notaries()

    def tearDown(self):
        Base.metadata.drop_all(self.engine, tables=TABLES)
        self.engine.dispose()
        self.tmp.cleanup()

    def _seed_notaries(self):
        db = self.Session()
        try:
            db.add_all(
                [
                    _notary(1, "notaria-uno"),
                    _notary(2, "notaria-dos"),
                ]
            )
            db.commit()
        finally:
            db.close()

    def test_parser_extracts_paragraphs_tables_and_stable_locations(self):
        parsed = NotarialDocumentParser().parse_bytes(build_docx_bytes(), filename="minuta.docx")

        locations = [block.location_key for block in parsed.blocks]
        texts = [block.text for block in parsed.blocks]

        self.assertIn("paragraph:1", locations)
        self.assertIn("paragraph:2", locations)
        self.assertIn("table:1/row:1/cell:1/paragraph:1", locations)
        self.assertIn("table:1/row:1/cell:2/paragraph:1", locations)
        self.assertIn("APARTAMENTO: 804", texts)
        self.assertTrue(parsed.parser_version.startswith("python-docx:"))

    @unittest.skipUnless(importlib.util.find_spec("unstructured"), "unstructured no esta instalado en el entorno local")
    def test_parser_reconciles_unstructured_semantics_when_available(self):
        parsed = NotarialDocumentParser().parse_bytes(build_docx_bytes(title="COMPRAVENTA INMUEBLE"), filename="minuta.docx")

        self.assertTrue(parsed.metadata["unstructured"]["enabled"])
        self.assertGreater(parsed.metadata["unstructured"]["reconciled_blocks"], 0)
        self.assertTrue(any(block.parser_source == "python-docx+unstructured" for block in parsed.blocks))

    def test_ingestion_stores_private_file_and_reuses_existing_document_by_hash(self):
        db = self.Session()
        try:
            content = build_docx_bytes()
            result = NotarialDocumentIngestionService(db, notary_id=1, storage=self.storage).ingest_batch(
                IngestBatchRequest(
                    name="Lote idempotente",
                    documents=[
                        DocumentUpload(filename="minuta-a.docx", content=content),
                        DocumentUpload(filename="minuta-b.docx", content=content),
                    ],
                )
            )

            self.assertEqual(result.total_documents, 2)
            self.assertEqual(result.unique_documents, 1)
            self.assertEqual(result.duplicate_documents, 1)
            self.assertEqual(db.query(NotarialDocument).count(), 1)
            self.assertEqual(db.query(NotarialDocumentBatchItem).count(), 2)
            self.assertGreater(db.query(NotarialDocumentBlock).count(), 0)
            stored_path = Path(result.documents[0].storage_path)
            self.assertTrue(stored_path.exists())
            self.assertTrue(str(stored_path).startswith(str(Path(self.tmp.name) / "storage")))
            self.assertEqual(result.documents[1].status, DocumentProcessingStatus.REUSED)
        finally:
            db.close()

    def test_ingestion_processes_synthetic_batch_equivalent_to_hundreds(self):
        db = self.Session()
        try:
            documents = [
                DocumentUpload(
                    filename=f"minuta-{index:03d}.docx",
                    content=build_docx_bytes(apartment=str(800 + index), include_table=index % 2 == 0),
                    metadata={"synthetic": True, "index": index},
                )
                for index in range(1, 201)
            ]

            result = NotarialDocumentIngestionService(db, notary_id=1, storage=self.storage).ingest_batch(
                IngestBatchRequest(name="Lote sintetico 200", source_type="synthetic", documents=documents)
            )

            self.assertEqual(result.total_documents, 200)
            self.assertEqual(result.error_documents, 0)
            self.assertEqual(result.unique_documents, 200)
            self.assertEqual(db.query(NotarialDocument).count(), 200)
            self.assertEqual(db.query(NotarialDocumentBatchItem).count(), 200)
            self.assertGreaterEqual(db.query(NotarialDocumentBlock).count(), 500)
        finally:
            db.close()

    def test_queue_and_process_batch_is_recoverable_by_batch_id(self):
        db = self.Session()
        try:
            service = NotarialDocumentIngestionService(db, notary_id=1, storage=self.storage)
            queued = service.queue_batch(
                IngestBatchRequest(
                    name="Lote async",
                    source_type="api_async",
                    documents=[
                        DocumentUpload(filename="async-1.docx", content=build_docx_bytes(apartment="1201")),
                        DocumentUpload(filename="async-2.docx", content=build_docx_bytes(apartment="1202", include_table=False)),
                    ],
                )
            )

            self.assertEqual(queued.status, BatchStatus.QUEUED)
            self.assertEqual(queued.total_documents, 2)
            self.assertEqual(db.query(NotarialDocumentBlock).count(), 0)

            processed = service.process_queued_batch(queued.batch_id)

            self.assertEqual(processed.status, BatchStatus.COMPLETED)
            self.assertEqual(processed.error_documents, 0)
            self.assertEqual(db.query(NotarialDocumentBlock).count(), 6)
            statuses = [item.status for item in db.query(NotarialDocumentBatchItem).order_by(NotarialDocumentBatchItem.item_index).all()]
            self.assertEqual(statuses, ["processed", "processed"])
        finally:
            db.close()

    def test_same_hash_is_deduplicated_per_notary_without_cross_tenant_leak(self):
        content = build_docx_bytes(apartment="1301")
        db = self.Session()
        try:
            first = NotarialDocumentIngestionService(db, notary_id=1, storage=self.storage).ingest_batch(
                IngestBatchRequest(name="Notaria 1", documents=[DocumentUpload(filename="mismo.docx", content=content)])
            )
            second = NotarialDocumentIngestionService(db, notary_id=2, storage=self.storage).ingest_batch(
                IngestBatchRequest(name="Notaria 2", documents=[DocumentUpload(filename="mismo.docx", content=content)])
            )

            self.assertEqual(first.unique_documents, 1)
            self.assertEqual(second.unique_documents, 1)
            self.assertEqual(db.query(NotarialDocument).count(), 2)
            self.assertNotEqual(first.documents[0].document_id, second.documents[0].document_id)
            self.assertIn("notary_1", first.documents[0].storage_path)
            self.assertIn("notary_2", second.documents[0].storage_path)
        finally:
            db.close()

    def test_vector_store_filters_by_notary_and_orders_by_distance_in_sqlite_fallback(self):
        db = self.Session()
        try:
            version = NotarialEmbeddingVersion(version_key="test-3d", provider="test", model_name="mini", dimensions=3)
            db.add(version)
            db.flush()
            store = NotarialVectorStore(db, dimensions=3)
            db.add_all(
                [
                    NotarialDocumentEmbedding(
                        notary_id=1,
                        embedding_version_id=version.id,
                        source_type="block",
                        source_id=1,
                        content_hash="a" * 64,
                        embedding_dimensions=3,
                        embedding=store.encode_vector([1.0, 0.0, 0.0]),
                    ),
                    NotarialDocumentEmbedding(
                        notary_id=1,
                        embedding_version_id=version.id,
                        source_type="block",
                        source_id=2,
                        content_hash="b" * 64,
                        embedding_dimensions=3,
                        embedding=store.encode_vector([0.0, 1.0, 0.0]),
                    ),
                    NotarialDocumentEmbedding(
                        notary_id=2,
                        embedding_version_id=version.id,
                        source_type="block",
                        source_id=3,
                        content_hash="c" * 64,
                        embedding_dimensions=3,
                        embedding=store.encode_vector([1.0, 0.0, 0.0]),
                    ),
                ]
            )
            db.commit()

            results = store.search(1, [0.9, 0.1, 0.0], embedding_version_id=version.id, limit=10)

            self.assertEqual([item.source_id for item in results], [1, 2])
            self.assertTrue(results[0].distance < results[1].distance)
        finally:
            db.close()

    def test_vector_store_rejects_incompatible_dimensions(self):
        db = self.Session()
        try:
            store = NotarialVectorStore(db, dimensions=3)
            with self.assertRaises(ValueError):
                store.encode_vector([1.0, 0.0])
        finally:
            db.close()

    def test_intelligence_engine_persists_classification_embeddings_rag_and_audit(self):
        db = self.Session()
        try:
            result = NotarialDocumentIngestionService(db, notary_id=1, storage=self.storage).ingest_batch(
                IngestBatchRequest(
                    name="Motor inteligencia",
                    documents=[
                        DocumentUpload(
                            filename="motor.docx",
                            content=build_docx_bytes(
                                title="COMPRAVENTA INMUEBLE",
                                apartment="301",
                                extra_lines=[
                                    "NOTARIA 16 DE MEDELLIN",
                                    "VENDEDOR C.C. 1.234.567",
                                    "COMPRADOR C.C. 9.876.543",
                                    "BANCOLOMBIA aprueba credito hipotecario",
                                    "VALOR $ 120.000.000,00",
                                ],
                            ),
                        ),
                        DocumentUpload(filename="motor-source.docx", content=build_docx_bytes(title="COMPRAVENTA INMUEBLE", apartment="302")),
                    ],
                )
            )
            document_id = result.documents[0].document_id

            engine = NotarialIntelligenceEngine(db, embedding_provider=HashEmbeddingProvider(dimensions=8))
            engine.run_document(1, result.documents[1].document_id)
            run = engine.run_document(1, document_id)

            document = db.get(NotarialDocument, document_id)
            self.assertEqual(run.document_id, document_id)
            self.assertEqual(run.classification["document_type"], "escritura_compraventa")
            self.assertIn("hipoteca", run.classification["acts"])
            self.assertEqual(document.document_type, "escritura_compraventa")
            self.assertEqual(document.bank_name, "Bancolombia")
            self.assertGreater(db.query(NotarialDocumentEmbedding).filter(NotarialDocumentEmbedding.document_id == document_id).count(), 0)
            self.assertEqual(db.query(NotarialDocumentFamilyMember).filter(NotarialDocumentFamilyMember.document_id == document_id).count(), 1)
            self.assertGreater(db.query(NotarialDocumentEntity).filter(NotarialDocumentEntity.document_id == document_id).count(), 0)
            evidence_types = {
                row.evidence_type
                for row in db.query(NotarialDocumentEvidence).filter(NotarialDocumentEvidence.document_id == document_id).all()
            }
            self.assertIn("document_classification", evidence_types)
            self.assertIn("hybrid_rag", evidence_types)
            decision = db.query(NotarialDocumentDecision).filter(NotarialDocumentDecision.document_id == document_id).one()
            hybrid = json.loads(decision.hybrid_decision_json)
            self.assertEqual(hybrid["selected"]["classification"]["document_type"], "escritura_compraventa")
            self.assertEqual(run.llm_error, "llm_disabled")
            metadata = json.loads(document.metadata_json)
            self.assertEqual(metadata["intelligence"]["decision_id"], decision.id)
        finally:
            db.close()

    def test_intelligence_engine_classifies_fixed_variable_and_optional_family_blocks(self):
        db = self.Session()
        try:
            result = NotarialDocumentIngestionService(db, notary_id=1, storage=self.storage).ingest_batch(
                IngestBatchRequest(
                    name="Familia variable",
                    documents=[
                        DocumentUpload(filename="familia-a.docx", content=build_docx_bytes(apartment="301", include_table=True)),
                        DocumentUpload(filename="familia-b.docx", content=build_docx_bytes(apartment="302", include_table=False)),
                    ],
                )
            )
            ids = [item.document_id for item in result.documents]
            engine = NotarialIntelligenceEngine(db, embedding_provider=HashEmbeddingProvider(dimensions=8))
            first = engine.run_document(1, ids[0])
            second = engine.run_document(1, ids[1])

            self.assertEqual(first.family_id, second.family_id)
            self.assertGreater(second.fixed_variable_counts["fixed"], 0)
            self.assertGreater(second.fixed_variable_counts["variable"], 0)
            self.assertGreater(second.fixed_variable_counts["optional"], 0)
            labels = {
                row.text: row.fixed_variable_label
                for row in db.query(NotarialDocumentBlock)
                .join(NotarialDocumentParseRun, NotarialDocumentParseRun.id == NotarialDocumentBlock.parse_run_id)
                .filter(NotarialDocumentBlock.document_id == ids[0], NotarialDocumentParseRun.is_active.is_(True))
                .all()
            }
            self.assertEqual(labels["ESCRITURA DE COMPRAVENTA"], "fixed")
            self.assertEqual(labels["APARTAMENTO: 301"], "variable")
            self.assertEqual(labels["MATRICULA"], "optional")
        finally:
            db.close()

    def test_intelligence_engine_persists_typed_llm_decision_and_token_audit(self):
        db = self.Session()
        try:
            result = NotarialDocumentIngestionService(db, notary_id=1, storage=self.storage).ingest_batch(
                IngestBatchRequest(name="LLM auditado", documents=[DocumentUpload(filename="llm.docx", content=build_docx_bytes(apartment="401"))])
            )
            document_id = result.documents[0].document_id
            provider = _FakeLLMProvider(
                {
                    "document_type": "escritura_compraventa",
                    "document_subtype": None,
                    "acts": ["compraventa"],
                    "fields": [
                        {
                            "field_code": "numero_apartamento",
                            "label": "Apartamento",
                            "value": "401",
                            "confidence": 0.91,
                            "evidence_block_ids": [],
                            "fixed_variable_label": "variable",
                            "requires_human_review": False,
                            "reason": "Detectado en bloque de apartamento.",
                        }
                    ],
                    "conflicts": [],
                    "confidence": 0.93,
                    "_usage": {"input_tokens": 123, "output_tokens": 45},
                }
            )

            run = NotarialIntelligenceEngine(
                db,
                embedding_provider=HashEmbeddingProvider(dimensions=8),
                llm_provider=provider,
            ).run_document(1, document_id, llm_mode=IntelligenceMode.ASSIST)

            decision = db.query(NotarialDocumentDecision).filter(NotarialDocumentDecision.document_id == document_id).one()
            llm_payload = json.loads(decision.llm_decision_json)
            metadata = json.loads(decision.metadata_json)
            self.assertEqual(run.llm_mode, "assist")
            self.assertIsNone(run.llm_error)
            self.assertEqual(llm_payload["fields"][0]["field_code"], "numero_apartamento")
            self.assertEqual(metadata["llm_audit"]["input_tokens"], 123)
            self.assertEqual(metadata["llm_audit"]["output_tokens"], 45)
            self.assertEqual(provider.calls, 1)
        finally:
            db.close()

    def test_intelligence_run_redelivery_is_idempotent_without_duplicate_outputs(self):
        db = self.Session()
        try:
            result = NotarialDocumentIngestionService(db, notary_id=1, storage=self.storage).ingest_batch(
                IngestBatchRequest(
                    name="Redelivery inteligencia",
                    documents=[
                        DocumentUpload(filename="base-redelivery.docx", content=build_docx_bytes(apartment="501")),
                        DocumentUpload(filename="fuente-redelivery.docx", content=build_docx_bytes(apartment="502")),
                    ],
                )
            )
            engine = NotarialIntelligenceEngine(db, embedding_provider=HashEmbeddingProvider(dimensions=8))
            engine.run_document(1, result.documents[1].document_id)
            run = engine.create_document_run(1, result.documents[0].document_id)

            first = engine.run_intelligence_run(run.id)
            second = engine.run_intelligence_run(run.id)

            self.assertEqual(first["decision_id"], second["decision_id"])
            self.assertEqual(db.query(NotarialDocumentDecision).filter(NotarialDocumentDecision.document_id == result.documents[0].document_id).count(), 1)
            self.assertEqual(db.query(NotarialDocumentEvidence).filter(NotarialDocumentEvidence.document_id == result.documents[0].document_id).count(), first["rag_hits"] + 1)
            self.assertEqual(db.get(NotarialIntelligenceRun, run.id).status, "completed")
        finally:
            db.close()

    def test_create_document_run_recovers_concurrent_run_key_race(self):
        db_path = Path(self.tmp.name) / "intelligence-run-key-race.db"
        db_engine = create_engine(
            f"sqlite:///{db_path}",
            connect_args={"check_same_thread": False, "timeout": 30},
        )
        Base.metadata.create_all(db_engine, tables=TABLES)
        Session = sessionmaker(bind=db_engine)
        db = Session()
        try:
            db.add(_notary(1, "notaria-run-key"))
            db.commit()
            result = NotarialDocumentIngestionService(db, notary_id=1, storage=self.storage).ingest_batch(
                IngestBatchRequest(name="Create run race", documents=[DocumentUpload(filename="race.docx", content=build_docx_bytes(apartment="551"))])
            )
            document_id = result.documents[0].document_id
        finally:
            db.close()

        barrier = threading.Barrier(2)
        run_ids: list[int] = []
        errors: list[BaseException] = []

        def create_run():
            session = Session()
            try:
                barrier.wait(timeout=5)
                run = NotarialIntelligenceEngine(session, embedding_provider=HashEmbeddingProvider(dimensions=8)).create_document_run(1, document_id)
                session.commit()
                run_ids.append(run.id)
            except BaseException as exc:
                errors.append(exc)
            finally:
                session.close()

        threads = [threading.Thread(target=create_run) for _ in range(2)]
        for thread in threads:
            thread.start()
        for thread in threads:
            thread.join(timeout=10)

        self.assertEqual(errors, [])
        self.assertFalse(any(thread.is_alive() for thread in threads))
        self.assertEqual(len(set(run_ids)), 1)
        db = Session()
        try:
            self.assertEqual(db.query(NotarialIntelligenceRun).count(), 1)
        finally:
            db.close()
            db_engine.dispose()

    def test_intelligence_run_claim_is_atomic_across_sessions(self):
        db_path = Path(self.tmp.name) / "intelligence-claim-race.db"
        db_engine = create_engine(
            f"sqlite:///{db_path}",
            connect_args={"check_same_thread": False, "timeout": 30},
        )
        Base.metadata.create_all(db_engine, tables=TABLES)
        Session = sessionmaker(bind=db_engine)
        db = Session()
        try:
            db.add(_notary(1, "notaria-claim"))
            db.commit()
            result = NotarialDocumentIngestionService(db, notary_id=1, storage=self.storage).ingest_batch(
                IngestBatchRequest(
                    name="Claim atomico",
                    documents=[
                        DocumentUpload(filename="claim-target.docx", content=build_docx_bytes(apartment="561")),
                        DocumentUpload(filename="claim-source.docx", content=build_docx_bytes(apartment="562")),
                    ],
                )
            )
            intelligence_engine = NotarialIntelligenceEngine(db, embedding_provider=HashEmbeddingProvider(dimensions=8))
            intelligence_engine.run_document(1, result.documents[1].document_id)
            run = intelligence_engine.create_document_run(1, result.documents[0].document_id)
            db.commit()
            run_id = run.id
        finally:
            db.close()

        entered = threading.Event()
        release = threading.Event()
        results: list[dict] = []
        errors: list[BaseException] = []

        class SlowEngine(NotarialIntelligenceEngine):
            def run_document(self, *args, **kwargs):
                entered.set()
                release.wait(timeout=5)
                return super().run_document(*args, **kwargs)

        def first_worker():
            session = Session()
            try:
                results.append(SlowEngine(session, embedding_provider=HashEmbeddingProvider(dimensions=8)).run_intelligence_run(run_id))
            except BaseException as exc:
                errors.append(exc)
            finally:
                session.close()

        def second_worker():
            session = Session()
            try:
                entered.wait(timeout=5)
                results.append(NotarialIntelligenceEngine(session, embedding_provider=HashEmbeddingProvider(dimensions=8)).run_intelligence_run(run_id))
            except BaseException as exc:
                errors.append(exc)
            finally:
                session.close()

        first = threading.Thread(target=first_worker)
        first.start()
        entered.wait(timeout=5)
        second = threading.Thread(target=second_worker)
        second.start()
        second.join(timeout=10)
        release.set()
        first.join(timeout=10)

        self.assertEqual(errors, [])
        self.assertFalse(first.is_alive())
        self.assertFalse(second.is_alive())
        self.assertTrue(any(item.get("claimed") is False for item in results))
        self.assertTrue(any(item.get("decision_id") for item in results))
        db = Session()
        try:
            stored = db.get(NotarialIntelligenceRun, run_id)
            self.assertEqual(stored.status, "completed")
            self.assertEqual(stored.attempts, 1)
            self.assertEqual(db.query(NotarialDocumentDecision).filter(NotarialDocumentDecision.document_id == result.documents[0].document_id).count(), 1)
        finally:
            db.close()
            db_engine.dispose()

    def test_rag_persists_target_and_source_trace_without_self_retrieval(self):
        db = self.Session()
        try:
            result = NotarialDocumentIngestionService(db, notary_id=1, storage=self.storage).ingest_batch(
                IngestBatchRequest(
                    name="RAG trazable",
                    documents=[
                        DocumentUpload(filename="target.docx", content=build_docx_bytes(apartment="601")),
                        DocumentUpload(filename="source.docx", content=build_docx_bytes(apartment="602")),
                    ],
                )
            )
            target_id = result.documents[0].document_id
            source_id = result.documents[1].document_id
            engine = NotarialIntelligenceEngine(db, embedding_provider=HashEmbeddingProvider(dimensions=8))
            engine.run_document(1, source_id)
            target_run = engine.run_document(1, target_id)

            evidences = db.query(NotarialDocumentEvidence).filter(NotarialDocumentEvidence.document_id == target_id, NotarialDocumentEvidence.evidence_type == "hybrid_rag").all()
            self.assertGreater(len(evidences), 0)
            payloads = [json.loads(row.payload_json) for row in evidences]
            self.assertTrue(all(item["target_document_id"] == target_id for item in payloads))
            self.assertTrue(all(item["target_parse_run_id"] == target_run.parse_run_id for item in payloads))
            self.assertTrue(all(item["source_document_id"] != target_id for item in payloads))
            self.assertIn(source_id, {item["source_document_id"] for item in payloads})
        finally:
            db.close()

    def test_rag_vector_scores_are_limited_to_filtered_candidate_blocks(self):
        db = self.Session()
        try:
            result = NotarialDocumentIngestionService(db, notary_id=1, storage=self.storage).ingest_batch(
                IngestBatchRequest(
                    name="RAG filtrado",
                    documents=[
                        DocumentUpload(filename="rag-target.docx", content=build_docx_bytes(apartment="611")),
                        DocumentUpload(filename="rag-source.docx", content=build_docx_bytes(apartment="612")),
                        DocumentUpload(filename="rag-wrong-type.docx", content=build_docx_bytes(title="PROMESA DE COMPRAVENTA", apartment="613")),
                    ],
                )
            )
            target_id, source_id, wrong_id = [item.document_id for item in result.documents]
            provider = HashEmbeddingProvider(dimensions=8)
            engine = NotarialIntelligenceEngine(db, embedding_provider=provider)
            for item in result.documents:
                engine.embedding_service.reindex_document(1, item.document_id)
            db.get(NotarialDocument, target_id).document_type = "escritura_compraventa"
            db.get(NotarialDocument, source_id).document_type = "escritura_compraventa"
            db.get(NotarialDocument, wrong_id).document_type = "promesa_compraventa"
            query_vector = provider.encode(["escritura_compraventa compraventa"])[0]
            encoded_query = NotarialVectorStore(db, dimensions=8).encode_vector(query_vector)
            for row in db.query(NotarialDocumentEmbedding).filter(NotarialDocumentEmbedding.document_id.in_([target_id, wrong_id])).all():
                row.embedding = encoded_query
            source_vector = provider.encode(["escritura_compraventa compraventa fuente"])[0]
            encoded_source = NotarialVectorStore(db, dimensions=8).encode_vector(source_vector)
            for row in db.query(NotarialDocumentEmbedding).filter(NotarialDocumentEmbedding.document_id == source_id).all():
                row.embedding = encoded_source
            db.commit()

            hits = engine.rag.search(1, "escritura_compraventa compraventa", document_type="escritura_compraventa", exclude_document_id=target_id, limit=8)

            self.assertGreater(len(hits), 0)
            self.assertEqual({hit.document_id for hit in hits}, {source_id})
            self.assertTrue(all(hit.document_id != target_id for hit in hits))
            self.assertTrue(all(hit.vector_score > 0 for hit in hits))
        finally:
            db.close()

    def test_structural_alignment_tolerates_shifted_paragraphs_and_table_variants(self):
        db = self.Session()
        try:
            result = NotarialDocumentIngestionService(db, notary_id=1, storage=self.storage).ingest_batch(
                IngestBatchRequest(
                    name="Alineacion estructural",
                    documents=[
                        DocumentUpload(filename="alineacion-a.docx", content=build_docx_bytes(apartment="701", include_table=True)),
                        DocumentUpload(filename="alineacion-b.docx", content=build_docx_bytes(apartment="702", include_table=True, extra_lines=["CLAUSULA INSERTADA VARIABLE"])),
                    ],
                )
            )
            engine = NotarialIntelligenceEngine(db, embedding_provider=HashEmbeddingProvider(dimensions=8))
            engine.run_document(1, result.documents[0].document_id)
            run = engine.run_document(1, result.documents[1].document_id)

            self.assertGreater(db.query(NotarialBlockAlignment).filter(NotarialBlockAlignment.run_id == run.run_id).count(), 0)
            labels = {
                row.text: row.fixed_variable_label
                for row in db.query(NotarialDocumentBlock)
                .filter(NotarialDocumentBlock.document_id == result.documents[1].document_id)
                .all()
            }
            self.assertEqual(labels["ESCRITURA DE COMPRAVENTA"], "fixed")
            self.assertEqual(labels["APARTAMENTO: 702"], "variable")
        finally:
            db.close()

    def test_vector_candidate_clustering_consolidates_three_near_duplicate_documents(self):
        db = self.Session()
        try:
            result = NotarialDocumentIngestionService(db, notary_id=1, storage=self.storage).ingest_batch(
                IngestBatchRequest(
                    name="Cluster tres",
                    documents=[
                        DocumentUpload(filename="cluster-a.docx", content=build_docx_bytes(apartment="801")),
                        DocumentUpload(filename="cluster-b.docx", content=build_docx_bytes(apartment="802")),
                        DocumentUpload(filename="cluster-c.docx", content=build_docx_bytes(apartment="803")),
                    ],
                )
            )
            engine = NotarialIntelligenceEngine(db, embedding_provider=HashEmbeddingProvider(dimensions=8))
            for item in result.documents:
                engine.run_document(1, item.document_id)

            near_cluster = db.query(NotarialDocumentCluster).filter(NotarialDocumentCluster.cluster_type == "near_duplicate").order_by(NotarialDocumentCluster.id.desc()).first()
            self.assertIsNotNone(near_cluster)
            self.assertGreaterEqual(db.query(NotarialDocumentClusterMember).filter(NotarialDocumentClusterMember.cluster_id == near_cluster.id).count(), 3)
            self.assertIn("pgvector+structural-rerank", near_cluster.algorithm)
        finally:
            db.close()

    def test_vector_candidate_clustering_finds_similar_high_id_document(self):
        db = self.Session()
        try:
            documents = [DocumentUpload(filename="cluster-target.docx", content=build_docx_bytes(apartment="821"))]
            documents.extend(
                DocumentUpload(filename=f"cluster-distractor-{index}.docx", content=build_docx_bytes(title="HIPOTECA", apartment=str(830 + index)))
                for index in range(8)
            )
            documents.append(DocumentUpload(filename="cluster-similar-high-id.docx", content=build_docx_bytes(apartment="822")))
            result = NotarialDocumentIngestionService(db, notary_id=1, storage=self.storage).ingest_batch(IngestBatchRequest(name="Cluster high id", documents=documents))
            target_id = result.documents[0].document_id
            similar_id = result.documents[-1].document_id
            provider = HashEmbeddingProvider(dimensions=8)
            engine = NotarialIntelligenceEngine(db, embedding_provider=provider)
            for item in result.documents:
                engine.embedding_service.reindex_document(1, item.document_id)
            vector_store = NotarialVectorStore(db, dimensions=8)
            target_vector = vector_store.encode_vector([1.0, 0.0, 0.0, 0.0, 0.0, 0.0, 0.0, 0.0])
            similar_vector = vector_store.encode_vector([0.99, 0.01, 0.0, 0.0, 0.0, 0.0, 0.0, 0.0])
            distractor_vector = vector_store.encode_vector([0.0, 1.0, 0.0, 0.0, 0.0, 0.0, 0.0, 0.0])
            for row in db.query(NotarialDocumentEmbedding).all():
                if row.document_id == target_id:
                    row.embedding = target_vector
                elif row.document_id == similar_id:
                    row.embedding = similar_vector
                else:
                    row.embedding = distractor_vector
            db.commit()

            candidate_ids = engine.family_cluster._candidate_document_ids(1, target_id, limit=3)
            clusters = engine.family_cluster.update_duplicate_clusters(1, db.get(NotarialDocument, target_id), near_duplicate_threshold=0.75)

            self.assertIn(similar_id, candidate_ids)
            self.assertNotEqual(candidate_ids, sorted(candidate_ids))
            near_clusters = [cluster for cluster in clusters if cluster.cluster_type == "near_duplicate"]
            self.assertTrue(near_clusters)
            member_ids = {
                row.document_id
                for row in db.query(NotarialDocumentClusterMember).filter(NotarialDocumentClusterMember.cluster_id == near_clusters[-1].id).all()
            }
            self.assertIn(similar_id, member_ids)
        finally:
            db.close()

    def test_http_intelligence_run_uses_outbox_and_exposes_status(self):
        original_task = route_module.process_intelligence_run_task

        class FakeTask:
            @staticmethod
            def delay(run_id: int):
                return type("FakeAsyncResult", (), {"id": f"intel-{run_id}"})()

        route_module.process_intelligence_run_task = FakeTask()
        self.addCleanup(lambda: setattr(route_module, "process_intelligence_run_task", original_task))
        client = self._build_api_client()
        ingest = client.post(
            "/api/v1/notarial-intelligence/batches/ingest",
            data={"name": "HTTP inteligencia", "source_type": "test"},
            files=[("files", ("intel.docx", build_docx_bytes(apartment="901"), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"))],
        )
        document_id = ingest.json()["documents"][0]["document_id"]

        response = client.post(f"/api/v1/notarial-intelligence/documents/{document_id}/intelligence/run", json={"llm_mode": "off"})

        self.assertEqual(response.status_code, 202)
        payload = response.json()
        self.assertEqual(payload["task_id"], f"intel-{payload['run_id']}")
        detail = client.get(payload["status_url"])
        self.assertEqual(detail.status_code, 200)
        self.assertEqual(detail.json()["run_id"], payload["run_id"])
        db = self.Session()
        try:
            publication = db.query(NotarialTaskPublication).filter(NotarialTaskPublication.target_type == "intelligence_run").one()
            run = db.get(NotarialIntelligenceRun, payload["run_id"])
            self.assertEqual(publication.task_id, payload["task_id"])
            self.assertEqual(run.task_id, payload["task_id"])
        finally:
            db.close()

    def test_http_intelligence_publication_failure_reconciles_without_duplicate_run(self):
        original_task = route_module.process_intelligence_run_task

        class FailingTask:
            @staticmethod
            def delay(run_id: int):
                raise ConnectionError("broker offline")

        class SuccessfulTask:
            @staticmethod
            def delay(run_id: int):
                return type("FakeAsyncResult", (), {"id": f"retry-intel-{run_id}"})()

        route_module.process_intelligence_run_task = FailingTask()
        self.addCleanup(lambda: setattr(route_module, "process_intelligence_run_task", original_task))
        client = self._build_api_client()
        ingest = client.post(
            "/api/v1/notarial-intelligence/batches/ingest",
            data={"name": "HTTP intel fail", "source_type": "test"},
            files=[("files", ("intel-fail.docx", build_docx_bytes(apartment="902"), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"))],
        )
        document_id = ingest.json()["documents"][0]["document_id"]
        failed = client.post(f"/api/v1/notarial-intelligence/documents/{document_id}/intelligence/run", json={"llm_mode": "off"})
        self.assertEqual(failed.status_code, 503)

        route_module.process_intelligence_run_task = SuccessfulTask()
        reconciled = client.post("/api/v1/notarial-intelligence/publications/reconcile")

        self.assertEqual(reconciled.status_code, 200)
        db = self.Session()
        try:
            self.assertEqual(db.query(NotarialIntelligenceRun).count(), 1)
            run = db.query(NotarialIntelligenceRun).one()
            self.assertEqual(run.task_id, f"retry-intel-{run.id}")
            self.assertEqual(db.query(NotarialTaskPublication).filter(NotarialTaskPublication.target_type == "intelligence_run").count(), 1)
        finally:
            db.close()

    def test_benchmark_persists_versioned_metrics(self):
        db = self.Session()
        try:
            result = NotarialDocumentIngestionService(db, notary_id=1, storage=self.storage).ingest_batch(
                IngestBatchRequest(name="Benchmark", documents=[DocumentUpload(filename="bench.docx", content=build_docx_bytes(apartment="1001"))])
            )
            NotarialIntelligenceEngine(db, embedding_provider=HashEmbeddingProvider(dimensions=8)).run_document(1, result.documents[0].document_id)

            benchmark = NotarialIntelligenceEngine(db, embedding_provider=HashEmbeddingProvider(dimensions=8)).benchmark_notary(1)

            self.assertIn("metrics", benchmark)
            self.assertEqual(benchmark["metrics"]["status"], "insufficient_ground_truth")
            self.assertIn("classification", benchmark["metrics"])
            self.assertIn("retrieval", benchmark["metrics"])
            self.assertEqual(benchmark["metrics"]["classification"]["status"], "insufficient_ground_truth")
            self.assertEqual(db.query(NotarialBenchmarkRun).count(), 1)
            self.assertEqual(db.query(NotarialBenchmarkRun).one().status, "insufficient_ground_truth")
            self.assertTrue(benchmark["corpus_version"].startswith("corpus-"))
        finally:
            db.close()

    def test_benchmark_uses_ground_truth_for_quality_metrics(self):
        db = self.Session()
        try:
            result = NotarialDocumentIngestionService(db, notary_id=1, storage=self.storage).ingest_batch(
                IngestBatchRequest(
                    name="Benchmark ground truth",
                    documents=[
                        DocumentUpload(filename="bench-target.docx", content=build_docx_bytes(apartment="1001")),
                        DocumentUpload(filename="bench-source.docx", content=build_docx_bytes(apartment="1002")),
                    ],
                )
            )
            target_id = result.documents[0].document_id
            source_id = result.documents[1].document_id
            engine = NotarialIntelligenceEngine(db, embedding_provider=HashEmbeddingProvider(dimensions=8))
            engine.run_document(1, source_id)
            engine.run_document(1, target_id)
            target = db.get(NotarialDocument, target_id)
            source = db.get(NotarialDocument, source_id)
            target.metadata_json = json.dumps(
                {
                    **json.loads(target.metadata_json),
                    "labels": {
                        "document_type": "escritura_compraventa",
                        "acts": ["compraventa"],
                        "field_codes": ["numero_apartamento"],
                        "relevant_document_ids": [source_id],
                        "cluster_label": "familia-compraventa",
                        "fixed_variable": {"APARTAMENTO: 1001": "variable"},
                    },
                },
                ensure_ascii=False,
                sort_keys=True,
            )
            source.metadata_json = json.dumps(
                {
                    **json.loads(source.metadata_json),
                    "labels": {
                        "document_type": "escritura_compraventa",
                        "acts": ["compraventa"],
                        "cluster_label": "familia-compraventa",
                    },
                },
                ensure_ascii=False,
                sort_keys=True,
            )
            db.commit()

            benchmark = engine.benchmark_notary(1)
            metrics = benchmark["metrics"]

            self.assertEqual(metrics["status"], "completed")
            self.assertEqual(metrics["classification"]["tp"], 4)
            self.assertEqual(metrics["classification"]["fp"], 0)
            self.assertEqual(metrics["classification"]["fn"], 0)
            self.assertEqual(metrics["classification"]["precision"], 1.0)
            self.assertEqual(metrics["classification"]["recall"], 1.0)
            self.assertEqual(metrics["classification"]["f1"], 1.0)
            self.assertEqual(metrics["retrieval"]["recall_at_8"], 1.0)
            self.assertEqual(metrics["retrieval"]["mrr"], 1.0)
            self.assertEqual(metrics["clustering"]["tp"], 1)
            self.assertEqual(metrics["fields"]["status"], "completed")
            self.assertEqual(metrics["fixed_variable"]["status"], "completed")
            self.assertEqual(db.query(NotarialBenchmarkRun).order_by(NotarialBenchmarkRun.id.desc()).first().status, "completed")
        finally:
            db.close()

    def test_human_review_approves_template_without_catalog_autoapproval(self):
        db = self.Session()
        try:
            db.add_all(
                [
                    FieldDefinition(field_code="VALOR_VENTA", display_name="Valor venta", status="approved", confidence=1.0),
                    FieldDefinition(field_code="NUMERO_APARTAMENTO", display_name="Apartamento", status="approved", confidence=1.0),
                ]
            )
            result = NotarialDocumentIngestionService(db, notary_id=1, storage=self.storage).ingest_batch(
                IngestBatchRequest(
                    name="Revision humana",
                    documents=[
                        DocumentUpload(
                            filename="review.docx",
                            content=build_docx_bytes(
                                apartment="1201",
                                extra_lines=["VALOR $ 150.000.000", "VALOR $ 150.000.000"],
                            ),
                        )
                    ],
                )
            )
            engine = NotarialIntelligenceEngine(db, embedding_provider=HashEmbeddingProvider(dimensions=8))
            run = engine.run_document(1, result.documents[0].document_id)
            block = (
                db.query(NotarialDocumentBlock)
                .filter(NotarialDocumentBlock.document_id == run.document_id, NotarialDocumentBlock.text == "VALOR $ 150.000.000")
                .order_by(NotarialDocumentBlock.id.asc())
                .first()
            )
            db.add(
                NotarialDocumentEntity(
                    notary_id=1,
                    document_id=run.document_id,
                    parse_run_id=run.parse_run_id,
                    block_id=block.id,
                    entity_type="money",
                    canonical_field_code="valor_venta",
                    text="VALOR $ 150.000.000",
                    normalized_text="150000000",
                    confidence=0.92,
                    source="synthetic-test",
                    requires_human_review=True,
                )
            )
            db.commit()
            decision = db.get(NotarialDocumentDecision, run.decision_id)
            client = self._build_api_client(role_code="protocolist")

            created = client.post(
                "/api/v1/notarial-intelligence/human-review/sessions",
                json={"decision_id": decision.id, "idempotency_key": "review-create-1"},
            )
            self.assertEqual(created.status_code, 200)
            session = created.json()["session"]
            fields = created.json()["fields"]
            money_field = next(item for item in fields if item["field_code"] == "VALOR_VENTA")

            accepted = client.post(
                f"/api/v1/notarial-intelligence/human-review/sessions/{session['id']}/fields/{money_field['id']}",
                json={
                    "action": "correct",
                    "proposed_field_code": "VALOR_VENTA",
                    "proposed_value": "150000000",
                    "apply_scope": "all",
                    "fixed_variable_label": "variable",
                    "idempotency_key": "field-money-1",
                },
            )
            self.assertEqual(accepted.status_code, 200)

            new_field = client.post(
                f"/api/v1/notarial-intelligence/human-review/sessions/{session['id']}/fields/{money_field['id']}",
                json={
                    "action": "propose_new",
                    "proposed_field_code": "CAMPO_NUEVO_SUPERVISADO",
                    "proposed_value": "observacion",
                    "apply_scope": "single",
                    "idempotency_key": "field-new-1",
                },
            )
            self.assertEqual(new_field.status_code, 200)
            self.assertTrue(any(item["status"] == "proposed_catalog_review" and item["is_new_field_proposal"] for item in new_field.json()["fields"]))
            self.assertEqual(db.query(FieldDefinition).filter(FieldDefinition.field_code == "CAMPO_NUEVO_SUPERVISADO").count(), 0)

            for field in new_field.json()["fields"]:
                if field["status"] == "pending":
                    response = client.post(
                        f"/api/v1/notarial-intelligence/human-review/sessions/{session['id']}/fields/{field['id']}",
                        json={"action": "accept", "fixed_variable_label": "variable", "idempotency_key": f"field-accept-{field['id']}"},
                    )
                    self.assertEqual(response.status_code, 200)
            proposed_rows = [
                item for item in client.get(f"/api/v1/notarial-intelligence/human-review/sessions/{session['id']}").json()["fields"] if item["status"] == "proposed_catalog_review"
            ]
            for field in proposed_rows:
                response = client.post(
                    f"/api/v1/notarial-intelligence/human-review/sessions/{session['id']}/fields/{field['id']}",
                    json={"action": "reject", "idempotency_key": f"field-reject-{field['id']}"},
                )
                self.assertEqual(response.status_code, 200)

            approved = client.post(
                f"/api/v1/notarial-intelligence/human-review/sessions/{session['id']}/approve",
                json={"template_name": "Compraventa revisada", "template_kind": "individual", "idempotency_key": "approve-session-1"},
            )
            self.assertEqual(approved.status_code, 200)
            payload = approved.json()
            self.assertEqual(payload["session"]["status"], "approved")
            self.assertEqual(payload["library_item"]["status"], "approved")
            self.assertIn("{{VALOR_VENTA}}", payload["version"]["placeholder_map"])

            library = client.get("/api/v1/notarial-intelligence/template-library?act_code=compraventa")
            self.assertEqual(library.status_code, 200)
            self.assertEqual(len(library.json()["items"]), 1)

            docx_response = client.get(f"/api/v1/notarial-intelligence/template-library/versions/{payload['version']['id']}/docx")
            self.assertEqual(docx_response.status_code, 200)
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                tmp.write(docx_response.content)
                tmp_path = Path(tmp.name)
            try:
                doc = Document(str(tmp_path))
                red_markers = [
                    run.text
                    for paragraph in doc.paragraphs
                    for run in paragraph.runs
                    if run.text == "{{VALOR_VENTA}}" and run.font.color.rgb is not None
                ]
                self.assertTrue(red_markers)
            finally:
                tmp_path.unlink(missing_ok=True)

            rollback = client.post(
                f"/api/v1/notarial-intelligence/template-library/{payload['library_item']['id']}/rollback",
                json={"target_version_id": payload["version"]["id"], "idempotency_key": "rollback-1"},
            )
            self.assertEqual(rollback.status_code, 200)
            self.assertEqual(rollback.json()["version"]["rollback_of_version_id"], payload["version"]["id"])
            self.assertGreaterEqual(db.query(NotarialHumanReviewAudit).count(), 4)
        finally:
            db.close()

    def test_llm_service_retries_transient_provider_errors_only_until_success(self):
        provider = _FlakyLLMProvider()
        service = NotarialLLMService(provider=provider, timeout_seconds=1, max_retries=1)

        result = service.analyze({"document": {"id": 1}, "evidence": []}, IntelligenceMode.SHADOW)

        self.assertIsNone(result.audit.error)
        self.assertEqual(result.audit.retries, 1)
        self.assertEqual(provider.calls, 2)
        self.assertEqual(result.decision.document_type, "escritura_compraventa")

    def test_llm_service_does_not_retry_invalid_provider_json(self):
        provider = _FakeLLMProvider({"document_type": "escritura_compraventa", "confidence": 2.0})
        service = NotarialLLMService(provider=provider, timeout_seconds=1, max_retries=2)

        result = service.analyze({"document": {"id": 1}, "evidence": []}, IntelligenceMode.SHADOW)

        self.assertIsNone(result.decision)
        self.assertIn("ValidationError", result.audit.error)
        self.assertEqual(result.audit.retries, 0)
        self.assertEqual(provider.calls, 1)

    def test_llm_service_rejects_unknown_field_code_and_foreign_evidence_id(self):
        invalid_code = _FakeLLMProvider(
            {
                "document_type": "escritura_compraventa",
                "fields": [
                    {
                        "field_code": "codigo_inventado",
                        "label": "Inventado",
                        "confidence": 0.8,
                        "evidence_block_ids": [1],
                        "reason": "No permitido.",
                    }
                ],
                "confidence": 0.8,
            }
        )
        result = NotarialLLMService(provider=invalid_code, timeout_seconds=1, max_retries=2).analyze(
            {"document": {"id": 1}, "evidence": [{"block_id": 1}]},
            IntelligenceMode.ASSIST,
            allowed_field_codes={"numero_apartamento"},
            allowed_evidence_block_ids={1},
        )
        self.assertIn("invalid_field_code", result.audit.error)
        self.assertEqual(invalid_code.calls, 1)

        invalid_evidence = _FakeLLMProvider(
            {
                "document_type": "escritura_compraventa",
                "fields": [
                    {
                        "field_code": "numero_apartamento",
                        "label": "Apartamento",
                        "confidence": 0.8,
                        "evidence_block_ids": [999],
                        "reason": "Bloque ajeno.",
                    }
                ],
                "confidence": 0.8,
            }
        )
        evidence_result = NotarialLLMService(provider=invalid_evidence, timeout_seconds=1, max_retries=2).analyze(
            {"document": {"id": 1}, "evidence": [{"block_id": 1}]},
            IntelligenceMode.ASSIST,
            allowed_field_codes={"numero_apartamento"},
            allowed_evidence_block_ids={1},
        )
        self.assertIn("invalid_evidence_block_ids", evidence_result.audit.error)
        self.assertEqual(invalid_evidence.calls, 1)

    def test_llm_modes_have_distinct_hybrid_semantics(self):
        db = self.Session()
        try:
            result = NotarialDocumentIngestionService(db, notary_id=1, storage=self.storage).ingest_batch(
                IngestBatchRequest(name="Modos LLM", documents=[DocumentUpload(filename="modo.docx", content=build_docx_bytes(apartment="1101"))])
            )
            document_id = result.documents[0].document_id
            provider = _FakeLLMProvider(
                {
                    "document_type": "escritura_compraventa",
                    "fields": [],
                    "conflicts": [],
                    "confidence": 0.95,
                }
            )
            engine = NotarialIntelligenceEngine(db, embedding_provider=HashEmbeddingProvider(dimensions=8), llm_provider=provider)
            shadow = engine.run_document(1, document_id, llm_mode=IntelligenceMode.SHADOW)
            assist = engine.run_document(1, document_id, llm_mode=IntelligenceMode.ASSIST)
            gated = engine.run_document(1, document_id, llm_mode=IntelligenceMode.GATED)

            decisions = db.query(NotarialDocumentDecision).filter(NotarialDocumentDecision.document_id == document_id).order_by(NotarialDocumentDecision.id).all()
            sources = [json.loads(row.hybrid_decision_json)["selected_source"] for row in decisions[-3:]]
            self.assertEqual(sources, ["deterministic_shadow_llm", "deterministic_with_llm_assist", "llm_gated_with_deterministic_evidence"])
            self.assertEqual({shadow.llm_mode, assist.llm_mode, gated.llm_mode}, {"shadow", "assist", "gated"})
        finally:
            db.close()

    def test_shared_circuit_breaker_state_blocks_second_instance(self):
        first = RedisCircuitBreaker("unit-shared-breaker", failure_threshold=1, recovery_seconds=60)
        second = RedisCircuitBreaker("unit-shared-breaker", failure_threshold=1, recovery_seconds=60)
        first.record_failure()

        with self.assertRaises(CircuitOpenError):
            second.before_call()

        first.record_success()

    def test_reparse_keeps_previous_parse_run_and_switches_active_on_success(self):
        db = self.Session()
        try:
            service = NotarialDocumentIngestionService(db, notary_id=1, storage=self.storage)
            result = service.ingest_batch(
                IngestBatchRequest(name="Versionado", documents=[DocumentUpload(filename="version.docx", content=build_docx_bytes(apartment="1401"))])
            )
            document_id = result.documents[0].document_id
            Path(result.documents[0].storage_path).write_bytes(build_docx_bytes(apartment="1402", include_table=False))

            reparsed = service.reparse_document(document_id)

            self.assertEqual(reparsed.block_count, 2)
            runs = db.query(NotarialDocumentParseRun).filter(NotarialDocumentParseRun.document_id == document_id).order_by(NotarialDocumentParseRun.parse_version).all()
            self.assertEqual([run.parse_version for run in runs], [1, 2])
            self.assertEqual([run.is_active for run in runs], [False, True])
            self.assertEqual(db.query(NotarialDocumentBlock).filter(NotarialDocumentBlock.document_id == document_id).count(), 6)
        finally:
            db.close()

    def test_http_ingests_docx_and_returns_batch_detail(self):
        client = self._build_api_client()
        response = client.post(
            "/api/v1/notarial-intelligence/batches/ingest",
            data={"name": "Lote HTTP", "source_type": "test"},
            files=[
                (
                    "files",
                    (
                        "minuta-http.docx",
                        build_docx_bytes(apartment="901"),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    ),
                )
            ],
        )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload["total_documents"], 1)
        self.assertEqual(payload["unique_documents"], 1)
        self.assertEqual(payload["documents"][0]["block_count"], 4)

        detail = client.get(f"/api/v1/notarial-intelligence/batches/{payload['batch_id']}")
        self.assertEqual(detail.status_code, 200)
        self.assertEqual(detail.json()["documents"][0]["filename"], "minuta_http.docx")

    def test_http_async_ingest_returns_accepted_batch_contract(self):
        original_task = route_module.process_queued_document_batch_task

        class FakeTask:
            @staticmethod
            def delay(batch_id: int, notary_id: int):
                return type("FakeAsyncResult", (), {"id": f"task-{batch_id}-{notary_id}"})()

        route_module.process_queued_document_batch_task = FakeTask()
        self.addCleanup(lambda: setattr(route_module, "process_queued_document_batch_task", original_task))
        client = self._build_api_client()

        response = client.post(
            "/api/v1/notarial-intelligence/batches/ingest",
            data={"name": "Lote async HTTP", "source_type": "test", "async_processing": "true"},
            files=[("files", ("async-http.docx", build_docx_bytes(apartment="1901"), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"))],
        )

        self.assertEqual(response.status_code, 202)
        payload = response.json()
        self.assertEqual(payload["status"], "queued")
        self.assertEqual(payload["task_id"], f"task-{payload['batch_id']}-1")
        self.assertIn(f"/notarial-intelligence/batches/{payload['batch_id']}", payload["status_url"])
        detail = client.get(f"/api/v1/notarial-intelligence/batches/{payload['batch_id']}")
        self.assertEqual(detail.status_code, 200)
        self.assertEqual(detail.json()["metadata"]["celery"]["task_id"], payload["task_id"])

    def test_http_async_ingest_marks_publication_failed_when_celery_publish_fails(self):
        original_task = route_module.process_queued_document_batch_task

        class FailingTask:
            @staticmethod
            def delay(batch_id: int, notary_id: int):
                raise ConnectionError("broker offline")

        route_module.process_queued_document_batch_task = FailingTask()
        self.addCleanup(lambda: setattr(route_module, "process_queued_document_batch_task", original_task))
        client = self._build_api_client()

        response = client.post(
            "/api/v1/notarial-intelligence/batches/ingest",
            data={"name": "Lote async fallido", "source_type": "test", "async_processing": "true"},
            files=[("files", ("async-fail.docx", build_docx_bytes(apartment="1902"), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"))],
        )

        self.assertEqual(response.status_code, 503)
        db = self.Session()
        try:
            batch = db.query(NotarialDocumentBatch).one()
            publication = db.query(NotarialTaskPublication).one()
            self.assertEqual(batch.status, BatchStatus.PUBLICATION_FAILED.value)
            self.assertIn("publication_failed", batch.error_message)
            self.assertEqual(publication.status, "publication_failed")
            self.assertEqual(publication.attempts, 1)
            self.assertIn("broker offline", publication.last_error)
        finally:
            db.close()

    def test_http_reconcile_republishes_failed_batch_without_duplicates(self):
        original_task = route_module.process_queued_document_batch_task

        class FailingTask:
            @staticmethod
            def delay(batch_id: int, notary_id: int):
                raise ConnectionError("broker offline")

        class SuccessfulTask:
            @staticmethod
            def delay(batch_id: int, notary_id: int):
                return type("FakeAsyncResult", (), {"id": f"retry-{batch_id}-{notary_id}"})()

        route_module.process_queued_document_batch_task = FailingTask()
        self.addCleanup(lambda: setattr(route_module, "process_queued_document_batch_task", original_task))
        client = self._build_api_client()

        response = client.post(
            "/api/v1/notarial-intelligence/batches/ingest",
            data={"name": "Reconcile batch", "source_type": "test", "async_processing": "true"},
            files=[("files", ("reconcile.docx", build_docx_bytes(apartment="1903"), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"))],
        )
        self.assertEqual(response.status_code, 503)

        route_module.process_queued_document_batch_task = SuccessfulTask()
        reconciled = client.post("/api/v1/notarial-intelligence/publications/reconcile")

        self.assertEqual(reconciled.status_code, 200)
        self.assertEqual(reconciled.json()["dispatched"], 1)
        db = self.Session()
        try:
            publication = db.query(NotarialTaskPublication).one()
            self.assertEqual(publication.status, "published")
            self.assertEqual(publication.attempts, 2)
            self.assertEqual(publication.task_id, "retry-1-1")
            self.assertEqual(db.query(NotarialDocument).count(), 1)
            self.assertEqual(db.query(NotarialDocumentBatchItem).count(), 1)
        finally:
            db.close()

    def test_reconcile_recovers_publishing_window_without_duplicating_records(self):
        db = self.Session()
        try:
            service = NotarialDocumentIngestionService(db, notary_id=1, storage=self.storage)
            queued = service.queue_batch(
                IngestBatchRequest(name="Ventana publishing", documents=[DocumentUpload(filename="window.docx", content=build_docx_bytes(apartment="1904"))])
            )
            publication = service.get_publication_for_target("batch", queued.batch_id)
            publication.status = "publishing"
            publication.attempts = 1
            db.commit()

            def fake_delay(batch_id: int, notary_id: int):
                return type("FakeAsyncResult", (), {"id": f"recovered-{batch_id}-{notary_id}"})()

            results = service.dispatch_recoverable_publications(
                {"notarial_document_intelligence.document.process_queued_batch": fake_delay}
            )

            self.assertEqual(results[0]["status"], "published")
            db.refresh(publication)
            self.assertEqual(publication.task_id, "recovered-1-1")
            self.assertEqual(publication.attempts, 2)
            self.assertEqual(db.query(NotarialDocumentBatch).count(), 1)
            self.assertEqual(db.query(NotarialDocument).count(), 1)
            self.assertEqual(db.query(NotarialDocumentBatchItem).count(), 1)
        finally:
            db.close()

    def test_http_ingests_zip_without_persisting_zip(self):
        client = self._build_api_client()
        zip_buffer = BytesIO()
        with zipfile.ZipFile(zip_buffer, "w") as archive:
            archive.writestr("lote/minuta-1.docx", build_docx_bytes(apartment="1001"))
            archive.writestr("lote/minuta-2.docx", build_docx_bytes(apartment="1002"))

        response = client.post(
            "/api/v1/notarial-intelligence/batches/ingest",
            data={"name": "Lote ZIP", "source_type": "zip"},
            files=[("files", ("lote.zip", zip_buffer.getvalue(), "application/zip"))],
        )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload["total_documents"], 2)
        self.assertEqual(payload["unique_documents"], 2)
        self.assertEqual(payload["error_documents"], 0)
        stored_files = list((Path(self.tmp.name) / "storage").rglob("*"))
        self.assertFalse(any(path.name == "lote.zip" for path in stored_files))

    def test_http_rejects_unsupported_zip_entry(self):
        client = self._build_api_client()
        zip_buffer = BytesIO()
        with zipfile.ZipFile(zip_buffer, "w") as archive:
            archive.writestr("lote/minuta.docx", build_docx_bytes(apartment="1003"))
            archive.writestr("lote/ignorar.txt", b"no docx")

        response = client.post(
            "/api/v1/notarial-intelligence/batches/ingest",
            data={"name": "Lote ZIP inseguro", "source_type": "zip"},
            files=[("files", ("lote.zip", zip_buffer.getvalue(), "application/zip"))],
        )

        self.assertEqual(response.status_code, 422)
        self.assertIn("Entrada no soportada", response.json()["detail"])

    def test_http_rejects_unsafe_zip_path(self):
        client = self._build_api_client()
        zip_buffer = BytesIO()
        with zipfile.ZipFile(zip_buffer, "w") as archive:
            archive.writestr("../minuta.docx", build_docx_bytes(apartment="1004"))

        response = client.post(
            "/api/v1/notarial-intelligence/batches/ingest",
            data={"name": "Lote ZIP ruta insegura", "source_type": "zip"},
            files=[("files", ("lote.zip", zip_buffer.getvalue(), "application/zip"))],
        )

        self.assertEqual(response.status_code, 422)
        self.assertIn("Ruta insegura", response.json()["detail"])

    def test_http_rejects_zip_bomb_ratio(self):
        client = self._build_api_client()
        zip_buffer = BytesIO()
        with zipfile.ZipFile(zip_buffer, "w", compression=zipfile.ZIP_DEFLATED) as archive:
            archive.writestr("lote/minuta.docx", b"0" * 1024 * 1024)

        response = client.post(
            "/api/v1/notarial-intelligence/batches/ingest",
            data={"name": "Lote ZIP bomb", "source_type": "zip"},
            files=[("files", ("lote.zip", zip_buffer.getvalue(), "application/zip"))],
        )

        self.assertEqual(response.status_code, 413)
        self.assertIn("ZIP bomb", response.json()["detail"])

    def test_http_rejects_invalid_docx(self):
        client = self._build_api_client()
        response = client.post(
            "/api/v1/notarial-intelligence/batches/ingest",
            data={"name": "DOCX invalido", "source_type": "test"},
            files=[("files", ("minuta.docx", b"no es docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"))],
        )

        self.assertEqual(response.status_code, 422)
        self.assertIn("DOCX invalido", response.json()["detail"])

    def test_http_reparse_document_returns_accepted_and_worker_rebuilds_blocks_from_storage(self):
        original_task = route_module.reparse_document_task

        class FakeTask:
            @staticmethod
            def delay(document_id: int, notary_id: int, parse_run_id: int):
                return type("FakeAsyncResult", (), {"id": f"reparse-{document_id}-{parse_run_id}"})()

        route_module.reparse_document_task = FakeTask()
        self.addCleanup(lambda: setattr(route_module, "reparse_document_task", original_task))
        client = self._build_api_client()
        response = client.post(
            "/api/v1/notarial-intelligence/batches/ingest",
            data={"name": "Lote reparse", "source_type": "test"},
            files=[
                (
                    "files",
                    (
                        "minuta-reparse.docx",
                        build_docx_bytes(apartment="1101"),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    ),
                )
            ],
        )
        self.assertEqual(response.status_code, 200)
        document = response.json()["documents"][0]
        Path(document["storage_path"]).write_bytes(build_docx_bytes(apartment="2202", include_table=False))

        reparse = client.post(f"/api/v1/notarial-intelligence/documents/{document['document_id']}/reparse")

        self.assertEqual(reparse.status_code, 202)
        accepted = reparse.json()
        self.assertEqual(accepted["status"], "queued")
        self.assertEqual(accepted["task_id"], f"reparse-{document['document_id']}-{accepted['parse_run_id']}")
        db = self.Session()
        try:
            service = NotarialDocumentIngestionService(db, notary_id=1, storage=self.storage)
            result = service.reparse_document(document["document_id"], parse_run_id=accepted["parse_run_id"])
            self.assertEqual(result.block_count, 2)
            texts = [
                row.text
                for row in db.query(NotarialDocumentBlock)
                .join(NotarialDocumentParseRun, NotarialDocumentParseRun.id == NotarialDocumentBlock.parse_run_id)
                .filter(
                    NotarialDocumentBlock.document_id == document["document_id"],
                    NotarialDocumentParseRun.is_active.is_(True),
                )
                .all()
            ]
            self.assertIn("APARTAMENTO: 2202", texts)
            self.assertNotIn("APARTAMENTO: 1101", texts)
        finally:
            db.close()

    def test_http_reparse_requires_role_permission(self):
        client = self._build_api_client(role_code="client")
        response = client.post(
            "/api/v1/notarial-intelligence/batches/ingest",
            data={"name": "Sin permiso", "source_type": "test"},
            files=[("files", ("sin-permiso.docx", build_docx_bytes(apartment="1701"), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"))],
        )

        self.assertEqual(response.status_code, 403)

    def test_http_requires_permission_in_active_notary_scope(self):
        client = self._build_api_client(notary_id=1, role_code="protocolist", assignment_notary_id=2)
        response = client.post(
            "/api/v1/notarial-intelligence/batches/ingest",
            data={"name": "Rol en otra notaria", "source_type": "test"},
            files=[("files", ("otra-notaria.docx", build_docx_bytes(apartment="1702"), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"))],
        )

        self.assertEqual(response.status_code, 403)

    def test_http_blocks_cross_notary_batch_and_reparse_access(self):
        client_1 = self._build_api_client(notary_id=1)
        response = client_1.post(
            "/api/v1/notarial-intelligence/batches/ingest",
            data={"name": "Privado", "source_type": "test"},
            files=[("files", ("privado.docx", build_docx_bytes(apartment="1501"), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"))],
        )
        self.assertEqual(response.status_code, 200)
        payload = response.json()

        client_2 = self._build_api_client(notary_id=2)
        detail = client_2.get(f"/api/v1/notarial-intelligence/batches/{payload['batch_id']}")
        reparse = client_2.post(f"/api/v1/notarial-intelligence/documents/{payload['documents'][0]['document_id']}/reparse")

        self.assertEqual(detail.status_code, 404)
        self.assertEqual(reparse.status_code, 404)

    def test_failed_parse_run_from_worker_is_persisted_after_batch_rollback(self):
        class FailingParser(NotarialDocumentParser):
            def parse_bytes(self, content: bytes, filename: str) -> ParsedDocument:
                raise ValueError("documento deterministico invalido")

        db = self.Session()
        try:
            service = NotarialDocumentIngestionService(db, notary_id=1, storage=self.storage)
            queued = service.queue_batch(
                IngestBatchRequest(name="Fallo worker", documents=[DocumentUpload(filename="fallo.docx", content=build_docx_bytes(apartment="1801"))])
            )
            failing_service = NotarialDocumentIngestionService(db, notary_id=1, storage=self.storage, parser=FailingParser())
            result = failing_service.process_queued_batch(queued.batch_id)

            self.assertEqual(result.status, BatchStatus.ERROR)
            run = db.query(NotarialDocumentParseRun).one()
            document = db.query(NotarialDocument).one()
            self.assertEqual(run.status, "error")
            self.assertIn("documento deterministico invalido", run.error_message)
            self.assertEqual(document.processing_status, DocumentProcessingStatus.ERROR.value)
            self.assertIn("last_reparse_error", document.metadata_json)
        finally:
            db.close()

    def test_failed_reparse_preserves_existing_active_parse_and_document_status(self):
        class FailingParser(NotarialDocumentParser):
            def parse_bytes(self, content: bytes, filename: str) -> ParsedDocument:
                raise ValueError("nuevo reparse invalido")

        db = self.Session()
        try:
            service = NotarialDocumentIngestionService(db, notary_id=1, storage=self.storage)
            result = service.ingest_batch(
                IngestBatchRequest(name="Activo previo", documents=[DocumentUpload(filename="activo.docx", content=build_docx_bytes(apartment="1851"))])
            )
            document_id = result.documents[0].document_id
            active_before = db.query(NotarialDocumentParseRun).filter(NotarialDocumentParseRun.document_id == document_id, NotarialDocumentParseRun.is_active.is_(True)).one()
            failing_service = NotarialDocumentIngestionService(db, notary_id=1, storage=self.storage, parser=FailingParser())

            with self.assertRaises(Exception):
                failing_service.reparse_document(document_id)

            document = db.get(NotarialDocument, document_id)
            active_after = db.query(NotarialDocumentParseRun).filter(NotarialDocumentParseRun.document_id == document_id, NotarialDocumentParseRun.is_active.is_(True)).one()
            failed_run = db.query(NotarialDocumentParseRun).filter(NotarialDocumentParseRun.document_id == document_id, NotarialDocumentParseRun.status == "error").one()
            self.assertEqual(document.processing_status, DocumentProcessingStatus.PARSED.value)
            self.assertEqual(active_after.id, active_before.id)
            self.assertFalse(failed_run.is_active)
            self.assertIn("nuevo reparse invalido", failed_run.error_message)
            self.assertIn("last_reparse_failed_at", document.metadata_json)
        finally:
            db.close()

    def test_transient_parse_error_is_persisted_and_bubbles_for_celery_retry(self):
        class TransientParser(NotarialDocumentParser):
            def parse_bytes(self, content: bytes, filename: str) -> ParsedDocument:
                raise TimeoutError("storage temporalmente no disponible")

        db = self.Session()
        try:
            service = NotarialDocumentIngestionService(db, notary_id=1, storage=self.storage)
            queued = service.queue_batch(
                IngestBatchRequest(name="Retry worker", documents=[DocumentUpload(filename="retry.docx", content=build_docx_bytes(apartment="1802"))])
            )
            retrying_service = NotarialDocumentIngestionService(db, notary_id=1, storage=self.storage, parser=TransientParser())

            with self.assertRaises(TransientDocumentProcessingError):
                retrying_service.process_queued_batch(queued.batch_id)

            item = db.query(NotarialDocumentBatchItem).one()
            run = db.query(NotarialDocumentParseRun).one()
            batch = db.query(NotarialDocumentBatch).one()
            self.assertEqual(item.status, "queued")
            self.assertEqual(run.status, "error")
            self.assertIn("storage temporalmente no disponible", run.error_message)
            self.assertEqual(batch.status, BatchStatus.RUNNING.value)
            self.assertIsNone(batch.finished_at)
        finally:
            db.close()

    def test_http_reparse_publish_failure_persists_failed_parse_run(self):
        original_task = route_module.reparse_document_task

        class FailingTask:
            @staticmethod
            def delay(document_id: int, notary_id: int, parse_run_id: int):
                raise ConnectionError("broker offline")

        route_module.reparse_document_task = FailingTask()
        self.addCleanup(lambda: setattr(route_module, "reparse_document_task", original_task))
        client = self._build_api_client()
        response = client.post(
            "/api/v1/notarial-intelligence/batches/ingest",
            data={"name": "Reparse publish fail", "source_type": "test"},
            files=[("files", ("publish-fail.docx", build_docx_bytes(apartment="1850"), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"))],
        )
        self.assertEqual(response.status_code, 200)
        document_id = response.json()["documents"][0]["document_id"]

        reparse = client.post(f"/api/v1/notarial-intelligence/documents/{document_id}/reparse")

        self.assertEqual(reparse.status_code, 503)
        db = self.Session()
        try:
            failed_run = db.query(NotarialDocumentParseRun).order_by(NotarialDocumentParseRun.parse_version.desc()).first()
            document = db.get(NotarialDocument, document_id)
            self.assertEqual(failed_run.status, "error")
            self.assertIn("publication_failed", failed_run.error_message)
            self.assertEqual(document.processing_status, DocumentProcessingStatus.PARSED.value)
            self.assertIn("last_reparse_error", document.metadata_json)
        finally:
            db.close()

    def test_supabase_path_upload_uses_file_stream_helper(self):
        original_upload = storage_module._upload_path_to_supabase
        captured = {}

        def fake_upload(bucket, object_key, source_path, content_type):
            captured["bucket"] = bucket
            captured["object_key"] = object_key
            captured["source_path"] = Path(source_path)
            captured["content_type"] = content_type

        storage_module._upload_path_to_supabase = fake_upload
        self.addCleanup(lambda: setattr(storage_module, "_upload_path_to_supabase", original_upload))
        source = Path(self.tmp.name) / "supabase-stream.docx"
        source.write_bytes(build_docx_bytes(apartment="1852"))
        storage = NotarialIntelligenceStorage(local_root=Path(self.tmp.name) / "unused", use_supabase=True)

        result = storage.store_document_path(1, "supabase-stream.docx", source, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        self.assertEqual(result.storage_backend, "supabase")
        self.assertEqual(captured["source_path"], source)
        self.assertIn("/notary_1/documents/", captured["object_key"])

    def test_concurrent_reparse_queue_assigns_unique_parse_versions(self):
        content = build_docx_bytes(apartment="1651")
        db_path = Path(self.tmp.name) / "parse-version-concurrency.db"
        engine = create_engine(
            f"sqlite:///{db_path}",
            connect_args={"check_same_thread": False, "timeout": 30},
        )
        Base.metadata.create_all(engine, tables=TABLES)
        Session = sessionmaker(bind=engine)
        seed = Session()
        try:
            seed.add(_notary(1, "notaria-versiones"))
            result = NotarialDocumentIngestionService(seed, notary_id=1, storage=self.storage).ingest_batch(
                IngestBatchRequest(name="Base", documents=[DocumentUpload(filename="base.docx", content=content)])
            )
            document_id = result.documents[0].document_id
        finally:
            seed.close()

        barrier = threading.Barrier(2)
        errors: list[BaseException] = []

        def queue_reparse() -> None:
            session = Session()
            try:
                barrier.wait(timeout=5)
                NotarialDocumentIngestionService(session, notary_id=1, storage=self.storage).queue_reparse_document(document_id)
            except BaseException as exc:
                errors.append(exc)
            finally:
                session.close()

        threads = [threading.Thread(target=queue_reparse) for _ in range(2)]
        for thread in threads:
            thread.start()
        for thread in threads:
            thread.join(timeout=10)

        self.assertEqual(errors, [])
        db = Session()
        try:
            versions = [
                row.parse_version
                for row in db.query(NotarialDocumentParseRun)
                .filter(NotarialDocumentParseRun.document_id == document_id)
                .order_by(NotarialDocumentParseRun.parse_version)
                .all()
            ]
            self.assertEqual(versions, [1, 2, 3])
        finally:
            db.close()
            engine.dispose()

    def test_http_rejects_zip_global_uncompressed_limit_across_multiple_zip_files(self):
        original_limit = route_module.MAX_ZIP_UNCOMPRESSED_BYTES
        docx_bytes = build_docx_bytes(apartment="2001")
        route_module.MAX_ZIP_UNCOMPRESSED_BYTES = len(docx_bytes) + 1
        self.addCleanup(lambda: setattr(route_module, "MAX_ZIP_UNCOMPRESSED_BYTES", original_limit))
        client = self._build_api_client()
        files = []
        for index in range(2):
            zip_buffer = BytesIO()
            with zipfile.ZipFile(zip_buffer, "w") as archive:
                archive.writestr(f"lote/minuta-{index}.docx", docx_bytes)
            files.append(("files", (f"lote-{index}.zip", zip_buffer.getvalue(), "application/zip")))

        response = client.post(
            "/api/v1/notarial-intelligence/batches/ingest",
            data={"name": "Lote ZIP global", "source_type": "zip"},
            files=files,
        )

        self.assertEqual(response.status_code, 413)
        self.assertIn("limite global", response.json()["detail"])

    def test_http_rejects_zip_entry_larger_than_single_file_limit(self):
        original_file_limit = route_module.MAX_FILE_BYTES
        original_ratio = route_module.MAX_ZIP_COMPRESSION_RATIO
        route_module.MAX_FILE_BYTES = 250
        route_module.MAX_ZIP_COMPRESSION_RATIO = 10_000
        self.addCleanup(lambda: setattr(route_module, "MAX_FILE_BYTES", original_file_limit))
        self.addCleanup(lambda: setattr(route_module, "MAX_ZIP_COMPRESSION_RATIO", original_ratio))
        client = self._build_api_client()
        zip_buffer = BytesIO()
        with zipfile.ZipFile(zip_buffer, "w", compression=zipfile.ZIP_DEFLATED) as archive:
            archive.writestr("lote/grande.docx", b"0" * 300)

        response = client.post(
            "/api/v1/notarial-intelligence/batches/ingest",
            data={"name": "Entrada grande", "source_type": "zip"},
            files=[("files", ("entrada-grande.zip", zip_buffer.getvalue(), "application/zip"))],
        )

        self.assertEqual(response.status_code, 413)
        self.assertIn("Entrada DOCX demasiado grande", response.json()["detail"])

    def test_postgresql_vector_search_uses_cosine_operator(self):
        captured = {}

        class FakeDialect:
            name = "postgresql"

        class FakeBind:
            dialect = FakeDialect()

        class FakeResult:
            def all(self):
                return [SimpleNamespace(id=1, source_type="block", source_id=7, distance=0.125)]

        class FakeSession:
            bind = FakeBind()

            def execute(self, statement, params):
                captured["sql"] = str(statement)
                captured["params"] = params
                return FakeResult()

        results = NotarialVectorStore(FakeSession(), dimensions=3).search(
            1,
            [1.0, 0.0, 0.0],
            embedding_version_id=9,
            limit=5,
        )

        self.assertEqual(results[0].distance, 0.125)
        self.assertIn("<=>", captured["sql"])
        self.assertNotIn("<->", captured["sql"])

    def test_concurrent_same_hash_creates_one_document_and_two_batch_items(self):
        content = build_docx_bytes(apartment="1601")
        db_path = Path(self.tmp.name) / "concurrency.db"
        engine = create_engine(
            f"sqlite:///{db_path}",
            connect_args={"check_same_thread": False, "timeout": 30},
        )
        Base.metadata.create_all(engine, tables=TABLES)
        Session = sessionmaker(bind=engine)
        seed = Session()
        try:
            seed.add_all([_notary(1, "notaria-uno-concurrente"), _notary(2, "notaria-dos-concurrente")])
            seed.commit()
        finally:
            seed.close()

        barrier = threading.Barrier(2)
        errors: list[BaseException] = []

        def ingest_one(name: str) -> None:
            session = Session()
            try:
                barrier.wait(timeout=5)
                NotarialDocumentIngestionService(session, notary_id=1, storage=self.storage).queue_batch(
                    IngestBatchRequest(name=name, documents=[DocumentUpload(filename=f"{name}.docx", content=content)])
                )
            except BaseException as exc:
                errors.append(exc)
            finally:
                session.close()

        threads = [threading.Thread(target=ingest_one, args=(f"concurrente-{index}",)) for index in range(2)]
        for thread in threads:
            thread.start()
        for thread in threads:
            thread.join(timeout=10)

        self.assertEqual(errors, [])
        self.assertFalse(any(thread.is_alive() for thread in threads))
        db = Session()
        try:
            self.assertEqual(db.query(NotarialDocument).filter(NotarialDocument.notary_id == 1).count(), 1)
            self.assertEqual(db.query(NotarialDocumentBatchItem).filter(NotarialDocumentBatchItem.notary_id == 1).count(), 2)
        finally:
            db.close()
            engine.dispose()

    def _build_api_client(self, notary_id: int | None = None, role_code: str = "protocolist", assignment_notary_id: int | None = None) -> TestClient:
        active_notary_id = notary_id or self.notary_id
        app = FastAPI()
        app.include_router(intelligence_router, prefix="/api/v1")

        def override_get_db():
            session = self.Session()
            try:
                yield session
            finally:
                session.close()

        def override_service():
            session = self.Session()
            try:
                yield NotarialDocumentIngestionService(session, notary_id=active_notary_id, storage=self.storage)
            finally:
                session.close()

        app.dependency_overrides[get_db] = override_get_db
        app.dependency_overrides[get_ingestion_service] = override_service
        app.dependency_overrides[get_current_user] = lambda: _user(active_notary_id, role_code, assignment_notary_id)
        client = TestClient(app)
        self.addCleanup(client.close)
        return client


class _FakeLLMProvider:
    provider_name = "fake-llm"
    model_name = "typed-json-fixture"
    prompt_version = "notarial-intelligence-test-v1"

    def __init__(self, response: dict):
        self.response = response
        self.calls = 0

    def complete_json(self, system_prompt: str, payload: dict, *, timeout_seconds: int) -> dict:
        self.calls += 1
        return dict(self.response)


class _FlakyLLMProvider(_FakeLLMProvider):
    def __init__(self):
        super().__init__(
            {
                "document_type": "escritura_compraventa",
                "document_subtype": None,
                "acts": ["compraventa"],
                "fields": [],
                "conflicts": [],
                "confidence": 0.88,
            }
        )
        self.model_name = "flaky-json-fixture"

    def complete_json(self, system_prompt: str, payload: dict, *, timeout_seconds: int) -> dict:
        self.calls += 1
        if self.calls == 1:
            raise ConnectionError("broker llm temporal")
        return dict(self.response)


def _user(notary_id: int, role_code: str = "protocolist", assignment_notary_id: int | None = None):
    role = SimpleNamespace(code=role_code)
    assignment = SimpleNamespace(role=role, notary_id=assignment_notary_id if assignment_notary_id is not None else notary_id)
    return type("UserStub", (), {"default_notary_id": notary_id, "role_assignments": [assignment]})()


def _notary(notary_id: int, slug: str) -> Notary:
    return Notary(
        id=notary_id,
        slug=slug,
        catalog_identity_key=f"{slug}-identity",
        commercial_name=slug,
        legal_name=slug,
        municipality="Medellin",
        notary_label=slug,
        city="Medellin",
    )


if __name__ == "__main__":
    unittest.main()
