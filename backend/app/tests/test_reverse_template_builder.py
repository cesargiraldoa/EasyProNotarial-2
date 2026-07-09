from __future__ import annotations

import unittest
from io import BytesIO

from docx import Document
from fastapi import FastAPI
from fastapi.testclient import TestClient

from app.modules.template_builder.router import router as template_builder_router


class ReverseTemplateBuilderTests(unittest.TestCase):
    def _build_docx(self) -> bytes:
        buffer = BytesIO()
        doc = Document()
        doc.add_paragraph(
            "EL COMPRADOR JUAN CAMILO VASQUEZ MIRA, identificado con C.C. 1.037.657.164, "
            "celular 3001234567 y correo juan@example.com."
        )
        doc.add_paragraph(
            "LA VENDEDORA MARIA ELENA GOMEZ RESTREPO transfiere el APARTAMENTO 804 con "
            "matricula inmobiliaria 001-1528731 por precio de $212.600.000."
        )
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "NIT vendedor"
        table.cell(0, 1).text = "900.123.456-7"
        doc.save(buffer)
        return buffer.getvalue()

    def test_endpoint_analyzes_single_docx_and_groups_candidates(self):
        app = FastAPI()
        app.include_router(template_builder_router, prefix="/api/v1")
        client = TestClient(app)

        response = client.post(
            "/api/v1/template-builder/reverse/single/analyze",
            files={
                "archivo": (
                    "minuta.docx",
                    self._build_docx(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload["mode"], "single")
        self.assertEqual(payload["filename"], "minuta.docx")
        self.assertGreaterEqual(payload["summary"]["total_candidates"], 7)

        by_type = {candidate["type"]: candidate for candidate in payload["candidates"]}
        self.assertEqual(by_type["matricula_inmobiliaria"]["text"], "001-1528731")
        self.assertEqual(by_type["money"]["text"], "$212.600.000")
        self.assertEqual(by_type["email"]["text"], "juan@example.com")
        self.assertEqual(by_type["property_unit"]["text"], "804")
        self.assertTrue(any(candidate["text"] == "1.037.657.164" for candidate in payload["candidates"]))
        self.assertTrue(any(candidate["text"] == "900.123.456-7" for candidate in payload["candidates"]))
        self.assertTrue(any("table 1 row 1 cell 2" in context["location"] for candidate in payload["candidates"] for context in candidate["contexts"]))
        self.assertTrue(all(candidate["status"] == "pending" for candidate in payload["candidates"]))

    def test_endpoint_rejects_non_docx(self):
        app = FastAPI()
        app.include_router(template_builder_router, prefix="/api/v1")
        client = TestClient(app)

        response = client.post(
            "/api/v1/template-builder/reverse/single/analyze",
            files={"archivo": ("minuta.txt", b"texto", "text/plain")},
        )

        self.assertEqual(response.status_code, 422)


if __name__ == "__main__":
    unittest.main()
