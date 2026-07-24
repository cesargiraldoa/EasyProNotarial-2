#!/usr/bin/env python3
"""Extrae texto limpio (con tildes) de un .docx de minuta notarial.

Uso:
    python extract_docx.py "ruta/minuta.docx" > /tmp/minuta.txt

Requiere python-docx (`pip install python-docx`). Para .doc VIEJOS (Word 97)
python-docx no sirve: conviértelos antes a .docx (Word: Guardar como → .docx) o
pide el .docx a la notaría — el texto extraído de .doc con otras herramientas
suele perder acentos y no sirve para texto legal verbatim.
"""
from __future__ import annotations

import sys


def main() -> int:
    if len(sys.argv) < 2:
        print("Uso: python extract_docx.py <archivo.docx>", file=sys.stderr)
        return 2
    path = sys.argv[1]
    if not path.lower().endswith(".docx"):
        print(
            "[AVISO] El archivo no es .docx. Si es .doc viejo, conviértelo a .docx "
            "primero (los .doc pierden acentos al extraerlos).",
            file=sys.stderr,
        )
    try:
        import docx  # python-docx
    except ImportError:
        print("Falta python-docx. Instala: pip install python-docx", file=sys.stderr)
        return 1

    document = docx.Document(path)
    out: list[str] = []

    def emit_paragraph_text(text: str) -> None:
        text = text.rstrip()
        if text.strip():
            out.append(text)

    # Párrafos del cuerpo.
    for paragraph in document.paragraphs:
        emit_paragraph_text(paragraph.text)

    # Texto dentro de tablas (algunas minutas ponen la carátula/SNR en tablas).
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            line = " | ".join(c for c in cells if c)
            if line:
                out.append(line)

    sys.stdout.write("\n".join(out) + "\n")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
